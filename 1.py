#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
AI PowerShell Orchestrator with Google Search Integration
Интегрирует LM Studio, PowerShell команды и поиск Google

"""

# Подавляем предупреждения PyTorch для чистого запуска
import warnings
warnings.filterwarnings("ignore", message="expandable_segments not supported on this platform")
warnings.filterwarnings("ignore", message="There is an imbalance between your GPUs")

import requests
import subprocess
import json
import time
import sys as _sys
import os
import base64
import io
import re
import threading
import tempfile
import shutil
import queue
import asyncio
from pathlib import Path
import math
import pyautogui
import mss
import queue
import logging
import argparse
from typing import Dict, Any, List, Union, Optional, TYPE_CHECKING, Tuple
import urllib.parse
from PIL import Image
from io import BytesIO
from collections import defaultdict
import asyncio
import telegram
from telegram import Update, Bot
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes
from dotenv import load_dotenv
import concurrent.futures

# Импорты для работы с электронной почтой
import imaplib
import smtplib
import email
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from email.mime.base import MIMEBase
from email import encoders
from email.header import decode_header
from email.utils import parseaddr, formataddr
import email.utils
import mimetypes

# Глобальные переменные для фоновой загрузки
_background_loader = None
_initialization_lock = threading.Lock()

class BackgroundInitializer:
    """Класс для фоновой инициализации тяжелых компонентов"""
    
    def __init__(self):
        self.loaded_components = {}
        self.loading_tasks = {}
        self.executor = concurrent.futures.ThreadPoolExecutor(max_workers=4)
        self._chromadb_manager = None
        self._easyocr_reader = None
        self._is_loading = set()
        
    def start_loading(self, component_name, loader_func, *args, **kwargs):
        """Запускает фоновую загрузку компонента"""
        if component_name not in self._is_loading and component_name not in self.loaded_components:
            self._is_loading.add(component_name)
            future = self.executor.submit(self._safe_load, component_name, loader_func, *args, **kwargs)
            self.loading_tasks[component_name] = future
            return future
        return None
    
    def _safe_load(self, component_name, loader_func, *args, **kwargs):
        """Безопасная загрузка компонента с обработкой ошибок"""
        try:
            result = loader_func(*args, **kwargs)
            self.loaded_components[component_name] = result
            self._is_loading.discard(component_name)
            return result
        except Exception as e:
            print(f"Ошибка загрузки {component_name}: {e}")
            self._is_loading.discard(component_name)
            return None
    
    def get_component(self, component_name, timeout=30):
        """Получает компонент, ждет завершения загрузки если нужно"""
        if component_name in self.loaded_components:
            return self.loaded_components[component_name]
        
        if component_name in self.loading_tasks:
            try:
                result = self.loading_tasks[component_name].result(timeout=timeout)
                return result
            except concurrent.futures.TimeoutError:
                print(f"Таймаут загрузки {component_name}")
                return None
        
        return None
    
    def is_loaded(self, component_name):
        """Проверяет, загружен ли компонент"""
        return component_name in self.loaded_components
    
    def shutdown(self):
        """Завершает работу загрузчика"""
        self.executor.shutdown(wait=True)

def get_background_loader():
    """Получает глобальный экземпляр фонового загрузчика"""
    global _background_loader
    if _background_loader is None:
        _background_loader = BackgroundInitializer()
    return _background_loader

# Функции для фоновой загрузки тяжелых компонентов
def load_chromadb(embedding_model="all-MiniLM-L6-v2"):
    """Загружает ChromaDB"""
    try:
        print("Загружаем ChromaDB...")
        # Ленивый импорт
        import chromadb
        from sentence_transformers import SentenceTransformer
        
        client = chromadb.PersistentClient(path="./chroma_db")
        collection = client.get_or_create_collection(
            name="ai_memories",
            metadata={"hnsw:space": "cosine"}
        )
        
        # Используем переданную модель
        model = SentenceTransformer(embedding_model)
        return {'client': client, 'collection': collection, 'model': model}
    except Exception as e:
        print(f"Ошибка загрузки ChromaDB: {e}")
        return None

def load_easyocr():
    """Загружает EasyOCR"""
    try:
        print("Загружаем EasyOCR...")
        import easyocr  # type: ignore
        reader = easyocr.Reader(['ru', 'en'])
        return reader
    except Exception as e:
        print(f"Ошибка загрузки EasyOCR: {e}")
        return None

def load_torch():
    """Загружает PyTorch"""
    try:
        print("Загружаем PyTorch...")
        import torch
        return torch
    except Exception as e:
        print(f"Ошибка загрузки PyTorch: {e}")
        return None

# Система плагинов
# Plugin system
try:
    from plugins import PluginManager, PluginError
    PLUGINS_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Система плагинов недоступна: {e}")
    PLUGINS_AVAILABLE = False
    PluginManager = None
    PluginError = Exception

# Помощь статическим анализаторам: явные объявления для опциональных внешних символов
from typing import Any as _Any
chromadb: _Any = None
Settings: _Any = None
SentenceTransformer: _Any = None
torch: _Any = None
_imageio: _Any = None
_pygame: _Any = None

# Импорты для работы с документами
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

try:
    import pandas as pd
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    pd = None
    openpyxl = None

# Импорты для работы с PDF
try:
    import PyPDF2  # type: ignore
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    PyPDF2 = None

# Импорты для генерации файлов
try:
    from reportlab.pdfgen import canvas  # type: ignore
    from reportlab.lib.pagesizes import letter, A4  # type: ignore
    from reportlab.lib.styles import getSampleStyleSheet  # type: ignore
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer  # type: ignore
    from reportlab.lib.units import inch  # type: ignore
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    canvas = None
    _letter = None
    _A4 = None
    getSampleStyleSheet = None
    SimpleDocTemplate = None
    Paragraph = None
    Spacer = None
    _inch = None

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    markdown = None

# Импорты для OCR - теперь ленивые
OCR_AVAILABLE = True  # Будем проверять при первом использовании

# Импорты для ChromaDB и векторного поиска - теперь ленивые
CHROMADB_AVAILABLE = True  # Будем проверять при первом использовании

# Импорты для Torch - теперь ленивые
TORCH_AVAILABLE = True  # Будем проверять при первом использовании

# Проверки доступности опциональных модулей
try:
    import torch as _torch
    TORCH_AVAILABLE = True
except Exception:
    TORCH_AVAILABLE = False

try:
    import diffusers as _diffusers
    DIFFUSERS_AVAILABLE = True
except Exception:
    DIFFUSERS_AVAILABLE = False

try:
    import imageio as _imageio
    IMAGEIO_AVAILABLE = True
except Exception:
    IMAGEIO_AVAILABLE = False

try:
    import pygame as _pygame
    PYGAME_AVAILABLE = True
except Exception:
    PYGAME_AVAILABLE = False

# Загружаем переменные окружения из .env файла
# override=True - перезаписываем существующие переменные
load_dotenv(override=True)

# Determine if running in web mode (show verbose console logs)
IS_WEB = any(arg == '--web' for arg in _sys.argv)

# Настройка логирования: всегда пишем подробный файл, но в консоль показываем INFO только в --web
log_file = "ai_orchestrator.log"
root_logger = logging.getLogger()
root_logger.setLevel(logging.DEBUG)

# File handler: keep full INFO logs for later inspection
file_handler = logging.FileHandler(log_file, mode='w', encoding='utf-8')
file_handler.setLevel(logging.INFO)
file_formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
file_handler.setFormatter(file_formatter)
root_logger.addHandler(file_handler)

# Console handler: verbose only for --web, otherwise warnings+ only
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.INFO)
console_handler.setFormatter(file_formatter)
root_logger.addHandler(console_handler)

logger = logging.getLogger(__name__)

# Disable chromadb telemetry / noisy chromadb logs entirely
try:
    chroma_logger = logging.getLogger('chromadb')
    chroma_logger.disabled = True
except Exception:
    pass

# Filter for telemetry messages: only applied to console handler so file logs keep full info
class TelemetryFilter(logging.Filter):
    def __init__(self, patterns=None):
        super().__init__()
        self.patterns = patterns or ["Failed to send telemetry event", "telemetry", "capture() takes"]

    def filter(self, record: logging.LogRecord) -> bool:
        try:
            msg = record.getMessage()
        except Exception:
            msg = ''
        for p in self.patterns:
            if p in msg:
                return False
        return True

# Apply telemetry filter to console only
try:
    console_handler.addFilter(TelemetryFilter())
except Exception:
    pass


class _StderrFilterWriter:
    """A file-like wrapper for stderr that filters lines containing given patterns."""
    def __init__(self, orig, patterns):
        self.orig = orig
        self.patterns = patterns

    def write(self, data):
        if not data:
            return
        for p in self.patterns:
            if p in data:
                return
        try:
            self.orig.write(data)
        except Exception:
            pass

    def flush(self):
        try:
            self.orig.flush()
        except Exception:
            pass
if TYPE_CHECKING:
    try:
        import chromadb  # type: ignore
        from chromadb.config import Settings  # type: ignore
    except Exception:
        pass
    # Stubs for optional heavy libraries so Pylance doesn't warn
    try:
        import imageio as _imageio  # type: ignore
    except Exception:
        pass
    try:
        import pygame as _pygame  # type: ignore
    except Exception:
        pass
    try:
        import diffusers as _diffusers  # type: ignore
    except Exception:
        pass
    try:
        import torch as _torch  # type: ignore
    except Exception:
        pass


from contextlib import contextmanager

@contextmanager
def suppress_stderr_patterns(patterns):
    orig = _sys.stderr
    try:
        _sys.stderr = _StderrFilterWriter(orig, patterns)
        yield
    finally:
        _sys.stderr = orig

### НОВОЕ: Класс для работы с векторным хранилищем ChromaDB ###
class ChromaDBManager:
    """
    Менеджер векторного хранилища ChromaDB для преодоления ограничений контекста
    и обучения на предпочтениях пользователя
    """
    
    def __init__(self, db_path: str = "chroma_db", embedding_model: str = "all-MiniLM-L6-v2", use_gpu: bool = True):
        """
        Инициализация ChromaDB менеджера
        
        Args:
            db_path: Путь к базе данных ChromaDB
            embedding_model: Модель для создания эмбеддингов (784 размерности)
            use_gpu: Использовать GPU для создания эмбеддингов (если доступен)
        """
        self.db_path = db_path
        self.embedding_model = embedding_model
        self.use_gpu = use_gpu
        self.client = None
        self.collection = None
        self.embedding_model_obj = None
        self.initialized = False
        
        # Создаем папку для базы данных если её нет
        os.makedirs(db_path, exist_ok=True)
        
        # Запускаем фоновую инициализацию ChromaDB
        self._start_background_initialization()
    
    def _start_background_initialization(self):
        """Запускает фоновую инициализацию ChromaDB"""
        loader = get_background_loader()
        loader.start_loading('chromadb', load_chromadb, self.embedding_model)
        
    def _ensure_initialized(self, timeout=30):
        """Обеспечивает инициализацию компонентов"""
        if self.initialized:
            return True
            
        loader = get_background_loader()
        chromadb_data = loader.get_component('chromadb', timeout=timeout)
        
        if chromadb_data:
            self.client = chromadb_data['client']
            self.collection = chromadb_data['collection']
            self.embedding_model_obj = chromadb_data['model']
            self.initialized = True
            return True
        else:
            # Fallback к синхронной инициализации
            return self._initialize_chromadb_sync()
    
    def _initialize_chromadb_sync(self):
        """Синхронная инициализация ChromaDB как fallback"""
        try:
            print("Синхронная инициализация ChromaDB...")
            import chromadb
            from chromadb.config import Settings
            from sentence_transformers import SentenceTransformer
            
            # Инициализируем клиент ChromaDB
            self.client = chromadb.PersistentClient(
                path=self.db_path,
                settings=Settings(
                    anonymized_telemetry=False,
                    allow_reset=True
                )
            )
            
            # Создаем или получаем коллекцию
            self.collection = self.client.get_or_create_collection(
                name="conversation_memory",
                metadata={"description": "Векторное хранилище диалогов и предпочтений пользователя"}
            )
            
            # Инициализируем модель эмбеддингов
            self.embedding_model_obj = SentenceTransformer(self.embedding_model)
            self.initialized = True
            return True
            
        except Exception as e:
            print(f"Ошибка синхронной инициализации ChromaDB: {e}")
            return False
    
    def get_gpu_info(self) -> Dict[str, Any]:
        """
        Получает информацию о GPU для ChromaDB
        
        Returns:
            Словарь с информацией о GPU
        """
        gpu_info = {
            "gpu_available": False,
            "gpu_name": None,
            "gpu_memory": None,
            "device_used": "cpu"
        }
        
        try:
            if torch.cuda.is_available():
                gpu_info["gpu_available"] = True
                gpu_info["gpu_name"] = torch.cuda.get_device_name(0)
                gpu_info["gpu_memory"] = torch.cuda.get_device_properties(0).total_memory / 1024**3  # GB
                gpu_info["device_used"] = "cuda" if self.use_gpu else "cpu"
                
                logger.info(f"🎮 GPU доступен: {gpu_info['gpu_name']}")
                logger.info(f"💾 GPU память: {gpu_info['gpu_memory']:.1f} GB")
            else:
                logger.info("💻 GPU недоступен, используется CPU")
                
        except Exception as e:
            logger.warning(f"⚠️ Ошибка получения информации о GPU: {e}")
            
        return gpu_info
    
    def add_conversation_memory(self, user_message: str, ai_response: str,
                               context: str = "", metadata: Optional[Dict[str, Any]] = None, 
                               force_add: bool = False) -> bool:
        """
        Добавляет диалог в векторное хранилище
        
        Args:
            user_message: Сообщение пользователя
            ai_response: Ответ ИИ
            context: Дополнительный контекст
            metadata: Дополнительные метаданные
            force_add: Принудительно добавить без проверки дубликатов
            
        Returns:
            True если успешно добавлено, False при ошибке или дубликате
        """
        if not self._ensure_initialized():
            return False
        
        try:
            # Проверяем дубликаты, если не принудительное добавление
            if not force_add:
                logger.debug(f"🔍 Проверяем дубликаты для сообщения: '{user_message[:50]}...'")
                similar_conversations = self.search_similar_conversations(
                    user_message, n_results=1, similarity_threshold=0.7
                )
                
                if similar_conversations and len(similar_conversations) > 0:
                    similarity = similar_conversations[0].get('similarity', 0)
                    logger.debug(f"🔍 Найден похожий диалог с similarity={similarity:.3f}")
                    
                    # Если similarity больше 0.7 (70%), считаем дубликатом
                    if similarity > 0.7:
                        logger.info(f"⚠️ Найден дубликат с similarity={similarity:.3f}, пропускаем добавление")
                        return False
                else:
                    logger.debug("🔍 Похожие диалоги не найдены, можно добавлять")
            else:
                logger.debug("🔄 Принудительное добавление, пропускаем проверку дубликатов")
            # Создаем уникальный ID для записи
            import uuid
            timestamp = int(time.time())
            unique_suffix = str(uuid.uuid4())[:8]  # Первые 8 символов UUID
            record_id = f"conv_{timestamp}_{unique_suffix}"
            
            # Объединяем текст для создания эмбеддинга
            combined_text = f"User: {user_message}\nAI: {ai_response}"
            if context:
                combined_text += f"\nContext: {context}"
            
            # Создаем эмбеддинг (если модель доступна)
            if not self.initialized or self.embedding_model_obj is None:
                logger.warning("⚠️ Эмбеддинговая модель не инициализирована, пропускаю добавление в ChromaDB")
                return False
            embedding = self.embedding_model_obj.encode(combined_text).tolist()
            
            # Подготавливаем метаданные
            record_metadata = {
                "timestamp": timestamp,
                "user_message": user_message,
                "ai_response": ai_response,
                "context": context,
                "type": "conversation"
            }
            
            if metadata:
                record_metadata.update(metadata)
            
            # Добавляем в коллекцию
            if self.collection is None:
                logger.warning("⚠️ Коллекция ChromaDB не инициализирована при попытке add")
                return False
            
            # Проверяем, существует ли уже такой ID (маловероятно с UUID, но проверим)
            try:
                existing = self.collection.get(ids=[record_id])
                if existing and existing.get('ids') and len(existing['ids']) > 0:
                    # ID уже существует, добавим еще один уникальный суффикс
                    record_id = f"conv_{timestamp}_{unique_suffix}_{hash(ai_response) % 1000}"
                    logger.info(f"🔄 ID уже существует, используем новый: {record_id}")
            except Exception:
                pass  # Нормально, ID не существует
            
            self.collection.add(
                embeddings=[embedding],
                documents=[combined_text],
                metadatas=[record_metadata],
                ids=[record_id]
            )
            
            # Дополнительная проверка - считаем количество записей
            try:
                total_count = self.collection.count()
                logger.info(f"💾 Добавлена запись в ChromaDB: {record_id} (всего записей: {total_count})")
            except Exception:
                logger.info(f"💾 Добавлена запись в ChromaDB: {record_id}")
            
            return True
            
        except Exception as e:
            logger.error(f"❌ Ошибка добавления в ChromaDB: {e}")
            return False
    
    def add_user_preference(self, preference_text: str, category: str = "general",
                           metadata: Optional[Dict[str, Any]] = None) -> bool:
        """
        Добавляет предпочтение пользователя в векторное хранилище
        
        Args:
            preference_text: Текст предпочтения
            category: Категория предпочтения
            metadata: Дополнительные метаданные
            
        Returns:
            True если успешно добавлено, False при ошибке
        """
        if not self._ensure_initialized():
            return False
        
        try:
            # Создаем уникальный ID
            timestamp = int(time.time())
            record_id = f"pref_{timestamp}_{hash(preference_text) % 10000}"
            
            # Создаем эмбеддинг (если модель доступна)
            if not self.initialized or self.embedding_model_obj is None:
                logger.warning("⚠️ Эмбеддинговая модель не инициализирована, пропускаю добавление предпочтения")
                return False
            embedding = self.embedding_model_obj.encode(preference_text).tolist()
            
            # Подготавливаем метаданные
            record_metadata = {
                "timestamp": timestamp,
                "preference_text": preference_text,
                "category": category,
                "type": "preference"
            }
            
            if metadata:
                record_metadata.update(metadata)
            
            # Добавляем в коллекцию
                if self.collection is None:
                    logger.warning("⚠️ Коллекция ChromaDB не инициализирована при попытке add preference")
                    return False
                self.collection.add(
                embeddings=[embedding],
                documents=[preference_text],
                metadatas=[record_metadata],
                ids=[record_id]
            )
            
            logger.info(f"💾 Добавлено предпочтение в ChromaDB: {record_id}")
            return True
            
        except Exception as e:
            logger.error(f"❌ Ошибка добавления предпочтения в ChromaDB: {e}")
            return False
    
    def search_similar_conversations(self, query: str, n_results: int = 5,
                                   similarity_threshold: Optional[float] = None) -> List[Dict[str, Any]]:
        """
        Ищет похожие диалоги в векторном хранилище
        
        Args:
            query: Поисковый запрос
            n_results: Количество результатов
            similarity_threshold: Порог схожести (автоматический если None)
            
        Returns:
            Список найденных диалогов с метаданными
        """
        if not self._ensure_initialized():
            return []
        
        try:
            # Создаем эмбеддинг для запроса
            if not self.initialized or self.embedding_model_obj is None:
                logger.warning("⚠️ Эмбеддинговая модель не инициализирована, поиск невозможен")
                return []
            
            logger.info(f"🔍 Ищем похожие диалоги для запроса: '{query}'")
            query_embedding = self.embedding_model_obj.encode(query).tolist()
            
            # Ищем похожие записи (если коллекция доступна)
            if self.collection is None:
                logger.warning("⚠️ Коллекция ChromaDB не доступна, поиск невозможен")
                return []
            
            # Проверяем общее количество записей перед поиском
            try:
                total_count = self.collection.count()
                logger.info(f"🔍 Всего записей в ChromaDB: {total_count}")
            except Exception as e:
                logger.warning(f"⚠️ Не удалось получить количество записей: {e}")
            
            # Увеличиваем количество результатов для лучшего поиска
            search_results = max(n_results * 3, 15)
            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=search_results,
                where={"type": "conversation"}  # type: ignore[arg-type]
            )
            
            # Анализируем результаты для определения адаптивного порога
            filtered_results = []
            found_count = 0
            
            # Защищаемся от отсутствия ключей или пустых результатов
            if isinstance(results, dict) and results:
                distances = results.get('distances')
                ids = results.get('ids')
                documents = results.get('documents')
                metadatas = results.get('metadatas')

                if distances and isinstance(distances, list) and distances and distances[0]:
                    logger.info(f"📊 Обработка {len(distances[0])} результатов поиска")
                    
                    # Вычисляем адаптивный порог если не задан
                    if similarity_threshold is None:
                        similarities = [1 - d for d in distances[0]]
                        if similarities:
                            max_sim = max(similarities)
                            avg_sim = sum(similarities) / len(similarities)
                            
                            # Адаптивный порог: берем результаты выше среднего, но не выше 0.5
                            if max_sim > 0.1:
                                adaptive_threshold = min(avg_sim + 0.1, 0.3, max_sim - 0.05)
                            else:
                                adaptive_threshold = -0.2  # Очень низкий порог для слабых совпадений
                            
                            logger.info(f"🎯 Адаптивный порог схожести: {adaptive_threshold:.3f} (макс: {max_sim:.3f}, средн: {avg_sim:.3f})")
                        else:
                            adaptive_threshold = 0.1
                    else:
                        adaptive_threshold = similarity_threshold
                    
                    for i, distance in enumerate(distances[0]):
                        # ChromaDB возвращает расстояния, конвертируем в схожесть
                        similarity = 1 - distance
                        
                        # Логируем для отладки
                        if i < 3:  # Показываем первые 3 результата
                            logger.info(f"   Результат {i+1}: схожесть={similarity:.3f}, расстояние={distance:.3f}")
                        
                        if similarity >= adaptive_threshold:
                            # Проверяем, что остальные структуры содержат нужные элементы
                            doc = None
                            meta = None
                            idv = None
                            try:
                                idv = ids[0][i] if ids and ids[0] and len(ids[0]) > i else None
                                doc = documents[0][i] if documents and documents[0] and len(documents[0]) > i else None
                                meta = metadatas[0][i] if metadatas and metadatas[0] and len(metadatas[0]) > i else None
                            except Exception as e:
                                logger.warning(f"⚠️ Ошибка извлечения данных для результата {i}: {e}")
                                continue

                            result = {
                                'id': idv,
                                'document': doc,
                                'metadata': meta,
                                'similarity': similarity,
                                'distance': distance
                            }
                            filtered_results.append(result)
                            found_count += 1
                            
                            # Ограничиваем количество результатов
                            if found_count >= n_results:
                                break
                    
                    # Если ничего не найдено даже с адаптивным порогом, берем лучшие результаты
                    if not filtered_results and distances[0]:
                        logger.info(f"⚠️ Ничего не найдено с порогом {adaptive_threshold:.3f}, берем {min(3, len(distances[0]))} лучших результата")
                        best_results = min(3, len(distances[0]))
                        for i in range(best_results):
                            distance = distances[0][i]
                            similarity = 1 - distance
                            
                            try:
                                idv = ids[0][i] if ids and ids[0] and len(ids[0]) > i else None
                                doc = documents[0][i] if documents and documents[0] and len(documents[0]) > i else None
                                meta = metadatas[0][i] if metadatas and metadatas[0] and len(metadatas[0]) > i else None
                                
                                result = {
                                    'id': idv,
                                    'document': doc,
                                    'metadata': meta,
                                    'similarity': similarity,
                                    'distance': distance
                                }
                                filtered_results.append(result)
                            except Exception as e:
                                logger.warning(f"⚠️ Ошибка извлечения лучшего результата {i}: {e}")
                else:
                    logger.warning("⚠️ Пустые результаты поиска в ChromaDB")
            else:
                logger.warning("⚠️ Некорректный формат результатов поиска")
            
            logger.info(f"✅ Найдено {len(filtered_results)} похожих диалогов")
            return filtered_results
            
        except Exception as e:
            logger.error(f"❌ Ошибка поиска в ChromaDB: {e}")
            return []
    
    def search_user_preferences(self, query: str, category: Optional[str] = None,
                               n_results: int = 3) -> List[Dict[str, Any]]:
        """
        Ищет предпочтения пользователя
        
        Args:
            query: Поисковый запрос
            category: Категория предпочтений (опционально)
            n_results: Количество результатов
            
        Returns:
            Список найденных предпочтений
        """
        if not self._ensure_initialized():
            return []
        
        try:
            # Создаем эмбеддинг для запроса
            if not self.initialized or self.embedding_model_obj is None:
                logger.warning("⚠️ Эмбеддинговая модель не инициализирована, поиск предпочтений невозможен")
                return []
            query_embedding = self.embedding_model_obj.encode(query).tolist()
            
            # Формируем условия поиска
            where_condition = {"type": "preference"}
            if category:
                where_condition["category"] = category
            
            # Ищем похожие записи
            if self.collection is None:
                logger.warning("⚠️ Коллекция ChromaDB не доступна, поиск предпочтений невозможен")
                return []
            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=n_results,
                where=where_condition  # type: ignore[arg-type]
            )
            
            # Формируем результат
            preferences = []
            if isinstance(results, dict) and results:
                docs = results.get('documents')
                distances = results.get('distances')
                ids = results.get('ids')
                metadatas = results.get('metadatas')

                if docs and isinstance(docs, list) and docs and docs[0]:
                    for i, document in enumerate(docs[0]):
                        try:
                            dist = distances[0][i] if distances and distances[0] and len(distances[0]) > i else None
                            sim = 1 - dist if dist is not None else 0.0
                            pref = {
                                'id': ids[0][i] if ids and ids[0] and len(ids[0]) > i else None,
                                'preference_text': document,
                                'metadata': metadatas[0][i] if metadatas and metadatas[0] and len(metadatas[0]) > i else {},
                                'similarity': sim
                            }
                            preferences.append(pref)
                        except Exception:
                            continue
            
            logger.info(f"🔍 Найдено {len(preferences)} предпочтений пользователя")
            return preferences
            
        except Exception as e:
            logger.error(f"❌ Ошибка поиска предпочтений в ChromaDB: {e}")
            return []
    
    def get_conversation_context(self, query: str, max_context_length: int = 2000) -> str:
        """
        Получает релевантный контекст из предыдущих диалогов
        
        Args:
            query: Текущий запрос пользователя
            max_context_length: Максимальная длина контекста
            
        Returns:
            Строка с релевантным контекстом
        """
        if not self.initialized:
            return ""
        
        try:
            # Ищем похожие диалоги с адаптивным порогом
            similar_conversations = self.search_similar_conversations(
                query, n_results=5  # Автоматический порог
            )
            
            if not similar_conversations:
                logger.info("📚 Релевантный контекст не найден")
                return ""
            
            logger.info(f"📚 Найдено {len(similar_conversations)} релевантных диалогов для контекста")
            
            # Формируем контекст
            context_parts = []
            current_length = 0
            
            for i, conv in enumerate(similar_conversations):
                # Извлекаем пользовательское сообщение из метаданных если доступно
                user_msg = ""
                ai_resp = ""
                
                if conv.get('metadata') and isinstance(conv['metadata'], dict):
                    user_msg = conv['metadata'].get('user_message', '')
                    ai_resp = conv['metadata'].get('ai_response', '')
                
                # Если метаданные недоступны, пытаемся извлечь из документа
                if not user_msg and conv.get('document'):
                    doc = conv['document']
                    if 'User:' in doc and 'AI:' in doc:
                        parts = doc.split('AI:', 1)
                        if len(parts) >= 2:
                            user_part = parts[0].replace('User:', '').strip()
                            user_msg = user_part
                
                if user_msg:
                    conv_text = f"Похожий запрос #{i+1} (схожесть: {conv['similarity']:.3f}):\n"
                    conv_text += f"Пользователь: {user_msg[:200]}{'...' if len(user_msg) > 200 else ''}\n"
                    
                    if ai_resp and len(ai_resp) < 300:  # Включаем короткие ответы
                        clean_ai_resp = ai_resp.replace('<think>', '').replace('</think>', '')
                        if len(clean_ai_resp) < 200:
                            conv_text += f"Ответ: {clean_ai_resp[:150]}{'...' if len(clean_ai_resp) > 150 else ''}\n"
                    
                    conv_text += "\n"
                    
                    if current_length + len(conv_text) <= max_context_length:
                        context_parts.append(conv_text)
                        current_length += len(conv_text)
                    else:
                        break
            
            context = "".join(context_parts)
            if context:
                logger.info(f"✅ Сформирован контекст длиной {len(context)} символов из {len(context_parts)} диалогов")
            return context
            
        except Exception as e:
            logger.error(f"❌ Ошибка получения контекста из ChromaDB: {e}")
            return ""
    
    def get_user_preferences_summary(self, query: Optional[str] = None) -> str:
        """
        Получает краткое резюме предпочтений пользователя
        
        Args:
            query: Контекстный запрос (опционально)
            
        Returns:
            Строка с резюме предпочтений
        """
        if not self.initialized:
            return ""
        
        try:
            # Ищем релевантные предпочтения
            if query:
                preferences = self.search_user_preferences(query, n_results=5)
            else:
                # Получаем все предпочтения
                results = self.collection.get(where={"type": "preference"})  # type: ignore[arg-type]
                preferences = []
                if isinstance(results, dict) and results:
                    docs = results.get('documents')
                    metadatas = results.get('metadatas')
                    if docs:
                        for i, doc in enumerate(docs):
                            pref_meta = metadatas[i] if metadatas and len(metadatas) > i else {}
                            preference = {
                                'preference_text': doc,
                                'metadata': pref_meta,
                                'similarity': 1.0  # Для общих предпочтений
                            }
                            preferences.append(preference)
            
            if not preferences:
                return ""
            
            # Формируем резюме
            summary_parts = ["Предпочтения пользователя:"]
            
            for pref in preferences[:3]:  # Ограничиваем 3 предпочтениями
                category = pref['metadata'].get('category', 'general')
                summary_parts.append(f"- {category}: {pref['preference_text'][:100]}...")
            
            summary = "\n".join(summary_parts)
            logger.info(f"📋 Сформировано резюме предпочтений длиной {len(summary)} символов")
            return summary
            
        except Exception as e:
            logger.error(f"❌ Ошибка получения резюме предпочтений: {e}")
            return ""
    
    def cleanup_old_records(self, days_to_keep: int = 30) -> int:
        """
        Удаляет старые записи из базы данных
        
        Args:
            days_to_keep: Количество дней для хранения записей
            
        Returns:
            Количество удаленных записей
        """
        if not self.initialized:
            return 0
        
        try:
            cutoff_timestamp = int(time.time()) - (days_to_keep * 24 * 60 * 60)
            
            # Получаем все записи
            if self.collection is None:
                logger.warning("⚠️ Коллекция ChromaDB не доступна, очистка записей пропущена")
                return 0
            results = self.collection.get()

            # Защита от отсутствия ключей/пустых результатов
            if not isinstance(results, dict) or not results:
                return 0

            ids = results.get('ids')
            metadatas = results.get('metadatas')
            if not ids:
                return 0

            # Находим записи для удаления
            ids_to_delete = []
            if metadatas:
                for i, metadata in enumerate(metadatas):
                    timestamp = metadata.get('timestamp', 0) if isinstance(metadata, dict) else 0
                    # Преобразуем timestamp в число если это строка
                    try:
                        timestamp_num = float(timestamp) if timestamp else 0
                    except (ValueError, TypeError):
                        timestamp_num = 0
                    
                    if timestamp_num < cutoff_timestamp:
                        # Защищаем доступ к ids
                        if ids and len(ids) > i:
                            ids_to_delete.append(ids[i])
            
            # Удаляем старые записи
            if ids_to_delete:
                if self.collection is None:
                    logger.warning("⚠️ Коллекция ChromaDB не доступна, удаление невозможно")
                    return 0
                self.collection.delete(ids=ids_to_delete)
                logger.info(f"🧹 Удалено {len(ids_to_delete)} старых записей из ChromaDB")
                return len(ids_to_delete)
            
            return 0
            
        except Exception as e:
            logger.error(f"❌ Ошибка очистки ChromaDB: {e}")
            return 0
    
    def get_database_stats(self) -> Dict[str, Any]:
        """
        Получает статистику базы данных
        
        Returns:
            Словарь со статистикой
        """
        if not self.initialized:
            return {"error": "ChromaDB не инициализирован"}
        
        try:
            if self.collection is None:
                logger.warning("⚠️ Коллекция ChromaDB не доступна, статистика недоступна")
                return {"error": "ChromaDB не инициализирован"}

            total_count = self.collection.count()
            # Подсчитываем по типам
            conversations = self.collection.get(where={"type": "conversation"})  # type: ignore[arg-type]
            preferences = self.collection.get(where={"type": "preference"})  # type: ignore[arg-type]

            conv_ids = conversations.get('ids') if isinstance(conversations, dict) else None
            pref_ids = preferences.get('ids') if isinstance(preferences, dict) else None

            stats = {
                "total_records": total_count,
                "conversations": len(conv_ids) if conv_ids else 0,
                "preferences": len(pref_ids) if pref_ids else 0,
                "database_path": self.db_path,
                "embedding_model": self.embedding_model
            }
            
            return stats
            
        except Exception as e:
            logger.error(f"❌ Ошибка получения статистики ChromaDB: {e}")
            return {"error": str(e)}

### НОВОЕ: Функция для сжатия изображений ###
def image_to_base64_balanced(image_path: str, max_size=(500, 500), palette_colors=12) -> str:
    """
    Конвертирует изображение в PNG base64 без ч/б и quantize, только ресайз (если нужно).
    """
    try:
        with Image.open(image_path) as img:
            img = img.convert("RGB")
            img.thumbnail(max_size, Image.Resampling.LANCZOS)
            buf = BytesIO()
            img.save(buf, format="PNG", optimize=True)
            return base64.b64encode(buf.getvalue()).decode("ascii")
    except Exception as e:
        logger.error(f"Ошибка кодирования (balanced) {image_path}: {e}")
        return ""

class ModelManager:
    """
    Класс для управления Stable Diffusion моделями и LoRA
    """
    
    def __init__(self, base_dir: str | None = None):
        if base_dir is None:
            base_dir = os.path.dirname(os.path.abspath(__file__))
        
        self.base_dir = base_dir
        self.stable_diff_dir = os.path.join(base_dir, "stable_diff")
        self.checkpoints_dir = os.path.join(self.stable_diff_dir, "checkpoints")
        self.lora_dir = os.path.join(self.stable_diff_dir, "lora")
        self.lora_config_path = os.path.join(self.lora_dir, "lora_config.json")
        
        # Кэш для конфигурации LoRA
        self._lora_config_cache = {}
        self._lora_config_last_modified = 0
        
        # Создаем папки если их нет
        self._ensure_directories()
        
        # Инициализируем конфигурацию LoRA
        self._init_lora_config()
    
    def _ensure_directories(self):
        """Создает необходимые папки если их нет"""
        os.makedirs(self.checkpoints_dir, exist_ok=True)
        os.makedirs(os.path.join(self.lora_dir, "sd"), exist_ok=True)
        os.makedirs(os.path.join(self.lora_dir, "sdxl"), exist_ok=True)
    
    def _init_lora_config(self):
        """Инициализирует конфигурацию LoRA"""
        if not os.path.exists(self.lora_config_path):
            self._generate_lora_config()
        else:
            self._scan_and_update_lora_config()
    
    def _scan_lora_files(self) -> Dict[str, List[str]]:
        """Сканирует папки LoRA и возвращает найденные файлы"""
        lora_files = {"sd": [], "sdxl": []}
        
        for model_type in ["sd", "sdxl"]:
            lora_type_dir = os.path.join(self.lora_dir, model_type)
            if os.path.exists(lora_type_dir):
                for file in os.listdir(lora_type_dir):
                    if file.lower().endswith(('.safetensors', '.ckpt', '.pt')):
                        lora_files[model_type].append(file)
        
        return lora_files
    
    def _generate_lora_config(self):
        """Генерирует базовую конфигурацию LoRA"""
        lora_files = self._scan_lora_files()
        config = {"loras": {}}
        
        for model_type, files in lora_files.items():
            for filename in files:
                lora_name = os.path.splitext(filename)[0]
                config["loras"][f"{model_type}_{lora_name}"] = {
                    "filename": filename,
                    "model_type": model_type,
                    "enabled": True,
                    "strength": 1.0,
                    "triggers": [],
                    "description": f"Auto-generated config for {filename}"
                }
        
        with open(self.lora_config_path, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
        
        logger.info(f"✅ Создан конфигурационный файл LoRA: {len(config['loras'])} файлов")
    
    def _scan_and_update_lora_config(self):
        """Сканирует LoRA файлы и обновляет конфигурацию новыми с анализом метаданных"""
        lora_files = self._scan_lora_files()
        
        try:
            with open(self.lora_config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
        except:
            config = {"loras": {}}
        
        if "loras" not in config:
            config["loras"] = {}
        
        # Добавляем новые LoRA файлы
        updated = False
        for model_type, files in lora_files.items():
            for filename in files:
                lora_name = os.path.splitext(filename)[0]
                lora_key = f"{model_type}_{lora_name}"
                
                if lora_key not in config["loras"]:
                    # Анализируем метаданные LoRA
                    lora_path = os.path.join(self.lora_dir, model_type, filename)
                    metadata = self.analyze_lora_metadata(lora_path)
                    
                    # Определяем тип модели из метаданных или используем папку
                    detected_model_type = metadata.get("model_type", model_type)
                    if detected_model_type != "unknown" and detected_model_type != model_type:
                        logger.warning(f"⚠️ LoRA {filename} в папке {model_type}/, но метаданные указывают на {detected_model_type}")
                        # Используем тип из метаданных как более точный
                        actual_model_type = detected_model_type
                        lora_key = f"{actual_model_type}_{lora_name}"
                    else:
                        actual_model_type = model_type
                    
                    # Создаем конфигурацию с метаданными
                    config["loras"][lora_key] = {
                        "filename": filename,
                        "model_type": actual_model_type,
                        "enabled": True,
                        "strength": metadata.get("preferred_weight", 1.0),
                        "triggers": metadata.get("triggers", [])[:3],  # Берем топ-3 триггера
                        "description": metadata.get("description", f"Auto-detected: {metadata.get('base_model', 'Unknown')} LoRA"),
                        "base_model": metadata.get("base_model", "Unknown"),
                        "resolution": metadata.get("resolution", "Unknown"),
                        "author": metadata.get("author", ""),
                        "metadata_analyzed": True
                    }
                    updated = True
                    
                    logger.info(f"📋 Создана конфигурация для {filename}")
                    logger.info(f"   🎯 Тип: {actual_model_type} ({metadata.get('base_model', 'Unknown')})")
                    if metadata.get("triggers"):
                        logger.info(f"   🔤 Триггеры: {', '.join(metadata['triggers'][:3])}")
        
        if updated:
            with open(self.lora_config_path, 'w', encoding='utf-8') as f:
                json.dump(config, f, ensure_ascii=False, indent=2)
            logger.info(f"✅ Обновлен конфигурационный файл LoRA с анализом метаданных")
    
    def get_lora_config(self, force_reload: bool = False) -> Dict:
        """Получает конфигурацию LoRA с кэшированием"""
        try:
            # Проверяем время модификации файла
            if os.path.exists(self.lora_config_path):
                mtime = os.path.getmtime(self.lora_config_path)
                
                # Если файл изменился или принудительная перезагрузка
                if force_reload or mtime > self._lora_config_last_modified:
                    with open(self.lora_config_path, 'r', encoding='utf-8') as f:
                        self._lora_config_cache = json.load(f)
                    self._lora_config_last_modified = mtime
                    logger.info("🔄 Перезагружена конфигурация LoRA")
                
                return self._lora_config_cache
            else:
                return {"loras": {}}
        except Exception as e:
            logger.error(f"❌ Ошибка загрузки конфигурации LoRA: {e}")
            return {"loras": {}}
    
    def analyze_lora_metadata(self, lora_path: str) -> Dict[str, Any]:
        """
        Анализирует метаданные LoRA файла для определения совместимости
        
        Args:
            lora_path: Путь к LoRA файлу
            
        Returns:
            Словарь с метаданными LoRA
        """
        try:
            from safetensors import safe_open
            
            # Результат анализа
            metadata = {
                "model_type": "unknown",
                "base_model": "unknown", 
                "resolution": "unknown",
                "triggers": [],
                "description": "",
                "author": "",
                "version": "",
                "activation_text": "",
                "preferred_weight": 1.0
            }
            
            # Анализируем расширение файла
            file_ext = os.path.splitext(lora_path)[1].lower()
            
            if file_ext == ".safetensors":
                # Читаем метаданные из safetensors
                with safe_open(lora_path, framework="pt") as f:
                    metadata_raw = f.metadata()
                    
                    if metadata_raw:
                        # Извлекаем информацию о базовой модели
                        if "ss_base_model_version" in metadata_raw:
                            base_version = metadata_raw["ss_base_model_version"]
                            if "xl" in base_version.lower():
                                metadata["model_type"] = "sdxl"
                                metadata["base_model"] = "SDXL"
                            else:
                                metadata["model_type"] = "sd"
                                metadata["base_model"] = "SD 1.5"
                        
                        # Разрешение обучения
                        if "ss_resolution" in metadata_raw:
                            metadata["resolution"] = metadata_raw["ss_resolution"]
                        elif "ss_bucket_info" in metadata_raw:
                            try:
                                bucket_info = json.loads(metadata_raw["ss_bucket_info"])
                                if "buckets" in bucket_info:
                                    resolutions = list(bucket_info["buckets"].keys())
                                    if resolutions:
                                        metadata["resolution"] = resolutions[0]
                            except:
                                pass
                        
                        # Извлекаем теги и триггеры
                        if "ss_tag_frequency" in metadata_raw:
                            try:
                                tag_freq = json.loads(metadata_raw["ss_tag_frequency"])
                                # Получаем самые частые теги как потенциальные триггеры
                                all_tags = {}
                                for dataset_tags in tag_freq.values():
                                    all_tags.update(dataset_tags)
                                
                                # Сортируем по частоте и берем топ-5
                                sorted_tags = sorted(all_tags.items(), key=lambda x: x[1], reverse=True)
                                metadata["triggers"] = [tag for tag, _ in sorted_tags[:5]]
                            except:
                                pass
                        
                        # Другие поля метаданных
                        metadata_mapping = {
                            "ss_dataset_dirs": "description",
                            "modelspec.architecture": "architecture",
                            "modelspec.implementation": "implementation",
                            "modelspec.title": "title"
                        }
                        
                        for key, target in metadata_mapping.items():
                            if key in metadata_raw:
                                metadata[target] = metadata_raw[key]
                        
                        # Пытаемся извлечь автора и описание из названия файла
                        filename = os.path.basename(lora_path)
                        if "_" in filename or "-" in filename:
                            parts = filename.replace("_", " ").replace("-", " ").split()
                            metadata["author"] = parts[0] if parts else ""
                        
                        logger.info(f"🔍 Проанализированы метаданные LoRA: {filename}")
                        logger.info(f"   📋 Базовая модель: {metadata['base_model']}")
                        logger.info(f"   📐 Разрешение: {metadata['resolution']}")
                        if metadata["triggers"]:
                            logger.info(f"   🎯 Найденные триггеры: {', '.join(metadata['triggers'][:3])}")
            
            elif file_ext in [".ckpt", ".pt"]:
                # Для старых форматов используем эвристический анализ
                filename = os.path.basename(lora_path).lower()
                
                # Определяем тип по имени файла
                if any(keyword in filename for keyword in ["sdxl", "xl", "illustrious", "pony"]):
                    metadata["model_type"] = "sdxl"
                    metadata["base_model"] = "SDXL"
                else:
                    metadata["model_type"] = "sd"
                    metadata["base_model"] = "SD 1.5"
                
                logger.info(f"🔍 Анализ LoRA по имени файла: {metadata['base_model']}")
            
            return metadata
            
        except ImportError:
            logger.warning("⚠️ safetensors не установлен, анализ метаданных недоступен")
            return {"model_type": "unknown", "error": "safetensors not available"}
        except Exception as e:
            logger.error(f"❌ Ошибка анализа метаданных LoRA {lora_path}: {e}")
            return {"model_type": "unknown", "error": str(e)}
    
    def get_model_path(self) -> str:
        """Получает путь к модели с приоритетом .env > stable_diff"""
        # Приоритет 1: переменная окружения
        env_path = os.getenv('STABLE_DIFFUSION_MODEL_PATH', '').strip()
        if env_path and os.path.exists(env_path):
            return env_path
        
        # Приоритет 2: папка stable_diff/checkpoints
        if os.path.exists(self.checkpoints_dir):
            for file in os.listdir(self.checkpoints_dir):
                if file.lower().endswith(('.safetensors', '.ckpt')):
                    model_path = os.path.join(self.checkpoints_dir, file)
                    logger.info(f"🔍 Автоопределена модель: {file}")
                    return model_path
        
        # Fallback: возвращаем путь из .env даже если файл не существует
        return env_path if env_path else ""
    
    def detect_model_type(self, model_path: str) -> str:
        """
        Определяет тип модели (sd/sdxl) по метаданным или имени файла
        
        Args:
            model_path: Путь к checkpoint файлу
            
        Returns:
            Тип модели: 'sd' или 'sdxl'
        """
        if not os.path.exists(model_path):
            logger.warning(f"⚠️ Файл модели не найден: {model_path}")
            return 'sd'  # По умолчанию SD 1.5
        
        file_ext = os.path.splitext(model_path)[1].lower()
        model_name = os.path.basename(model_path).lower()
        
        # Сначала пытаемся анализировать метаданные
        if file_ext == ".safetensors":
            try:
                metadata = self.analyze_checkpoint_metadata(model_path)
                detected_type = metadata.get("model_type", "unknown")
                
                if detected_type != "unknown":
                    logger.info(f"🔍 Тип модели определен по метаданным: {detected_type}")
                    return detected_type
                    
            except Exception as e:
                logger.warning(f"⚠️ Ошибка анализа метаданных checkpoint: {e}")
        
        # Резервный анализ по имени файла
        if any(keyword in model_name for keyword in ['sdxl', 'xl', 'illustrious', 'pony']):
            logger.info(f"🔍 Тип модели определен по имени файла: sdxl")
            return 'sdxl'
        else:
            logger.info(f"🔍 Тип модели определен по имени файла: sd")
            return 'sd'
    
    def analyze_checkpoint_metadata(self, checkpoint_path: str) -> Dict[str, Any]:
        """
        Анализирует метаданные checkpoint файла
        
        Args:
            checkpoint_path: Путь к checkpoint файлу
            
        Returns:
            Словарь с метаданными checkpoint
        """
        try:
            from safetensors import safe_open
            
            metadata = {
                "model_type": "unknown",
                "architecture": "unknown",
                "base_model": "unknown",
                "resolution": "unknown",
                "model_name": "",
                "author": "",
                "description": "",
                "version": ""
            }
            
            file_ext = os.path.splitext(checkpoint_path)[1].lower()
            
            if file_ext == ".safetensors":
                with safe_open(checkpoint_path, framework="pt") as f:
                    metadata_raw = f.metadata()
                    tensor_keys = list(f.keys())
                    
                    logger.info(f"🔍 Найдено {len(tensor_keys)} тензоров в checkpoint")
                    if metadata_raw:
                        logger.info(f"🔍 Найдено {len(metadata_raw)} записей метаданных")
                    
                    # Анализируем ключи тензоров для определения архитектуры
                    sdxl_indicators = [
                        "conditioner.embedders.1.model.transformer.resblocks",
                        "conditioner.embedders.0.transformer.text_model",
                        "first_stage_model.encoder.down.0.block.0.norm1.weight",
                        "model.diffusion_model.input_blocks.4.1.transformer_blocks.0.attn2.to_k.weight"
                    ]
                    
                    sd_indicators = [
                        "cond_stage_model.transformer.text_model.encoder.layers",
                        "first_stage_model.encoder.down.0.block.0.norm1.weight",
                        "model.diffusion_model.input_blocks.1.1.transformer_blocks.0.attn1.to_q.weight"
                    ]
                    
                    # Ищем характерные ключи для SDXL
                    sdxl_score = 0
                    sd_score = 0
                    
                    for key in tensor_keys[:100]:  # Проверяем первые 100 ключей
                        for indicator in sdxl_indicators:
                            if indicator in key:
                                sdxl_score += 1
                                break
                        
                        for indicator in sd_indicators:
                            if indicator in key and "conditioner.embedders.1" not in key:
                                sd_score += 1
                                break
                    
                    # Дополнительная проверка по размерам моделей
                    try:
                        # Проверяем размер текстового энкодера
                        text_encoder_keys = [k for k in tensor_keys if "text_model.embeddings.token_embedding.weight" in k]
                        if text_encoder_keys:
                            tensor = f.get_tensor(text_encoder_keys[0])
                            vocab_size = tensor.shape[0]
                            logger.info(f"🔍 Размер словаря текстового энкодера: {vocab_size}")
                            
                            if vocab_size > 50000:  # SDXL обычно имеет больший словарь
                                sdxl_score += 2
                            else:
                                sd_score += 2
                    except:
                        pass
                    
                    # Проверяем размеры UNet
                    try:
                        unet_keys = [k for k in tensor_keys if "model.diffusion_model.input_blocks.0.0.weight" in k]
                        if unet_keys:
                            tensor = f.get_tensor(unet_keys[0])
                            input_channels = tensor.shape[1]
                            logger.info(f"🔍 Входные каналы UNet: {input_channels}")
                            
                            if input_channels == 4:  # Стандартно для обеих архитектур
                                # Проверяем другие размеры
                                output_channels = tensor.shape[0]
                                if output_channels >= 320:
                                    sdxl_score += 1
                    except:
                        pass
                    
                    logger.info(f"🔍 Счет определения: SDXL={sdxl_score}, SD={sd_score}")
                    
                    # Определяем тип модели на основе счета
                    if sdxl_score > sd_score:
                        metadata["model_type"] = "sdxl"
                        metadata["architecture"] = "SDXL"
                        metadata["base_model"] = "SDXL"
                        metadata["resolution"] = "1024x1024"
                    elif sd_score > 0:
                        metadata["model_type"] = "sd"
                        metadata["architecture"] = "SD 1.5"
                        metadata["base_model"] = "SD 1.5"
                        metadata["resolution"] = "512x512"
                    
                    # Извлекаем метаданные из заголовка файла
                    if metadata_raw:
                        # Стандартные поля
                        standard_fields = {
                            "modelspec.title": "model_name",
                            "modelspec.description": "description", 
                            "modelspec.author": "author",
                            "modelspec.implementation": "implementation",
                            "modelspec.architecture": "architecture_info"
                        }
                        
                        for raw_key, meta_key in standard_fields.items():
                            if raw_key in metadata_raw:
                                metadata[meta_key] = metadata_raw[raw_key]
                        
                        # Ищем другие полезные поля
                        for key, value in metadata_raw.items():
                            if "title" in key.lower() and not metadata.get("model_name"):
                                metadata["model_name"] = value
                            elif "description" in key.lower() and not metadata.get("description"):
                                metadata["description"] = value
                            elif "author" in key.lower() and not metadata.get("author"):
                                metadata["author"] = value
                    
                    logger.info(f"🔍 Финальное определение типа: {metadata['model_type']}")
                    if metadata["model_type"] != "unknown":
                        logger.info(f"   📋 Архитектура: {metadata['architecture']}")
                        logger.info(f"   📐 Разрешение: {metadata['resolution']}")
                        
                        if metadata.get("model_name"):
                            logger.info(f"   📝 Название: {metadata['model_name']}")
                    
            return metadata
            
        except ImportError:
            logger.warning("⚠️ safetensors не установлен, анализ checkpoint метаданных недоступен")
            return {"model_type": "unknown", "error": "safetensors not available"}
        except Exception as e:
            logger.error(f"❌ Ошибка анализа метаданных checkpoint {checkpoint_path}: {e}")
            return {"model_type": "unknown", "error": str(e)}
    
    def get_active_loras(self, model_type: str) -> List[Dict]:
        """Получает список активных LoRA для указанного типа модели"""
        config = self.get_lora_config()
        active_loras = []
        
        for lora_key, lora_config in config.get("loras", {}).items():
            if (lora_config.get("enabled", False) and 
                lora_config.get("model_type") == model_type):
                active_loras.append(lora_config)
        
        return active_loras
    
    def apply_lora_triggers(self, prompt: str, model_type: str) -> str:
        """Добавляет триггер-слова LoRA к промпту"""
        active_loras = self.get_active_loras(model_type)
        triggers = []
        
        for lora in active_loras:
            lora_triggers = lora.get("triggers", [])
            if lora_triggers:
                triggers.extend(lora_triggers)
        
        if triggers:
            trigger_text = ", ".join(triggers)
            enhanced_prompt = f"{prompt}, {trigger_text}"
            logger.info(f"🎯 Добавлены LoRA триггеры: {trigger_text}")
            return enhanced_prompt
        
        return prompt
    
    def analyze_all_loras(self) -> Dict[str, Dict[str, Any]]:
        """
        Анализирует метаданные всех LoRA файлов в системе
        
        Returns:
            Словарь с результатами анализа всех LoRA
        """
        results = {}
        lora_files = self._scan_lora_files()
        
        logger.info("🔍 Запускаю анализ метаданных всех LoRA файлов...")
        
        for model_type, files in lora_files.items():
            for filename in files:
                lora_path = os.path.join(self.lora_dir, model_type, filename)
                lora_key = f"{model_type}_{os.path.splitext(filename)[0]}"
                
                logger.info(f"📋 Анализирую: {filename}")
                metadata = self.analyze_lora_metadata(lora_path)
                
                results[lora_key] = {
                    "filename": filename,
                    "path": lora_path,
                    "folder_type": model_type,
                    "detected_type": metadata.get("model_type", "unknown"),
                    "base_model": metadata.get("base_model", "Unknown"),
                    "resolution": metadata.get("resolution", "Unknown"),
                    "triggers": metadata.get("triggers", []),
                    "author": metadata.get("author", ""),
                    "description": metadata.get("description", ""),
                    "compatible": metadata.get("model_type", model_type) == model_type,
                    "analysis_success": "error" not in metadata
                }
                
                # Предупреждение о несоответствии
                if (metadata.get("model_type", "unknown") != "unknown" and 
                    metadata.get("model_type") != model_type):
                    logger.warning(f"⚠️ {filename}: в папке {model_type}/, но предназначен для {metadata.get('model_type')}")
        
        logger.info(f"✅ Анализ завершен: {len(results)} LoRA файлов")
        return results
    
    def update_lora_metadata(self, force_update: bool = False) -> bool:
        """
        Обновляет метаданные существующих LoRA в конфигурации
        
        Args:
            force_update: Принудительно обновить все LoRA (даже уже проанализированные)
            
        Returns:
            True если конфигурация была обновлена
        """
        try:
            config = self.get_lora_config(force_reload=True)
            if "loras" not in config:
                config["loras"] = {}
            
            updated = False
            
            for lora_key, lora_config in config["loras"].items():
                # Пропускаем уже проанализированные LoRA (если не force_update)
                if not force_update and lora_config.get("metadata_analyzed", False):
                    continue
                
                filename = lora_config.get("filename")
                model_type = lora_config.get("model_type", "sd")
                
                if not filename:
                    continue
                
                # Ищем файл в соответствующей папке
                lora_path = os.path.join(self.lora_dir, model_type, filename)
                
                if not os.path.exists(lora_path):
                    logger.warning(f"⚠️ LoRA файл не найден: {lora_path}")
                    continue
                
                logger.info(f"🔍 Обновляю метаданные для {filename}")
                
                # Анализируем метаданные
                metadata = self.analyze_lora_metadata(lora_path)
                
                # Определяем актуальный тип модели
                detected_type = metadata.get("model_type", model_type)
                if detected_type != "unknown" and detected_type != model_type:
                    logger.warning(f"⚠️ LoRA {filename} в папке {model_type}/, но метаданные указывают на {detected_type}")
                    actual_model_type = detected_type
                    
                    # Создаем новый ключ с правильным типом
                    new_lora_key = f"{actual_model_type}_{os.path.splitext(filename)[0]}"
                    if new_lora_key != lora_key:
                        logger.info(f"🔄 Перемещаю конфигурацию: {lora_key} -> {new_lora_key}")
                        # Копируем в новый ключ
                        config["loras"][new_lora_key] = lora_config.copy()
                        # Удаляем старый ключ
                        del config["loras"][lora_key]
                        lora_key = new_lora_key
                        lora_config = config["loras"][lora_key]
                else:
                    actual_model_type = model_type
                
                # Сохраняем пользовательские настройки
                user_enabled = lora_config.get("enabled", True)
                user_strength = lora_config.get("strength", 1.0)
                user_triggers = lora_config.get("triggers", [])
                
                # Обновляем конфигурацию с метаданными
                config["loras"][lora_key].update({
                    "model_type": actual_model_type,
                    "enabled": user_enabled,  # Сохраняем пользовательскую настройку
                    "strength": user_strength,  # Сохраняем пользовательскую силу
                    "triggers": user_triggers if user_triggers else metadata.get("triggers", [])[:3],
                    "description": metadata.get("description", f"Auto-detected: {metadata.get('base_model', 'Unknown')} LoRA"),
                    "base_model": metadata.get("base_model", "Unknown"),
                    "resolution": metadata.get("resolution", "Unknown"),
                    "author": metadata.get("author", ""),
                    "metadata_analyzed": True
                })
                
                updated = True
                
                logger.info(f"✅ Обновлены метаданные для {filename}")
                logger.info(f"   🎯 Тип: {actual_model_type} ({metadata.get('base_model', 'Unknown')})")
                if metadata.get("triggers") and not user_triggers:
                    logger.info(f"   🔤 Новые триггеры: {', '.join(metadata['triggers'][:3])}")
            
            if updated:
                with open(self.lora_config_path, 'w', encoding='utf-8') as f:
                    json.dump(config, f, ensure_ascii=False, indent=2)
                logger.info(f"✅ Конфигурация LoRA обновлена с метаданными")
                return True
            else:
                logger.info(f"📋 Все LoRA уже имеют актуальные метаданные")
                return False
                
        except Exception as e:
            logger.error(f"❌ Ошибка обновления метаданных LoRA: {e}")
            return False

class PromptLoader:
    """
    Класс для динамической загрузки системных промптов и модулей из .md файлов
    """
    
    def __init__(self, base_dir: str | None = None):
        if base_dir is None:
            base_dir = os.path.dirname(os.path.abspath(__file__))
        self.prompts_dir = os.path.join(base_dir, "promts")
        self.base_prompt_file = os.path.join(self.prompts_dir, "PROMPT_SYSTEM.md")
        
        # Карта команд к файлам модулей
        self.module_commands = {
            "get_image_generation_help": "image_generation_module.md",
            "get_email_module_help": "email_module.md", 
            "get_pc_control_help": "pc_control_module.md",
            "get_file_processing_help": "file_processing_module.md",
            "get_error_handling_help": "error_handling_module.md",
            "get_additional_modules_help": "additional_modules.md",
            "get_search_help": "additional_modules.md",
            "get_media_analysis_help": "additional_modules.md",
            "get_plugins_help": "additional_modules.md",
            "get_memory_help": "additional_modules.md",
            "get_speech_help": "additional_modules.md",
            "get_workflow_help": "additional_modules.md",
            "get_strategy_help": "additional_modules.md"
        }
        
        # Кэш загруженных модулей
        self._module_cache = {}
        self._base_prompt_cache = None
    
    def load_base_prompt(self) -> str:
        """
        Загружает базовый системный промпт из PROMPT_SYSTEM.md
        """
        if self._base_prompt_cache is not None:
            return self._base_prompt_cache
            
        try:
            if not os.path.exists(self.base_prompt_file):
                logger.error(f"Файл базового промпта не найден: {self.base_prompt_file}")
                return self._get_fallback_prompt()
                
            with open(self.base_prompt_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Извлекаем только секцию "БАЗОВЫЙ УНИВЕРСАЛЬНЫЙ ПРОМПТ"
            base_section = self._extract_base_section(content)
            self._base_prompt_cache = base_section
            return base_section
            
        except Exception as e:
            logger.error(f"Ошибка загрузки базового промпта: {e}")
            return self._get_fallback_prompt()
    
    def _extract_base_section(self, content: str) -> str:
        """
        Извлекает секцию базового промпта из полного файла
        """
        lines = content.split('\n')
        base_lines = []
        in_base_section = False
        
        for line in lines:
            # Начало базовой секции
            if "## БАЗОВЫЙ УНИВЕРСАЛЬНЫЙ ПРОМПТ" in line:
                in_base_section = True
                continue
            
            # Конец базовой секции (следующий заголовок уровня 2)
            if in_base_section and line.startswith("## ") and "БАЗОВЫЙ" not in line:
                break
                
            if in_base_section:
                base_lines.append(line)
        
        if base_lines:
            return '\n'.join(base_lines).strip()
        else:
            # Если секция не найдена, возвращаем весь контент до первого модуля
            for i, line in enumerate(lines):
                if line.startswith("## МОДУЛЬ:"):
                    return '\n'.join(lines[:i]).strip()
            return content.strip()
    
    def load_module(self, command: str) -> str:
        """
        Загружает модуль по команде (например, get_image_generation_help)
        """
        if command in self._module_cache:
            return self._module_cache[command]
            
        if command not in self.module_commands:
            logger.warning(f"Неизвестная команда модуля: {command}")
            return f"Модуль для команды '{command}' не найден."
        
        module_file = self.module_commands[command]
        module_path = os.path.join(self.prompts_dir, module_file)
        
        try:
            if not os.path.exists(module_path):
                logger.error(f"Файл модуля не найден: {module_path}")
                return f"Файл модуля {module_file} не найден."
                
            with open(module_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Для additional_modules.md извлекаем конкретную секцию
            if module_file == "additional_modules.md":
                content = self._extract_specific_module(content, command)
            
            self._module_cache[command] = content
            return content
            
        except Exception as e:
            logger.error(f"Ошибка загрузки модуля {module_file}: {e}")
            return f"Ошибка загрузки модуля: {e}"
    
    def _extract_specific_module(self, content: str, command: str) -> str:
        """
        Извлекает конкретную секцию из additional_modules.md
        """
        # Карта команд к заголовкам секций
        section_map = {
            "get_additional_modules_help": "ВСЕ МОДУЛИ",  # Возвращает весь файл
            "get_search_help": "МОДУЛЬ: ИНТЕРНЕТ ПОИСК",
            "get_media_analysis_help": "МОДУЛЬ: ВИДЕО И АУДИО АНАЛИЗ",
            "get_plugins_help": "МОДУЛЬ: ПЛАГИНЫ",
            "get_memory_help": "МОДУЛЬ: ВЕКТОРНАЯ ПАМЯТЬ CHROMADB",
            "get_speech_help": "МОДУЛЬ: ОЗВУЧКА",
            "get_workflow_help": "МОДУЛЬ: ЦЕПОЧКИ ДЕЙСТВИЙ",
            "get_strategy_help": "МОДУЛЬ: СТРАТЕГИЧЕСКОЕ МЫШЛЕНИЕ"
        }
        
        target_section = section_map.get(command)
        if not target_section:
            return content
            
        # Для get_additional_modules_help возвращаем весь файл
        if command == "get_additional_modules_help":
            return content
        
        lines = content.split('\n')
        section_lines = []
        in_target_section = False
        
        for line in lines:
            if target_section in line:
                in_target_section = True
                section_lines.append(line)
                continue
                
            if in_target_section:
                # Конец секции - следующий заголовок уровня 2
                if line.startswith("## МОДУЛЬ:") and target_section not in line:
                    break
                section_lines.append(line)
        
        return '\n'.join(section_lines).strip() if section_lines else content
    
    def _get_fallback_prompt(self) -> str:
        """
        Возвращает резервный промпт если файлы не найдены
        """
        return """Ты — Нейро, интеллектуальный программный оркестратор.

ВСЕГДА отвечай в формате JSON с одним из доступных действий:
- "powershell" — выполнение команд PowerShell
- "search" — поиск в интернете
- "send_email" — отправка письма
- "get_emails" — получение писем
- "reply_email" — ответ на письмо
- "search_emails" — поиск писем
- "generate_image" — генерация изображения
- "response" — финальный ответ пользователю
- "move_mouse", "left_click", "right_click", "scroll_up", "scroll_down" — управление мышью
- "type_text" — ввод текста
- "take_screenshot" — создание скриншота

КРИТИЧЕСКИ ВАЖНЫЕ ПРАВИЛА:
1. НЕ ВЫДУМЫВАЙ результаты операций
2. ВСЕГДА ЭКРАНИРУЙ обратные слэши в путях (\\\\)
3. ИСПОЛЬЗУЙ UTF-8 для русского текста
4. СТРОЙ ЦЕПОЧКИ действий для сложных задач

Если нужна дополнительная информация о конкретных инструментах, используй команды:
- get_image_generation_help
- get_email_module_help
- get_pc_control_help
- get_file_processing_help
- get_error_handling_help"""

    def is_module_command(self, message: str) -> bool:
        """
        Проверяет, является ли сообщение командой для загрузки модуля
        """
        return message.strip() in self.module_commands
    
    def get_available_commands(self) -> List[str]:
        """
        Возвращает список доступных команд модулей
        """
        return list(self.module_commands.keys())

class AIOrchestrator:
    def extract_video_frames(self, video_path: str, fps: int = 1) -> list:
        """
        Извлекает по одному кадру на каждую секунду видео (fps=1).
        Возвращает список кортежей (таймкод, base64 PNG).
        """
        frames = []
        temp_dir = tempfile.mkdtemp()
        try:
            # Получаем длительность видео через ffprobe
            cmd = [
                'ffprobe', '-v', 'error', '-show_entries', 'format=duration',
                '-of', 'default=noprint_wrappers=1:nokey=1', video_path
            ]
            result = subprocess.run(cmd, capture_output=True, text=True)
            duration = float(result.stdout.strip()) if result.returncode == 0 else 0
            if duration == 0:
                return []
            # Извлекаем кадры с помощью ffmpeg
            # -vf fps=1: по одному кадру в секунду
            frame_pattern = os.path.join(temp_dir, 'frame_%05d.png')
            cmd = [
                'ffmpeg', '-i', video_path, '-vf', f'fps={fps}', '-q:v', '2', frame_pattern, '-hide_banner', '-loglevel', 'error'
            ]
            subprocess.run(cmd, check=True)
            # Собираем кадры и таймкоды
            total_frames = int(math.ceil(duration))
            for i in range(1, total_frames + 1):
                frame_path = os.path.join(temp_dir, f'frame_{i:05d}.png')
                if not os.path.exists(frame_path):
                    continue
                # Таймкод в формате [HH:MM:SS]
                sec = i - 1
                h = sec // 3600
                m = (sec % 3600) // 60
                s = sec % 60
                timecode = f"[{h:02}:{m:02}:{s:02}]"
                # base64 через существующую функцию
                b64 = image_to_base64_balanced(frame_path)
                frames.append((timecode, b64))
            return frames
        except Exception as e:
            self.logger.error(f"Ошибка извлечения кадров: {e}")
            return []
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)
    def download_youtube_video(self, url: str, out_dir: Optional[str] = None) -> Optional[str]:
        """
        Скачивает видео с YouTube по ссылке (использует yt-dlp)
        Возвращает путь к mp4-файлу или пустую строку
        """
        if out_dir is None:
            out_dir = os.path.join(os.path.dirname(__file__), "Video")
        os.makedirs(out_dir, exist_ok=True)
        out_path = os.path.join(out_dir, "yt_video.%(ext)s")
        
        # Проверяем наличие cookies
        cookies_path = self.get_youtube_cookies_path()
        use_cookies = False
        
        if cookies_path and self.check_cookies_validity(cookies_path):
            use_cookies = True
            self.logger.info("🍪 Использую cookies для аутентификации YouTube")
        else:
            self.logger.info("ℹ️ Cookies не найдены или невалидны, использую базовые параметры")
            if not cookies_path:
                self.suggest_cookies_update()
        
        # Базовые параметры для yt-dlp
        base_cmd = [
            "yt-dlp",
            "--force-ipv4",
            "--user-agent", "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "--extractor-args", "youtube:player_client=android",  # Используем Android клиент
            "--no-check-certificate",  # Игнорируем SSL ошибки
            "--prefer-insecure",  # Предпочитаем HTTP
            "--geo-bypass",  # Обход геоблокировки
            "--geo-bypass-country", "US",  # Страна для обхода
            "-f", "bestvideo[ext=mp4]+bestaudio[ext=m4a]/mp4/best[ext=mp4]/best",
            "-o", out_path
        ]
        
        # Добавляем cookies если доступны
        if cookies_path:
            base_cmd.extend(["--cookies", str(cookies_path)])  # type: ignore[arg-type]
        
        # Добавляем URL в конец
        cmd = base_cmd + [url]
        
        try:
            self.logger.info(f"Скачиваю видео с YouTube: {url}")
            # Логируем команду в одну строку для избежания обрезания
            cmd_str = " ".join(cmd)
            self.logger.info(f"Команда: {cmd_str}")
            
            # Запускаем с таймаутом
            result = subprocess.run(cmd, check=True, capture_output=True, text=True, timeout=300)
            
            if result.stdout:
                self.logger.info(f"yt-dlp stdout: {result.stdout}")
            if result.stderr:
                self.logger.warning(f"yt-dlp stderr: {result.stderr}")
            
            # Найти скачанный файл
            for fname in os.listdir(out_dir):
                if fname.startswith("yt_video") and fname.endswith('.mp4'):
                    self.logger.info(f"✅ Видео успешно скачано: {fname}")
                    return os.path.join(out_dir, fname)
            
            self.logger.warning("⚠️ Файл не найден после скачивания")
            return ""
            
        except subprocess.TimeoutExpired:
            self.logger.error("❌ Таймаут скачивания видео (5 минут)")
            return ""
        except subprocess.CalledProcessError as e:
            self.logger.error(f"❌ Ошибка yt-dlp: {e}")
            if e.stderr:
                self.logger.error(f"stderr: {e.stderr}")
            return ""
        except Exception as e:
            self.logger.error(f"❌ Неожиданная ошибка скачивания видео: {e}")
            
            # Пробуем альтернативный метод с другими параметрами
            self.logger.info("🔄 Пробую альтернативный метод скачивания...")
            try:
                alt_cmd = [
                    "yt-dlp",
                    "--force-ipv4",
                    "--user-agent", "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                    "--extractor-args", "youtube:player_client=web",
                    "--no-check-certificate",
                    "--geo-bypass",
                    "--geo-bypass-country", "US",
                    "-f", "best[ext=mp4]/best",
                    "-o", out_path
                ]
                
                # Добавляем cookies если доступны
                if cookies_path:
                    alt_cmd.extend(["--cookies", str(cookies_path)])  # type: ignore[arg-type]
                
                alt_cmd.append(url)
                
                # Логируем команду в одну строку
                alt_cmd_str = " ".join(alt_cmd)
                self.logger.info(f"Альтернативная команда: {alt_cmd_str}")
                result = subprocess.run(alt_cmd, check=True, capture_output=True, text=True, timeout=300)
                
                # Найти скачанный файл
                for fname in os.listdir(out_dir):
                    if fname.startswith("yt_video") and fname.endswith('.mp4'):
                        self.logger.info(f"✅ Видео успешно скачано альтернативным методом: {fname}")
                        return os.path.join(out_dir, fname)
                        
            except Exception as alt_e:
                self.logger.error(f"❌ Альтернативный метод также не сработал: {alt_e}")
                
                # Пробуем третий метод с максимально простыми параметрами
                self.logger.info("🔄 Пробую третий метод (минимальные параметры)...")
                try:
                    simple_cmd = [
                        "yt-dlp",
                        "--force-ipv4",
                        "--user-agent", "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                        "--no-check-certificate",
                        "-f", "best",
                        "-o", out_path
                    ]
                    # Добавляем cookies если доступны
                    if use_cookies and cookies_path:
                        simple_cmd.extend(["--cookies", str(cookies_path)])  # type: ignore[arg-type]

                    simple_cmd.append(url)
                    
                    # Логируем команду в одну строку
                    simple_cmd_str = " ".join(simple_cmd)
                    self.logger.info(f"Третий метод: {simple_cmd_str}")
                    result = subprocess.run(simple_cmd, check=True, capture_output=True, text=True, timeout=300)
                    
                    # Найти скачанный файл
                    for fname in os.listdir(out_dir):
                        if fname.startswith("yt_video") and fname.endswith('.mp4'):
                            self.logger.info(f"✅ Видео успешно скачано третьим методом: {fname}")
                            return os.path.join(out_dir, fname)
                            
                except Exception as simple_e:
                    self.logger.error(f"❌ Третий метод также не сработал: {simple_e}")
            
            return ""
    
    def check_vpn_status(self) -> bool:
        """
        Проверяет статус VPN соединения
        """
        try:
            import requests
            # Пробуем получить IP адрес
            response = requests.get("https://ifconfig.me", timeout=10)
            if response.status_code == 200:
                ip = response.text.strip()
                self.logger.info(f"🌐 Текущий IP адрес: {ip}")
                
                # Проверяем, не из РФ ли IP
                ru_ips = ["185.", "31.", "46.", "37.", "95.", "178.", "79.", "5.", "176.", "195."]
                if any(ip.startswith(prefix) for prefix in ru_ips):
                    self.logger.warning("⚠️ IP адрес похож на российский. VPN может не работать корректно.")
                    return False
                else:
                    self.logger.info("✅ IP адрес не из РФ. VPN работает.")
                    return True
            else:
                self.logger.warning(f"⚠️ Не удалось проверить IP: {response.status_code}")
                return False
                
        except Exception as e:
            self.logger.error(f"❌ Ошибка проверки VPN: {e}")
            return False

    def get_youtube_info(self, url: str) -> dict:
        """
        Получает информацию о YouTube видео без скачивания
        """
        try:
            import json
            # Проверяем наличие cookies
            cookies_path = self.get_youtube_cookies_path()
            use_cookies = False
            
            if cookies_path and self.check_cookies_validity(cookies_path):
                use_cookies = True
                self.logger.info("🍪 Использую cookies для получения информации о видео")
            
            # Базовые параметры для yt-dlp
            base_cmd = [
                "yt-dlp",
                "--force-ipv4",
                "--user-agent", "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                "--extractor-args", "youtube:player_client=android",
                "--no-check-certificate",
                "--geo-bypass",
                "--dump-json"
            ]
            
            # Добавляем cookies если доступны
            if use_cookies:
                base_cmd.extend(["--cookies", str(cookies_path)])  # type: ignore[arg-type]
            
            # Добавляем URL в конец
            cmd = base_cmd + [url]
            
            self.logger.info("📋 Получаю информацию о YouTube видео...")
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            
            if result.returncode == 0 and result.stdout:
                try:
                    import json
                    info = json.loads(result.stdout)
                    title = info.get('title', 'Неизвестное видео')
                    duration = info.get('duration', 0)
                    uploader = info.get('uploader', 'Неизвестный автор')
                    
                    self.logger.info(f"✅ Информация получена: {title} ({duration}с) от {uploader}")
                    return {
                        'title': title,
                        'duration': duration,
                        'uploader': uploader,
                        'success': True
                    }
                except json.JSONDecodeError:
                    self.logger.error("❌ Ошибка парсинга JSON информации о видео")
                    return {'success': False, 'error': 'JSON parse error'}
            else:
                self.logger.error(f"❌ Не удалось получить информацию: {result.stderr}")
                
                # Пробуем альтернативный метод без Android клиента
                self.logger.info("🔄 Пробую альтернативный метод получения информации...")
                try:
                    alt_cmd = [
                        "yt-dlp",
                        "--force-ipv4",
                        "--user-agent", "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                        "--extractor-args", "youtube:player_client=web",
                        "--no-check-certificate",
                        "--geo-bypass",
                        "--dump-json"
                    ]
                    
                    # Добавляем cookies если доступны
                    if use_cookies:
                        alt_cmd.extend(["--cookies", str(cookies_path)])  # type: ignore[arg-type]
                    
                    alt_cmd.append(url)
                    
                    self.logger.info("🔄 Альтернативная команда для получения информации...")
                    alt_result = subprocess.run(alt_cmd, capture_output=True, text=True, timeout=60)
                    
                    if alt_result.returncode == 0 and alt_result.stdout:
                        try:
                            import json
                            info = json.loads(alt_result.stdout)
                            title = info.get('title', 'Неизвестное видео')
                            duration = info.get('duration', 0)
                            uploader = info.get('uploader', 'Неизвестный автор')
                            
                            self.logger.info(f"✅ Информация получена альтернативным методом: {title} ({duration}с) от {uploader}")
                            return {
                                'title': title,
                                'duration': duration,
                                'uploader': uploader,
                                'success': True
                            }
                        except json.JSONDecodeError:
                            self.logger.error("❌ Ошибка парсинга JSON альтернативным методом")
                            return {'success': False, 'error': 'JSON parse error (alt method)'}
                    else:
                        self.logger.error(f"❌ Альтернативный метод также не сработал: {alt_result.stderr}")
                        return {'success': False, 'error': result.stderr}
                        
                except Exception as alt_e:
                    self.logger.error(f"❌ Ошибка альтернативного метода: {alt_e}")
                    return {'success': False, 'error': result.stderr}
                
        except Exception as e:
            self.logger.error(f"❌ Ошибка получения информации о видео: {e}")
            return {'success': False, 'error': str(e)}

    def check_youtube_accessibility(self, url: str) -> bool:
        """
        Проверяет доступность YouTube ссылки различными методами
        """
        try:
            # Проверяем наличие cookies
            cookies_path = self.get_youtube_cookies_path()
            use_cookies = False
            
            if cookies_path and self.check_cookies_validity(cookies_path):
                use_cookies = True
                self.logger.info("🍪 Использую cookies для проверки доступности")
            
            # Базовые параметры для yt-dlp
            base_cmd = [
                "yt-dlp",
                "--force-ipv4",
                "--user-agent", "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                "--extractor-args", "youtube:player_client=android",
                "--no-check-certificate",
                "--geo-bypass",
                "--list-formats"
            ]
            
            # Добавляем cookies если доступны
            if use_cookies:
                base_cmd.extend(["--cookies", str(cookies_path)])  # type: ignore[arg-type]
            
            # Добавляем URL в конец
            test_cmd = base_cmd + [url]
            
            self.logger.info("🔍 Проверяю доступность YouTube ссылку...")
            result = subprocess.run(test_cmd, capture_output=True, text=True, timeout=60)
            
            if result.returncode == 0:
                self.logger.info("✅ YouTube ссылка доступна")
                return True
            else:
                self.logger.warning(f"⚠️ YouTube ссылка недоступна: {result.stderr}")
                
                # Пробуем альтернативный метод с web клиентом
                self.logger.info("🔄 Пробую альтернативный метод проверки...")
                try:
                    alt_test_cmd = [
                        "yt-dlp",
                        "--force-ipv4",
                        "--user-agent", "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                        "--extractor-args", "youtube:player_client=web",
                        "--no-check-certificate",
                        "--geo-bypass",
                        "--list-formats"
                    ]
                    
                    # Добавляем cookies если доступны
                    if use_cookies:
                        alt_test_cmd.extend(["--cookies", str(cookies_path)])  # type: ignore[arg-type]
                    
                    alt_test_cmd.append(url)
                    
                    alt_result = subprocess.run(alt_test_cmd, capture_output=True, text=True, timeout=60)
                    
                    if alt_result.returncode == 0:
                        self.logger.info("✅ YouTube ссылка доступна через альтернативный метод")
                        return True
                    else:
                        self.logger.warning(f"⚠️ YouTube ссылка недоступна и через альтернативный метод: {alt_result.stderr}")
                        return False
                        
                except Exception as alt_e:
                    self.logger.error(f"❌ Ошибка альтернативной проверки: {alt_e}")
                    return False
                
        except Exception as e:
            self.logger.error(f"❌ Ошибка проверки доступности YouTube: {e}")
            return False

    def _auto_load_brain_model(self):
        """Автоматически загружает модель мозга при инициализации"""
        try:
            # Проверяем, запущена ли модель через прямой запрос к API
            try:
                resp = requests.get(f"{self.lm_studio_url}/v1/models", timeout=10)
                if resp.status_code == 200:
                    data = resp.json()
                    model_loaded = False
                    for m in data.get("data", []):
                        if self.brain_model in m.get("id", "") and m.get("isLoaded", False):
                            model_loaded = True
                            # Сохраняем короткий ID модели для API вызовов
                            self.brain_model_id = m.get("id")
                            self.logger.info(f"✅ Модель мозга уже загружена: {os.path.basename(self.brain_model)} (ID: {self.brain_model_id})")
                            return
                else:
                    self.logger.warning(f"⚠️ Не удалось проверить статус моделей: {resp.status_code}")
            except Exception as e:
                self.logger.warning(f"⚠️ Ошибка проверки статуса моделей: {e}")
            
            # Если модель не загружена, пытаемся загрузить
            self.logger.info(f"🧠 Автоматически загружаю модель мозга: {os.path.basename(self.brain_model)}")
            
            # Пытаемся загрузить модель через API
            payload = {
                "model": self.brain_model,
                "load": True
            }
            
            try:
                resp = requests.post(f"{self.lm_studio_url}/v1/models/load", json=payload, timeout=30)
                if resp.status_code == 200:
                    self.logger.info("✅ Модель мозга успешно загружена через API")
                    # Получаем короткий ID модели после загрузки
                    self._update_brain_model_id()
                else:
                    self.logger.warning(f"⚠️ Не удалось загрузить модель через API: {resp.status_code}")
                    # Пробуем запустить через LM Studio
                    self.launch_model(self.brain_model)
                    self.logger.info("🔄 Запускаю модель через LM Studio...")
                    # Пытаемся получить ID модели после запуска
                    self._update_brain_model_id()
            except Exception as e:
                self.logger.warning(f"⚠️ Ошибка API загрузки модели: {e}")
                # Пробуем запустить через LM Studio
                self.launch_model(self.brain_model)
                self.logger.info("🔄 Запускаю модель через LM Studio...")
                # Пытаемся получить ID модели после запуска
                self._update_brain_model_id()
                
        except Exception as e:
            self.logger.error(f"❌ Ошибка автозагрузки модели мозга: {e}")
    
    def _update_brain_model_id(self):
        """Обновляет короткий ID модели мозга из API"""
        try:
            resp = requests.get(f"{self.lm_studio_url}/v1/models", timeout=10)
            if resp.status_code == 200:
                data = resp.json()
                for m in data.get("data", []):
                    if self.brain_model in m.get("id", "") and m.get("isLoaded", False):
                        self.brain_model_id = m.get("id")
                        self.logger.info(f"✅ Обновлен ID модели мозга: {self.brain_model_id}")
                        return
                self.logger.warning("⚠️ Не удалось найти загруженную модель для получения ID")
            else:
                self.logger.warning(f"⚠️ Не удалось получить список моделей для обновления ID: {resp.status_code}")
        except Exception as e:
            self.logger.warning(f"⚠️ Ошибка обновления ID модели мозга: {e}")
    
    def _start_background_loading(self):
        """Запускает фоновую загрузку тяжелых компонентов"""
        loader = get_background_loader()
        
        # Запускаем загрузку EasyOCR в фоне
        loader.start_loading('easyocr', load_easyocr)
        
        # Запускаем загрузку PyTorch в фоне (если нужно)
        loader.start_loading('torch', load_torch)
        
        self.logger.info("🚀 Запущена фоновая загрузка компонентов")
    
    def _ensure_ocr_initialized(self):
        """Обеспечивает инициализацию OCR"""
        if self.ocr_reader is not None:
            return True
            
        loader = get_background_loader()
        ocr_reader = loader.get_component('easyocr', timeout=30)
        
        # Проверяем, что reader не None (успешно загружен)
        if ocr_reader is not None:
            self.ocr_reader = ocr_reader
            self.logger.info("✅ EasyOCR загружен из фонового потока")
            return True
        else:
            # Fallback к синхронной инициализации
            return self._initialize_ocr_sync()
    
    def _initialize_ocr_sync(self):
        """Синхронная инициализация OCR как fallback"""
        try:
            self.logger.info("📖 Синхронная инициализация EasyOCR...")
            import easyocr  # type: ignore
            self.ocr_reader = easyocr.Reader(['ru', 'en'])
            self.logger.info("✅ EasyOCR инициализирован синхронно")
            return True
        except Exception as e:
            self.logger.error(f"❌ Ошибка синхронной инициализации EasyOCR: {e}")
            return False
    
    def _reconnect_brain_model(self):
        """Переподключается к модели мозга, если соединение потеряно"""
        try:
            self.logger.info("🔄 Переподключение к модели мозга...")
            
            # Сначала пытаемся загрузить модель
            self._auto_load_brain_model()
            
            # Ждем немного для загрузки
            time.sleep(3)
            
            # Проверяем, что модель доступна
            response = requests.get(f"{self.lm_studio_url}/v1/models", timeout=10)
            if response.status_code == 200:
                models = response.json().get("data", [])
                if any(self.brain_model in m.get("id", "") for m in models):
                    self.logger.info("✅ Переподключение к модели мозга успешно")
                    return True
            
            self.logger.warning("⚠️ Модель мозга недоступна после переподключения")
            return False
            
        except Exception as e:
            self.logger.error(f"❌ Ошибка переподключения к модели мозга: {e}")
            return False
    
    def _initialize_ocr(self):
        """Инициализирует OCR для распознавания текста на изображениях (теперь ленивая загрузка)"""
        # OCR теперь загружается в фоне, здесь ничего не делаем
        pass

    def _ensure_chromadb_initialized(self):
        """Ленивая инициализация ChromaDB - вызывается только при первом использовании"""
        if not self._chromadb_initialized:
            try:
                self.logger.info("🔄 Инициализирую ChromaDB...")
                self.chromadb_manager = ChromaDBManager(
                    db_path=self._chromadb_config["db_path"],
                    use_gpu=self._chromadb_config["use_gpu"]
                )
                self._chromadb_initialized = True
                self.logger.info("✅ ChromaDB инициализирован")
            except Exception as e:
                self.logger.error(f"❌ Ошибка инициализации ChromaDB: {e}")
                self.chromadb_manager = None
    
    def _check_ffmpeg(self):
        """Проверяет наличие ffmpeg в системе для конвертации аудио"""
        try:
            result = subprocess.run(['ffmpeg', '-version'], capture_output=True, text=True, timeout=10)
            if result.returncode == 0:
                self.logger.info("✅ ffmpeg найден в системе")
            else:
                self.logger.warning("⚠️ ffmpeg найден, но не может быть запущен")
        except (subprocess.CalledProcessError, FileNotFoundError):
            self.logger.warning("⚠️ ffmpeg не найден в системе. Установите ffmpeg для конвертации аудио.")
            self.logger.info("💡 Скачайте с https://ffmpeg.org/download.html")
        except Exception as e:
            self.logger.warning(f"⚠️ Ошибка проверки ffmpeg: {e}")
    
    def is_model_running(self, model_name: str) -> bool:
        """
        Проверяет, запущена ли модель в LM Studio через /v1/models
        """
        try:
            resp = requests.get(f"{self.lm_studio_url}/v1/models")
            if resp.status_code == 200:
                data = resp.json()
                for m in data.get("data", []):
                    if model_name in m.get("id", "") and m.get("isLoaded", False):
                        return True
            return False
        except Exception as e:
            self.logger.error(f"Ошибка проверки модели {model_name}: {e}")
            return False

    def get_model_context_info(self) -> Dict[str, int]:
        """
        Получает информацию о максимальном контексте модели из LM Studio API
        Возвращает словарь с max_context и safe_context
        """
        try:
            resp = requests.get(f"{self.lm_studio_url}/v1/models", timeout=10)
            if resp.status_code == 200:
                data = resp.json()
                
                # Ищем нашу модель по ключевым словам
                target_model = None
                search_terms = ["huihui-qwen3-4b-thinking", "qwen3-4b", "thinking"]
                
                for m in data.get("data", []):
                    model_id = m.get("id", "").lower()
                    for term in search_terms:
                        if term.lower() in model_id:
                            target_model = m
                            self.logger.info(f"🎯 Найдена модель: {m.get('id')}")
                            break
                    if target_model:
                        break
                
                if target_model:
                    # Пытаемся получить информацию через запрос к модели
                    context_info = self._get_context_info_via_chat(target_model.get("id"))
                    if context_info:
                        return context_info
                    
                    # Если не удалось получить через чат, сохраняем ID модели для использования
                    if not hasattr(self, 'brain_model_id') or not self.brain_model_id:
                        self.brain_model_id = target_model.get("id")
                        self.logger.info(f"✅ Сохранен ID модели мозга: {self.brain_model_id}")
                
            # Если не удалось получить информацию, используем значения по умолчанию
            self.logger.warning("⚠️ Не удалось получить информацию о контексте модели, используем значения по умолчанию")
            return {
                "max_context": 262144,
                "safe_context": 32768
            }
        except Exception as e:
            self.logger.warning(f"⚠️ Ошибка получения информации о контексте модели: {e}")
            return {
                "max_context": 262144,
                "safe_context": 32768
            }

    def _get_context_info_via_chat(self, model_id: str) -> Optional[Dict[str, int]]:
        """
        Пытается получить информацию о контексте через запрос к модели
        """
        try:
            # Простой запрос для получения информации о модели
            payload = {
                "model": model_id,
                "messages": [{"role": "user", "content": "Hello"}],
                "max_tokens": 1,
                "temperature": 0
            }
            
            resp = requests.post(f"{self.lm_studio_url}/v1/chat/completions", json=payload, timeout=10)
            
            if resp.status_code == 200:
                data = resp.json()
                
                # Проверяем поле stats
                stats = data.get("stats", {})
                if stats:
                    # Ищем информацию о контексте в stats
                    context_length = None
                    if "context_length" in stats:
                        context_length = stats["context_length"]
                    elif "max_context" in stats:
                        context_length = stats["max_context"]
                    elif "max_tokens" in stats:
                        context_length = stats["max_tokens"]
                    
                    if context_length:
                        safe_context = max(context_length // 8, 32768)
                        self.logger.info(f"✅ Найден context_length в stats: {context_length}")
                        return {
                            "max_context": context_length,
                            "safe_context": safe_context
                        }
                
                # Проверяем другие поля на предмет информации о контексте
                for key, value in data.items():
                    if isinstance(value, dict) and ("context" in key.lower() or "token" in key.lower()):
                        self.logger.debug(f"🔍 Проверяем поле {key}: {value}")
                
                self.logger.debug("❌ Информация о контексте не найдена в ответе модели")
                return None
            else:
                self.logger.warning(f"❌ Ошибка запроса к модели: {resp.status_code}")
                return None
                
        except Exception as e:
            self.logger.warning(f"❌ Ошибка получения информации через чат: {e}")
            return None

    def _initialize_dynamic_context(self):
        """
        Инициализирует динамические параметры контекста на основе информации о модели
        """
        try:
            context_info = self.get_model_context_info()
            self.max_context_length = context_info["max_context"]
            self.safe_context_length = context_info["safe_context"]
            self.logger.info(f"📊 Контекст инициализирован: максимум {self.max_context_length:,}, безопасный {self.safe_context_length:,}")
        except Exception as e:
            self.logger.warning(f"⚠️ Ошибка инициализации динамического контекста: {e}")
            # Оставляем значения по умолчанию
            self.max_context_length = 262144
            self.safe_context_length = 32768

    def _initialize_email_config(self):
        """
        Инициализирует конфигурацию почтовой системы
        """
        try:
            # Загружаем настройки почты из переменных окружения
            self.email_config = {
                'gmail': {
                    'smtp_server': os.getenv('GMAIL_SMTP_SERVER', 'smtp.gmail.com'),
                    'smtp_port': int(os.getenv('GMAIL_SMTP_PORT', '587')),
                    'imap_server': os.getenv('GMAIL_IMAP_SERVER', 'imap.gmail.com'),
                    'imap_port': int(os.getenv('GMAIL_IMAP_PORT', '993')),
                    'email': os.getenv('GMAIL_EMAIL'),
                    'app_password': os.getenv('GMAIL_APP_PASSWORD')
                },
                'outlook': {
                    'smtp_server': os.getenv('OUTLOOK_SMTP_SERVER', 'smtp-mail.outlook.com'),
                    'smtp_port': int(os.getenv('OUTLOOK_SMTP_PORT', '587')),
                    'imap_server': os.getenv('OUTLOOK_IMAP_SERVER', 'outlook.office365.com'),
                    'imap_port': int(os.getenv('OUTLOOK_IMAP_PORT', '993')),
                    'email': os.getenv('OUTLOOK_EMAIL'),
                    'app_password': os.getenv('OUTLOOK_APP_PASSWORD')
                },
                'yandex': {
                    'smtp_server': os.getenv('YANDEX_SMTP_SERVER', 'smtp.yandex.ru'),
                    'smtp_port': int(os.getenv('YANDEX_SMTP_PORT', '587')),
                    'imap_server': os.getenv('YANDEX_IMAP_SERVER', 'imap.yandex.ru'),
                    'imap_port': int(os.getenv('YANDEX_IMAP_PORT', '993')),
                    'email': os.getenv('YANDEX_EMAIL'),
                    'app_password': os.getenv('YANDEX_APP_PASSWORD')
                },
                'mail_ru': {
                    'smtp_server': os.getenv('MAIL_RU_SMTP_SERVER', 'smtp.mail.ru'),
                    'smtp_port': int(os.getenv('MAIL_RU_SMTP_PORT', '587')),
                    'imap_server': os.getenv('MAIL_RU_IMAP_SERVER', 'imap.mail.ru'),
                    'imap_port': int(os.getenv('MAIL_RU_IMAP_PORT', '993')),
                    'email': os.getenv('MAIL_RU_EMAIL'),
                    'app_password': os.getenv('MAIL_RU_APP_PASSWORD')
                }
            }

            # Проверяем какие провайдеры настроены
            self.available_email_providers = []
            for provider, config in self.email_config.items():
                if config['email'] and config['app_password']:
                    self.available_email_providers.append(provider)
            
            if self.available_email_providers:
                self.logger.info(f"📧 Почтовая система инициализирована. Доступные провайдеры: {', '.join(self.available_email_providers)}")
            else:
                self.logger.warning("⚠️ Почтовые провайдеры не настроены. Проверьте переменные окружения.")
                
        except Exception as e:
            self.logger.error(f"❌ Ошибка инициализации почтовой системы: {e}")
            self.email_config = {}
            self.available_email_providers = []

    def send_email(self, provider: str, to_email: str, subject: str, body: str, attachments: Optional[List[str]] = None, reply_to: Optional[str] = None):
        """
        Отправляет email через указанного провайдера
        
        Args:
            provider: провайдер (gmail, outlook, yandex, mail_ru)
            to_email: получатель
            subject: тема письма
            body: текст письма
            attachments: список файлов для прикрепления
            reply_to: ID письма для ответа (опционально)
        """
        try:
            if provider not in self.available_email_providers:
                return f"❌ Провайдер {provider} не настроен или недоступен"
            
            config = self.email_config[provider]
            
            # Создаем сообщение
            msg = MIMEMultipart()
            msg['From'] = config['email']
            msg['To'] = to_email
            msg['Subject'] = subject
            
            if reply_to:
                msg['In-Reply-To'] = reply_to
                msg['References'] = reply_to
            
            # Добавляем текст
            msg.attach(MIMEText(body, 'plain', 'utf-8'))
            
            # Добавляем вложения
            if attachments:
                for file_path in attachments:
                    if os.path.exists(file_path):
                        with open(file_path, 'rb') as attachment:
                            part = MIMEApplication(attachment.read(), Name=os.path.basename(file_path))
                            part['Content-Disposition'] = f'attachment; filename="{os.path.basename(file_path)}"'
                            msg.attach(part)
            
            # Отправляем письмо
            with smtplib.SMTP(config['smtp_server'], config['smtp_port']) as server:
                server.starttls()
                server.login(config['email'], config['app_password'])
                server.send_message(msg)
            
            self.logger.info(f"✅ Письмо отправлено через {provider} на {to_email}")
            return f"✅ Письмо успешно отправлено на {to_email}"
            
        except Exception as e:
            error_msg = f"❌ Ошибка отправки письма через {provider}: {e}"
            self.logger.error(error_msg)
            return error_msg

    def get_emails(self, provider: str, folder: str = 'INBOX', limit: int = 10, search_criteria: str = 'ALL'):
        """
        Получает список писем из почтового ящика
        
        Args:
            provider: провайдер (gmail, outlook, yandex, mail_ru)
            folder: папка (INBOX, SENT, DRAFT и т.д.)
            limit: количество писем для получения
            search_criteria: критерии поиска (ALL, UNSEEN, FROM "email", SUBJECT "text" и т.д.)
        """
        try:
            if provider not in self.available_email_providers:
                return f"❌ Провайдер {provider} не настроен или недоступен"
            
            config = self.email_config[provider]
            emails = []
            
            with imaplib.IMAP4_SSL(config['imap_server'], config['imap_port']) as imap:
                imap.login(config['email'], config['app_password'])
                imap.select(folder)
                
                # Поиск писем
                status, messages = imap.search(None, search_criteria)
                if status != 'OK':
                    return f"❌ Ошибка поиска писем: {status}"
                
                message_ids = messages[0].split()
                
                # Ограничиваем количество и берем самые новые
                message_ids = message_ids[-limit:] if len(message_ids) > limit else message_ids
                message_ids.reverse()  # Самые новые сначала
                
                for msg_id in message_ids:
                    status, msg_data = imap.fetch(msg_id, '(RFC822)')
                    if status == 'OK' and msg_data and msg_data[0] and len(msg_data[0]) > 1:
                        email_body = msg_data[0][1]
                        if isinstance(email_body, bytes):
                            email_message = email.message_from_bytes(email_body)
                        
                            # Извлекаем основную информацию
                            email_info = {
                                'id': msg_id.decode(),
                                'from': email_message.get('From'),
                                'to': email_message.get('To'),
                                'subject': email_message.get('Subject'),
                                'date': email_message.get('Date'),
                                'body': self._extract_email_body(email_message),
                                'attachments': self._get_email_attachments_info(email_message)
                            }
                            emails.append(email_info)
            
            self.logger.info(f"📧 Получено {len(emails)} писем из {folder} ({provider})")
            return emails
            
        except Exception as e:
            error_msg = f"❌ Ошибка получения писем из {provider}: {e}"
            self.logger.error(error_msg)
            return error_msg

    def reply_to_email(self, provider: str, original_email_id: str, reply_text: str, attachments: Optional[List[str]] = None):
        """
        Отвечает на письмо
        
        Args:
            provider: провайдер
            original_email_id: ID оригинального письма
            reply_text: текст ответа
            attachments: вложения
        """
        try:
            if provider not in self.available_email_providers:
                return f"❌ Провайдер {provider} не настроен или недоступен"
            
            config = self.email_config[provider]
            
            # Получаем оригинальное письмо
            with imaplib.IMAP4_SSL(config['imap_server'], config['imap_port']) as imap:
                imap.login(config['email'], config['app_password'])
                imap.select('INBOX')
                
                status, msg_data = imap.fetch(original_email_id, '(RFC822)')
                if status == 'OK' and msg_data and msg_data[0] and len(msg_data[0]) > 1:
                    email_body = msg_data[0][1]
                    if isinstance(email_body, bytes):
                        original_message = email.message_from_bytes(email_body)
                        
                        # Формируем ответ
                        original_from = original_message.get('From')
                        original_subject = original_message.get('Subject', '')
                        reply_subject = f"Re: {original_subject}" if not original_subject.startswith('Re:') else original_subject
                        message_id = original_message.get('Message-ID')
                        
                        # Проверяем, что у нас есть адрес получателя
                        if not original_from:
                            return "❌ Не удалось получить адрес отправителя оригинального письма"
                        
                        # Отправляем ответ
                        return self.send_email(
                            provider=provider,
                            to_email=original_from,
                            subject=reply_subject,
                            body=reply_text,
                            attachments=attachments,
                            reply_to=message_id
                        )
                else:
                    return f"❌ Не удалось получить оригинальное письмо с ID {original_email_id}"
            
        except Exception as e:
            error_msg = f"❌ Ошибка ответа на письмо: {e}"
            self.logger.error(error_msg)
            return error_msg

    def _extract_email_body(self, email_message):
        """
        Извлекает текст письма из объекта email
        """
        try:
            if email_message.is_multipart():
                for part in email_message.walk():
                    content_type = part.get_content_type()
                    if content_type == "text/plain":
                        body = part.get_payload(decode=True)
                        if body:
                            return body.decode('utf-8', errors='ignore')
            else:
                body = email_message.get_payload(decode=True)
                if body:
                    return body.decode('utf-8', errors='ignore')
            return "Не удалось извлечь текст письма"
        except Exception as e:
            return f"Ошибка извлечения текста: {e}"

    def _get_email_attachments_info(self, email_message):
        """
        Получает информацию о вложениях письма
        """
        attachments = []
        try:
            if email_message.is_multipart():
                for part in email_message.walk():
                    if part.get_content_disposition() == 'attachment':
                        filename = part.get_filename()
                        if filename:
                            attachments.append({
                                'filename': filename,
                                'size': len(part.get_payload(decode=True)) if part.get_payload(decode=True) else 0
                            })
        except Exception as e:
            self.logger.warning(f"Ошибка получения информации о вложениях: {e}")
        return attachments

    def search_emails(self, provider: str, query: str, folder: str = 'INBOX', limit: int = 20):
        """
        Поиск писем по различным критериям
        
        Args:
            provider: провайдер
            query: поисковый запрос (может быть текстом для поиска в теме/тексте)
            folder: папка для поиска
            limit: максимальное количество результатов
        """
        try:
            # Формируем IMAP критерии поиска
            search_criteria = f'(OR SUBJECT "{query}" BODY "{query}")'
            
            emails = self.get_emails(provider, folder, limit, search_criteria)
            
            if isinstance(emails, list):
                self.logger.info(f"🔍 Найдено {len(emails)} писем по запросу '{query}'")
                return emails
            else:
                return emails  # Возвращаем сообщение об ошибке
                
        except Exception as e:
            error_msg = f"❌ Ошибка поиска писем: {e}"
            self.logger.error(error_msg)
            return error_msg

    def _trim_context_if_needed(self):
        """
        Обрезает контекст если он превышает безопасные лимиты
        Использует self.current_context_length (total_tokens) для проверки
        """
        if self.current_context_length > self.max_context_length:
            # Критический лимит - агрессивная обрезка
            self.conversation_history = self.conversation_history[-2:]  # Оставляем только 2 последних сообщения
            self.logger.warning(f"Критическое превышение контекста ({self.current_context_length:,} > {self.max_context_length:,}) - агрессивная обрезка истории")
        elif self.current_context_length > self.safe_context_length:
            # Превышение безопасного лимита - аккуратная обрезка
            self.conversation_history = self.conversation_history[-5:]  # Оставляем только 5 последних сообщений
            self.logger.warning(f"Превышение безопасного контекста ({self.current_context_length:,} > {self.safe_context_length:,}) - аккуратная обрезка истории")
        elif self.current_context_length > self.safe_context_length * 0.8:
            # Приближение к безопасному лимиту - профилактическая обрезка
            self.conversation_history = self.conversation_history[-10:]  # Оставляем только 10 последних сообщений
            self.logger.info(f"Приближение к безопасному лимиту ({self.current_context_length:,} > {self.safe_context_length * 0.8:,}) - профилактическая обрезка истории")

    def launch_model(self, model_path: str):
        """
        Запускает модель через LM Studio (локально, subprocess)
        """
        try:
            # threading уже импортирован в начале файла
            lmstudio_exe = os.getenv("LMSTUDIO_EXE", r"C:\Program Files\LM Studio\LM Studio.exe")
            self.logger.info(f"Запускаю модель: {model_path}")
            threading.Thread(target=lambda: os.system(f'"{lmstudio_exe}" --model "{model_path}"'), daemon=True).start()
        except Exception as e:
            self.logger.error(f"Ошибка запуска модели: {e}")

    def ask_qwen(self, question: str) -> Optional[str]:
        """Запрос к Qwen для генерации промтов изображений
        Использует основной мозг для генерации промтов
        """
        # Используем основной мозг для генерации промтов
        image_model = self.brain_model_id if hasattr(self, 'brain_model_id') and self.brain_model_id else self.brain_model
        payload = {
            "model": image_model,
            "messages": [
                {"role": "system", "content": "Ты — ассистент для генерации идеальных промтов для Stable Diffusion. Твоя задача — создать идеальный промт для генерации изображения на основе запроса пользователя. ВАЖНО: prompt и negative_prompt должны быть ТОЛЬКО на английском языке, иначе будет ошибка! ВСЕГДА включай negative_prompt - это обязательное поле! Формируй промт и настройки строго в формате JSON: {\"prompt\":..., \"negative_prompt\":..., \"params\":{...}}. Пример negative_prompt: '(worst quality, low quality, normal quality:1.4), (deformed, distorted, disfigured:1.3), poorly drawn, bad anatomy'. Не добавляй ничего лишнего!"},
                {"role": "user", "content": f"Вопрос: {question}\n\nВАЖНО: prompt и negative_prompt должны быть ТОЛЬКО на английском языке! ОБЯЗАТЕЛЬНО включи negative_prompt в JSON!"}
            ],
            "temperature": 0.2,
            "max_tokens": 1024,
            "stream": False
        }
        try:
            resp = requests.post(f"{self.lm_studio_url}/v1/chat/completions", json=payload, headers={"Content-Type": "application/json"})
            if resp.status_code == 200:
                result = resp.json()
                return result["choices"][0]["message"]["content"].strip()
            else:
                self.logger.error(f"Ошибка Qwen: {resp.status_code} - {resp.text}")
                return None
        except Exception as e:
            self.logger.error(f"Ошибка запроса к Qwen: {e}")
            return None

    def get_youtube_cookies_path(self) -> Optional[str]:
        """
        Получает путь к файлу cookies для YouTube
        Возвращает путь к файлу или None если файл не найден
        """
        cookies_file = "youtube_cookies.txt"
        
        # Сначала ищем в текущей рабочей директории
        cookies_path = os.path.join(os.getcwd(), cookies_file)
        if os.path.exists(cookies_path) and os.path.getsize(cookies_path) > 0:
            self.logger.info(f"🍪 Найден файл cookies в рабочей директории: {cookies_file}")
            return cookies_path
        
        # Затем ищем в директории скрипта
        cookies_path = os.path.join(os.path.dirname(__file__), cookies_file)
        if os.path.exists(cookies_path) and os.path.getsize(cookies_path) > 0:
            self.logger.info(f"🍪 Найден файл cookies в директории скрипта: {cookies_file}")
            return cookies_path
        
        # Если файл не найден нигде
        self.logger.info(f"ℹ️ Файл cookies не найден: {cookies_file}")
        return None

    def check_cookies_validity(self, cookies_path: str) -> bool:
        """
        Проверяет валидность файла cookies
        """
        try:
            with open(cookies_path, 'r', encoding='utf-8') as f:
                content = f.read()
                
            # Проверяем базовую структуру
            if not content.strip():
                return False
                
            # Проверяем наличие YouTube доменов
            youtube_domains = ['youtube.com', '.youtube.com', 'google.com', '.google.com']
            has_youtube = any(domain in content for domain in youtube_domains)
            
            if not has_youtube:
                self.logger.warning("⚠️ В файле cookies не найдены домены YouTube")
                return False
                
            # Проверяем формат (должен содержать табуляции)
            if '\t' not in content:
                self.logger.warning("⚠️ Неверный формат файла cookies (отсутствуют табуляции)")
                return False
                
            self.logger.info("✅ Файл cookies валиден")
            return True
            
        except Exception as e:
            self.logger.error(f"❌ Ошибка проверки cookies: {e}")
            return False

    def suggest_cookies_update(self):
        """
        Предлагает пользователю обновить cookies
        """
        self.logger.info("💡 Для улучшения работы с YouTube рекомендуется:")
        self.logger.info("   1. Запустить: python extract_chrome_cookies.py")
        self.logger.info("   2. Закрыть Chrome перед извлечением")
        self.logger.info("   3. Войти в YouTube через VPN")
        self.logger.info("   4. Cookies обновляются каждые 2-3 месяца")

    def generate_image_stable_diffusion(self, prompt: str, negative_prompt: str, params: dict) -> Optional[str]:
        """Генерация изображения через прямую интеграцию со Stable Diffusion"""
        start_time = time.time()
        
        # Логируем полученные параметры
        self.logger.info(f"🔧 Получены параметры генерации: prompt='{prompt[:50]}...', negative_prompt='{negative_prompt}'")
        
        # Горячая перезагрузка конфигурации LoRA
        self.model_manager.get_lora_config(force_reload=True)
        
        # Автоматически включаем генерацию изображений при необходимости
        if not getattr(self, 'use_image_generation', False):
            self.logger.info("🔧 Автоматически включаю генерацию изображений")
            self.use_image_generation = True
            # Запускаем таймер автоматического выключения
            self.auto_disable_tools("image_generation")
        
        # Получаем путь к модели через ModelManager (приоритет .env > stable_diff)
        model_path = self.model_manager.get_model_path()
        if not model_path:
            self.logger.error("❌ Не удалось найти Stable Diffusion модель")
            return None
        
        if not os.path.exists(model_path):
            self.logger.error(f"❌ Модель не найдена: {model_path}")
            return None
        
        # Определяем тип модели
        model_type = self.model_manager.detect_model_type(model_path)
        self.logger.info(f"🔍 Определен тип модели: {model_type} для {os.path.basename(model_path)}")
        
        # Применяем LoRA триггеры к промпту
        enhanced_prompt = self.model_manager.apply_lora_triggers(prompt, model_type)
        
        # Параметры по умолчанию (будут обновлены в зависимости от типа модели)
        default_params = {
            "seed": -1,
            "steps": 30,
            "width": 1024,  # Временно, будет обновлено ниже
            "height": 1024,  # Временно, будет обновлено ниже
            "cfg": 7.0,
            "sampler_name": "dpmpp_2m",
            "scheduler": "karras"
        }
        
        # Обновляем параметры пользовательскими значениями
        gen_params = default_params.copy()
        gen_params.update(params)
        
        # Исправляем seed если он -1
        if gen_params["seed"] == -1:
            import random
            gen_params["seed"] = random.randint(0, 2**32 - 1)
            self.logger.info(f"🎲 Сгенерирован случайный seed: {gen_params['seed']}")
        
        # Определяем тип модели для корректировки размеров
        model_path = os.getenv("STABLE_DIFFUSION_MODEL_PATH")
        if model_path:
            model_name = os.path.basename(model_path).lower()
            is_sdxl = any(keyword in model_name for keyword in ['xl', 'sdxl', 'illustrious', 'pony'])
            
            # Обновляем размеры по умолчанию в зависимости от типа модели (только если не заданы пользователем)
            if not params.get("width") and not params.get("height"):
                if is_sdxl:
                    # SDXL модели работают лучше с 1024x1024 (уже установлено по умолчанию)
                    pass
                else:
                    # SD 1.5 модели работают лучше с 512x512
                    gen_params["width"] = 512
                    gen_params["height"] = 512
                    self.logger.info("📐 Автоматически установил размеры для SD 1.5 модели: 512x512")
        
        self.logger.info(f"🔧 Параметры генерации: {gen_params}")
        
        try:
            # Устанавливаем необходимые зависимости
            self._install_diffusers_dependencies()
            
            # Импортируем необходимые библиотеки (рекомендованные подмодули для совместимости с Pylance)
            from diffusers.pipelines.stable_diffusion.pipeline_stable_diffusion import StableDiffusionPipeline  # type: ignore
            from diffusers.pipelines.stable_diffusion_xl.pipeline_stable_diffusion_xl import StableDiffusionXLPipeline  # type: ignore
            from diffusers.schedulers.scheduling_dpmsolver_multistep import DPMSolverMultistepScheduler  # type: ignore
            import torch
            
            self.logger.info(f"📦 Загружаю модель: {model_path}")
            
            # Загружаем соответствующий pipeline
            is_sdxl = (model_type == 'sdxl')
            if is_sdxl:
                self.logger.info("🎯 Обнаружена SDXL модель, использую StableDiffusionXLPipeline")
                pipe = StableDiffusionXLPipeline.from_single_file(
                    model_path,
                    torch_dtype=torch.float16,
                    use_safetensors=True
                )
            else:
                self.logger.info("🎯 Обнаружена SD 1.5 модель, использую StableDiffusionPipeline")
                pipe = StableDiffusionPipeline.from_single_file(
                    model_path,
                    torch_dtype=torch.float16,
                    use_safetensors=True
                )
            
            # Перемещаем на GPU если доступен
            if torch.cuda.is_available():
                pipe = pipe.to("cuda")
                self.logger.info("🚀 Модель перемещена на GPU")
            else:
                self.logger.warning("⚠️ GPU недоступен, использую CPU")
            
            # Загружаем активные LoRA для данного типа модели
            active_loras = self.model_manager.get_active_loras(model_type)
            if active_loras:
                self.logger.info(f"🎭 Найдено {len(active_loras)} активных LoRA для типа {model_type}")
                
                # Проверяем доступность PEFT
                peft_available = False
                try:
                    import peft
                    peft_available = True
                    self.logger.info(f"✅ PEFT версии {peft.__version__} доступен")
                except ImportError as e:
                    self.logger.warning(f"⚠️ PEFT не установлен: {e}")
                    self.logger.warning("   LoRA файлы в формате safetensors могут не работать")
                except Exception as e:
                    self.logger.warning(f"⚠️ Ошибка импорта PEFT: {e}")
                    self.logger.warning("   LoRA файлы в формате safetensors могут не работать")
                
                loaded_loras = []
                for lora in active_loras:
                    try:
                        lora_filename = lora.get('filename', '')
                        lora_strength = lora.get('strength', 1.0)
                        
                        # Определяем путь к LoRA файлу
                        lora_path = os.path.join(self.model_manager.lora_dir, model_type, lora_filename)
                        
                        if os.path.exists(lora_path):
                            # Проверяем формат файла
                            file_ext = os.path.splitext(lora_filename)[1].lower()
                            
                            if file_ext == '.safetensors':
                                # Современный формат - используем правильный метод
                                adapter_name = os.path.splitext(lora_filename)[0]
                                
                                if not peft_available:
                                    self.logger.warning(f"⚠️ Пропускаю LoRA {lora_filename} - PEFT не доступен")
                                    continue
                                
                                try:
                                    # Метод 1: Пробуем загрузить с указанием папки и имени файла
                                    try:
                                        pipe.load_lora_weights(os.path.dirname(lora_path), weight_name=lora_filename, adapter_name=adapter_name)
                                        loaded_loras.append((adapter_name, lora_strength))
                                        self.logger.info(f"✅ Загружена LoRA (метод 1/folder): {lora_filename} (сила: {lora_strength})")
                                    except Exception as e1:
                                        # Метод 2: Пробуем загрузить напрямую с именем адаптера
                                        try:
                                            pipe.load_lora_weights(lora_path, adapter_name=adapter_name)
                                            loaded_loras.append((adapter_name, lora_strength))
                                            self.logger.info(f"✅ Загружена LoRA (метод 2/direct): {lora_filename} (сила: {lora_strength})")
                                        except Exception as e2:
                                            # Метод 3: Пробуем загрузить без имени адаптера (legacy)
                                            try:
                                                pipe.load_lora_weights(lora_path)
                                                loaded_loras.append((lora_filename, lora_strength))
                                                self.logger.info(f"✅ Загружена LoRA (метод 3/legacy): {lora_filename} (сила: {lora_strength})")
                                            except Exception as e3:
                                                # Все методы не сработали
                                                self.logger.error(f"❌ Не удалось загрузить safetensors LoRA {lora_filename}")
                                                self.logger.error(f"   Метод 1 (folder): {str(e1)[:100]}...")
                                                self.logger.error(f"   Метод 2 (direct): {str(e2)[:100]}...")  
                                                self.logger.error(f"   Метод 3 (legacy): {str(e3)[:100]}...")
                                                
                                                # Дополнительная диагностика
                                                if "PEFT" in str(e3):
                                                    self.logger.error("   💡 Рекомендация: убедитесь что PEFT установлен: pip install peft")
                                                    self.logger.error("   📋 Или попробуйте конвертировать LoRA в формат .ckpt")
                                                continue
                                        
                                except Exception as e:
                                    self.logger.error(f"❌ Общая ошибка загрузки safetensors LoRA {lora_filename}: {e}")
                                    continue
                            else:
                                # Старый формат (.ckpt, .pt)
                                try:
                                    pipe.load_lora_weights(lora_path)
                                    loaded_loras.append((lora_filename, lora_strength))
                                    self.logger.info(f"✅ Загружена LoRA (legacy): {lora_filename} (сила: {lora_strength})")
                                except Exception as e:
                                    self.logger.error(f"❌ Не удалось загрузить legacy LoRA {lora_filename}: {e}")
                                    continue
                        else:
                            self.logger.warning(f"⚠️ LoRA файл не найден: {lora_path}")
                    
                    except Exception as e:
                        self.logger.error(f"❌ Ошибка загрузки LoRA {lora.get('filename', 'unknown')}: {e}")
                        continue
                
                # Применяем силу LoRA если есть загруженные адаптеры
                if loaded_loras:
                    try:
                        if hasattr(pipe, 'set_adapters') and len(loaded_loras) > 0:
                            adapter_names = [name for name, _ in loaded_loras]
                            adapter_weights = [weight for _, weight in loaded_loras]
                            pipe.set_adapters(adapter_names, adapter_weights=adapter_weights)
                            self.logger.info(f"⚙️ Настроены веса адаптеров: {dict(loaded_loras)}")
                    except Exception as e:
                        self.logger.warning(f"⚠️ Не удалось настроить веса адаптеров: {e}")
                        
            else:
                self.logger.info(f"📝 Нет активных LoRA для типа модели {model_type}")
            
            # Сохраняем pipeline для последующей выгрузки
            self.current_pipeline = pipe
            
            # Настраиваем scheduler
            if gen_params["sampler_name"] == "dpmpp_2m":
                pipe.scheduler = DPMSolverMultistepScheduler.from_config(pipe.scheduler.config)
                self.logger.info("⚙️ Использую DPMSolverMultistepScheduler")
            
            # Генерируем изображение
            self.logger.info(f"🎨 Генерирую изображение: {enhanced_prompt[:50]}...")

            result = pipe(
                prompt=enhanced_prompt,
                negative_prompt=negative_prompt,
                num_inference_steps=gen_params["steps"],
                guidance_scale=gen_params["cfg"],
                width=gen_params["width"],
                height=gen_params["height"],
                generator=torch.Generator(device="cuda" if torch.cuda.is_available() else "cpu").manual_seed(gen_params["seed"])
            )

            # Получаем изображение: результат pipe может быть объектом с атрибутом images или кортежем (image, extras)
            image = None
            try:
                imgs = getattr(result, 'images', None)
                if imgs:
                    image = imgs[0]
                elif isinstance(result, (tuple, list)) and len(result) > 0:
                    image = result[0]
            except Exception:
                image = None

            if image is None:
                raise RuntimeError('Не удалось получить изображение из результата pipeline')

            # Нормализация: если image — numpy array или torch tensor, конвертируем в PIL.Image
            try:
                import numpy as _np
                try:
                    import torch as _torch
                except Exception:
                    _torch = None

                if hasattr(image, 'save'):
                    img_to_save = image
                elif _np and isinstance(image, _np.ndarray):
                    img_to_save = Image.fromarray(image.astype('uint8'))
                elif _torch and _torch.is_tensor(image):
                    arr = image.detach().cpu().numpy()
                    img_to_save = Image.fromarray(arr.astype('uint8'))
                else:
                    # Попытка сконвертировать общим способом
                    img_to_save = Image.fromarray(_np.array(image).astype('uint8'))
            except Exception:
                # В крайнем случае пытаемся работать напрямую — пусть вызов .save выбросит понятную ошибку
                img_to_save = image
            # Подсказка для статического анализатора: гарантируем, что img_to_save рассматривается как PIL.Image
            try:
                from typing import cast
                img_to_save = cast(Image.Image, img_to_save)
            except Exception:
                pass
            
            # Сохраняем изображение
            output_dir = os.path.join(os.path.dirname(__file__), "Images", "generated")
            os.makedirs(output_dir, exist_ok=True)
            
            filename = f"ConsoleTest_{gen_params['seed']}.png"
            output_path = os.path.join(output_dir, filename)
            
            # Сохраняем PIL.Image
            try:
                # Если уже PIL.Image
                if isinstance(img_to_save, Image.Image):
                    img_to_save.save(output_path)
                    self.logger.info(f"💾 Изображение сохранено: {output_path}")
                else:
                    # Пытаемся сохранить оригинальный объект как PIL
                    if isinstance(image, Image.Image):
                        image.save(output_path)
                        self.logger.info(f"💾 Изображение сохранено (fallback): {output_path}")
                    else:
                        import numpy as _np
                        Image.fromarray(_np.array(image).astype('uint8')).save(output_path)
                        self.logger.info(f"💾 Изображение сохранено (converted fallback): {output_path}")
                        
                # Выгружаем pipeline сразу после сохранения для освобождения VRAM
                self._unload_current_pipeline()
                
                # Проверяем наличие модели RealESRGAN и применяем апскейл если доступен
                if self._is_realesrgan_available():
                    self.logger.info("🔍 Модель RealESRGAN найдена, применяю апскейл...")
                    upscaled_path = self.upscale_image_realesrgan(output_path)
                    if upscaled_path and os.path.exists(upscaled_path):
                        self.logger.info("✨ Использую увеличенное изображение")
                        # Обновляем img_to_save для финального base64
                        img_to_save = Image.open(upscaled_path)
                        # Также показываем пользователю увеличенную версию
                        output_path = upscaled_path
                    else:
                        self.logger.warning("⚠️ Апскейл не удался, использую оригинальное изображение")
                else:
                    self.logger.info("ℹ️ Модель RealESRGAN не найдена, пропускаю апскейл")
                    
            except Exception:
                self.logger.error(f"❌ Не удалось сохранить изображение ни одним из способов")
            
            # Автоматически открываем изображение (отключено, т.к. открывается через show_image_base64_temp)
            # try:
            #     subprocess.run(["start", output_path], shell=True, check=True)
            #     self.logger.info("🖼️ Изображение автоматически открыто")
            # except Exception as e:
            #     self.logger.warning(f"⚠️ Не удалось открыть изображение: {e}")
            
            # Конвертируем в base64
            buf = BytesIO()
            try:
                if isinstance(img_to_save, Image.Image):
                    img_to_save.save(buf, format="PNG")
                elif isinstance(image, Image.Image):
                    image.save(buf, format="PNG")
                else:
                    import numpy as _np
                    Image.fromarray(_np.array(image).astype('uint8')).save(buf, format="PNG")
            except Exception:
                self.logger.error("❌ Не удалось сконвертировать изображение в буфер PNG")
            img_b64 = base64.b64encode(buf.getvalue()).decode("ascii")
            
            return img_b64
            
        except Exception as e:
            self.logger.error(f"❌ Ошибка генерации изображения: {e}")
            import traceback
            traceback.print_exc()
            return None
        finally:
            # Записываем метрику производительности
            response_time = time.time() - start_time
            self.add_performance_metric("image_generation", response_time)
            self.logger.info(f"🎨 Изображение сгенерировано за {response_time:.2f} сек")

    def _unload_current_pipeline(self):
        """Выгружает текущий pipeline для экономии VRAM"""
        try:
            if hasattr(self, 'current_pipeline') and self.current_pipeline is not None:
                self.logger.info("🔄 Выгружаю pipeline для экономии VRAM...")
                
                # Перемещаем на CPU
                if hasattr(self.current_pipeline, 'to'):
                    self.current_pipeline.to('cpu')
                
                # Удаляем pipeline
                del self.current_pipeline
                self.current_pipeline = None
                
                # Принудительная очистка памяти GPU
                try:
                    import torch
                    if torch.cuda.is_available():
                        torch.cuda.empty_cache()
                        self.logger.info("🧹 Очищен кэш CUDA")
                except Exception as e:
                    self.logger.warning(f"⚠️ Не удалось очистить CUDA кэш: {e}")
                
                self.logger.info("✅ Pipeline выгружен")
        except Exception as e:
            self.logger.warning(f"⚠️ Ошибка выгрузки pipeline: {e}")

    def _is_realesrgan_available(self) -> bool:
        """Проверяет доступность модели RealESRGAN"""
        try:
            base_dir = os.path.dirname(os.path.abspath(__file__))
            model_path = os.path.join(base_dir, "stable_diff", "RealESRGAN_x4.pth")
            return os.path.exists(model_path)
        except Exception:
            return False

    def upscale_image_realesrgan(self, image_path: str, output_path: Optional[str] = None) -> Optional[str]:
        """
        Увеличивает изображение в 4 раза с помощью RealESRGAN
        
        Args:
            image_path: Путь к исходному изображению
            output_path: Путь для сохранения результата (опционально)
        
        Returns:
            Путь к увеличенному изображению или None в случае ошибки
        """
        try:
            self.logger.info(f"📈 Начинаю апскейл изображения: {os.path.basename(image_path)}")
            
            # Путь к модели RealESRGAN
            base_dir = os.path.dirname(os.path.abspath(__file__))
            model_path = os.path.join(base_dir, "stable_diff", "RealESRGAN_x4.pth")
            
            if not os.path.exists(model_path):
                self.logger.info(f"ℹ️ Модель RealESRGAN не найдена: {model_path}")
                self.logger.info("💡 Поместите файл RealESRGAN_x4.pth в папку stable_diff для включения апскейла")
                return None
            
            # Проверяем исходное изображение
            if not os.path.exists(image_path):
                self.logger.error(f"❌ Исходное изображение не найдено: {image_path}")
                return None
            
            # Определяем выходной путь
            if output_path is None:
                base_name = os.path.splitext(os.path.basename(image_path))[0]
                output_dir = os.path.dirname(image_path)
                output_path = os.path.join(output_dir, f"{base_name}_upscaled_4x.png")
            
            # Устанавливаем Real-ESRGAN если нужно
            self._install_realesrgan_dependencies()
            
            # Импортируем библиотеки
            try:
                import cv2
                import torch
                import numpy as np
                from PIL import Image
                
                # Пытаемся импортировать RealESRGAN
                try:
                    from realesrgan import RealESRGANer
                    from basicsr.archs.rrdbnet_arch import RRDBNet
                except ImportError:
                    # Если realesrgan не установлен, пытаемся использовать базовую реализацию
                    self.logger.warning("⚠️ realesrgan пакет не найден, использую альтернативный метод")
                    return self._upscale_image_alternative(image_path, output_path)
                
                # Настраиваем модель
                model = RRDBNet(
                    num_in_ch=3, 
                    num_out_ch=3, 
                    num_feat=64, 
                    num_block=23, 
                    num_grow_ch=32, 
                    scale=4
                )
                
                # Создаем upsampler
                upsampler = RealESRGANer(
                    scale=4,
                    model_path=model_path,
                    model=model,
                    tile=0,
                    tile_pad=10,
                    pre_pad=0,
                    half=torch.cuda.is_available()
                )
                
                # Загружаем изображение
                img = cv2.imread(image_path, cv2.IMREAD_COLOR)
                if img is None:
                    raise ValueError(f"Не удалось загрузить изображение: {image_path}")
                
                self.logger.info(f"📐 Исходный размер: {img.shape[1]}x{img.shape[0]}")
                
                # Выполняем апскейл
                self.logger.info("🚀 Выполняю апскейл...")
                output, _ = upsampler.enhance(img, outscale=4)
                
                # Сохраняем результат
                cv2.imwrite(output_path, output)
                
                self.logger.info(f"📐 Результирующий размер: {output.shape[1]}x{output.shape[0]}")
                self.logger.info(f"💾 Апскейл сохранен: {output_path}")
                
                return output_path
                
            except Exception as e:
                self.logger.error(f"❌ Ошибка в процессе апскейла: {e}")
                return self._upscale_image_alternative(image_path, output_path)
                
        except Exception as e:
            self.logger.error(f"❌ Общая ошибка апскейла: {e}")
            return None
    
    def _upscale_image_alternative(self, image_path: str, output_path: str) -> Optional[str]:
        """
        Альтернативный метод апскейла с помощью простого бикубического интерполирования
        """
        try:
            self.logger.info("🔄 Использую альтернативный метод апскейла...")
            
            from PIL import Image
            
            # Загружаем изображение
            with Image.open(image_path) as img:
                original_size = img.size
                new_size = (original_size[0] * 4, original_size[1] * 4)
                
                # Увеличиваем с помощью бикубической интерполяции
                upscaled = img.resize(new_size, Image.Resampling.LANCZOS)
                
                # Сохраняем результат
                upscaled.save(output_path, "PNG")
                
                self.logger.info(f"📐 Увеличено с {original_size} до {new_size}")
                self.logger.info(f"💾 Альтернативный апскейл сохранен: {output_path}")
                
                return output_path
                
        except Exception as e:
            self.logger.error(f"❌ Ошибка альтернативного апскейла: {e}")
            return None
    
    def _install_realesrgan_dependencies(self):
        """Устанавливает зависимости для RealESRGAN"""
        try:
            # Проверяем установлен ли basicsr
            try:
                import basicsr
                self.logger.debug("✅ basicsr уже установлен")
            except ImportError:
                self.logger.info("📦 Устанавливаю basicsr...")
                subprocess.run([_sys.executable, '-m', 'pip', 'install', 'basicsr'], 
                             check=True, capture_output=True)
            
            # Проверяем установлен ли realesrgan
            try:
                import realesrgan
                self.logger.debug("✅ realesrgan уже установлен")
            except ImportError:
                self.logger.info("📦 Устанавливаю realesrgan...")
                subprocess.run([_sys.executable, '-m', 'pip', 'install', 'realesrgan'], 
                             check=True, capture_output=True)
                             
        except Exception as e:
            self.logger.warning(f"⚠️ Не удалось установить зависимости RealESRGAN: {e}")

    def generate_video_stable_diffusion(self, prompt: str, negative_prompt: str, params: dict) -> Optional[str]:
        """Генерация видео через прямую интеграцию со Stable Diffusion"""
        start_time = time.time()
        
        # Автоматически включаем генерацию изображений при необходимости
        if not getattr(self, 'use_image_generation', False):
            self.logger.info("🔧 Автоматически включаю генерацию изображений")
            self.use_image_generation = True
            # Запускаем таймер автоматического выключения
            self.auto_disable_tools("image_generation")
        
        # Параметры по умолчанию для видео
        default_params = {
            "seed": -1,
            "steps": 20,
            "width": 512,
            "height": 512,
            "cfg": 7.0,
            "num_frames": 24,
            "fps": 8,
            "key_frames": 4
        }
        
        # Обновляем параметры пользовательскими значениями
        gen_params = default_params.copy()
        gen_params.update(params)
        
        # Исправляем seed если он -1
        if gen_params["seed"] == -1:
            import random
            gen_params["seed"] = random.randint(0, 2**32 - 1)
            self.logger.info(f"🎲 Сгенерирован случайный seed: {gen_params['seed']}")
        
        self.logger.info(f"🔧 Параметры генерации видео: {gen_params}")
        
        try:
            # Устанавливаем необходимые зависимости
            self._install_diffusers_dependencies()
            
            # Импортируем необходимые библиотеки (рекомендованные подмодули для совместимости с Pylance)
            from diffusers.pipelines.stable_diffusion.pipeline_stable_diffusion import StableDiffusionPipeline  # type: ignore
            from diffusers.pipelines.stable_diffusion_xl.pipeline_stable_diffusion_xl import StableDiffusionXLPipeline  # type: ignore
            from diffusers.schedulers.scheduling_dpmsolver_multistep import DPMSolverMultistepScheduler  # type: ignore
            import torch
            from PIL import Image
            import numpy as np
            import imageio  # type: ignore
            
            # Путь к модели из .env файла
            model_path = os.getenv("STABLE_DIFFUSION_MODEL_PATH")
            if not model_path:
                self.logger.error("❌ STABLE_DIFFUSION_MODEL_PATH не указан в .env файле")
                return None
            
            # Проверяем существование модели
            if not os.path.exists(model_path):
                self.logger.error(f"❌ Модель не найдена: {model_path}")
                return None
            
            self.logger.info(f"📦 Загружаю модель: {model_path}")
            
            # Определяем тип модели по имени файла (SDXL модели обычно содержат xl, sdxl, illustrious в названии)
            model_name = os.path.basename(model_path).lower()
            is_sdxl = any(keyword in model_name for keyword in ['xl', 'sdxl', 'illustrious', 'pony'])
            
            # Загружаем соответствующий pipeline
            if is_sdxl:
                self.logger.info("🎯 Обнаружена SDXL модель, использую StableDiffusionXLPipeline")
                pipe = StableDiffusionXLPipeline.from_single_file(
                    model_path,
                    torch_dtype=torch.float16,
                    use_safetensors=True
                )
            else:
                self.logger.info("🎯 Обнаружена SD 1.5 модель, использую StableDiffusionPipeline")
                pipe = StableDiffusionPipeline.from_single_file(
                    model_path,
                    torch_dtype=torch.float16,
                    use_safetensors=True
                )
            
            # Перемещаем на GPU если доступен
            if torch.cuda.is_available():
                pipe = pipe.to("cuda")
                self.logger.info("🚀 Модель перемещена на GPU")
            else:
                self.logger.warning("⚠️ GPU недоступен, использую CPU")
            
            # Сохраняем pipeline для последующей выгрузки
            self.current_pipeline = pipe
            
            # Настраиваем scheduler
            pipe.scheduler = DPMSolverMultistepScheduler.from_config(pipe.scheduler.config)
            self.logger.info("⚙️ Использую DPMSolverMultistepScheduler")
            
            # Параметры генерации
            generation_config = {
                "width": gen_params["width"],
                "height": gen_params["height"],
                "num_inference_steps": gen_params["steps"],
                "guidance_scale": gen_params["cfg"],
                "num_images_per_prompt": 1
            }
            
            self.logger.info(f"🎬 Генерирую {gen_params['num_frames']} кадров для видео...")
            
            frames = []
            key_frames = gen_params["key_frames"]
            
            # Создаем вариации промпта для ключевых кадров
            key_prompts = [
                prompt,
                self._add_dynamic_elements(prompt, 1, key_frames),
                self._add_dynamic_elements(prompt, 2, key_frames),
                self._add_dynamic_elements(prompt, 3, key_frames)
            ]
            
            # Генерируем ключевые кадры
            for i in range(key_frames):
                seed = gen_params["seed"] + i * 50  # Разные seed'ы
                generator = torch.Generator(device="cuda" if torch.cuda.is_available() else "cpu").manual_seed(seed)
                
                with torch.no_grad():
                    result = pipe(
                        prompt=key_prompts[i],
                        negative_prompt=negative_prompt,
                        generator=generator,
                        **generation_config
                    )
                
                # Безопасно извлекаем изображение из результата pipeline
                frame_img = None
                try:
                    imgs = getattr(result, 'images', None)
                    if imgs:
                        frame_img = imgs[0]
                    elif isinstance(result, (tuple, list)) and len(result) > 0:
                        frame_img = result[0]
                except Exception:
                    frame_img = None

                if frame_img is None:
                    raise RuntimeError('Не удалось получить кадр из результата pipeline')

                frames.append(frame_img)
                self.logger.info(f"  ✅ Ключевой кадр {i+1} готов")
            
            # Создаем интерполированные кадры между ключевыми кадрами
            frames_per_segment = gen_params["num_frames"] // (key_frames - 1)
            
            for segment in range(key_frames - 1):
                img1 = np.array(frames[segment])
                img2 = np.array(frames[segment + 1])
                
                for i in range(frames_per_segment):
                    # Вычисляем коэффициент интерполяции
                    t = i / frames_per_segment
                    
                    # Используем более плавную интерполяцию (ease-in-out)
                    t_smooth = 3 * t * t - 2 * t * t * t
                    
                    # Интерполяция между двумя изображениями
                    interpolated_array = img1 * (1 - t_smooth) + img2 * t_smooth
                    
                    # Конвертируем обратно в PIL Image
                    interpolated_image = Image.fromarray(interpolated_array.astype(np.uint8))
                    frames.append(interpolated_image)
                    
                    frame_num = segment * frames_per_segment + i + 1
                    self.logger.info(f"  ✅ Кадр {frame_num}/{gen_params['num_frames']} готов (сегмент {segment+1}, интерполяция: {t_smooth:.2f})")
            
            # Добавляем последний ключевой кадр если нужно
            if len(frames) < gen_params["num_frames"]:
                frames.append(frames[-1])
                self.logger.info(f"  ✅ Добавлен финальный кадр")
            
            frames = frames[:gen_params["num_frames"]]  # Убеждаемся, что возвращаем нужное количество кадров
            
            # Создаем папку для выходных файлов
            output_dir = os.path.join(os.path.dirname(__file__), "Videos", "generated")
            os.makedirs(output_dir, exist_ok=True)
            
            # Сохраняем кадры
            self.logger.info("💾 Сохраняю кадры...")
            for i, frame in enumerate(frames):
                frame_path = os.path.join(output_dir, f"video_frame_{i:03d}.png")
                try:
                    if hasattr(frame, 'save'):
                        try:
                            from typing import cast
                            frame = cast(Image.Image, frame)
                        except Exception:
                            pass
                        frame.save(frame_path)
                    else:
                        # Попытка привести numpy array / tensor к PIL Image
                        import numpy as _np
                        try:
                            import torch as _torch
                        except Exception:
                            _torch = None

                        if _torch and _torch.is_tensor(frame):
                            arr = frame.detach().cpu().numpy()
                            Image.fromarray(arr.astype('uint8')).save(frame_path)
                        elif isinstance(frame, _np.ndarray):
                            Image.fromarray(frame.astype('uint8')).save(frame_path)
                        else:
                            Image.fromarray(_np.array(frame).astype('uint8')).save(frame_path)

                    self.logger.info(f"  💾 Кадр {i+1} сохранен: {frame_path}")
                except Exception as e:
                    self.logger.warning(f"⚠️ Не удалось сохранить кадр {i+1}: {e}")
            
            # Создаем видео
            video_path = os.path.join(output_dir, f"ConsoleVideo_{gen_params['seed']}.mp4")
            self.logger.info(f"🎬 Создаю видео: {video_path}")
            
            # Конвертируем PIL изображения в numpy массивы
            video_frames = []
            for frame in frames:
                frame_array = np.array(frame)
                video_frames.append(frame_array)
            
            # Создаем видео с высоким качеством
            imageio.mimsave(video_path, video_frames, fps=gen_params["fps"], quality=8)
            
            self.logger.info(f"✅ Видео создано: {video_path}")
            
            # Автоматически открываем видео
            try:
                subprocess.run(["start", video_path], shell=True, check=True)
                self.logger.info("🎬 Видео автоматически открыто")
            except Exception as e:
                self.logger.warning(f"⚠️ Не удалось открыть видео: {e}")
            
            return video_path
            
        except Exception as e:
            self.logger.error(f"❌ Ошибка генерации видео: {e}")
            import traceback
            traceback.print_exc()
            return None
        finally:
            # Записываем метрику производительности
            response_time = time.time() - start_time
            self.add_performance_metric("video_generation", response_time)
            self.logger.info(f"🎬 Видео сгенерировано за {response_time:.2f} сек")

    def _add_dynamic_elements(self, prompt, frame_index, total_frames):
        """Добавляет динамические элементы к промпту в зависимости от номера кадра"""
        
        # Базовые динамические элементы для разных типов промптов
        dynamic_elements = {
            "pose": [
                "slight head turn", "head turning", "looking to the side", "looking up", "looking down",
                "slight body movement", "body turning", "arm movement", "hand gesture", "finger movement",
                "eye movement", "blinking", "mouth movement", "smile change", "expression change"
            ],
            "lighting": [
                "slight lighting change", "light shift", "shadow movement", "highlight change",
                "ambient light variation", "light intensity change", "color temperature shift"
            ],
            "camera": [
                "slight camera movement", "camera angle change", "zoom effect", "perspective shift",
                "depth change", "focus adjustment", "blur variation"
            ],
            "motion": [
                "motion blur", "movement lines", "wind effect", "hair movement", "clothing movement",
                "particle effects", "energy flow", "magical effects", "sparkle effects"
            ]
        }
        
        # Определяем тип промпта
        prompt_lower = prompt.lower()
        
        # Выбираем подходящие динамические элементы
        if any(word in prompt_lower for word in ["anime", "girl", "boy", "character", "person"]):
            # Для персонажей добавляем движения и выражения
            elements = dynamic_elements["pose"] + dynamic_elements["motion"]
        elif any(word in prompt_lower for word in ["landscape", "nature", "scenery", "background"]):
            # Для пейзажей добавляем изменения освещения и камеры
            elements = dynamic_elements["lighting"] + dynamic_elements["camera"]
        else:
            # Для остальных используем все элементы
            elements = dynamic_elements["pose"] + dynamic_elements["lighting"] + dynamic_elements["camera"] + dynamic_elements["motion"]
        
        # Выбираем элемент в зависимости от номера кадра
        if elements:
            # Равномерно распределяем элементы по кадрам
            element_index = int((frame_index / total_frames) * len(elements))
            selected_element = elements[element_index % len(elements)]
            
            # Добавляем элемент к промпту
            enhanced_prompt = f"{prompt}, {selected_element}"
            
            # Добавляем интенсивность изменения в зависимости от прогресса
            progress = frame_index / total_frames
            if progress > 0.5:
                enhanced_prompt += ", subtle animation"
            
            return enhanced_prompt
        
        return prompt
    
    def _install_diffusers_dependencies(self):
        """Устанавливает необходимые зависимости для diffusers"""
        try:
            import diffusers
            import torch
            import peft  # Проверяем PEFT также
            self.logger.info("✅ diffusers, torch и PEFT уже установлены")
            
            # Убеждаемся что PEFT backend включен для diffusers
            os.environ["USE_PEFT_BACKEND"] = "1"
            return
        except ImportError:
            self.logger.info("📦 Устанавливаю зависимости для diffusers...")
            
            try:
                subprocess.run([_sys.executable, "-m", "pip", "install", "diffusers", "transformers", "torch", "torchvision", "accelerate", "safetensors", "peft"], 
                             check=True, capture_output=True)
                self.logger.info("✅ Зависимости установлены успешно")
                # Устанавливаем переменную окружения для PEFT
                os.environ["USE_PEFT_BACKEND"] = "1"
            except subprocess.CalledProcessError as e:
                self.logger.error(f"❌ Ошибка установки зависимостей: {e}")
                raise

    def show_image_base64_temp(self, b64img: str):
        """Показать изображение из base64 через универсальный метод Windows"""
        try:
            # В веб-режиме отключаем всплывающее окно показа
            if not getattr(self, 'show_images_locally', True):
                return
            
            # Создаем временный файл для показа
            import tempfile
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                tmp_path = tmp_file.name
                img_data = base64.b64decode(b64img)
                tmp_file.write(img_data)
            
            # Открываем через универсальный метод Windows
            try:
                subprocess.run(["start", tmp_path], shell=True, check=True)
                self.logger.info("🖼️ Изображение автоматически открыто")
            except Exception as e:
                self.logger.warning(f"⚠️ Не удалось открыть изображение: {e}")
                
        except Exception as e:
            self.logger.error(f"Ошибка показа изображения: {e}")

    def find_new_audio(self) -> Optional[str]:
        """Находит новый аудиофайл для обработки"""
        audio_extensions = ['.mp3', '.wav', '.m4a', '.flac', '.ogg', '.aac']
        
        # Ищем в папке Audio
        audio_dir = os.path.join(self.base_dir, 'Audio')
        if os.path.exists(audio_dir):
            for file in os.listdir(audio_dir):
                if any(file.lower().endswith(ext) for ext in audio_extensions):
                    # Проверяем, что файл не помечен как обработанный
                    if '.used' not in file and not file.endswith('.used'):
                        file_path = os.path.join(audio_dir, file)
                        # Проверяем, что файл действительно существует и не пустой
                        if os.path.exists(file_path) and os.path.getsize(file_path) > 0:
                            return file_path
        
        return ""

    def mark_audio_used(self, audio_path: str):
        """Удаляет аудиофайл после обработки"""
        try:
            if os.path.exists(audio_path):
                # Удаляем файл полностью
                os.remove(audio_path)
                self.logger.info(f"✅ Аудиофайл удален после обработки: {os.path.basename(audio_path)}")
        except Exception as e:
            self.logger.error(f"❌ Ошибка при удалении аудиофайла: {e}")

    def transcribe_audio_whisper(self, audio_path: str, lang: str = "ru", use_separator: bool = True) -> Optional[str]:
        """
        Распознаёт аудио через whisper-cli. Если use_separator=True, предварительно выделяет вокал через audio-separator.
        Возвращает текст транскрипта (выводит только один раз при получении).
        """
        start_time = time.time()
        
        # Автоматически включаем audio модель при необходимости
        if not getattr(self, 'use_audio', False):
            self.logger.info("🔧 Автоматически включаю audio модель")
            self.use_audio = True
            # Запускаем таймер автоматического выключения
            self.auto_disable_tools("audio")
        
        # Проверяем и загружаем whisper модель если нужно
        if not self.check_whisper_setup():
            return "[Whisper error] Проблемы с настройкой Whisper. Проверьте наличие whisper-cli.exe и модели."
        
        try:
            base_dir = os.path.dirname(os.path.abspath(__file__))
            exe_path = os.path.join(base_dir, "Release", "whisper-cli.exe")
            model_path = os.path.join(base_dir, "models", "whisper-large-v3-q8_0.gguf")
            
            # Проверяем существование whisper-cli.exe
            if not os.path.exists(exe_path):
                return "[Whisper error] Не найден whisper-cli.exe в папке Release"
            
            # Проверяем существование модели
            if not os.path.exists(model_path):
                return "[Whisper error] Не найдена модель whisper в папке models"
            
            audio_for_whisper = audio_path
            
            # Используем audio separator если включен
            if use_separator:
                try:
                    from audio_separator.separator import Separator
                    self.logger.info("🎵 Использую audio-separator для выделения вокала...")
                    out_dir = os.path.join(base_dir, "separated")
                    os.makedirs(out_dir, exist_ok=True)
                    separator = Separator(output_dir=out_dir)
                    separator.load_model(model_filename='htdemucs_ft.yaml')
                    output_files = separator.separate(audio_path)
                    vocals_path = None
                    for file_path in output_files:
                        if '(Vocals)' in os.path.basename(file_path):
                            vocals_path = file_path  # audio-separator возвращает полный путь
                            self.logger.info(f"[SUCCESS] Вокал найден: {vocals_path}")
                            break
                    if not vocals_path:
                        self.logger.warning("⚠️ Не удалось найти файл с голосом после разделения дорожек, использую оригинал")
                    else:
                        audio_for_whisper = vocals_path
                except ImportError:
                    self.logger.warning("⚠️ Не установлена библиотека audio-separator. Пытаюсь установить автоматически...")
                    try:
                        import subprocess
                        subprocess.run([_sys.executable, "-m", "pip", "install", "audio-separator"], 
                                     capture_output=True, check=True)
                        self.logger.info("✅ audio-separator успешно установлен")
                        # Повторно пытаемся импортировать
                        from audio_separator.separator import Separator
                        self.logger.info("🎵 Использую audio-separator для выделения вокала...")
                        out_dir = os.path.join(base_dir, "separated")
                        os.makedirs(out_dir, exist_ok=True)
                        separator = Separator(output_dir=out_dir)
                        separator.load_model(model_filename='htdemucs_ft.yaml')
                        output_files = separator.separate(audio_path)
                        vocals_path = None
                        for file_path in output_files:
                            if '(Vocals)' in os.path.basename(file_path):
                                vocals_path = file_path  # audio-separator возвращает полный путь
                                self.logger.info(f"[SUCCESS] Вокал найден: {vocals_path}")
                                break
                        if not vocals_path:
                            self.logger.warning("⚠️ Не удалось найти файл с голосом после разделения дорожек, использую оригинал")
                        else:
                            audio_for_whisper = vocals_path
                    except Exception as install_error:
                        self.logger.warning(f"⚠️ Не удалось установить audio-separator: {install_error}")
                        self.logger.info("ℹ️ Продолжаю без разделения дорожек")
                except Exception as e:
                    self.logger.warning(f"⚠️ Ошибка audio-separator: {e}, использую оригинал")
            
            # Конвертируем аудио в WAV формат для Whisper (если это не уже WAV)
            if not audio_for_whisper.lower().endswith('.wav'):
                wav_path = self.convert_audio_to_wav(audio_for_whisper)
                if wav_path:
                    audio_for_whisper = wav_path
                    self.logger.info(f"✅ Аудио конвертировано в WAV: {os.path.basename(wav_path)}")
                else:
                    self.logger.warning("⚠️ Не удалось конвертировать в WAV, использую оригинал")
            else:
                self.logger.info("✅ Аудио уже в WAV формате")
            
            # Переименовать используемый файл в .used.расширение
            base_used, ext_used = os.path.splitext(audio_for_whisper)
            used_path = base_used + ".used" + ext_used
            try:
                if os.path.exists(audio_for_whisper):
                    os.rename(audio_for_whisper, used_path)
                    self.logger.info(f"✅ Аудиофайл переименован в: {os.path.basename(used_path)}")
                else:
                    self.logger.warning(f"⚠️ Аудиофайл не найден для переименования: {audio_for_whisper}")
                    used_path = audio_for_whisper
            except Exception as e:
                self.logger.error(f"Ошибка переименования аудио после whisper: {e}")
                # Если не удалось переименовать, используем оригинальный файл
                used_path = audio_for_whisper
            
            # Теперь используем used_path для whisper
            cmd = [exe_path, "--model", model_path]
            if lang:
                cmd += ["--language", lang]
            cmd.append(used_path)
            self.logger.info(f"[INFO] Запуск Whisper: {' '.join(cmd)}")
            import subprocess
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=300, encoding="utf-8", errors="replace")
            transcript = result.stdout.strip() if result.stdout else ""
            if transcript:
                self.logger.info("\n=== ТРАНСКРИПТ АУДИО ===\n" + transcript)
                return transcript
            
            # Очистка временных файлов если был separator
            if use_separator and 'separated' in audio_for_whisper:
                try:
                    separated_dir = os.path.dirname(audio_for_whisper)
                    if os.path.exists(separated_dir):
                        shutil.rmtree(separated_dir)
                        self.logger.info("🧹 Временные файлы audio-separator очищены")
                except Exception as e:
                    self.logger.warning(f"⚠️ Не удалось очистить временные файлы: {e}")
            
            err = result.stderr.strip() if result.stderr else ""
            return f"[Whisper error] Не удалось получить транскрипт. STDERR: {err}"
        except Exception as e:
            error_msg = f"Исключение whisper-cli: {str(e)}"
            self.logger.error(error_msg)
            return f"[Whisper error] {error_msg}"
        finally:
            # Записываем метрику производительности
            response_time = time.time() - start_time
            self.add_performance_metric("whisper_transcription", response_time)
            self.logger.info(f"🎤 Whisper обработал за {response_time:.2f} сек")

    def convert_audio_to_wav(self, audio_path: str) -> Optional[str]:
        """
        Конвертирует аудиофайл в WAV формат для Whisper.
        Возвращает путь к WAV файлу или None при ошибке.
        """
        try:
            if not audio_path or not os.path.exists(audio_path):
                return None
            
            # Если уже WAV, не конвертируем
            if audio_path.lower().endswith('.wav'):
                return audio_path
            
            # Проверяем наличие ffmpeg
            try:
                subprocess.run(['ffmpeg', '-version'], capture_output=True, check=True)
            except (subprocess.CalledProcessError, FileNotFoundError):
                self.logger.warning("⚠️ ffmpeg не найден в системе. Установите ffmpeg для конвертации аудио.")
                return None
            
            # Создаем временную папку для конвертации
            temp_dir = os.path.join(os.path.dirname(audio_path), "temp_convert")
            os.makedirs(temp_dir, exist_ok=True)
            
            # Имя выходного WAV файла
            base_name = os.path.splitext(os.path.basename(audio_path))[0]
            wav_path = os.path.join(temp_dir, f"{base_name}.wav")
            
            # Команда для конвертации через ffmpeg
            cmd = [
                'ffmpeg', '-i', audio_path,
                '-acodec', 'pcm_s16le',  # 16-bit PCM
                '-ar', '16000',          # 16kHz sample rate (оптимально для Whisper)
                '-ac', '1',              # моно
                '-y',                    # перезаписать существующий файл
                wav_path
            ]
            
            self.logger.info(f"🔄 Конвертирую аудио в WAV: {' '.join(cmd)}")
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            
            if result.returncode == 0 and os.path.exists(wav_path):
                self.logger.info(f"✅ Конвертация успешна: {os.path.basename(wav_path)}")
                return wav_path
            else:
                self.logger.error(f"❌ Ошибка конвертации: {result.stderr}")
                return None
                
        except Exception as e:
            self.logger.error(f"❌ Ошибка конвертации аудио в WAV: {e}")
            return None

    def check_whisper_setup(self) -> bool:
        """
        Проверяет настройку Whisper: наличие whisper-cli.exe и модели.
        Возвращает True если всё готово, False если есть проблемы.
        """
        try:
            base_dir = os.path.dirname(os.path.abspath(__file__))
            exe_path = os.path.join(base_dir, "Release", "whisper-cli.exe")
            model_path = os.path.join(base_dir, "models", "whisper-large-v3-q8_0.gguf")
            
            # Проверяем whisper-cli.exe
            if not os.path.exists(exe_path):
                self.logger.error(f"❌ Не найден whisper-cli.exe в папке Release: {exe_path}")
                self.logger.info("💡 Скачайте whisper.cpp с https://github.com/ggerganov/whisper.cpp")
                return False
            
            # Проверяем модель
            if not os.path.exists(model_path):
                self.logger.warning(f"⚠️ Не найдена модель whisper в папке models: {model_path}")
                self.logger.info("🔄 Пытаюсь автоматически скачать модель...")
                if self.download_whisper_model():
                    self.logger.info("✅ Модель whisper успешно загружена")
                else:
                    self.logger.error("❌ Не удалось загрузить модель whisper")
                    self.logger.info("💡 Скачайте модель whisper-large-v3-q8_0.gguf вручную")
                    return False
            
            # Проверяем права на выполнение
            try:
                result = subprocess.run([exe_path, "--help"], capture_output=True, text=True, timeout=10)
                if result.returncode != 0:
                    self.logger.warning("⚠️ whisper-cli.exe не может быть запущен")
                    return False
            except Exception as e:
                self.logger.warning(f"⚠️ Ошибка запуска whisper-cli.exe: {e}")
                return False
            
            self.logger.info("✅ Whisper настройка проверена успешно")
            return True
            
        except Exception as e:
            self.logger.error(f"❌ Ошибка проверки настройки Whisper: {e}")
            return False

    def download_whisper_model(self) -> bool:
        """
        Автоматически скачивает модель whisper-large-v3-q8_0.gguf.
        Возвращает True если успешно, False если ошибка.
        """
        try:
            base_dir = os.path.dirname(os.path.abspath(__file__))
            models_dir = os.path.join(base_dir, "models")
            os.makedirs(models_dir, exist_ok=True)
            
            model_name = "whisper-large-v3-q8_0.gguf"
            model_path = os.path.join(models_dir, model_name)
            
            # URL для скачивания модели (используем Hugging Face)
            model_url = "https://huggingface.co/ggerganov/whisper.cpp/resolve/main/ggml-large-v3-q8_0.bin"
            
            self.logger.info(f"📥 Скачиваю модель whisper: {model_name}")
            self.logger.info(f"🔗 URL: {model_url}")
            
            # Скачиваем модель
            response = requests.get(model_url, stream=True, timeout=300)
            response.raise_for_status()
            
            total_size = int(response.headers.get('content-length', 0))
            downloaded = 0
            
            with open(model_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    if chunk:
                        f.write(chunk)
                        downloaded += len(chunk)
                        if total_size > 0:
                            percent = (downloaded / total_size) * 100
                            self.logger.info(f"📊 Прогресс: {percent:.1f}% ({downloaded}/{total_size} байт)")
            
            self.logger.info(f"✅ Модель скачана: {model_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"❌ Ошибка скачивания модели whisper: {e}")
            return False

    def download_youtube_audio(self, url: str, out_dir: Optional[str] = None) -> str:
        """
        Скачивает аудиодорожку с YouTube по ссылке (использует yt-dlp)
        Возвращает путь к аудиофайлу или пустую строку
        """
        # subprocess уже импортирован в начале файла
        if out_dir is None:
            out_dir = os.path.join(os.path.dirname(__file__), "Audio")
        os.makedirs(out_dir, exist_ok=True)
        out_path = os.path.join(out_dir, "yt_audio.%(ext)s")
        # Проверяем наличие cookies
        cookies_path = self.get_youtube_cookies_path()
        use_cookies = False
        
        if cookies_path and self.check_cookies_validity(cookies_path):
            use_cookies = True
            self.logger.info("🍪 Использую cookies для аутентификации YouTube")
        else:
            self.logger.info("ℹ️ Cookies не найдены или невалидны, использую базовые параметры")
        
        # Базовые параметры для yt-dlp
        base_cmd = [
            "yt-dlp",
            "--force-ipv4",
            "--user-agent", "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "--extractor-args", "youtube:player_client=android",  # Используем Android клиент
            "--no-check-certificate",  # Игнорируем SSL ошибки
            "--prefer-insecure",  # Предпочитаем HTTP
            "--geo-bypass",  # Обход геоблокировки
            "--geo-bypass-country", "US",  # Страна для обхода
            "-f", "bestaudio[ext=m4a]/bestaudio/best",
            "--extract-audio", "--audio-format", "wav",  # Сразу в WAV для Whisper
            "-o", out_path
        ]

        # Добавляем cookies если доступны
        if use_cookies:
            base_cmd.extend(["--cookies", str(cookies_path)])  # type: ignore[arg-type]
        
        # Добавляем URL в конец
        cmd = base_cmd + [url]
        
        try:
            self.logger.info(f"Скачиваю аудио с YouTube: {url}")
            # Логируем команду в одну строку для избежания обрезания
            cmd_str = " ".join(cmd)
            self.logger.info(f"Команда: {cmd_str}")
            
            # Запускаем с таймаутом
            result = subprocess.run(cmd, check=True, capture_output=True, text=True, timeout=300)
            
            if result.stdout:
                self.logger.info(f"yt-dlp stdout: {result.stdout}")
            if result.stderr:
                self.logger.warning(f"yt-dlp stderr: {result.stderr}")
            
            # Найти скачанный файл
            for fname in os.listdir(out_dir):
                if fname.startswith("yt_audio") and fname.endswith(('.wav', '.m4a', '.mp3', '.ogg', '.flac')):
                    self.logger.info(f"✅ Аудио успешно скачано: {fname}")
                    return os.path.join(out_dir, fname)
            
            self.logger.warning("⚠️ Аудиофайл не найден после скачивания")
            return ""
            
        except subprocess.TimeoutExpired:
            self.logger.error("❌ Таймаут скачивания аудио (5 минут)")
            return ""
        except subprocess.CalledProcessError as e:
            self.logger.error(f"❌ Ошибка yt-dlp: {e}")
            if e.stderr:
                self.logger.error(f"stderr: {e.stderr}")
            return ""
        except Exception as e:
            self.logger.error(f"❌ Неожиданная ошибка скачивания аудио: {e}")
            
            # Пробуем альтернативный метод с другими параметрами
            self.logger.info("🔄 Пробую альтернативный метод скачивания...")
            try:
                alt_cmd = [
                    "yt-dlp",
                    "--force-ipv4",
                    "--user-agent", "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                    "--extractor-args", "youtube:player_client=web",
                    "--no-check-certificate",
                    "--geo-bypass",
                    "--geo-bypass-country", "US",
                    "-f", "bestaudio",
                    "--extract-audio", "--audio-format", "wav",  # Сразу в WAV для Whisper
                    "-o", out_path
                ]

                # Добавляем cookies если доступны
                if use_cookies:
                    alt_cmd.extend(["--cookies", str(cookies_path)])  # type: ignore[arg-type]

                alt_cmd.append(url)
                
                # Логируем команду в одну строку
                alt_cmd_str = " ".join(alt_cmd)
                self.logger.info(f"Альтернативная команда: {alt_cmd_str}")
                result = subprocess.run(alt_cmd, check=True, capture_output=True, text=True, timeout=300)
                
                # Найти скачанный файл
                for fname in os.listdir(out_dir):
                    if fname.startswith("yt_audio") and fname.endswith(('.wav', '.m4a', '.mp3', '.ogg', '.flac')):
                        self.logger.info(f"✅ Аудио успешно скачано альтернативным методом: {fname}")
                        return os.path.join(out_dir, fname)
                        
            except Exception as alt_e:
                self.logger.error(f"❌ Альтернативный метод также не сработал: {alt_e}")
                
                # Пробуем третий метод с максимально простыми параметрами
                self.logger.info("🔄 Пробую третий метод (минимальные параметры)...")
                try:
                    simple_cmd = [
                        "yt-dlp",
                        "--force-ipv4",
                        "--user-agent", "Mozilla/0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                        "--no-check-certificate",
                        "-f", "bestaudio",
                        "--extract-audio", "--audio-format", "mp3",
                        "-o", out_path
                    ]
                    
                    # Добавляем cookies если доступны
                    if use_cookies:
                        simple_cmd.extend(["--cookies", str(cookies_path)])  # type: ignore[arg-type]
                    
                    simple_cmd.append(url)
                    
                    self.logger.info(f"Третий метод: {' '.join(simple_cmd)}")
                    result = subprocess.run(simple_cmd, check=True, capture_output=True, text=True, timeout=300)
                    
                    # Найти скачанный файл
                    for fname in os.listdir(out_dir):
                        if fname.startswith("yt_audio") and fname.endswith(('.m4a', '.mp3', '.wav', '.ogg', '.flac')):
                            self.logger.info(f"✅ Аудио успешно скачано третьим методом: {fname}")
                            return os.path.join(out_dir, fname)
                            
                except Exception as simple_e:
                    self.logger.error(f"❌ Третий метод также не сработал: {simple_e}")
            
            return ""
    def find_new_image(self) -> str:
        """
        Находит первое новое изображение (png или jpg) в папке Photos, игнорируя файлы с .used перед расширением
        """
        photos_dir = os.path.join(os.path.dirname(__file__), "Photos")
        if not os.path.exists(photos_dir):
            return ""
        for fname in os.listdir(photos_dir):
            lower = fname.lower()
            if lower.endswith(('.png', '.jpg', '.jpeg')) and '.used' not in os.path.splitext(lower)[0]:
                return os.path.join(photos_dir, fname)
        return ""

    def mark_image_used(self, image_path: str):
        """
        Переименовывает изображение, чтобы нейросеть его больше не использовала
        """
        if not image_path:
            return
        base, ext = os.path.splitext(image_path)
        new_path = base + ".used" + ext
        try:
            os.rename(image_path, new_path)
        except Exception as e:
            self.logger.error(f"Ошибка переименования изображения: {e}")
    def extract_think_content(self, text: str) -> Optional[str]:
        """
        Извлекает содержимое из блока <think> или альтернативных маркеров размышлений.
        Поддерживает форматы:
        - <think>...</think> 
        - <|begin_of_thought|>...<|end_of_thought|>
        - BEGIN_OF_THOUGHT...END_OF_THOUGHT

        Args:
            text: Входной текст для поиска блока размышлений

        Returns:
            Optional[str]: Содержимое блока размышлений или None если блок не найден
        """
        # Паттерны для поиска (case-insensitive)
        patterns = [
            r'<think>(.*?)</think>',
            r'<\|begin_of_thought\|>(.*?)<\|end_of_thought\|>',
            r'BEGIN_OF_THOUGHT(.*?)END_OF_THOUGHT'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
            if match:
                content = match.group(1).strip()
                if content:  # Проверяем что контент не пустой
                    return content
        
        return None

    def extract_first_json(self, text: str, allow_json_in_think: bool = False) -> str:
        """
        Извлекает первый корректный JSON-блок из текста.
        
        Args:
            text: Входной текст для поиска JSON
            allow_json_in_think: Искать ли JSON внутри блока think
            
        Returns:
            str: Найденный JSON или исходный текст если JSON не найден
        """
        # Сначала ищем think-блок и удаляем его из текста если он есть
        think_content = None
        clean_text = text
        
        patterns = [
            r'<think>.*?</think>',
            r'<\|begin_of_thought\|>.*?<\|end_of_thought\|>',
            r'BEGIN_OF_THOUGHT.*?END_OF_THOUGHT'
        ]
        
        for pattern in patterns:
            think_match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
            if think_match:
                think_content = self.extract_think_content(think_match.group(0))
                clean_text = re.sub(pattern, '', text, flags=re.DOTALL | re.IGNORECASE).strip()
                break
        
        # Сначала пробуем найти чистый JSON (без обрамлений)
        json_in_text = self._extract_json_from_text(clean_text)
        if json_in_text:
            return json_in_text
        
        # Если не нашли, пробуем искать JSON с обрамлением и удалять его
        json_with_wrapper = re.search(r'```(?:json)?\s*(.*?)\s*```', clean_text, re.DOTALL)
        if json_with_wrapper:
            potential_json = json_with_wrapper.group(1).strip()
            json_in_wrapper = self._extract_json_from_text(potential_json)
            if json_in_wrapper:
                return json_in_wrapper
            
        # Если разрешено и есть think-контент - ищем JSON там
        if allow_json_in_think and think_content:
            json_in_think = self._extract_json_from_text(think_content)
            if json_in_think:
                return json_in_think
        
        return text  # если не найдено, вернуть исходное
    
    def _extract_json_from_text(self, text: str) -> str:
        """
        Вспомогательная функция для извлечения JSON из текста
        """
        # re уже импортирован в начале файла
        stack = []
        start = None
        
        for i, c in enumerate(text):
            if c == '{':
                if not stack:
                    start = i
                stack.append('{')
            elif c == '}':
                if stack:
                    stack.pop()
                    if not stack and start is not None:
                        return text[start:i+1]
        
        return ""

    def _smart_json_parse(self, s: str):
        """Умный парсер JSON с несколькими попытками автокоррекции.

        Возвращает tuple (data_or_none, fixes_list).
        """
        logger.info(f"🔍 Парсинг JSON: {s[:200]}...")
        try:
            return json.loads(s), []
        except Exception as e:
            fixes = [f"Первый парсинг не удался: {e}"]

        # Попытка закрыть скобки
        open_braces = s.count('{')
        close_braces = s.count('}')
        if open_braces > close_braces:
            s2 = s + '}' * (open_braces - close_braces)
            fixes.append(f"Добавлено {open_braces - close_braces} }} для баланса скобок")
            try:
                return json.loads(s2), fixes
            except Exception:
                pass
        elif close_braces > open_braces:
            s2 = re.sub(r'}+$', '', s)
            fixes.append(f"Удалены лишние закрывающие скобки")
            try:
                return json.loads(s2), fixes
            except Exception:
                pass

        # Заменяем одинарные кавычки на двойные если это безопасно
        if "'" in s and '"' not in s:
            s2 = s.replace("'", '"')
            try:
                return json.loads(s2), fixes+['Заменены одинарные кавычки на двойные']
            except Exception:
                pass

        # Удаляем лишние запятые перед закрывающей скобкой
        s3 = re.sub(r',\s*([}\]])', r'\1', s)
        try:
            return json.loads(s3), fixes+['Удалены лишние запятые']
        except Exception:
            pass

        # Оборачиваем ключи в кавычки (грубая попытка)
        s4 = re.sub(r'([,{]\s*)([a-zA-Z0-9_]+)\s*:', r'\1"\2":', s3)
        try:
            return json.loads(s4), fixes+['Добавлены кавычки к ключам']
        except Exception:
            pass

        # Исправляем незакрытые строки
        s5 = re.sub(r'([^\"])\s*$', r'\1"', s4)
        try:
            return json.loads(s5), fixes+['Исправлены незакрытые строки']
        except Exception:
            pass

        # Финальная попытка с очисткой от непечатаемых символов
        s6 = re.sub(r'[^\x20-\x7E]', '', s5)
        try:
            return json.loads(s6), fixes+['Очищены непечатаемые символы']
        except Exception as e2:
            fixes.append(f"Не удалось распарсить даже после исправлений: {e2}")

        return None, fixes
    def __init__(self, lm_studio_url: str = "http://localhost:1234", 
                 google_api_key: str = "", google_cse_id: str = ""):
        """
        Инициализация оркестратора
        
        Args:
            lm_studio_url: URL сервера LM Studio
            google_api_key: API ключ Google Custom Search
            google_cse_id: ID поисковой системы Google CSE
        """
        self.lm_studio_url = lm_studio_url.rstrip("/")
        self.google_api_key = google_api_key
        self.google_cse_id = google_cse_id
        # unify logger usage for instance methods
        self.logger = logger
        self.conversation_history: List[Dict[str, Any]] = []
        self.brain_model = "J:/models-LM Studio/mradermacher/Huihui-Qwen3-4B-Thinking-2507-abliterated-GGUF/Huihui-Qwen3-4B-Thinking-2507-abliterated.Q4_K_S.gguf"
        self.brain_model_id = None  # Короткий ID модели для API вызовов
        self.use_separator = True  # По умолчанию True, чтобы убрать предупреждение Pylance
        self.use_image_generation = False  # По умолчанию отключена генерация изображений
        # Тумблеры функционала (визуал и аудио)
        self.use_vision = False
        self.use_audio = False
        self.use_ocr = False  # По умолчанию отключен OCR
        # Управление локальным показом изображений (для веб-режима можно отключить)
        self.show_images_locally = True
        # Хранилище последнего сгенерированного изображения (base64) и ответа
        self.last_generated_image_b64 = None
        self.last_final_response = ""
        
        # Хранилище последнего сгенерированного файла для Telegram
        self.last_generated_file_path = None
        self.last_generated_file_name = None
        
        # Динамическое управление контекстом
        self.max_context_length = 262144  # Максимальный контекст (временно)
        self.safe_context_length = 32768   # Безопасный контекст (временно)
        self.current_context_length = 0    # Текущий размер контекста
        
        # Метрики производительности
        self.performance_metrics = []  # Список метрик производительности
        
        # Счетчик попыток для предотвращения зацикливания
        self.retry_count = 0
        self.max_retries = 3
        
        # Постоянная голосовая запись
        self.continuous_recording = False
        self.audio_queue = queue.Queue()
        self.recording_thread = None
        
        # Таймеры для автоматического выключения инструментов
        self.tool_timers = {}
        self.auto_disable_delay = 300  # Выключать инструменты через 5 минут после использования
        
        # Автоматически запускаем модель мозга при инициализации
        self._auto_load_brain_model()
        
        # Инициализируем динамические параметры контекста после загрузки модели
        self._initialize_dynamic_context()
        
        # Инициализируем базовую директорию
        self.base_dir = os.path.dirname(os.path.abspath(__file__))
        
        # Запускаем фоновую загрузку тяжелых компонентов
        self._start_background_loading()
        
        # Инициализируем ChromaDB для векторного хранилища (ленивая инициализация)
        self.chromadb_manager = None
        self._chromadb_initialized = False
        self._chromadb_config = {
            "db_path": os.path.join(self.base_dir, "chroma_db"),
            "use_gpu": True
        }
        
        # OCR будет инициализирован в фоне
        self.ocr_reader = None
        
        # Проверяем наличие ffmpeg для конвертации аудио
        self._check_ffmpeg()
        
        # Telegram Bot настройки
        self.telegram_bot_token = ""
        self.telegram_allowed_user_id = ""
        
        # Инициализируем систему плагинов
        # Initialize plugin system
        self.plugin_manager = None
        if PLUGINS_AVAILABLE and PluginManager is not None:
            try:
                self.plugin_manager = PluginManager(plugins_dir="plugins")
                self.plugin_manager.load_all_plugins(orchestrator=self)
                logger.info("✅ Система плагинов инициализирована")
            except Exception as e:
                logger.error(f"❌ Ошибка инициализации плагинов: {e}")
                self.plugin_manager = None
        else:
            logger.warning("⚠️ Система плагинов недоступна")
        
        # Инициализируем почтовые настройки
        self._initialize_email_config()

        # Инициализируем загрузчик промптов
        self.prompt_loader = PromptLoader(self.base_dir)
        
        # Инициализируем менеджер моделей и LoRA
        self.model_manager = ModelManager(self.base_dir)
        
        # Загружаем базовый системный промпт из файла
        self.system_prompt = self.prompt_loader.load_base_prompt()

    def list_folder_contents(self, folder_name: str) -> str:
        """
        Получение списка файлов в указанной папке
        
        Args:
            folder_name: Имя папки (Audio, Photos, Video, Excel, Docx, PDF)
        
        Returns:
            Строка со списком файлов или сообщение об ошибке
        """
        try:
            folder_path = os.path.join(self.base_dir, folder_name)
            
            if not os.path.exists(folder_path):
                return f"Папка {folder_name} не существует"
            
            files = os.listdir(folder_path)
            if not files:
                return f"Папка {folder_name} пуста"
            
            # Группируем файлы по типу
            file_types = {
                'Документы DOCX': [],
                'Таблицы Excel': [],
                'Документы PDF': [],
                'Изображения': [],
                'Аудио': [],
                'Видео': [],
                'Другие': []
            }
            
            for file in files:
                file_lower = file.lower()
                if file_lower.endswith(('.docx', '.doc')):
                    file_types['Документы DOCX'].append(file)
                elif file_lower.endswith(('.xlsx', '.xls', '.csv')):
                    file_types['Таблицы Excel'].append(file)
                elif file_lower.endswith('.pdf'):
                    file_types['Документы PDF'].append(file)
                elif file_lower.endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp')):
                    file_types['Изображения'].append(file)
                elif file_lower.endswith(('.mp3', '.wav', '.ogg', '.flac', '.m4a')):
                    file_types['Аудио'].append(file)
                elif file_lower.endswith(('.mp4', '.avi', '.mkv', '.mov', '.wmv')):
                    file_types['Видео'].append(file)
                else:
                    file_types['Другие'].append(file)
            
            result = f"Содержимое папки {folder_name}:\n"
            for file_type, file_list in file_types.items():
                if file_list:
                    result += f"\n{file_type}:\n"
                    for file in sorted(file_list):
                        result += f"  - {file}\n"
            
            return result
            
        except Exception as e:
            return f"Ошибка при чтении папки {folder_name}: {str(e)}"

    def extract_docx_content(self, file_path: str) -> tuple[str, str]:
        """
        Извлечение текста из DOCX файла
        
        Args:
            file_path: Путь к DOCX файлу
        
        Returns:
            Кортеж (текст, сообщение_об_ошибке)
        """
        try:
            from docx import Document
            
            if not os.path.exists(file_path):
                return "", f"Файл {file_path} не найден"
            
            doc = Document(file_path)
            text_content = []
            
            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            full_text = '\n'.join(text_content)
            
            if not full_text.strip():
                return "", "Документ не содержит текста"
            
            return full_text, ""
            
        except ImportError:
            return "", "Библиотека python-docx не установлена. Установите: pip install python-docx"
        except Exception as e:
            return "", f"Ошибка при чтении DOCX файла: {str(e)}"

    def extract_excel_content(self, file_path: str) -> tuple[str, str]:
        """
        Извлечение данных из Excel файла
        
        Args:
            file_path: Путь к Excel файлу
        
        Returns:
            Кортеж (данные_в_текстовом_формате, сообщение_об_ошибке)
        """
        try:
            import pandas as pd
            
            if not os.path.exists(file_path):
                return "", f"Файл {file_path} не найден"
            
            # Читаем все листы Excel файла
            excel_data = pd.read_excel(file_path, sheet_name=None)
            
            content_parts = []
            
            for sheet_name, df in excel_data.items():
                content_parts.append(f"=== Лист: {sheet_name} ===\n")
                
                # Проверяем, есть ли данные
                if df.empty:
                    content_parts.append("Лист пуст\n")
                    continue
                
                # Конвертируем DataFrame в текстовое представление
                content_parts.append(df.to_string(index=False))
                content_parts.append("\n")
            
            full_content = '\n'.join(content_parts)
            
            if not full_content.strip():
                return "", "Excel файл не содержит данных"
            
            return full_content, ""
            
        except ImportError:
            return "", "Библиотеки pandas/openpyxl не установлены. Установите: pip install pandas openpyxl"
        except Exception as e:
            return "", f"Ошибка при чтении Excel файла: {str(e)}"

    def extract_pdf_content(self, file_path: str) -> tuple[str, str]:
        """
        Извлечение текста из PDF файла
        
        Args:
            file_path: Путь к PDF файлу
        
        Returns:
            Кортеж (текст, сообщение_об_ошибке)
        """
        try:
            if not PDF_AVAILABLE or PyPDF2 is None:
                return "", "Библиотека PyPDF2 не установлена. Установите: pip install PyPDF2"
            
            if not os.path.exists(file_path):
                return "", f"Файл {file_path} не найден"
            
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Проверяем, есть ли страницы
                if len(pdf_reader.pages) == 0:
                    return "", "PDF файл не содержит страниц"
                
                # Извлекаем текст со всех страниц
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"=== Страница {page_num} ===\n{page_text.strip()}")
                    except Exception as e:
                        text_content.append(f"=== Страница {page_num} ===\n[Ошибка извлечения текста: {str(e)}]")
            
            full_text = '\n\n'.join(text_content)
            
            if not full_text.strip():
                return "", "PDF файл не содержит извлекаемого текста"
            
            return full_text, ""
            
        except Exception as e:
            return "", f"Ошибка при чтении PDF файла: {str(e)}"

    def rag_process_large_content(self, content: str, max_tokens: int = 4000) -> str:
        """
        RAG-обработка больших документов с разделением на части
        
        Args:
            content: Содержимое документа
            max_tokens: Максимальное количество токенов на часть
        
        Returns:
            Обработанное содержимое (сжатое или разделенное)
        """
        try:
            # Приблизительная оценка токенов (1 токен ≈ 4 символа для русского текста)
            estimated_tokens = len(content) // 4
            
            if estimated_tokens <= max_tokens:
                return content
            
            # Если документ слишком большой, разделяем на части
            chunk_size = max_tokens * 4  # Размер части в символах
            chunks = []
            
            # Разделяем по предложениям, чтобы сохранить смысл
            sentences = content.split('.')
            current_chunk = ""
            
            for sentence in sentences:
                if len(current_chunk + sentence) < chunk_size:
                    current_chunk += sentence + "."
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                        current_chunk = sentence + "."
                    else:
                        # Если одно предложение больше chunk_size, берем как есть
                        chunks.append(sentence.strip())
            
            # Добавляем последнюю часть
            if current_chunk:
                chunks.append(current_chunk.strip())
            
            # Возвращаем первые 3 части с указанием общего количества
            if len(chunks) <= 3:
                if len(chunks) > 1:
                    return f"=== ДОКУМЕНТ РАЗДЕЛЕН НА {len(chunks)} ЧАСТЕЙ ===\n\n" + '\n\n=== ЧАСТЬ ДОКУМЕНТА ===\n\n'.join(chunks)
                else:
                    return chunks[0] if chunks else content[:max_tokens * 4]
            else:
                result = f"=== ДОКУМЕНТ РАЗДЕЛЕН НА {len(chunks)} ЧАСТЕЙ ===\n\n"
                result += '\n\n=== ЧАСТЬ ДОКУМЕНТА ===\n\n'.join(chunks[:3])
                result += f"\n\n[ПОКАЗАНЫ ПЕРВЫЕ 3 ЧАСТИ ИЗ {len(chunks)}. СПРОСИТЕ, ЕСЛИ НУЖНО БОЛЬШЕ ИНФОРМАЦИИ]"
                return result
                
        except Exception as e:
            logger.error(f"Ошибка RAG обработки: {e}")
            return content[:max_tokens * 4] + "\n\n[СОДЕРЖИМОЕ ОБРЕЗАНО ДУЕ К РАЗМЕРУ]"

    def process_document_request(self, file_path: str) -> str:
        """
        Обработка запроса на работу с документом
        
        Args:
            file_path: Путь к файлу
        
        Returns:
            Обработанное содержимое документа
        """
        try:
            file_lower = file_path.lower()
            
            if file_lower.endswith(('.docx', '.doc')):
                content, error = self.extract_docx_content(file_path)
                if error:
                    return f"Ошибка при обработке DOCX: {error}"
                
                # RAG обработка для больших документов
                processed_content = self.rag_process_large_content(content)
                return f"Содержимое DOCX документа:\n\n{processed_content}"
                
            elif file_lower.endswith(('.xlsx', '.xls')):
                content, error = self.extract_excel_content(file_path)
                if error:
                    return f"Ошибка при обработке Excel: {error}"
                
                # RAG обработка для больших таблиц
                processed_content = self.rag_process_large_content(content)
                return f"Содержимое Excel файла:\n\n{processed_content}"
                
            elif file_lower.endswith('.pdf'):
                content, error = self.extract_pdf_content(file_path)
                if error:
                    return f"Ошибка при обработке PDF: {error}"
                
                # RAG обработка для больших PDF документов
                processed_content = self.rag_process_large_content(content)
                return f"Содержимое PDF документа:\n\n{processed_content}"
                
            elif file_lower.endswith('.csv'):
                # Для CSV используем pandas
                try:
                    import pandas as pd
                    df = pd.read_csv(file_path)
                    content = df.to_string(index=False)
                    processed_content = self.rag_process_large_content(content)
                    return f"Содержимое CSV файла:\n\n{processed_content}"
                except Exception as e:
                    return f"Ошибка при чтении CSV: {str(e)}"
                    
            elif file_lower.endswith(('.txt', '.md')):
                # Текстовые файлы
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    processed_content = self.rag_process_large_content(content)
                    file_type = "Markdown" if file_lower.endswith('.md') else "текстового"
                    return f"Содержимое {file_type} файла:\n\n{processed_content}"
                except UnicodeDecodeError:
                    # Попробуем другие кодировки
                    for encoding in ['cp1251', 'latin1']:
                        try:
                            with open(file_path, 'r', encoding=encoding) as f:
                                content = f.read()
                            processed_content = self.rag_process_large_content(content)
                            return f"Содержимое текстового файла (кодировка {encoding}):\n\n{processed_content}"
                        except:
                            continue
                    return f"Ошибка: не удалось определить кодировку файла {file_path}"
                except Exception as e:
                    return f"Ошибка при чтении текстового файла: {str(e)}"
                    
            elif file_lower.endswith(('.rtf')):
                # RTF файлы
                try:
                    # Простое извлечение текста из RTF (базовое)
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    # Удаляем RTF команды (очень упрощенно)
                    import re
                    content = re.sub(r'\\[a-z]+\d*', '', content)  # Убираем команды типа \par, \b1 и т.д.
                    content = re.sub(r'[{}]', '', content)  # Убираем фигурные скобки
                    content = content.strip()
                    processed_content = self.rag_process_large_content(content)
                    return f"Содержимое RTF файла:\n\n{processed_content}"
                except Exception as e:
                    return f"Ошибка при обработке RTF: {str(e)}"
                    
            elif file_lower.endswith('.json'):
                # JSON файлы
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    # Преобразуем JSON в читаемый формат
                    content = json.dumps(data, indent=2, ensure_ascii=False)
                    processed_content = self.rag_process_large_content(content)
                    return f"Содержимое JSON файла:\n\n{processed_content}"
                except json.JSONDecodeError as e:
                    return f"Ошибка в формате JSON: {str(e)}"
                except Exception as e:
                    return f"Ошибка при обработке JSON: {str(e)}"
                    
            elif file_lower.endswith(('.xml', '.html', '.htm')):
                # XML/HTML файлы
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    
                    # Для HTML попробуем извлечь только текст
                    if file_lower.endswith(('.html', '.htm')):
                        try:
                            from bs4 import BeautifulSoup
                            soup = BeautifulSoup(content, 'html.parser')
                            # Удаляем скрипты и стили
                            for script in soup(["script", "style"]):
                                script.decompose()
                            content = soup.get_text()
                            # Очищаем лишние пробелы
                            content = '\n'.join(line.strip() for line in content.splitlines() if line.strip())
                        except ImportError:
                            # Если BeautifulSoup не установлен, используем простое удаление тегов
                            import re
                            content = re.sub('<[^<]+?>', '', content)
                    
                    processed_content = self.rag_process_large_content(content)
                    file_type = "HTML" if file_lower.endswith(('.html', '.htm')) else "XML"
                    return f"Содержимое {file_type} файла:\n\n{processed_content}"
                except Exception as e:
                    return f"Ошибка при обработке XML/HTML: {str(e)}"
                    
            else:
                return f"Неподдерживаемый формат файла: {file_path}"
                
        except Exception as e:
            return f"Ошибка при обработке документа: {str(e)}"

    def generate_docx_file(self, content: str, filename: str) -> str:
        """
        Генерация DOCX файла
        
        Args:
            content: Содержимое документа
            filename: Полное имя файла (с расширением или без)
        
        Returns:
            Сообщение о результате создания файла
        """
        try:
            if not DOCX_AVAILABLE or Document is None:
                return "Ошибка: Библиотека python-docx не установлена"
            
            # Создаем документ
            doc = Document()
            
            # Разделяем контент на параграфы
            paragraphs = content.split('\n')
            
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    # Проверяем, является ли строка заголовком (начинается с #)
                    if paragraph_text.strip().startswith('#'):
                        # Убираем # и создаем заголовок
                        title_text = paragraph_text.strip().lstrip('#').strip()
                        heading = doc.add_heading(title_text, level=1)
                    else:
                        doc.add_paragraph(paragraph_text.strip())
            
            # Создаем папку если не существует
            os.makedirs(os.path.join(self.base_dir, "output"), exist_ok=True)
            
            # Сохраняем файл - убираем расширение если есть и добавляем .docx
            base_name = filename.replace('.docx', '').replace('.doc', '')
            output_path = os.path.join(self.base_dir, "output", f"{base_name}.docx")
            doc.save(output_path)
            
            return f"Документ успешно создан: {output_path}"
            
        except Exception as e:
            return f"Ошибка при создании DOCX файла: {str(e)}"

    def generate_excel_file(self, content: str, filename: str) -> str:
        """
        Генерация Excel файла
        
        Args:
            content: Содержимое в формате таблицы (разделители - табуляция или запятые)
            filename: Полное имя файла (с расширением или без)
        
        Returns:
            Сообщение о результате создания файла
        """
        try:
            if not EXCEL_AVAILABLE or pd is None:
                return "Ошибка: Библиотеки pandas/openpyxl не установлены"
            
            # Пытаемся разобрать контент как табличные данные
            lines = content.strip().split('\n')
            if not lines:
                return "Ошибка: Пустое содержимое для Excel файла"
            
            # Определяем разделитель (табуляция или запятая)
            delimiter = '\t' if '\t' in lines[0] else ','
            
            # Создаем DataFrame
            import io
            data_string = '\n'.join(lines)
            df = pd.read_csv(io.StringIO(data_string), delimiter=delimiter)
            
            # Сохраняем файл - убираем расширение если есть и добавляем .xlsx
            base_name = filename.replace('.xlsx', '').replace('.xls', '')
            output_path = os.path.join(self.base_dir, "output", f"{base_name}.xlsx")
            df.to_excel(output_path, index=False)
            
            return f"Excel файл успешно создан: {output_path}"
            
        except Exception as e:
            return f"Ошибка при создании Excel файла: {str(e)}"

    def generate_markdown_file(self, content: str, filename: str) -> str:
        """
        Генерация Markdown файла
        
        Args:
            content: Содержимое в формате Markdown
            filename: Полное имя файла (с расширением или без)
        
        Returns:
            Сообщение о результате создания файла
        """
        try:
            # Сохраняем файл - убираем расширение если есть и добавляем .md
            base_name = filename.replace('.md', '').replace('.markdown', '')
            output_path = os.path.join(self.base_dir, "output", f"{base_name}.md")
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return f"Markdown файл успешно создан: {output_path}"
            
        except Exception as e:
            return f"Ошибка при создании Markdown файла: {str(e)}"

    def generate_pdf_file(self, content: str, filename: str) -> str:
        """
        Генерация PDF файла
        
        Args:
            content: Содержимое документа
            filename: Полное имя файла (с расширением или без)
        
        Returns:
            Сообщение о результате создания файла
        """
        try:
            if not REPORTLAB_AVAILABLE:
                return "Ошибка: Библиотека reportlab не установлена"
            
            # Убираем расширение если есть и добавляем .pdf
            base_name = filename.replace('.pdf', '')
            os.makedirs(os.path.join(self.base_dir, "output"), exist_ok=True)
            output_path = os.path.join(self.base_dir, "output", f"{base_name}.pdf")
            
            # Создаем PDF документ - импорт во время выполнения
            from reportlab.lib.pagesizes import A4  # type: ignore
            from reportlab.lib.styles import getSampleStyleSheet  # type: ignore
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer  # type: ignore
            
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Разделяем контент на параграфы
            paragraphs = content.split('\n')
            
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    # Проверяем, является ли строка заголовком
                    if paragraph_text.strip().startswith('#'):
                        title_text = paragraph_text.strip().lstrip('#').strip()
                        p = Paragraph(title_text, styles['Heading1'])
                    else:
                        p = Paragraph(paragraph_text.strip(), styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 12))
            
            # Строим документ
            doc.build(story)
            
            return f"PDF файл успешно создан: {output_path}"
            
        except Exception as e:
            return f"Ошибка при создании PDF файла: {str(e)}"

    def generate_txt_file(self, content: str, filename: str) -> str:
        """
        Генерация простого текстового файла
        
        Args:
            content: Содержимое файла
            filename: Полное имя файла (с расширением или без)
        
        Returns:
            Сообщение о результате создания файла
        """
        try:
            # Убираем расширение если есть и добавляем .txt
            base_name = filename.replace('.txt', '')
            output_path = os.path.join(self.base_dir, "output", f"{base_name}.txt")
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return f"Текстовый файл успешно создан: {output_path}"
            
        except Exception as e:
            return f"Ошибка при создании текстового файла: {str(e)}"

    def generate_json_file(self, content: str, filename: str) -> str:
        """
        Генерация JSON файла
        
        Args:
            content: Содержимое файла в формате JSON (строка)
            filename: Полное имя файла (с расширением или без)
        
        Returns:
            Сообщение о результате создания файла
        """
        try:
            # Убираем расширение если есть и добавляем .json
            base_name = filename.replace('.json', '')
            output_path = os.path.join(self.base_dir, "output", f"{base_name}.json")
            
            # Проверяем валидность JSON
            import json
            try:
                json.loads(content)  # Проверяем что content - валидный JSON
            except json.JSONDecodeError:
                # Если не JSON, оборачиваем в кавычки как строку
                content = json.dumps(content, ensure_ascii=False, indent=2)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return f"JSON файл успешно создан: {output_path}"
            
        except Exception as e:
            return f"Ошибка при создании JSON файла: {str(e)}"

    def generate_csv_file(self, content: str, filename: str) -> str:
        """
        Генерация CSV файла
        
        Args:
            content: Содержимое файла в формате CSV (строка с разделителями)
            filename: Полное имя файла (с расширением или без)
        
        Returns:
            Сообщение о результате создания файла
        """
        try:
            # Убираем расширение если есть и добавляем .csv
            base_name = filename.replace('.csv', '')
            output_path = os.path.join(self.base_dir, "output", f"{base_name}.csv")
            
            with open(output_path, 'w', encoding='utf-8', newline='') as f:
                f.write(content)
            
            return f"CSV файл успешно создан: {output_path}"
            
        except Exception as e:
            return f"Ошибка при создании CSV файла: {str(e)}"

    def generate_html_file(self, content: str, filename: str) -> str:
        """
        Генерация HTML файла
        
        Args:
            content: Содержимое файла в формате HTML
            filename: Полное имя файла (с расширением или без)
        
        Returns:
            Сообщение о результате создания файла
        """
        try:
            # Убираем расширение если есть и добавляем .html
            base_name = filename.replace('.html', '').replace('.htm', '')
            output_path = os.path.join(self.base_dir, "output", f"{base_name}.html")
            
            # Если контент не содержит HTML структуру, добавляем базовую
            if not content.strip().lower().startswith('<!doctype') and not content.strip().lower().startswith('<html'):
                content = f"""<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{base_name}</title>
</head>
<body>
{content}
</body>
</html>"""
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return f"HTML файл успешно создан: {output_path}"
            
        except Exception as e:
            return f"Ошибка при создании HTML файла: {str(e)}"

    def generate_xml_file(self, content: str, filename: str) -> str:
        """
        Генерация XML файла
        
        Args:
            content: Содержимое файла в формате XML
            filename: Полное имя файла (с расширением или без)
        
        Returns:
            Сообщение о результате создания файла
        """
        try:
            # Убираем расширение если есть и добавляем .xml
            base_name = filename.replace('.xml', '')
            output_path = os.path.join(self.base_dir, "output", f"{base_name}.xml")
            
            # Если контент не содержит XML декларацию, добавляем
            if not content.strip().startswith('<?xml'):
                content = f'<?xml version="1.0" encoding="UTF-8"?>\n{content}'
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return f"XML файл успешно создан: {output_path}"
            
        except Exception as e:
            return f"Ошибка при создании XML файла: {str(e)}"

    def generate_bat_file(self, content: str, filename: str) -> str:
        """
        Создает .bat файл с командами для Windows
        
        Args:
            content: Содержимое bat файла (команды)
            filename: Имя файла (с расширением .bat или без)
            
        Returns:
            Сообщение о результате создания файла
        """
        try:
            # Обеспечиваем правильное расширение
            if not filename.lower().endswith('.bat'):
                filename += '.bat'
            
            # Путь для сохранения в папку output
            output_path = os.path.join(self.base_path, "output", filename)
            
            # Убеждаемся что папка output существует
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Добавляем @echo off в начало, если его нет
            if not content.strip().startswith('@echo off'):
                content = '@echo off\n' + content.strip()
            
            # Добавляем pause в конец, если его нет
            if not content.strip().endswith('pause'):
                content = content.strip() + '\npause'
            
            with open(output_path, 'w', encoding='cp1251') as f:  # cp1251 для Windows bat файлов
                f.write(content)
            
            self.logger.info(f"📄 BAT файл создан: {filename}")
            return f"BAT файл успешно создан: {output_path}"
            
        except Exception as e:
            return f"Ошибка при создании BAT файла: {str(e)}"

    def run_bat_file(self, file_path: str, working_dir: Optional[str] = None) -> Dict[str, Any]:
        """
        Запускает .bat файл
        
        Args:
            file_path: Путь к .bat файлу
            working_dir: Рабочая директория для выполнения (опционально)
            
        Returns:
            Результат выполнения с выводом и кодом возврата
        """
        try:
            # Разрешаем путь к файлу
            resolved_path = self.resolve_path(file_path)
            
            if not os.path.exists(resolved_path):
                return {
                    "success": False,
                    "error": f"BAT файл не найден: {resolved_path}",
                    "output": "",
                    "return_code": -1
                }
            
            # Определяем рабочую директорию
            if working_dir is None:
                working_dir = os.path.dirname(resolved_path)
            else:
                working_dir = self.resolve_path(working_dir)
            
            self.logger.info(f"🚀 Запускаю BAT файл: {os.path.basename(resolved_path)}")
            
            # Запускаем bat файл
            result = subprocess.run(
                [resolved_path],
                cwd=working_dir,
                capture_output=True,
                text=True,
                encoding='cp1251',  # Кодировка для Windows
                timeout=300  # Таймаут 5 минут
            )
            
            success = result.returncode == 0
            
            if success:
                self.logger.info(f"✅ BAT файл выполнен успешно")
            else:
                self.logger.warning(f"⚠️ BAT файл завершился с кодом: {result.returncode}")
            
            return {
                "success": success,
                "output": result.stdout,
                "error": result.stderr,
                "return_code": result.returncode,
                "working_dir": working_dir
            }
            
        except subprocess.TimeoutExpired:
            return {
                "success": False,
                "error": "Превышено время ожидания выполнения BAT файла (5 минут)",
                "output": "",
                "return_code": -1
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"Ошибка при выполнении BAT файла: {str(e)}",
                "output": "",
                "return_code": -1
            }

    def generate_file(self, content: str, filename: str, file_format: str) -> bool:
        """
        Универсальный метод генерации файлов
        
        Args:
            content: Содержимое файла
            filename: Полное имя файла с расширением
            file_format: Формат файла (docx, excel, md, pdf, txt, json, csv, html, xml, bat)
        
        Returns:
            True если файл создан успешно, False иначе
        """
        try:
            # Создаем папку output если её нет
            output_dir = os.path.join(self.base_dir, "output")
            os.makedirs(output_dir, exist_ok=True)
            
            format_lower = file_format.lower()
            
            if format_lower in ['docx', 'doc', 'word']:
                result = self.generate_docx_file(content, filename)
                return "успешно создан" in result.lower()
            elif format_lower in ['excel', 'xlsx', 'xls']:
                result = self.generate_excel_file(content, filename)
                return "успешно создан" in result.lower()
            elif format_lower in ['md', 'markdown']:
                result = self.generate_markdown_file(content, filename)
                return "успешно создан" in result.lower()
            elif format_lower in ['pdf']:
                result = self.generate_pdf_file(content, filename)
                return "успешно создан" in result.lower()
            elif format_lower in ['txt', 'text']:
                result = self.generate_txt_file(content, filename)
                return "успешно создан" in result.lower()
            elif format_lower in ['json']:
                result = self.generate_json_file(content, filename)
                return "успешно создан" in result.lower()
            elif format_lower in ['csv']:
                result = self.generate_csv_file(content, filename)
                return "успешно создан" in result.lower()
            elif format_lower in ['html', 'htm']:
                result = self.generate_html_file(content, filename)
                return "успешно создан" in result.lower()
            elif format_lower in ['xml']:
                result = self.generate_xml_file(content, filename)
                return "успешно создан" in result.lower()
            elif format_lower in ['bat', 'batch']:
                result = self.generate_bat_file(content, filename)
                return "успешно создан" in result.lower()
            else:
                logger.error(f"Неподдерживаемый формат файла: {file_format}")
                return False
                
        except Exception as e:
            logger.error(f"Ошибка при генерации файла: {str(e)}")
            return False

    def extract_text_from_image(self, image_path: str) -> Tuple[str, str]:
        """
        Извлекает текст из изображения с помощью OCR
        
        Args:
            image_path: Путь к изображению
        
        Returns:
            Tuple[str, str]: (extracted_text, error_message)
        """
        try:
            if not self._ensure_ocr_initialized():
                return "", "OCR не может быть инициализирован"
            
            if not os.path.exists(image_path):
                return "", f"Файл {image_path} не найден"
            
            logger.info(f"📖 Извлекаю текст из изображения: {image_path}")
            
            # Выполняем OCR (дополнительная проверка для типизации)
            if self.ocr_reader is None:
                return "", "OCR reader не инициализирован после попытки загрузки"
            results = self.ocr_reader.readtext(image_path)
            
            if not results:
                return "", "Текст на изображении не обнаружен"
            
            # Объединяем весь найденный текст
            extracted_text = []
            for (bbox, text, confidence) in results:
                # Фильтруем результаты с низкой уверенностью
                if confidence > 0.3:  # Порог уверенности 30%
                    extracted_text.append(text.strip())
            
            if not extracted_text:
                return "", "Текст не обнаружен с достаточной уверенностью"
            
            final_text = '\n'.join(extracted_text)
            logger.info(f"✅ Извлечено {len(extracted_text)} блоков текста")
            
            return final_text, ""
            
        except Exception as e:
            error_msg = f"Ошибка OCR: {str(e)}"
            logger.error(error_msg)
            return "", error_msg

    def extract_text_from_image_object(self, image_obj) -> Tuple[str, str]:
        """
        Извлекает текст из PIL.Image объекта с помощью OCR
        
        Args:
            image_obj: PIL.Image объект
        
        Returns:
            Tuple[str, str]: (extracted_text, error_message)
        """
        try:
            if not self._ensure_ocr_initialized():
                return "", "OCR не может быть инициализирован"
            
            import numpy as np
            
            # Конвертируем PIL Image в numpy array для EasyOCR
            image_array = np.array(image_obj)
            
            logger.info(f"📖 Извлекаю текст из изображения (объект)")
            
            # Выполняем OCR (дополнительная проверка для типизации)
            if self.ocr_reader is None:
                return "", "OCR reader не инициализирован после попытки загрузки"
            results = self.ocr_reader.readtext(image_array)
            
            if not results:
                return "", "Текст на изображении не обнаружен"
            
            # Объединяем весь найденный текст
            extracted_text = []
            for (bbox, text, confidence) in results:
                # Фильтруем результаты с низкой уверенностью
                if confidence > 0.3:  # Порог уверенности 30%
                    extracted_text.append(text.strip())
            
            if not extracted_text:
                return "", "Текст не обнаружен с достаточной уверенностью"
            
            final_text = '\n'.join(extracted_text)
            logger.info(f"✅ Извлечено {len(extracted_text)} блоков текста")
            
            return final_text, ""
            
        except Exception as e:
            error_msg = f"Ошибка OCR объекта: {str(e)}"
            logger.error(error_msg)
            return "", error_msg

    def should_use_ocr_on_image(self, vision_description: str) -> bool:
        """
        Определяет, нужно ли применять OCR на основе описания изображения.
        Теперь использует интеллектуальную логику вместо простого поиска ключевых слов.
        
        Args:
            vision_description: Описание изображения от vision модели
            
        Returns:
            bool: True если нужно применить OCR, False если нет
        """
        return self.should_use_ocr_intelligently(vision_description)

    def process_image_with_smart_ocr(self, image_path: str, vision_description: str = "", force_ocr: bool = False) -> Tuple[str, str, str]:
        """
        Обрабатывает изображение с умным применением OCR
        
        Args:
            image_path: Путь к изображению
            vision_description: Описание от vision модели (опционально)
            force_ocr: Принудительно применить OCR
        
        Returns:
            Tuple[str, str, str]: (extracted_text, description, error_message)
        """
        extracted_text = ""
        ocr_error = ""
        
        # Решаем, применять ли OCR
        should_ocr = force_ocr or self.should_use_ocr_on_image(vision_description)
        
        if should_ocr:
            logger.info("🔍 Применяю OCR к изображению")
            text, error = self.extract_text_from_image(image_path)
            if text:
                extracted_text = f"\n\n[Текст с изображения]:\n{text}"
                logger.info(f"✅ Извлечен текст: {len(text)} символов")
            elif error:
                ocr_error = error
                logger.warning(f"⚠️ OCR не удался: {error}")
        else:
            logger.info("⏭️ OCR не применяется - текст не обнаружен в описании")
        
        # Объединяем описание и извлеченный текст
        full_description = vision_description
        if extracted_text:
            full_description += extracted_text
        
        return extracted_text, full_description, ocr_error

    def auto_disable_tools(self, tool_name: Optional[str] = None):
        """Автоматически выключает инструмент через заданное время после использования"""
        import threading
        import time
        import gc
        
        def disable_tool(tool_name):
            time.sleep(self.auto_disable_delay)
            
            if tool_name == 'image_generation':
                if hasattr(self, 'use_image_generation') and self.use_image_generation:
                    self.use_image_generation = False
                    
                    # Выгружаем pipeline из памяти
                    if hasattr(self, 'current_pipeline') and self.current_pipeline is not None:
                        try:
                            # Освобождаем GPU память
                            if hasattr(self.current_pipeline, 'to'):
                                self.current_pipeline.to('cpu')
                            del self.current_pipeline
                            self.current_pipeline = None
                        except Exception as e:
                            self.logger.warning(f"⚠️ Ошибка при выгрузке pipeline: {e}")
                    
                    # Принудительная очистка GPU памяти
                    try:
                        import torch
                        if torch.cuda.is_available():
                            torch.cuda.empty_cache()
                            torch.cuda.synchronize()
                    except Exception as e:
                        self.logger.warning(f"⚠️ Ошибка при очистке GPU памяти: {e}")
                    
                    # Сборка мусора
                    gc.collect()
                    self.logger.info(f"🔧 Автоматически выключил {tool_name} и освободил память")
                    
            elif tool_name == 'vision':
                if hasattr(self, 'use_vision') and self.use_vision:
                    self.use_vision = False
                    
                    # Выгружаем vision модели если они есть
                    vision_attrs = ['vision_model', 'vision_processor', 'vision_pipeline']
                    for attr in vision_attrs:
                        if hasattr(self, attr):
                            try:
                                model = getattr(self, attr)
                                if model is not None and hasattr(model, 'to'):
                                    model.to('cpu')
                                delattr(self, attr)
                            except Exception as e:
                                self.logger.warning(f"⚠️ Ошибка при выгрузке {attr}: {e}")
                    
                    # Очистка GPU памяти
                    try:
                        import torch
                        if torch.cuda.is_available():
                            torch.cuda.empty_cache()
                    except Exception:
                        pass
                    
                    gc.collect()
                    self.logger.info(f"🔧 Автоматически выключил {tool_name} и освободил память")
                    
            elif tool_name == 'audio':
                if hasattr(self, 'use_audio') and self.use_audio:
                    self.use_audio = False
                    
                    # Выгружаем audio модели если они есть
                    audio_attrs = ['whisper_model', 'audio_model', 'tts_model']
                    for attr in audio_attrs:
                        if hasattr(self, attr):
                            try:
                                model = getattr(self, attr)
                                if model is not None and hasattr(model, 'to'):
                                    model.to('cpu')
                                delattr(self, attr)
                            except Exception as e:
                                self.logger.warning(f"⚠️ Ошибка при выгрузке {attr}: {e}")
                    
                    # Очистка GPU памяти
                    try:
                        import torch
                        if torch.cuda.is_available():
                            torch.cuda.empty_cache()
                    except Exception:
                        pass
                    
                    gc.collect()
                    self.logger.info(f"🔧 Автоматически выключил {tool_name} и освободил память")
        
        # Если указан конкретный инструмент, запускаем таймер только для него
        if tool_name:
            # Отменяем предыдущий таймер если он есть
            if tool_name in self.tool_timers and self.tool_timers[tool_name].is_alive():
                self.tool_timers[tool_name].cancel() if hasattr(self.tool_timers[tool_name], 'cancel') else None
                
            timer = threading.Thread(target=disable_tool, args=(tool_name,), daemon=True)
            self.tool_timers[tool_name] = timer
            timer.start()
            self.logger.info(f"⏰ Запустил таймер автоматического выключения для {tool_name}")
        else:
            # Запускаем таймеры для всех активных инструментов
            for tool in ['image_generation', 'vision', 'audio']:
                if tool not in self.tool_timers or not self.tool_timers[tool].is_alive():
                    timer = threading.Thread(target=disable_tool, args=(tool,), daemon=True)
                    self.tool_timers[tool] = timer
                    timer.start()
                    self.logger.info(f"⏰ Запустил таймер автоматического выключения для {tool}")
                
    def _log(self, message: str, level: str = "INFO"):
        """Логирование с временной меткой в файл и консоль"""
        timestamp = time.strftime("%H:%M:%S")
        formatted_message = f"[{timestamp}] {level}: {message}"
        
        # Логируем в файл
        if level == "ERROR":
            logger.error(message)
        elif level == "WARNING":
            logger.warning(message)
        else:
            logger.info(message)
        
        # Выводим в консоль
        print(formatted_message)
    
    def get_context_info(self) -> str:
        """Возвращает информацию о текущем использовании контекста"""
        return f"Контекст: {self.current_context_length:,} токенов / {self.safe_context_length:,} (безопасный) / {self.max_context_length:,} (максимум)"

    def add_performance_metric(self, action: str, response_time: float, context_length: int = 0):
        """Добавляет метрику производительности"""
        metric = {
            "timestamp": time.time(),
            "action": action,
            "response_time": response_time,
            "context_length": context_length
        }
        self.performance_metrics.append(metric)
        
        # Ограничиваем количество метрик
        if len(self.performance_metrics) > 100:
            self.performance_metrics.pop(0)
    
    def get_performance_stats(self) -> Dict[str, Any]:
        """Возвращает статистику производительности"""
        if not self.performance_metrics:
            return {"total_actions": 0, "avg_response_time": 0, "recent_metrics": []}
        
        total_actions = len(self.performance_metrics)
        avg_response_time = sum(m["response_time"] for m in self.performance_metrics) / total_actions
        recent_metrics = self.performance_metrics[-10:]  # Последние 10 метрик
        
        return {
            "total_actions": total_actions,
            "avg_response_time": round(avg_response_time, 3),
            "recent_metrics": recent_metrics
        }

    def take_screenshot(self) -> str:
        """
        Делает скриншот основного монитора и возвращает base64
        """
        try:
            # mss уже импортирован в начале файла
            with mss.mss() as sct:
                # Скриншот основного монитора
                monitor = sct.monitors[1]  # 0 - все мониторы, 1 - основной
                screenshot = sct.grab(monitor)
                # Конвертируем в PIL Image
                img = Image.frombytes("RGB", screenshot.size, screenshot.rgb)
                # Сжимаем для экономии места
                img.thumbnail((1280, 720), Image.Resampling.LANCZOS)
                # Конвертируем в base64
                buf = BytesIO()
                img.save(buf, format="PNG", optimize=True)
                return base64.b64encode(buf.getvalue()).decode("ascii")
        except ImportError:
            logger.warning("mss не установлен, используем pyautogui")
            try:
                # pyautogui уже импортирован в начале файла
                screenshot = pyautogui.screenshot()
                screenshot.thumbnail((1280, 720), Image.Resampling.LANCZOS)
                buf = BytesIO()
                screenshot.save(buf, format="PNG", optimize=True)
                return base64.b64encode(buf.getvalue()).decode("ascii")
            except ImportError:
                logger.error("pyautogui не установлен")
                return ""
        except Exception as e:
            logger.error(f"Ошибка создания скриншота: {e}")
            return ""

    def resolve_path(self, path: str) -> str:
        """
        Разрешает относительный путь к изображению
        """
        if os.path.isabs(path):
            return path
        
        # Проверяем в базовой директории
        full_path = os.path.join(self.base_dir, path)
        if os.path.exists(full_path):
            return full_path
            
        # Проверяем в папке Photos
        photos_path = os.path.join(self.base_dir, "Photos", path)
        if os.path.exists(photos_path):
            return photos_path
            
        # Проверяем в папке Images
        images_path = os.path.join(self.base_dir, "Images", path)
        if os.path.exists(images_path):
            return images_path
            
        return path  # Возвращаем исходный путь если не найден

    def analyze_image_with_vision(self, image_path: str) -> str:
        """
        Анализирует изображение с помощью vision модели
        """
        try:
            # Конвертируем изображение в base64
            image_b64 = image_to_base64_balanced(image_path)
            if not image_b64:
                return "Ошибка: не удалось загрузить изображение"
            
            # Автоматически включаем vision если не включен
            if not getattr(self, 'use_vision', False):
                self.logger.info("🔍 Автоматически включаю анализ изображений...")
                self.use_vision = True
                self.auto_disable_tools("vision")
            
            # Вызываем vision модель
            return self.call_vision_model(image_b64)
            
        except Exception as e:
            self.logger.error(f"Ошибка анализа изображения: {e}")
            return f"Ошибка анализа изображения: {str(e)}"

    def move_mouse(self, x: int, y: int) -> Dict[str, Any]:
        """Перемещение мыши в координаты (x, y)"""
        try:
            # pyautogui уже импортирован в начале файла
            pyautogui.moveTo(x, y, duration=0.2)
            return {"success": True, "message": f"Мышь перемещена в ({x}, {y})"}
        except ImportError:
            return {"success": False, "error": "pyautogui не установлен"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def left_click(self, x: int, y: int) -> Dict[str, Any]:
        """Клик левой кнопкой мыши по координатам (x, y)"""
        try:
            # pyautogui уже импортирован в начале файла
            pyautogui.click(x, y)
            return {"success": True, "message": f"ЛКМ клик в ({x}, {y})"}
        except ImportError:
            return {"success": False, "error": "pyautogui не установлен"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def right_click(self, x: int, y: int) -> Dict[str, Any]:
        """Клик правой кнопкой мыши по координатам (x, y)"""
        try:
            # pyautogui уже импортирован в начале файла
            pyautogui.rightClick(x, y)
            return {"success": True, "message": f"ПКМ клик в ({x}, {y})"}
        except ImportError:
            return {"success": False, "error": "pyautogui не установлен"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def scroll(self, pixels: int) -> Dict[str, Any]:
        """Прокрутка колесиком мыши. Положительные значения - вверх, отрицательные - вниз"""
        try:
            # pyautogui уже импортирован в начале файла
            pyautogui.scroll(pixels)
            direction = "вверх" if pixels > 0 else "вниз"
            return {"success": True, "message": f"Прокрутка {direction} на {abs(pixels)} пикселей"}
        except ImportError:
            return {"success": False, "error": "pyautogui не установлен"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def mouse_down(self, x: int, y: int) -> Dict[str, Any]:
        """Зажать левую кнопку мыши в координатах (x, y)"""
        try:
            # pyautogui уже импортирован в начале файла
            pyautogui.moveTo(x, y)
            pyautogui.mouseDown(button='left')
            return {"success": True, "message": f"ЛКМ зажата в ({x}, {y})"}
        except ImportError:
            return {"success": False, "error": "pyautogui не установлен"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def mouse_up(self, x: int, y: int) -> Dict[str, Any]:
        """Отпустить левую кнопку мыши в координатах (x, y)"""
        try:
            # pyautogui уже импортирован в начале файле
            pyautogui.moveTo(x, y)
            pyautogui.mouseUp(button='left')
            return {"success": True, "message": f"ЛКМ отпущена в ({x}, {y})"}
        except ImportError:
            return {"success": False, "error": "pyautogui не установлен"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def drag_and_drop(self, x1: int, y1: int, x2: int, y2: int) -> Dict[str, Any]:
        """Перетащить мышью из (x1, y1) в (x2, y2)"""
        try:
            import pyautogui
            pyautogui.dragTo(x2, y2, duration=0.5, button='left')
            return {"success": True, "message": f"Перетаскивание из ({x1}, {y1}) в ({x2}, {y2})"}
        except ImportError:
            return {"success": False, "error": "pyautogui не установлен"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def type_text(self, text: str) -> Dict[str, Any]:
        """Ввести текст"""
        try:
            import pyautogui
            pyautogui.typewrite(text, interval=0.05)
            return {"success": True, "message": f"Введён текст: {text}"}
        except ImportError:
            return {"success": False, "error": "pyautogui не установлен"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def start_continuous_recording(self):
        """Запуск постоянной голосовой записи"""
        if self.continuous_recording:
            return
        
        self.continuous_recording = True
        self.recording_thread = threading.Thread(target=self._continuous_recording_worker, daemon=True)
        self.recording_thread.start()
        logger.info("Постоянная голосовая запись запущена")

    def stop_continuous_recording(self):
        """Остановка постоянной голосовой записи"""
        self.continuous_recording = False
        if self.recording_thread:
            self.recording_thread.join(timeout=2)
        logger.info("Постоянная голосовая запись остановлена")

    def _continuous_recording_worker(self):
        """Воркер для постоянной голосовой записи (заглушка - нужна реализация через веб-интерфейс)"""
        # Эта функция будет вызываться из веб-интерфейса через API
        while self.continuous_recording:
            try:
                # Проверяем очередь на наличие аудиочанков
                if not self.audio_queue.empty():
                    audio_data = self.audio_queue.get_nowait()
                    # Обрабатываем аудио
                    self._process_audio_chunk(audio_data)
                time.sleep(0.1)
            except Exception as e:
                logger.error(f"Ошибка в continuous recording worker: {e}")

    def _process_audio_chunk(self, audio_data: bytes):
        """Обработка чанка аудио из постоянной записи"""
        try:
            # Сохраняем чанк во временный файл
            temp_dir = os.path.join(os.path.dirname(__file__), "temp_audio")
            os.makedirs(temp_dir, exist_ok=True)
            temp_file = os.path.join(temp_dir, f"chunk_{int(time.time())}.wav")
            
            with open(temp_file, 'wb') as f:
                f.write(audio_data)
            
            # Распознаём аудио
            transcript = self.transcribe_audio_whisper(temp_file, use_separator=False)
            
            if transcript and not transcript.startswith("[Whisper error]"):
                # Проверяем, содержит ли текст команду или имя "Алиса"
                if self._is_valid_command(transcript):
                    logger.info(f"Получена команда из голоса: {transcript}")
                    # Делаем скриншот для контекста
                    screenshot_b64 = self.take_screenshot()
                    vision_desc = ""
                    if screenshot_b64:
                        vision_desc = self.call_vision_model(screenshot_b64)
                    
                    # Формируем запрос для мозга
                    brain_input = f"[Скриншот экрана]: {vision_desc}\n\nГолосовая команда: {transcript}"
                    
                    # Отправляем в мозг
                    ai_response = self.call_brain_model(brain_input)
                    self.process_ai_response(ai_response)
                else:
                    # Игнорируем бессмысленные фразы
                    pass
            
            # Удаляем временный файл
            try:
                os.remove(temp_file)
            except Exception:
                pass
                
        except Exception as e:
            logger.error(f"Ошибка обработки аудиочанка: {e}")

    def _is_valid_command(self, text: str) -> bool:
        """Всегда возвращает True — фильтрация отключена, все команды проходят к нейросети"""
        return True

    def call_vision_model(self, image_base64: str) -> str:
        """
        Отправка изображения только в vision-модель ("глаза")
        Возвращает описание изображения (текст).
        """
        start_time = time.time()
        
        # Автоматически включаем vision модель при необходимости
        if not getattr(self, 'use_vision', False):
            logger.info("🔧 Автоматически включаю vision модель")
            self.use_vision = True
            # Запускаем таймер автоматического выключения
            self.auto_disable_tools("vision")
        
        try:
            payload = {
                "model": "moondream2-llamafile",
                "messages": [
                    {"role": "user", "content": [
                        {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{image_base64}"}}
                    ]}
                ],
                "temperature": 0.0,
                "max_tokens": 2048,
                "stream": False
            }
            logger.info("Отправляю изображение в vision-модель (глаза)")
            response = requests.post(
                f"{self.lm_studio_url}/v1/chat/completions",
                json=payload,
                headers={"Content-Type": "application/json"}
            )
            if response.status_code == 200:
                result = response.json()
                return result["choices"][0]["message"]["content"].strip()
            else:
                error_msg = f"Ошибка vision-модели: {response.status_code} - {response.text}"
                logger.error(error_msg)
                return f"[Vision error] {error_msg}"
        except Exception as e:
            error_msg = f"Исключение vision: {str(e)}"
            logger.error(error_msg)
            return f"[Vision error] {error_msg}"
        finally:
            # Записываем метрику производительности
            response_time = time.time() - start_time
            self.add_performance_metric("vision_processing", response_time)
            logger.info(f"👁️ Vision обработал за {response_time:.2f} сек")

    def call_brain_model(self, user_message: str, vision_desc: str = "") -> str:
        """
        Отправка текстового запроса (и опционально описания изображения) в "мозг" (текстовая модель)
        """
        start_time = time.time()
        # Инициализируем переменные для использования в except блоке
        processed_message = user_message
        messages = []
        
        try:
            # Проверяем, является ли сообщение командой для загрузки модуля
            if self.prompt_loader.is_module_command(user_message.strip()):
                module_content = self.prompt_loader.load_module(user_message.strip())
                logger.info(f"📚 Загружен модуль для команды: {user_message.strip()}")
                return module_content
            
            # Обрабатываем сообщение плагинами (hook on_message_received)
            processed_message = user_message
            if self.plugin_manager:
                processed_message = self.plugin_manager.call_hook_message_received(user_message, self)
            
            # Улучшаем промпт с помощью памяти ChromaDB
            enhanced_system_prompt = self.enhance_prompt_with_memory(processed_message, self.system_prompt)
            
            # Добавляем информацию о доступных плагинах в системный промпт
            if self.plugin_manager:
                plugin_info = self._get_plugin_info_for_prompt()
                if plugin_info:
                    enhanced_system_prompt += f"\n\n{plugin_info}"
            
            messages: List[Dict[str, Any]] = [
                {"role": "system", "content": enhanced_system_prompt}
            ]
            
            # Добавляем историю разговора с управлением контекстом
            messages.extend(self.conversation_history)
            
            # Добавляем описание изображения, если есть
            if vision_desc:
                messages.append({"role": "user", "content": vision_desc})
            # Добавляем основной запрос пользователя
            messages.append({"role": "user", "content": user_message})
            
            # Динамическое управление контекстом - подсчитываем длину
            total_length = sum(len(str(msg.get("content", ""))) for msg in messages)
            self.current_context_length = total_length
            
            # Обрезаем контекст если необходимо (используем предварительную оценку)
            if total_length > self.safe_context_length:
                self.conversation_history = self.conversation_history[-5:]  # Оставляем только 5 последних сообщений
                logger.warning(f"Превышение безопасного контекста ({total_length:,} символов > {self.safe_context_length:,}) - аккуратная обрезка истории")
            
            # Пересобираем сообщения после обрезки
            messages = [{"role": "system", "content": enhanced_system_prompt}]
            messages.extend(self.conversation_history)
            if vision_desc:
                messages.append({"role": "user", "content": vision_desc})
            messages.append({"role": "user", "content": user_message})
            
            payload = {
                "model": self.brain_model_id if hasattr(self, 'brain_model_id') and self.brain_model_id else self.brain_model,
                "messages": messages,
                "temperature": 0.1,
                "max_tokens": 32767,
                "stream": False
            }
            logger.info(f"Отправляю запрос в мозг: {user_message[:100]}...")
            response = requests.post(
                f"{self.lm_studio_url}/v1/chat/completions",
                json=payload,
                headers={"Content-Type": "application/json"}
            )
            if response.status_code == 200:
                result = response.json()
                ai_response = result["choices"][0]["message"]["content"].strip()
                
                # Извлекаем информацию о токенах из ответа модели
                usage_info = result.get("usage", {})
                prompt_tokens = usage_info.get("prompt_tokens", 0)
                completion_tokens = usage_info.get("completion_tokens", 0)
                total_tokens = usage_info.get("total_tokens", 0)
                
                # Обновляем текущий размер контекста на основе total_tokens
                if total_tokens > 0:
                    self.current_context_length = total_tokens
                    logger.info(f"📊 Реальные токены: prompt={prompt_tokens}, completion={completion_tokens}, total={total_tokens}")
                    
                    # Обрезаем контекст на основе реальных токенов
                    self._trim_context_if_needed()
                
                # Добавляем в историю разговора (если ответ не пустой)
                if ai_response and ai_response != "{}":
                    self.conversation_history.append({"role": "user", "content": processed_message})
                    self.conversation_history.append({"role": "assistant", "content": ai_response})
                    
                    # Автоматически сохраняем диалог в ChromaDB
                    self.auto_save_conversation(processed_message, ai_response, vision_desc)
                    
                    # Извлекаем предпочтения пользователя из диалога
                    self.extract_preferences_from_response(processed_message, ai_response)
                
                # Обрабатываем ответ плагинами (hook on_response_generated)
                final_response = ai_response
                if self.plugin_manager:
                    final_response = self.plugin_manager.call_hook_response_generated(ai_response, self)
                
                return final_response
            else:
                error_msg = f"Ошибка brain-модели: {response.status_code} - {response.text}"
                logger.error(error_msg)
                
                # Попытка переподключения при ошибке
                if response.status_code in [404, 500, 503]:
                    logger.info("🔄 Попытка переподключения к модели мозга...")
                    if self._reconnect_brain_model():
                        # Повторяем запрос после переподключения
                        try:
                            retry_response = requests.post(
                                f"{self.lm_studio_url}/v1/chat/completions",
                                json=payload,
                                headers={"Content-Type": "application/json"}
                            )
                            if retry_response.status_code == 200:
                                result = retry_response.json()
                                ai_response = result["choices"][0]["message"]["content"].strip()
                                logger.info("✅ Повторный запрос после переподключения успешен")
                                
                                # Добавляем в историю разговора
                                if ai_response and ai_response != "{}":
                                    self.conversation_history.append({"role": "user", "content": processed_message})
                                    self.conversation_history.append({"role": "assistant", "content": ai_response})
                                    self.auto_save_conversation(processed_message, ai_response, vision_desc)
                                    self.extract_preferences_from_response(processed_message, ai_response)
                                
                                # Обрабатываем ответ плагинами
                                final_response = ai_response
                                if self.plugin_manager:
                                    final_response = self.plugin_manager.call_hook_response_generated(ai_response, self)
                                
                                return final_response
                        except Exception as retry_e:
                            logger.error(f"❌ Ошибка повторного запроса: {retry_e}")
                
                return f"[Brain error] {error_msg}"
        except Exception as e:
            error_msg = f"Исключение brain: {str(e)}"
            logger.error(error_msg)
            
            # Попытка переподключения при исключении (может быть связано с обрывом соединения)
            if "Connection" in str(e) or "timeout" in str(e).lower() or "refused" in str(e).lower():
                logger.info("🔄 Попытка переподключения к модели мозга из-за проблем соединения...")
                if self._reconnect_brain_model():
                    # Повторяем запрос после переподключения
                    try:
                        payload = {
                            "model": self.brain_model_id if hasattr(self, 'brain_model_id') and self.brain_model_id else self.brain_model,
                            "messages": messages,
                            "temperature": 0.1,
                            "max_tokens": 32767,
                            "stream": False
                        }
                        retry_response = requests.post(
                            f"{self.lm_studio_url}/v1/chat/completions",
                            json=payload,
                            headers={"Content-Type": "application/json"}
                        )
                        if retry_response.status_code == 200:
                            result = retry_response.json()
                            ai_response = result["choices"][0]["message"]["content"].strip()
                            logger.info("✅ Повторный запрос после переподключения успешен")
                            
                            # Добавляем в историю разговора
                            if ai_response and ai_response != "{}":
                                self.conversation_history.append({"role": "user", "content": processed_message})
                                self.conversation_history.append({"role": "assistant", "content": ai_response})
                                self.auto_save_conversation(processed_message, ai_response, vision_desc)
                                self.extract_preferences_from_response(processed_message, ai_response)
                            
                            # Обрабатываем ответ плагинами
                            final_response = ai_response
                            if self.plugin_manager:
                                final_response = self.plugin_manager.call_hook_response_generated(ai_response, self)
                            
                            return final_response
                    except Exception as retry_e:
                        logger.error(f"❌ Ошибка повторного запроса после исключения: {retry_e}")
            
            return f"[Brain error] {error_msg}"
        finally:
            # Записываем метрику производительности
            response_time = time.time() - start_time
            self.add_performance_metric("brain_response", response_time, self.current_context_length)
            logger.info(f"🧠 Мозг ответил за {response_time:.2f} сек")

    def execute_powershell(self, command: str) -> Dict[str, Any]:
        """
        Выполнение PowerShell команды
        
        Args:
            command: PowerShell команда
        
        Returns:
            Словарь с результатом выполнения
        """
        try:
            orig_command = command
            # Автоисправление: заменяем '&&' на PowerShell-совместимый синтаксис
            if '&&' in command:
                parts = [p.strip() for p in command.split('&&')]
                # Если первая часть cd, делаем push-location, затем вторую команду
                if parts[0].lower().startswith('cd '):
                    dir_path = parts[0][3:].strip().strip('"\'')
                    command = f"Push-Location '{dir_path}'; {parts[1]} ; Pop-Location"
                else:
                    # Просто объединяем через ';'
                    command = ' ; '.join(parts)
                logger.info(f"PowerShell: автоисправлен '&&' -> ';' или Push-Location: {command}")
            logger.info(f"Выполняю PowerShell: {command}")
            # Выполняем команду PowerShell с декодированием cp1251 и защитой
            result = subprocess.run(
                ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", command],
                capture_output=True,
                text=True,
                encoding='cp1251',
                errors='replace',
                timeout=60
            )
            success = result.returncode == 0
            # Защита от None
            output = (result.stdout if success else result.stderr) or ""
            logger.info(f"PowerShell результат (код: {result.returncode}): {output[:200]}...")
            return {
                "success": success,
                "returncode": result.returncode,
                "output": output,
                "error": (result.stderr or "") if not success else ""
            }
        except subprocess.TimeoutExpired:
            error_msg = "Команда превысила лимит времени выполнения (60 сек)"
            logger.error(error_msg)
            return {"success": False, "returncode": -1, "output": "", "error": error_msg}
        except Exception as e:
            error_msg = f"Ошибка выполнения PowerShell: {str(e)}"
            logger.error(error_msg)
            return {"success": False, "returncode": -1, "output": "", "error": error_msg}

    def google_search(self, query: str, num_results: int = 10) -> List[Dict[str, str]]:
        """
        Поиск в Google Custom Search API
        
        Args:
            query: Поисковый запрос
            num_results: Количество результатов для парсинга (по умолчанию 10)
            
        Returns:
            Список результатов поиска
        """
        try:
            if not self.google_api_key or not self.google_cse_id:
                return [{"error": "Google API ключ или CSE ID не настроены"}]
            
            logger.info(f"Выполняю поиск Google: {query}")
            
            # Кодируем запрос для URL
            encoded_query = urllib.parse.quote(query)
            
            # Формируем URL для Google Custom Search API (максимум 10 результатов)
            url = f"https://www.googleapis.com/customsearch/v1?key={self.google_api_key}&cx={self.google_cse_id}&q={encoded_query}&num=10"
            
            response = requests.get(url, timeout=10)
            
            if response.status_code == 200:
                data = response.json()
                
                if "items" not in data:
                    return [{"error": "Результаты не найдены"}]
                
                # Берем первые num_results результатов для парсинга (максимум 10)
                actual_results = min(num_results, 10)
                search_results = []
                for i, item in enumerate(data["items"][:actual_results]):
                    result = {
                        "title": item.get("title", ""),
                        "url": item.get("link", ""),
                        "snippet": item.get("snippet", "")
                    }
                    
                    # Пытаемся получить содержимое страницы
                    try:
                        page_response = requests.get(result["url"], timeout=5, headers={
                            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
                        })
                        if page_response.status_code == 200:
                            # Берем первые 2000 символов текста для полного анализа
                            content = page_response.text[:2000]
                            result["content"] = content
                        else:
                            result["content"] = "Не удалось получить содержимое страницы"
                    except:
                        result["content"] = "Ошибка при получении содержимого страницы"
                    
                    search_results.append(result)
                    logger.info(f"Получен результат {i+1}: {result['title']}")
                
                logger.info(f"Поиск завершен: найдено {len(search_results)} результатов")
                return search_results
            else:
                error_msg = f"Ошибка Google Search API: {response.status_code}"
                logger.error(error_msg)
                return [{"error": error_msg}]
                
        except Exception as e:
            error_msg = f"Ошибка поиска Google: {str(e)}"
            logger.error(error_msg)
            return [{"error": error_msg}]

    def process_ai_response(self, ai_response: str) -> bool:
        """Light wrapper to keep `process_ai_response` simple for static analysers.

        Реализует прокси к подробной реализации `_process_ai_response_impl`.
        Это уменьшает сложность видимой функции для Pylance.
        """
        return self._process_ai_response_impl(ai_response)

    # --- Разделённые обработчики действий (уменьшают сложность основного метода) ---
    def _is_english_simple(self, s: str) -> bool:
        if not s:
            return False
        allowed_chars = 0
        total_chars = len(s)
        for c in s:
            if (c.isascii() and (c.isalpha() or c.isdigit()) or
                c in '.,;:-_=+!@#$%^&*()[]{}<>?/\\|`~\'\" ' or
                c.isspace()):
                allowed_chars += 1
        return allowed_chars / total_chars > 0.85

    def _handle_powershell(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        command = action_data.get("command", "")
        description = action_data.get("description", "")
        logger.info(f"\n🔧 ВЫПОЛНЕНИЕ КОМАНДЫ: {description}")
        logger.info(f"📝 Команда: {command}")
        result = self.execute_powershell(command)
        if result["success"]:
            feedback = f"Команда '{command}' выполнена успешно. Результат: {result['output']}"
        else:
            feedback = f"Команда '{command}' завершилась с ошибкой: {result.get('error','') }"
        follow_up = self.call_brain_model(feedback)
        return follow_up

    def _handle_search(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        query = action_data.get("query", "")
        description = action_data.get("description", "")
        logger.info(f"\n🔍 ПОИСК В ИНТЕРНЕТЕ: {description}")
        logger.info(f"🔎 Запрос: {query}")
        search_results = self.google_search(query)
        results_text = "Результаты поиска:\n"
        for i, result in enumerate(search_results, 1):
            if "error" in result:
                results_text += f"{i}. Ошибка: {result['error']}\n"
            else:
                results_text += f"{i}. {result['title']}\n"
                results_text += f"   URL: {result['url']}\n"
                results_text += f"   Описание: {result['snippet']}\n"
                results_text += f"   Содержимое: {result.get('content', '')}\n\n"
        logger.info("✅ Поиск завершен")
        follow_up = self.call_brain_model(f"Результаты поиска по запросу '{query}': {results_text}")
        return follow_up

    def _handle_take_screenshot(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        logger.info(f"\n📸 СОЗДАНИЕ СКРИНШОТА")
        if not getattr(self, 'use_vision', False):
            logger.info("🔧 Автоматически включаю vision модель")
            self.use_vision = True
            self.auto_disable_tools("vision")
        screenshot_b64 = self.take_screenshot()
        if screenshot_b64:
            vision_desc = self.call_vision_model(screenshot_b64)
            feedback = f"Скриншот экрана получен. Описание от vision-модели: {vision_desc}"
        else:
            feedback = "Не удалось создать скриншот"
        follow_up = self.call_brain_model(feedback)
        return follow_up

    def _handle_move_mouse(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        x = action_data.get("x", 0)
        y = action_data.get("y", 0)
        description = action_data.get("description", "")
        logger.info(f"\n🖱️ ПЕРЕМЕЩЕНИЕ МЫШИ: {description}")
        result = self.move_mouse(x, y)
        feedback = f"Мышь перемещена в координаты ({x}, {y})" if result.get("success") else f"Ошибка перемещения мыши: {result.get('error','') }"
        follow_up = self.call_brain_model(feedback)
        return follow_up

    def _handle_left_click(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        x = action_data.get("x", 0)
        y = action_data.get("y", 0)
        result = self.left_click(x, y)
        feedback = f"Клик ЛКМ выполнен в координатах ({x}, {y})" if result.get("success") else f"Ошибка клика: {result.get('error','') }"
        follow_up = self.call_brain_model(feedback)
        return follow_up

    def _handle_right_click(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        x = action_data.get("x", 0)
        y = action_data.get("y", 0)
        result = self.right_click(x, y)
        feedback = f"Клик ПКМ выполнен в координатах ({x}, {y})" if result.get("success") else f"Ошибка клика ПКМ: {result.get('error','') }"
        follow_up = self.call_brain_model(feedback)
        return follow_up

    def _handle_scroll(self, action: str, action_data: Dict[str, Any]) -> Union[bool, str]:
        pixels = action_data.get("pixels", 100)
        if action == "scroll_down":
            pixels = -pixels
        result = self.scroll(pixels)
        feedback = f"Прокрутка выполнена: {result.get('message','') }" if result.get("success") else f"Ошибка прокрутки: {result.get('error','') }"
        follow_up = self.call_brain_model(feedback)
        return follow_up

    def _handle_mouse_down(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        x = action_data.get("x", 0)
        y = action_data.get("y", 0)
        result = self.mouse_down(x, y)
        feedback = f"ЛКМ зажата в координатах ({x}, {y})" if result.get("success") else f"Ошибка зажатия ЛКМ: {result.get('error','') }"
        follow_up = self.call_brain_model(feedback)
        return follow_up

    def _handle_mouse_up(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        x = action_data.get("x", 0)
        y = action_data.get("y", 0)
        result = self.mouse_up(x, y)
        feedback = f"ЛКМ отпущена в координатах ({x}, {y})" if result.get("success") else f"Ошибка отпускания ЛКМ: {result.get('error','') }"
        follow_up = self.call_brain_model(feedback)
        return follow_up

    def _handle_drag_and_drop(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        x1 = action_data.get("x1", 0)
        y1 = action_data.get("y1", 0)
        x2 = action_data.get("x2", 0)
        y2 = action_data.get("y2", 0)
        result = self.drag_and_drop(x1, y1, x2, y2)
        feedback = f"Перетаскивание выполнено от ({x1}, {y1}) к ({x2}, {y2})" if result.get("success") else f"Ошибка перетаскивания: {result.get('error','') }"
        follow_up = self.call_brain_model(feedback)
        return follow_up

    def _handle_type_text(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        text = action_data.get("text", "")
        result = self.type_text(text)
        feedback = f"Текст введён: {text}" if result.get("success") else f"Ошибка ввода текста: {result.get('error','') }"
        follow_up = self.call_brain_model(feedback)
        return follow_up

    def _handle_generate_image(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        if not getattr(self, 'use_image_generation', False):
            logger.info("🔧 Автоматически включаю генерацию изображений")
            self.use_image_generation = True
            self.auto_disable_tools("image_generation")
        if not getattr(self, 'use_image_generation', False):
            logger.error("❌ Генерация изображений отключена")
            follow_up = self.call_brain_model("Генерация изображений отключена. Предложи альтернативный способ помочь пользователю.")
            return follow_up

        description = action_data.get("description", "")
        text = action_data.get("text", "")
        style = action_data.get("style", "")
        negative_prompt = action_data.get("negative_prompt", "")
        prompt = text.strip() if text else (description or "")
        if style and prompt:
            prompt += f", {style}"

        params = {}
        if not isinstance(params, dict):
            params = {}

        if not prompt:
            follow_up = self.call_brain_model("Нейросеть вернула пустой промт для генерации изображения. Попроси её создать описание.")
            return follow_up

        if not self._is_english_simple(prompt):
            follow_up = self.call_brain_model(f"Нейросеть вернула промт не на английском языке: {prompt}. Попроси её создать промт на английском.")
            return follow_up

        neg = negative_prompt.strip() if negative_prompt else ""
        if not neg or not self._is_english_simple(neg):
            neg = "(worst quality, low quality, normal quality:1.4)"
            self.logger.info(f"⚠️ Используется fallback negative_prompt: {neg}")
        else:
            self.logger.info(f"✅ Используется negative_prompt из JSON: {neg}")
        
        self.logger.info(f"🎨 Передаем в генерацию - prompt: {prompt[:50]}..., negative_prompt: {neg}")

        # default params and validation (kept simple here)
        default_params = {"seed": -1, "steps": 30, "width": 1024, "height": 1024, "cfg": 4.0}
        gen_params = default_params.copy()

        img_b64 = self.generate_image_stable_diffusion(prompt, neg, gen_params)
        if img_b64:
            self.last_generated_image_b64 = img_b64
            self.show_image_base64_temp(img_b64)
            final_msg = f"✅ Изображение успешно сгенерировано по вашему описанию: {description}\n🎨 Использованный промт: {prompt}"
            self.last_final_response = final_msg
            logger.info(final_msg)
            return False
        else:
            follow_up = self.call_brain_model(f"Не удалось сгенерировать изображение по описанию: {description}.")
            return follow_up

    def _handle_generate_video(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        if not getattr(self, 'use_image_generation', False):
            logger.info("🔧 Автоматически включаю генерацию изображений")
            self.use_image_generation = True
            self.auto_disable_tools("image_generation")
        if not getattr(self, 'use_image_generation', False):
            follow_up = self.call_brain_model("Генерация изображений отключена. Предложи альтернативный способ помочь пользователю.")
            return follow_up

        description = action_data.get("description", "")
        text = action_data.get("text", "")
        style = action_data.get("style", "")
        negative_prompt = action_data.get("negative_prompt", "")
        prompt = text.strip() if text else (description or "")
        if style and prompt:
            prompt += f", {style}"

        if not prompt:
            follow_up = self.call_brain_model("Нейросеть вернула пустой промт для генерации видео. Попроси её создать описание.")
            return follow_up

        # negative prompt handling
        neg = negative_prompt.strip() if negative_prompt else ""
        fallback_negative = "(worst quality, low quality, normal quality:1.4), (deformed, distorted, disfigured:1.3), poorly drawn, bad anatomy, text, watermark"
        if not neg or not self._is_english_simple(neg):
            neg = fallback_negative
            self.logger.info(f"⚠️ Используется fallback negative_prompt для видео: {neg}")
        else:
            self.logger.info(f"✅ Используется negative_prompt из JSON для видео: {neg}")
        
        self.logger.info(f"🎬 Передаем в генерацию видео - prompt: {prompt[:50]}..., negative_prompt: {neg[:50]}...")

        # default video params and validation
        default_params = {"seed": -1, "steps": 20, "width": 512, "height": 512, "cfg": 7.0, "num_frames": 24, "fps": 8}
        params = action_data.get("params", {}) or {}
        gen_params = default_params.copy()
        if isinstance(params, dict):
            for key, value in params.items():
                if key in default_params:
                    try:
                        if key in ["seed", "steps", "width", "height", "num_frames", "fps"]:
                            gen_params[key] = int(value)
                        elif key == "cfg":
                            gen_params[key] = float(value)
                        else:
                            gen_params[key] = value
                    except (ValueError, TypeError):
                        logger.warning(f"⚠️ Неверный параметр {key}={value}, используется значение по умолчанию")

        # basic bounds
        if gen_params["steps"] < 1 or gen_params["steps"] > 100:
            gen_params["steps"] = 20
        if gen_params["width"] < 64 or gen_params["width"] > 2048:
            gen_params["width"] = 512
        if gen_params["height"] < 64 or gen_params["height"] > 2048:
            gen_params["height"] = 512

        video_path = self.generate_video_stable_diffusion(prompt, neg, gen_params)
        if video_path:
            final_msg = f"✅ Видео успешно сгенерировано по вашему описанию: {description}\n📁 Путь к видео: {video_path}"
            self.last_final_response = final_msg
            logger.info(final_msg)
            return False
        else:
            follow_up = self.call_brain_model(f"Не удалось сгенерировать видео по описанию: {description}.")
            return follow_up

    def _handle_speak(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        text_to_speak = action_data.get("text", "")
        voice = action_data.get("voice", "male")
        language = action_data.get("language", "ru")
        if not text_to_speak:
            follow_up = self.call_brain_model("Текст для озвучки пустой. Укажите текст в поле 'text'.")
            return follow_up
        audio_path = self.text_to_speech(text_to_speak, voice, language)
        if audio_path:
            follow_up = self.call_brain_model(f"Текст успешно озвучен: {text_to_speak}. Аудиофайл сохранен: {os.path.basename(audio_path)}")
        else:
            follow_up = self.call_brain_model(f"Не удалось озвучить текст: {text_to_speak}.")
        return follow_up

    def _handle_response(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        # Поддерживаем и "text" и "content" для совместимости
        content = action_data.get("text", action_data.get("content", ""))
        self.last_final_response = content
        logger.info(f"\n🤖 ФИНАЛЬНЫЙ ОТВЕТ:")
        logger.info(content)
        
        # Если есть сгенерированный файл, уведомляем об этом в интерактивном режиме
        if (hasattr(self, 'last_generated_file_path') and self.last_generated_file_path and 
            getattr(self, 'show_images_locally', True)):
            logger.info(f"\n📄 Создан файл: {self.last_generated_file_name}")
            logger.info(f"📂 Расположение: {self.last_generated_file_path}")
            # Очищаем после показа в интерактивном режиме
            self.last_generated_file_path = None
            self.last_generated_file_name = None
        
        return False

    def _handle_list_files(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для просмотра содержимого папок
        """
        folder = action_data.get("folder", "")
        description = action_data.get("description", f"Просмотр содержимого папки {folder}")
        
        logger.info(f"\n📁 ПРОСМОТР ПАПКИ: {description}")
        logger.info(f"📂 Папка: {folder}")
        
        result = self.list_folder_contents(folder)
        
        logger.info(f"📋 Результат:\n{result}")
        
        follow_up = self.call_brain_model(f"Содержимое папки '{folder}': {result}")
        return follow_up

    def _handle_process_document(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для обработки документов (DOCX, Excel, PDF)
        """
        file_path = action_data.get("file_path", "")
        description = action_data.get("description", f"Обработка документа {file_path}")
        
        logger.info(f"\n📄 ОБРАБОТКА ДОКУМЕНТА: {description}")
        logger.info(f"📁 Файл: {file_path}")
        
        # Проверяем, является ли путь относительным и добавляем базовую директорию
        if not os.path.isabs(file_path):
            # Проверяем, не начинается ли путь уже с папки
            if (file_path.startswith("Docx/") or file_path.startswith("Docx\\") or
                file_path.startswith("Excel/") or file_path.startswith("Excel\\") or
                file_path.startswith("PDF/") or file_path.startswith("PDF\\")):
                full_path = os.path.join(self.base_dir, file_path)
            else:
                # Определяем папку по расширению файла
                file_lower = file_path.lower()
                if file_lower.endswith(('.docx', '.doc')):
                    full_path = os.path.join(self.base_dir, "Docx", file_path)
                elif file_lower.endswith(('.xlsx', '.xls', '.csv')):
                    full_path = os.path.join(self.base_dir, "Excel", file_path)
                elif file_lower.endswith('.pdf'):
                    full_path = os.path.join(self.base_dir, "PDF", file_path)
                else:
                    full_path = os.path.join(self.base_dir, file_path)
        else:
            full_path = file_path
        
        result = self.process_document_request(full_path)
        
        logger.info(f"📋 Результат обработки:\n{result[:500]}...")
        
        follow_up = self.call_brain_model(f"Результат обработки документа '{file_path}': {result}")
        return follow_up

    def _handle_generate_file(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для генерации файлов (DOCX, Excel, PDF, Markdown)
        """
        content = action_data.get("content", "")
        filename = action_data.get("filename", "")
        file_type = action_data.get("file_type", "").lower()
        description = action_data.get("description", f"Генерация файла {filename}")
        
        logger.info(f"\n📝 ГЕНЕРАЦИЯ ФАЙЛА: {description}")
        logger.info(f"📁 Имя файла: {filename}")
        logger.info(f"📄 Тип файла: {file_type}")
        
        if not content:
            follow_up = self.call_brain_model("Ошибка: не указано содержимое для генерации файла")
            return follow_up
        
        if not filename:
            follow_up = self.call_brain_model("Ошибка: не указано имя файла")
            return follow_up
        
        # Определяем путь в папку output
        output_path = os.path.join(self.base_dir, "output", filename)
        
        try:
            success = self.generate_file(content, output_path, file_type)
            if success:
                logger.info(f"✅ Файл успешно создан: {output_path}")
                
                # Сохраняем информацию о последнем сгенерированном файле для Telegram
                self.last_generated_file_path = output_path
                self.last_generated_file_name = filename
                
                follow_up = self.call_brain_model(f"Файл '{filename}' успешно создан в папке output")
            else:
                logger.error(f"❌ Ошибка при создании файла: {output_path}")
                follow_up = self.call_brain_model(f"Ошибка при создании файла '{filename}'")
        except Exception as e:
            logger.error(f"❌ Исключение при создании файла: {e}")
            follow_up = self.call_brain_model(f"Ошибка при создании файла '{filename}': {str(e)}")
        
        return follow_up

    def _handle_extract_text(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для извлечения текста из изображений с помощью OCR
        """
        image_path = action_data.get("image_path", "")
        description = action_data.get("description", f"Извлечение текста из {image_path}")
        force_ocr = action_data.get("force_ocr", False)
        
        logger.info(f"\n📖 ИЗВЛЕЧЕНИЕ ТЕКСТА OCR: {description}")
        logger.info(f"📁 Изображение (исходный путь): {image_path}")
        
        # Проверяем, является ли путь относительным и добавляем базовую директорию
        if not os.path.isabs(image_path):
            logger.info(f"🔍 Путь относительный, проверяю логику...")
            # Проверяем, не начинается ли путь уже с папки Photos
            if image_path.startswith("Photos/") or image_path.startswith("Photos\\"):
                logger.info(f"🔍 Путь начинается с Photos/, используем как есть")
                full_path = os.path.join(self.base_dir, image_path)
            else:
                logger.info(f"🔍 Путь не начинается с Photos/, определяю по расширению")
                # Определяем папку по расширению файла
                file_lower = image_path.lower()
                if file_lower.endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp')):
                    logger.info(f"🔍 Это изображение, добавляю папку Photos")
                    full_path = os.path.join(self.base_dir, "Photos", image_path)
                else:
                    logger.info(f"🔍 Не изображение, добавляю к base_dir")
                    full_path = os.path.join(self.base_dir, image_path)
        else:
            logger.info(f"🔍 Путь абсолютный, используем как есть")
            full_path = image_path
        
        logger.info(f"🔍 Полный путь к файлу: {full_path}")
        
        # Сначала получим описание от vision модели если доступна
        vision_description = ""
        if getattr(self, 'use_vision', False):
            try:
                # Конвертируем изображение в base64 для vision модели
                with open(full_path, 'rb') as img_file:
                    img_b64 = base64.b64encode(img_file.read()).decode('ascii')
                    vision_description = self.call_vision_model(img_b64)
                    logger.info(f"🔍 Vision описание: {vision_description[:100]}...")
            except Exception as e:
                logger.warning(f"⚠️ Не удалось получить vision описание: {e}")
        
        # Применяем умный OCR
        extracted_text, full_description, ocr_error = self.process_image_with_smart_ocr(
            full_path, vision_description, force_ocr
        )
        
        if extracted_text:
            result_text = f"Извлеченный текст из изображения '{image_path}':{extracted_text}"
        elif ocr_error:
            result_text = f"Ошибка при обработке изображения '{image_path}':\n\n{ocr_error}\n\nОписание изображения (если доступно):\n{vision_description}"
        else:
            result_text = f"Описание изображения '{image_path}':\n\n{vision_description}"
        
        logger.info(f"📋 Результат OCR:\n{result_text}")
        
        follow_up = self.call_brain_model(result_text)
        return follow_up

    def _handle_analyze_image(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для анализа изображения с возможностью интеллектуального OCR.
        
        Args:
            action_data: Данные действия с полями:
                - image_path: Путь к изображению  
                - check_for_text: Нужно ли проверять наличие текста (bool)
                - description: Описание задачи (optional)
        
        Returns:
            str: Follow-up для модели с результатами анализа
        """
        try:
            image_path = action_data.get("image_path", "").strip()
            check_for_text = action_data.get("check_for_text", False)
            description = action_data.get("description", "")
            
            if not image_path:
                logger.error("❌ Не указан путь к изображению")
                return self.call_brain_model("❌ Ошибка: не указан путь к изображению")
                
            # Разрешаем путь
            full_path = self.resolve_path(image_path)
            if not os.path.exists(full_path):
                logger.error(f"❌ Изображение не найдено: {full_path}")
                return self.call_brain_model(f"❌ Изображение не найдено: {image_path}")
            
            logger.info(f"🔍 Анализирую изображение: {image_path}")
            if description:
                logger.info(f"📝 Задача: {description}")
            
            # Получаем vision описание
            vision_description = ""
            try:
                vision_description = self.analyze_image_with_vision(full_path)
                logger.info(f"👁️ Vision описание: {vision_description}")
            except Exception as e:
                logger.warning(f"⚠️ Не удалось получить vision описание: {e}")
                vision_description = f"Ошибка анализа изображения: {e}"
            
            # Если запрошена проверка текста, принимаем решение об OCR
            extracted_text = ""
            ocr_info = ""
            
            if check_for_text:
                logger.info("🔍 Проверяю наличие текста на изображении...")
                
                # Интеллектуальное решение об OCR на основе vision описания
                should_use_ocr = self.should_use_ocr_intelligently(vision_description, description)
                
                if should_use_ocr:
                    logger.info("✅ Применяю OCR для извлечения текста")
                    try:
                        extracted_text = self.extract_text_from_image(full_path)
                        if extracted_text:
                            ocr_info = f"\n\n📋 Извлеченный текст:\n{extracted_text}"
                        else:
                            ocr_info = "\n\n❌ OCR не обнаружил текст на изображении"
                    except Exception as e:
                        logger.warning(f"⚠️ Ошибка OCR: {e}")
                        ocr_info = f"\n\n❌ Ошибка при извлечении текста: {e}"
                else:
                    logger.info("❌ OCR не требуется для данного изображения")
                    ocr_info = "\n\n📝 OCR не применялся - на изображении не обнаружено значимого текста"
            
            # Формируем результат анализа
            result_text = f"Анализ изображения '{image_path}':\n\n📸 Описание:\n{vision_description}{ocr_info}"
            
            logger.info(f"✅ Анализ завершен")
            
            follow_up = self.call_brain_model(result_text)
            return follow_up
            
        except Exception as e:
            logger.error(f"❌ Ошибка при анализе изображения: {e}")
            return self.call_brain_model(f"❌ Ошибка при анализе изображения: {e}")

    def _handle_plugin_action(self, action: str, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для действий плагинов.
        
        Args:
            action: Строка действия в формате "plugin:plugin_name:action_name"
            action_data: Данные действия
        
        Returns:
            str: Follow-up для модели с результатами выполнения
        """
        try:
            # Проверяем доступность системы плагинов
            if not self.plugin_manager:
                logger.error("❌ Система плагинов недоступна")
                return self.call_brain_model("❌ Система плагинов недоступна")
            
            # Парсим action в формате "plugin:plugin_name:action_name"
            parts = action.split(":", 2)
            if len(parts) != 3 or parts[0] != "plugin":
                logger.error(f"❌ Неверный формат действия плагина: {action}")
                return self.call_brain_model(f"❌ Неверный формат действия плагина. Ожидается 'plugin:plugin_name:action_name'")
            
            plugin_name = parts[1]
            plugin_action = parts[2]
            plugin_data = action_data.get("data", {})
            
            logger.info(f"🔌 Выполняется действие плагина: {plugin_name}.{plugin_action}")
            
            # Выполняем действие плагина
            result = self.plugin_manager.execute_plugin_action(
                plugin_name=plugin_name,
                action=plugin_action,
                data=plugin_data,
                orchestrator=self
            )
            
            # Формируем ответ
            result_text = f"Результат выполнения плагина '{plugin_name}', действие '{plugin_action}':\n\n{result}"
            
            logger.info(f"✅ Действие плагина выполнено успешно")
            
            follow_up = self.call_brain_model(result_text)
            return follow_up
            
        except PluginError as e:
            logger.error(f"❌ Ошибка плагина: {e}")
            return self.call_brain_model(f"❌ Ошибка плагина: {e}")
        except Exception as e:
            logger.error(f"❌ Неожиданная ошибка при выполнении действия плагина: {e}")
            return self.call_brain_model(f"❌ Ошибка при выполнении действия плагина: {e}")

    def _handle_get_help(self, action: str, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для команд get_*_help - загружает соответствующий модуль.
        
        Args:
            action: Команда типа "get_image_generation_help"
            action_data: Данные действия
        
        Returns:
            str: Follow-up для модели с загруженным модулем
        """
        try:
            logger.info(f"📚 Загружаю модуль для команды: {action}")
            
            # Загружаем модуль через PromptLoader
            module_content = self.prompt_loader.load_module(action)
            
            if module_content is None:
                logger.warning(f"❌ Модуль для команды {action} не найден")
                return self.call_brain_model(f"❌ Модуль для команды {action} не найден. Доступные команды: {', '.join(self.prompt_loader.module_commands.keys())}")
            
            logger.info(f"✅ Модуль загружен, размер: {len(module_content)} символов")
            
            # Отправляем загруженный модуль как контекст для модели
            follow_up_prompt = f"""
Загружен модуль по запросу {action}:

{module_content}

Теперь ты можешь использовать эту информацию для ответа пользователю или выполнения соответствующих действий.
"""
            
            follow_up = self.call_brain_model(follow_up_prompt)
            return follow_up
            
        except Exception as e:
            logger.error(f"❌ Ошибка при загрузке модуля {action}: {e}")
            return self.call_brain_model(f"❌ Ошибка при загрузке модуля {action}: {e}")

    def _handle_send_email(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для отправки email
        """
        provider = action_data.get("provider", "")
        to_email = action_data.get("to_email", "")
        subject = action_data.get("subject", "")
        body = action_data.get("body", "")
        attachments = action_data.get("attachments", [])
        description = action_data.get("description", f"Отправка письма на {to_email}")
        
        logger.info(f"\n📧 ОТПРАВКА EMAIL: {description}")
        logger.info(f"📨 Провайдер: {provider}")
        logger.info(f"📧 Получатель: {to_email}")
        logger.info(f"📝 Тема: {subject}")
        
        if not provider:
            # Используем первый доступный провайдер
            if self.available_email_providers:
                provider = self.available_email_providers[0]
                logger.info(f"🔧 Автоматически выбран провайдер: {provider}")
            else:
                result = "❌ Почтовые провайдеры не настроены"
                logger.error(result)
                follow_up = self.call_brain_model(result)
                return follow_up
        
        result = self.send_email(provider, to_email, subject, body, attachments)
        logger.info(f"📧 Результат отправки: {result}")
        
        follow_up = self.call_brain_model(f"Результат отправки письма: {result}")
        return follow_up

    def _handle_get_emails(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для получения списка писем
        """
        provider = action_data.get("provider", "")
        folder = action_data.get("folder", "INBOX")
        limit = action_data.get("limit", 10)
        search_criteria = action_data.get("search_criteria", "ALL")
        description = action_data.get("description", f"Получение писем из {folder}")
        
        logger.info(f"\n📧 ПОЛУЧЕНИЕ ПИСЕМ: {description}")
        logger.info(f"📨 Провайдер: {provider}")
        logger.info(f"📁 Папка: {folder}")
        logger.info(f"🔢 Лимит: {limit}")
        
        if not provider:
            # Используем первый доступный провайдер
            if self.available_email_providers:
                provider = self.available_email_providers[0]
                logger.info(f"🔧 Автоматически выбран провайдер: {provider}")
            else:
                result = "❌ Почтовые провайдеры не настроены"
                logger.error(result)
                follow_up = self.call_brain_model(result)
                return follow_up
        
        emails = self.get_emails(provider, folder, limit, search_criteria)
        
        if isinstance(emails, list):
            # Форматируем список писем для отправки в модель
            emails_summary = f"Получено {len(emails)} писем из папки '{folder}':\n\n"
            for i, email_info in enumerate(emails, 1):
                emails_summary += f"{i}. От: {email_info.get('from', 'Неизвестно')}\n"
                emails_summary += f"   Тема: {email_info.get('subject', 'Без темы')}\n"
                emails_summary += f"   Дата: {email_info.get('date', 'Неизвестно')}\n"
                emails_summary += f"   ID: {email_info.get('id', 'Неизвестно')}\n"
                
                # Обрезаем текст письма до 200 символов для краткости
                body = email_info.get('body', '')
                if len(body) > 200:
                    body = body[:200] + "..."
                emails_summary += f"   Текст: {body}\n\n"
            
            logger.info(f"📧 Получено {len(emails)} писем")
            follow_up = self.call_brain_model(emails_summary)
        else:
            # Ошибка получения писем
            logger.error(f"❌ Ошибка: {emails}")
            follow_up = self.call_brain_model(f"Ошибка получения писем: {emails}")
        
        return follow_up

    def _handle_reply_email(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для ответа на письмо
        """
        provider = action_data.get("provider", "")
        original_email_id = action_data.get("original_email_id", "")
        reply_text = action_data.get("reply_text", "")
        attachments = action_data.get("attachments", [])
        description = action_data.get("description", f"Ответ на письмо {original_email_id}")
        
        logger.info(f"\n📧 ОТВЕТ НА ПИСЬМО: {description}")
        logger.info(f"📨 Провайдер: {provider}")
        logger.info(f"🆔 ID оригинального письма: {original_email_id}")
        
        if not provider:
            # Используем первый доступный провайдер
            if self.available_email_providers:
                provider = self.available_email_providers[0]
                logger.info(f"🔧 Автоматически выбран провайдер: {provider}")
            else:
                result = "❌ Почтовые провайдеры не настроены"
                logger.error(result)
                follow_up = self.call_brain_model(result)
                return follow_up
        
        result = self.reply_to_email(provider, original_email_id, reply_text, attachments)
        logger.info(f"📧 Результат ответа: {result}")
        
        follow_up = self.call_brain_model(f"Результат ответа на письмо: {result}")
        return follow_up

    def _handle_search_emails(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для поиска писем
        """
        provider = action_data.get("provider", "")
        query = action_data.get("query", "")
        folder = action_data.get("folder", "INBOX")
        limit = action_data.get("limit", 20)
        description = action_data.get("description", f"Поиск писем: {query}")
        
        logger.info(f"\n🔍 ПОИСК ПИСЕМ: {description}")
        logger.info(f"📨 Провайдер: {provider}")
        logger.info(f"🔎 Запрос: {query}")
        logger.info(f"📁 Папка: {folder}")
        
        if not provider:
            # Используем первый доступный провайдер
            if self.available_email_providers:
                provider = self.available_email_providers[0]
                logger.info(f"🔧 Автоматически выбран провайдер: {provider}")
            else:
                result = "❌ Почтовые провайдеры не настроены"
                logger.error(result)
                follow_up = self.call_brain_model(result)
                return follow_up
        
        emails = self.search_emails(provider, query, folder, limit)
        
        if isinstance(emails, list):
            # Форматируем результаты поиска
            search_summary = f"Найдено {len(emails)} писем по запросу '{query}':\n\n"
            for i, email_info in enumerate(emails, 1):
                search_summary += f"{i}. От: {email_info.get('from', 'Неизвестно')}\n"
                search_summary += f"   Тема: {email_info.get('subject', 'Без темы')}\n"
                search_summary += f"   Дата: {email_info.get('date', 'Неизвестно')}\n"
                search_summary += f"   ID: {email_info.get('id', 'Неизвестно')}\n"
                
                # Обрезаем текст письма до 200 символов
                body = email_info.get('body', '')
                if len(body) > 200:
                    body = body[:200] + "..."
                search_summary += f"   Текст: {body}\n\n"
            
            logger.info(f"🔍 Найдено {len(emails)} писем")
            follow_up = self.call_brain_model(search_summary)
        else:
            # Ошибка поиска
            logger.error(f"❌ Ошибка поиска: {emails}")
            follow_up = self.call_brain_model(f"Ошибка поиска писем: {emails}")
        
        return follow_up

    def _handle_run_bat_file(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для запуска BAT файлов
        """
        file_path = action_data.get("file_path", "")
        working_dir = action_data.get("working_dir", None)
        description = action_data.get("description", f"Запуск BAT файла: {file_path}")
        
        logger.info(f"\n🚀 ЗАПУСК BAT ФАЙЛА: {description}")
        logger.info(f"📄 Файл: {file_path}")
        if working_dir:
            logger.info(f"📁 Рабочая директория: {working_dir}")
        
        if not file_path:
            error_msg = "❌ Не указан путь к BAT файлу"
            logger.error(error_msg)
            follow_up = self.call_brain_model(f"Ошибка: {error_msg}")
            return follow_up
        
        try:
            result = self.run_bat_file(file_path, working_dir)
            
            if result["success"]:
                # Успешное выполнение
                output_info = f"✅ BAT файл выполнен успешно!\n"
                output_info += f"📄 Файл: {os.path.basename(file_path)}\n"
                output_info += f"📁 Рабочая директория: {result.get('working_dir', 'не указана')}\n"
                output_info += f"🔢 Код возврата: {result.get('return_code', 0)}\n"
                
                if result.get("output"):
                    output_info += f"\n📝 Вывод:\n{result['output']}"
                
                logger.info("✅ BAT файл выполнен успешно")
                follow_up = self.call_brain_model(output_info)
            else:
                # Ошибка выполнения
                error_info = f"❌ Ошибка выполнения BAT файла!\n"
                error_info += f"📄 Файл: {os.path.basename(file_path)}\n"
                error_info += f"🔢 Код возврата: {result.get('return_code', -1)}\n"
                error_info += f"❌ Ошибка: {result.get('error', 'Неизвестная ошибка')}\n"
                
                if result.get("output"):
                    error_info += f"\n📝 Вывод:\n{result['output']}"
                
                logger.error(f"❌ Ошибка выполнения BAT файла: {result.get('error')}")
                follow_up = self.call_brain_model(error_info)
                
        except Exception as e:
            error_msg = f"❌ Исключение при запуске BAT файла: {str(e)}"
            logger.error(error_msg)
            follow_up = self.call_brain_model(f"Ошибка: {error_msg}")
        
        return follow_up

    def _get_plugin_info_for_prompt(self) -> str:
        """
        Формирует информацию о доступных плагинах для добавления в системный промпт.
        
        Returns:
            str: Текст с информацией о плагинах
        """
        if not self.plugin_manager:
            return ""
        
        try:
            loaded_plugins = self.plugin_manager.get_loaded_plugins()
            if not loaded_plugins:
                return ""
            
            plugin_info_parts = ["ДОСТУПНЫЕ ПЛАГИНЫ:"]
            
            for plugin_name, plugin in loaded_plugins.items():
                try:
                    info = plugin.get_plugin_info()
                    actions = plugin.get_available_actions()
                    
                    plugin_desc = f"\n🔌 {info.get('name', plugin_name)} v{info.get('version', '1.0')}"
                    plugin_desc += f"\n   Описание: {info.get('description', 'Нет описания')}"
                    plugin_desc += f"\n   Автор: {info.get('author', 'Неизвестен')}"
                    plugin_desc += f"\n   Действия: {', '.join(actions)}"
                    plugin_desc += f"\n   Формат вызова: plugin:{plugin_name}:action_name"
                    
                    plugin_info_parts.append(plugin_desc)
                except Exception as e:
                    logger.warning(f"Ошибка получения информации о плагине {plugin_name}: {e}")
                    plugin_info_parts.append(f"\n🔌 {plugin_name} (ошибка получения информации)")
            
            return "\n".join(plugin_info_parts)
            
        except Exception as e:
            logger.error(f"Ошибка формирования информации о плагинах: {e}")
            return ""

    def should_use_ocr_intelligently(self, vision_description: str, task_description: str = "") -> bool:
        """
        Интеллектуальное решение о необходимости OCR на основе контекста.
        
        Args:
            vision_description: Описание изображения от vision модели
            task_description: Описание задачи пользователя
            
        Returns:
            bool: True если OCR нужен, False если нет
        """
        vision_lower = vision_description.lower()
        task_lower = task_description.lower()
        
        # Явные индикаторы необходимости OCR
        text_indicators = [
            "text", "writing", "words", "letters", "document", "page", "book", "sign", 
            "label", "caption", "title", "heading", "paragraph", "sentence", "line",
            "screen", "display", "interface", "menu", "button", "dialog", "window",
            "newspaper", "article", "magazine", "poster", "banner", "billboard",
            "form", "table", "chart", "graph", "spreadsheet", "invoice", "receipt",
            "card", "certificate", "license", "passport", "id", "ticket",
            "текст", "надпись", "слова", "буквы", "документ", "страница", "книга",
            "вывеска", "подпись", "заголовок", "строка", "экран", "интерфейс",
            "меню", "кнопка", "окно", "газета", "статья", "плакат", "форма",
            "таблица", "график", "чек", "карта", "сертификат", "билет"
        ]
        
        # Индикаторы отсутствия текста
        no_text_indicators = [
            "landscape", "nature", "animal", "person", "face", "building", "car",
            "food", "flower", "tree", "sky", "mountain", "water", "art", "painting",
            "photo", "picture", "image", "scenery", "portrait", "selfie",
            "пейзаж", "природа", "животное", "человек", "лицо", "здание", "машина",
            "еда", "цветок", "дерево", "небо", "гора", "вода", "искусство", "картина",
            "фото", "изображение", "портрет"
        ]
        
        # Проверяем явные запросы OCR в задаче
        if any(keyword in task_lower for keyword in ["текст", "text", "прочита", "read", "извлеч", "extract"]):
            return True
        
        # Считаем индикаторы в vision описании
        text_score = sum(1 for indicator in text_indicators if indicator in vision_lower)
        no_text_score = sum(1 for indicator in no_text_indicators if indicator in vision_lower)
        
        # Решение на основе баланса индикаторов
        if text_score >= 2:  # Несколько индикаторов текста
            return True
        elif text_score >= 1 and no_text_score == 0:  # Есть индикатор текста, нет противопоказаний
            return True
        elif no_text_score >= 2:  # Явно не текстовое изображение
            return False
        else:
            # Граничный случай - OCR может быть полезен
            return text_score > no_text_score

    def _process_ai_response_impl(self, ai_response: str) -> bool:
        """
        Подробная реализация обработки ответа AI (перенесена из original `process_ai_response`).

        Args:
            ai_response: JSON ответ от AI

        Returns:
            True если нужно продолжать диалог, False если завершить
        """
        # Итеративный обработчик цепочек follow_up: избегаем рекурсии.
        next_input: Optional[str] = ai_response
        attempts = 0
        while next_input is not None and attempts <= self.max_retries:
            attempts += 1
            try:
                json_str = self.extract_first_json(next_input, allow_json_in_think=True)
                if not json_str or json_str == next_input:
                    logger.info("💬 Модель вернула текстовый ответ без JSON")
                    if len(next_input.strip()) > 5 and not next_input.strip().startswith('{'):
                        logger.info("💬 Использую текстовый ответ как финальный")
                        self.last_final_response = next_input.strip()
                        return False
                    think_content = self.extract_think_content(next_input)
                    if think_content:
                        json_in_think = self._extract_json_from_text(think_content)
                        if json_in_think:
                            logger.info("🔍 Найден JSON внутри <think> блока, использую его")
                            json_str = json_in_think
                        else:
                            feedback = "Модель вернула неполный ответ. Пожалуйста, сформулируй полный ответ или действие."
                            next_input = self.call_brain_model(feedback)
                            continue
                    else:
                        feedback = "Модель вернула неполный ответ. Пожалуйста, сформулируй полный ответ или действие."
                        next_input = self.call_brain_model(feedback)
                        continue

                action_data, fixes = self._smart_json_parse(json_str)
                if fixes:
                    logger.warning(f"⚠️ Исправления JSON: {'; '.join(fixes)}")
                if not action_data:
                    logger.error(f"❌ Не удалось распарсить JSON даже после исправлений:\n{json_str}")
                    self.retry_count += 1
                    if self.retry_count > self.max_retries:
                        logger.warning(f"🔄 Достигнут лимит попыток ({self.max_retries}), завершаю диалог")
                        self.retry_count = 0
                        self.last_final_response = "Извините, возникла проблема с обработкой запроса. Попробуйте переформулировать вопрос."
                        return False
                    think_content = self.extract_think_content(next_input)
                    if think_content:
                        logger.info("💭 Обнаружен блок размышлений - модель размышляет, но не генерирует действие")
                        logger.info(f"💭 Содержимое: {think_content[:200]}...")
                        if len(think_content) > 20 and any(word in think_content.lower() for word in ['привет', 'hello', 'здравствуй']):
                            logger.info("💭 Обнаружено приветствие в размышлениях, завершаю диалог")
                            self.retry_count = 0
                            self.last_final_response = "Привет! Я Нейро, ваш AI-помощник. Чем могу помочь?"
                            return False
                        if self.retry_count >= 2 and 'правил' in think_content.lower():
                            logger.info("💭 Модель зацикливается на правилах, принудительно завершаю")
                            self.retry_count = 0
                            self.last_final_response = "Привет! Как дела? Чем могу помочь?"
                            return False
                        next_input = self.call_brain_model("Пожалуйста, дай простой дружелюбный ответ или конкретное действие. Не анализируй правила.")
                        continue
                    else:
                        next_input = self.call_brain_model("Пожалуйста, дай простой ответ или конкретное действие.")
                        continue

                if 'action' not in action_data:
                    logger.warning("⚠️ В JSON отсутствует ключ 'action', добавляю action: 'unknown'")
                    action_data['action'] = 'unknown'

                if action_data == {} or (len(action_data) == 1 and 'action' in action_data and action_data['action'] == 'unknown'):
                    logger.info("💭 Модель вернула пустой JSON - возможно, она размышляет")
                    think_content = self.extract_think_content(next_input)
                    if think_content and len(think_content) > 20:
                        logger.info("💭 Использую содержимое размышлений как ответ")
                        self.last_final_response = think_content
                        return False
                    next_input = self.call_brain_model("Модель вернула пустой JSON. Пожалуйста, сформулируй конкретный ответ или действие.")
                    continue

                action = action_data.get("action")
                self.retry_count = 0

                # Вызов соответствующего хендлера и обработка результата (следующий ввод или завершение)
                handler_result = None
                if action == "powershell":
                    handler_result = self._handle_powershell(action_data)
                elif action == "search":
                    handler_result = self._handle_search(action_data)
                elif action == "take_screenshot":
                    handler_result = self._handle_take_screenshot(action_data)
                elif action == "move_mouse":
                    handler_result = self._handle_move_mouse(action_data)
                elif action == "left_click":
                    handler_result = self._handle_left_click(action_data)
                elif action == "right_click":
                    handler_result = self._handle_right_click(action_data)
                elif action in ["scroll_up", "scroll_down"]:
                    handler_result = self._handle_scroll(action, action_data)
                elif action == "mouse_down":
                    handler_result = self._handle_mouse_down(action_data)
                elif action == "mouse_up":
                    handler_result = self._handle_mouse_up(action_data)
                elif action == "drag_and_drop":
                    handler_result = self._handle_drag_and_drop(action_data)
                elif action == "type_text":
                    handler_result = self._handle_type_text(action_data)
                elif action == "generate_image":
                    handler_result = self._handle_generate_image(action_data)
                elif action == "generate_video":
                    handler_result = self._handle_generate_video(action_data)
                elif action == "speak":
                    handler_result = self._handle_speak(action_data)
                elif action == "list_files":
                    handler_result = self._handle_list_files(action_data)
                elif action == "process_document":
                    handler_result = self._handle_process_document(action_data)
                elif action == "generate_file":
                    handler_result = self._handle_generate_file(action_data)
                elif action == "extract_text":
                    handler_result = self._handle_extract_text(action_data)
                elif action == "analyze_image":
                    handler_result = self._handle_analyze_image(action_data)
                elif action == "response":
                    handler_result = self._handle_response(action_data)
                elif action == "send_email":
                    handler_result = self._handle_send_email(action_data)
                elif action == "get_emails":
                    handler_result = self._handle_get_emails(action_data)
                elif action == "reply_email":
                    handler_result = self._handle_reply_email(action_data)
                elif action == "search_emails":
                    handler_result = self._handle_search_emails(action_data)
                elif action == "run_bat_file":
                    handler_result = self._handle_run_bat_file(action_data)
                elif action.startswith("plugin:"):
                    handler_result = self._handle_plugin_action(action, action_data)
                elif action.startswith("get_") and action.endswith("_help"):
                    handler_result = self._handle_get_help(action, action_data)
                else:
                    logger.warning(f"❓ Неизвестное действие: {action}")
                    return False

                # Обработка результата хендлера: False => завершить, str => новый ввод для итерации
                if handler_result is False:
                    return False
                elif isinstance(handler_result, str):
                    next_input = handler_result
                    continue
                else:
                    # Если хендлер вернул None или неожиданный тип — завершаем
                    logger.error("❌ Хендлер вернул некорректный результат, завершаю")
                    return False

            except json.JSONDecodeError as e:
                logger.error(f"❌ Ошибка парсинга JSON ответа AI: {e}")
                logger.info(f"📝 Ответ AI: {next_input}")
                return False
            except Exception as e:
                logger.error(f"❌ Ошибка обработки ответа AI: {str(e)}")
                return False

        logger.warning("🔄 Превышен цикл попыток обработки follow_up или достигнут лимит. Завершаю.")
        return False

    def run_interactive(self):
        """Запуск интерактивного режима (глаза, аудио, мозг)"""
        # Показываем сообщения только в консольном режиме
        if getattr(self, 'show_images_locally', True):
            logger.info("🚀 AI PowerShell Оркестратор запущен!")
            logger.info("💡 Используйте новые команды для работы с файлами:")
            logger.info("   - list_files: просмотр содержимого папок (Audio, Photos, Video, Excel, Docx, PDF)")
            logger.info("   - process_document: обработка всех типов документов (DOCX, DOC, PDF, XLSX, XLS, CSV, TXT, MD, RTF)")
            logger.info("   - extract_text: OCR распознавание текста с изображений (с поддержкой русского и английского)")
            logger.info("   - generate_file: создание файлов (DOCX, Excel, PDF, Markdown) в папке output")
            logger.info("📱 Telegram бот поддерживает:")
            logger.info("   - Изображения: автоматический OCR + vision анализ")
            logger.info("   - Документы: полная обработка всех поддерживаемых форматов")
            logger.info("   - Аудио: транскрипция через Whisper")
            logger.info("   - Видео: анализ кадров + OCR + извлечение аудио")
            logger.info(f"🧠 Модель: {os.path.basename(self.brain_model)}")
            logger.info(f"📊 {self.get_context_info()}")
            logger.info("💻 Доступные команды: 'stats' (метрики), 'reset' (сброс), 'logs' (логи), 'export' (экспорт), 'memory' (память), 'gpu' (видеокарта), 'search' (поиск), 'preferences' (предпочтения), 'cleanup' (очистка), 'unload' (выгрузка моделей), 'exit' (выход)")
            logger.info("="*60)

        vision_desc = ""
        audio_text = ""
        while True:
            try:
                # АВТОМАТИЧЕСКАЯ ОБРАБОТКА ФАЙЛОВ ОТКЛЮЧЕНА - ТЕПЕРЬ ТОЛЬКО ПО ЗАПРОСУ
                # Автоматический поиск и обработка изображений/аудио удалена
                # Используйте новые действия: list_files и process_document
                
                vision_desc = ""
                audio_text = ""

                # 3. Запрашиваем у пользователя текстовый вопрос
                try:
                    if getattr(self, 'show_images_locally', True):
                        user_input = input("\n👤 Ваш вопрос (или Enter для пропуска, либо вставьте ссылку на YouTube): ").strip()
                    else:
                        # В веб-режиме не запрашиваем ввод
                        user_input = ""
                except EOFError:
                    # Если ввод из файла/pipe, используем пустую строку
                    user_input = ""
                    if getattr(self, 'show_images_locally', True):
                        logger.info("📝 Ввод из файла/pipe, продолжаю...")
                if user_input.lower() in ['exit', 'quit', 'выход']:
                    if getattr(self, 'show_images_locally', True):
                        logger.info("👋 До свидания!")
                    break
                if user_input.lower() in ['stats', 'метрики', 'статистика']:
                    # Показываем метрики производительности только в консольном режиме
                    if getattr(self, 'show_images_locally', True):
                        stats = self.get_performance_stats()
                        logger.info("\n📊 МЕТРИКИ ПРОИЗВОДИТЕЛЬНОСТИ:")
                        logger.info(f"   Всего действий: {stats['total_actions']}")
                        logger.info(f"   Среднее время ответа: {stats['avg_response_time']} сек")
                        if stats['recent_metrics']:
                            logger.info("   Последние действия:")
                            for metric in stats['recent_metrics'][-5:]:  # Показываем последние 5
                                timestamp = time.strftime("%H:%M:%S", time.localtime(metric['timestamp']))
                                logger.info(f"     [{timestamp}] {metric['action']}: {metric['response_time']:.2f} сек")
                        logger.info(f"   {self.get_context_info()}")
                    continue
                if user_input.lower() in ['reset', 'сброс', 'очистить']:
                    # Сбрасываем метрики и историю
                    self.performance_metrics.clear()
                    self.conversation_history.clear()
                    self.current_context_length = 0
                    if getattr(self, 'show_images_locally', True):
                        logger.info("🔄 Метрики и история сброшены")
                    continue
                if user_input.lower() in ['logs', 'логи']:
                    # Показываем последние записи из лог-файла только в консольном режиме
                    if getattr(self, 'show_images_locally', True):
                        try:
                            with open("ai_orchestrator.log", "r", encoding="utf-8") as f:
                                lines = f.readlines()
                                logger.info("\n📝 ПОСЛЕДНИЕ ЗАПИСИ В ЛОГЕ:")
                                for line in lines[-10:]:  # Показываем последние 10 строк
                                    logger.info(f"   {line.strip()}")
                        except Exception as e:
                            logger.error(f"Ошибка чтения лог-файла: {e}")
                    continue
                if user_input.lower() in ['export', 'экспорт']:
                    # Экспортируем метрики в JSON файл только в консольном режиме
                    if getattr(self, 'show_images_locally', True):
                        try:
                            stats = self.get_performance_stats()
                            export_data = {
                                "export_time": time.strftime("%Y-%m-%d %H:%M:%S"),
                                "performance_stats": stats,
                                "context_info": {
                                    "current": self.current_context_length,
                                    "safe": self.safe_context_length,
                                    "max": self.max_context_length
                                }
                            }
                            filename = f"metrics_export_{int(time.time())}.json"
                            with open(filename, "w", encoding="utf-8") as f:
                                json.dump(export_data, f, ensure_ascii=False, indent=2)
                            logger.info(f"📊 Метрики экспортированы в {filename}")
                        except Exception as e:
                            logger.error(f"Ошибка экспорта метрик: {e}")
                    continue
                if user_input.lower() in ['memory', 'память', 'mem']:
                    # Показываем статистику памяти ChromaDB
                    if getattr(self, 'show_images_locally', True):
                        try:
                            stats = self.get_memory_stats()
                            logger.info("\n🧠 СТАТИСТИКА ПАМЯТИ CHROMADB:")
                            if "error" not in stats:
                                logger.info(f"   Всего записей: {stats['total_records']}")
                                logger.info(f"   Диалогов: {stats['conversations']}")
                                logger.info(f"   Предпочтений: {stats['preferences']}")
                                logger.info(f"   Путь к БД: {stats['database_path']}")
                                logger.info(f"   Модель эмбеддингов: {stats['embedding_model']}")
                            else:
                                logger.error(f"   Ошибка получения статистики: {stats['error']}")
                        except Exception as e:
                            logger.error(f"Ошибка получения статистики памяти: {e}")
                    continue
                if user_input.lower() in ['gpu', 'видеокарта', 'gpuinfo']:
                    # Показываем информацию о GPU для ChromaDB
                    if getattr(self, 'show_images_locally', True):
                        try:
                            gpu_info = self.get_gpu_info()
                            logger.info("\n🎮 ИНФОРМАЦИЯ О GPU ДЛЯ CHROMADB:")
                            if "error" not in gpu_info:
                                logger.info(f"   GPU доступен: {'Да' if gpu_info['gpu_available'] else 'Нет'}")
                                if gpu_info['gpu_available']:
                                    logger.info(f"   Название GPU: {gpu_info['gpu_name']}")
                                    logger.info(f"   Память GPU: {gpu_info['gpu_memory']:.1f} GB")
                                logger.info(f"   Используемое устройство: {gpu_info['device_used']}")
                            else:
                                logger.error(f"   Ошибка получения информации о GPU: {gpu_info['error']}")
                        except Exception as e:
                            logger.error(f"Ошибка получения информации о GPU: {e}")
                    continue
                if user_input.lower() in ['cleanup', 'очистка', 'clean']:
                    # Очищаем старые записи из памяти
                    if getattr(self, 'show_images_locally', True):
                        try:
                            days = input("🗑️ Введите количество дней для хранения записей (по умолчанию 30): ").strip()
                            days_to_keep = int(days) if days.isdigit() else 30
                            deleted_count = self.cleanup_old_memory(days_to_keep)
                            logger.info(f"🧹 Удалено {deleted_count} старых записей из памяти")
                        except Exception as e:
                            logger.error(f"Ошибка очистки памяти: {e}")
                    continue
                if user_input.lower() in ['unload', 'выгрузка', 'unload_models']:
                    # Принудительная выгрузка всех моделей кроме мозга
                    logger.info("🔧 Принудительная выгрузка всех моделей...")
                    try:
                        import gc
                        import torch
                        
                        # Выключаем все системы
                        self.use_image_generation = False
                        self.use_vision = False
                        self.use_audio = False
                        
                        # Выгружаем pipeline для генерации изображений
                        if hasattr(self, 'current_pipeline') and self.current_pipeline is not None:
                            try:
                                if hasattr(self.current_pipeline, 'to'):
                                    self.current_pipeline.to('cpu')
                                del self.current_pipeline
                                self.current_pipeline = None
                                logger.info("✅ Pipeline для генерации изображений выгружен")
                            except Exception as e:
                                logger.warning(f"⚠️ Ошибка при выгрузке pipeline: {e}")
                        
                        # Выгружаем другие модели
                        model_attrs = ['vision_model', 'vision_processor', 'vision_pipeline', 
                                     'whisper_model', 'audio_model', 'tts_model']
                        for attr in model_attrs:
                            if hasattr(self, attr):
                                try:
                                    model = getattr(self, attr)
                                    if model is not None and hasattr(model, 'to'):
                                        model.to('cpu')
                                    delattr(self, attr)
                                    logger.info(f"✅ {attr} выгружен")
                                except Exception as e:
                                    logger.warning(f"⚠️ Ошибка при выгрузке {attr}: {e}")
                        
                        # Очистка GPU памяти
                        if torch.cuda.is_available():
                            torch.cuda.empty_cache()
                            torch.cuda.synchronize()
                            logger.info("✅ GPU память очищена")
                        
                        # Сборка мусора
                        gc.collect()
                        logger.info("✅ Все модели выгружены, память освобождена")
                        
                    except Exception as e:
                        logger.error(f"Ошибка при выгрузке моделей: {e}")
                    continue
                if user_input.lower() in ['search', 'поиск', 'find']:
                    # Поиск похожих диалогов в памяти
                    if getattr(self, 'show_images_locally', True):
                        try:
                            query = input("🔍 Введите поисковый запрос: ").strip()
                            if query:
                                results = self.search_similar_conversations(query, n_results=3)
                                if results:
                                    logger.info(f"\n🔍 НАЙДЕНО {len(results)} ПОХОЖИХ ДИАЛОГОВ:")
                                    for i, result in enumerate(results, 1):
                                        logger.info(f"   {i}. Схожесть: {result['similarity']:.2f}")
                                        logger.info(f"      ID: {result['id']}")
                                        logger.info(f"      Текст: {result['document'][:100]}...")
                                        logger.info(f"      Время: {time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(result['metadata']['timestamp']))}")
                                        logger.info("")
                                else:
                                    logger.info("🔍 Похожих диалогов не найдено")
                        except Exception as e:
                            logger.error(f"Ошибка поиска: {e}")
                    continue
                if user_input.lower() in ['preferences', 'предпочтения', 'prefs']:
                    # Показываем предпочтения пользователя
                    if getattr(self, 'show_images_locally', True):
                        try:
                            query = input("👤 Введите контекст для поиска предпочтений (или Enter для всех): ").strip()
                            preferences = self.get_user_preferences(query if query else None)
                            if preferences:
                                logger.info(f"\n👤 ПРЕДПОЧТЕНИЯ ПОЛЬЗОВАТЕЛЯ:")
                                logger.info(preferences)
                            else:
                                logger.info("👤 Предпочтения не найдены")
                        except Exception as e:
                            logger.error(f"Ошибка получения предпочтений: {e}")
                    continue
                if not user_input:
                    continue

                # 4. Перехват: если есть YouTube-ссылка, скачиваем видео и аудио, обрабатываем аудио, затем формируем brain_input: вопрос пользователя (без ссылки, вместо неё название ролика), затем текст из аудио, и только после этого отправляем brain_input в мозг
                # re уже импортирован в начале файла
                yt_url_match = re.search(r'https?://(?:www\.)?(?:youtube\.com|youtu\.be)/\S+', user_input)
                yt_processed = False  # Флаг для отслеживания обработки YouTube
                if yt_url_match:
                    yt_url = yt_url_match.group(0)
                    logger.info("🔗 Обнаружена ссылка на YouTube, скачиваю видео...")
                    
                    # Проверяем VPN статус перед скачиванием
                    if not self.check_vpn_status():
                        logger.warning("⚠️ VPN может не работать корректно. YouTube может быть недоступен.")
                        user_input_no_url = re.sub(r'https?://\S+', '[YouTube недоступен - проверьте VPN]', user_input).strip()
                        brain_input = f"{user_input_no_url}\n\n[ОШИБКА]: Не удалось скачать YouTube видео. Проверьте VPN соединение."
                        ai_response = self.call_brain_model(brain_input)
                        continue
                    
                    # Проверяем доступность YouTube ссылки
                    if not self.check_youtube_accessibility(yt_url):
                        logger.error("❌ YouTube ссылка недоступна. Проверьте VPN или ссылку.")
                        user_input_no_url = re.sub(r'https?://\S+', '[YouTube недоступен]', user_input).strip()
                        brain_input = f"{user_input_no_url}\n\n[ОШИБКА]: YouTube ссылка недоступна. Проверьте VPN или ссылку."
                        ai_response = self.call_brain_model(brain_input)
                        continue
                    
                    # Получаем информацию о видео
                    video_info = self.get_youtube_info(yt_url)
                    if video_info.get('success'):
                        video_title = video_info['title']
                        logger.info(f"📹 Название видео: {video_title}")
                    else:
                        video_title = "Неизвестное видео"
                        logger.warning(f"⚠️ Не удалось получить название: {video_info.get('error', 'Неизвестная ошибка')}")
                    yt_video = self.download_youtube_video(yt_url)
                    if yt_video:
                        logger.info(f"✅ Видео скачано: {yt_video}")
                        # video_title уже получен выше из get_youtube_info
                    else:
                        logger.error("❌ Не удалось скачать видео с YouTube")
                        video_title = "Видео не скачано"
                    logger.info("🔗 Скачиваю аудиодорожку для транскрипции...")
                    yt_audio = self.download_youtube_audio(yt_url)
                    if yt_audio:
                        logger.info(f"✅ Аудио скачано: {yt_audio}")
                        # Определяем язык по названию видео, а не по имени файла
                        if "english" in video_title.lower() or "eng" in video_title.lower():
                            lang = "en"
                        elif "рус" in video_title.lower() or "russian" in video_title.lower():
                            lang = "ru"
                        else:
                            # Автоматически определяем язык по названию видео
                            # Для Rick Astley и подобных - английский
                            if any(word in video_title.lower() for word in ["rick", "astley", "never", "gonna", "give", "you", "up"]):
                                lang = "en"
                                logger.info("🌐 Автоматически определен язык: английский (по названию видео)")
                            else:
                                # По умолчанию используем русский
                                lang = "ru"
                                logger.info("🌐 Используется язык по умолчанию: русский")
                        audio_text = self.transcribe_audio_whisper(yt_audio, lang=lang, use_separator=getattr(self, 'use_separator', True))
                        # --- VISION ПОКАДРОВО ---
                        vision_frames_desc = ""
                        if getattr(self, 'use_vision', False) and yt_video:
                            logger.info("🖼️ Извлекаю кадры из видео для vision...")
                            frames = self.extract_video_frames(yt_video, fps=1)
                            frame_results = []  # [(timecode, desc)]
                            for idx, (timecode, b64) in enumerate(frames):
                                if not b64:
                                    continue
                                vision_prompt = "Describe absolutely everything you see in the image, including all small details, their positions, and any visible text. Be as detailed as possible."
                                desc = self.call_vision_model(b64 + "\n" + vision_prompt)
                                
                                # Добавляем OCR к кадру видео
                                if getattr(self, 'use_ocr', False):
                                    try:
                                        # Декодируем base64 для OCR
                                        from PIL import Image
                                        
                                        image_data = base64.b64decode(b64)
                                        image = Image.open(io.BytesIO(image_data))
                                        
                                        # Извлекаем текст с помощью OCR
                                        ocr_text, ocr_error = self.extract_text_from_image_object(image)
                                        if ocr_text and ocr_text.strip():
                                            desc += f"\n[OCR TEXT]: {ocr_text.strip()}"
                                            logger.info(f"[OCR][{timecode}] Извлечен текст: {ocr_text.strip()}")
                                        elif ocr_error:
                                            logger.warning(f"[OCR][{timecode}] Ошибка OCR: {ocr_error}")
                                    except Exception as e:
                                        logger.warning(f"[OCR][{timecode}] Ошибка OCR: {e}")
                                
                                frame_results.append((timecode, desc))
                                logger.info(f"[VISION][{timecode}] {desc}")
                            # Группировка одинаковых описаний
                            # collections.defaultdict уже импортирован в начале файла
                            grouped = []  # [(list_of_timecodes, desc)]
                            prev_desc = None
                            prev_times = []
                            for i, (tc, desc) in enumerate(frame_results):
                                if prev_desc is None:
                                    prev_desc = desc
                                    prev_times = [tc]
                                elif desc == prev_desc:
                                    prev_times.append(tc)
                                else:
                                    grouped.append((prev_times, prev_desc))
                                    prev_desc = desc
                                    prev_times = [tc]
                            if prev_times:
                                grouped.append((prev_times, prev_desc))
                            # Формируем текст с диапазонами и списками
                            def format_timecodes(times):
                                def tc_to_sec(tc):
                                    h, m, s = map(int, tc.strip('[]').split(':'))
                                    return h*3600 + m*60 + s
                                if len(times) == 1:
                                    return times[0]
                                secs = [tc_to_sec(t) for t in times]
                                sorted_pairs = sorted(zip(secs, times))
                                secs_sorted, times_sorted = zip(*sorted_pairs)
                                # Проверяем, есть ли диапазон подряд
                                ranges = []
                                start = end = secs_sorted[0]
                                start_tc = times_sorted[0]
                                for i in range(1, len(secs_sorted)):
                                    if secs_sorted[i] == end + 1:
                                        end = secs_sorted[i]
                                    else:
                                        if start != end:
                                            ranges.append(f"[{start_tc}-{times_sorted[i-1]}]")
                                        else:
                                            ranges.append(f"{start_tc}")
                                        start = end = secs_sorted[i]
                                        start_tc = times_sorted[i]
                                if start != end:
                                    ranges.append(f"[{start_tc}-{times_sorted[-1]}]")
                                else:
                                    ranges.append(f"{start_tc}")
                                return ', '.join(ranges)
                            for times, desc in grouped:
                                vision_frames_desc += f"{format_timecodes(times)}: {desc}\n"
                        # Формируем brain_input: вопрос пользователя (без ссылки, вместо неё название ролика), ответы vision по кадрам, затем текст из аудио (с таймкодами)
                        user_input_no_url = re.sub(r'https?://\S+', f'[Видео]: {video_title}' if video_title else '', user_input).strip()
                        brain_input = ""
                        brain_input += user_input_no_url
                        if vision_frames_desc:
                            brain_input += f"\n[Покадровое описание видео]:\n{vision_frames_desc}"
                        if audio_text:
                            brain_input += f"\n[Текст из аудио]:\n{audio_text}"
                        # Отправляем в мозг и обрабатываем ответ
                        ai_response = self.call_brain_model(brain_input)
                        # Показываем информацию о контексте
                        logger.info(f"📊 {self.get_context_info()}")
                        continue_dialog = self.process_ai_response(ai_response)
                        if not continue_dialog:
                            logger.info("\n" + "="*60)
                        yt_processed = True  # Отмечаем, что YouTube обработан
                        continue  # пропускаем стандартный brain_input ниже
                    else:
                        logger.error("❌ Не удалось скачать аудио с YouTube")

                # 5. Собираем итоговый запрос для мозга (только если не было YouTube обработки)
                if not yt_processed:  # Используем флаг вместо yt_url_match
                    brain_input = ""
                    if vision_desc:
                        brain_input += f"[Описание изображения]:\n{vision_desc}\n"
                    if audio_text:
                        brain_input += f"[Текст из аудио]:\n{audio_text}\n"
                    brain_input += user_input

                    # 6. Отправляем запрос в мозг
                    ai_response = self.call_brain_model(brain_input)
                    # Показываем информацию о контексте
                    logger.info(f"📊 {self.get_context_info()}")
                    # Обрабатываем ответ AI
                    continue_dialog = self.process_ai_response(ai_response)
                    if not continue_dialog:
                        logger.info("\n" + "="*60)

            except Exception as e:
                if isinstance(e, KeyboardInterrupt):
                    logger.info("\n👋 Программа прервана пользователем")
                    break
                logger.error(f"❌ Неожиданная ошибка: {str(e)}")
                # Убираем дублирование обработки аудио - это уже делается выше
                audio_text = ""

                # 3. Запрашиваем у пользователя текстовый вопрос
                try:
                    user_input = input("\n👤 Ваш вопрос (или Enter для пропуска, либо вставьте ссылку на YouTube): ").strip()
                except EOFError:
                    # Если ввод из файла/pipe, используем пустую строку
                    user_input = ""
                    logger.info("📝 Ввод из файла/pipe, продолжаю...")
                if user_input.lower() in ['exit', 'quit', 'выход']:
                    logger.info("👋 До свидания!")
                    break
                if user_input.lower() in ['stats', 'метрики', 'статистика']:
                    # Показываем метрики производительности
                    stats = self.get_performance_stats()
                    logger.info("\n📊 МЕТРИКИ ПРОИЗВОДИТЕЛЬНОСТИ:")
                    logger.info(f"   Всего действий: {stats['total_actions']}")
                    logger.info(f"   Среднее время ответа: {stats['avg_response_time']} сек")
                    if stats['recent_metrics']:
                        logger.info("   Последние действия:")
                        for metric in stats['recent_metrics'][-5:]:  # Показываем последние 5
                            timestamp = time.strftime("%H:%M:%S", time.localtime(metric['timestamp']))
                            logger.info(f"     [{timestamp}] {metric['action']}: {metric['response_time']:.2f} сек")
                    logger.info(f"   {self.get_context_info()}")
                    continue
                if user_input.lower() in ['reset', 'сброс', 'очистить']:
                    # Сбрасываем метрики и историю
                    self.performance_metrics.clear()
                    self.conversation_history.clear()
                    self.current_context_length = 0
                    logger.info("🔄 Метрики и история сброшены")
                    continue
                if user_input.lower() in ['logs', 'логи']:
                    # Показываем последние записи из лог-файла
                    try:
                        with open("ai_orchestrator.log", "r", encoding="utf-8") as f:
                            lines = f.readlines()
                            logger.info("\n📝 ПОСЛЕДНИЕ ЗАПИСИ В ЛОГЕ:")
                            for line in lines[-10:]:  # Показываем последние 10 строк
                                logger.info(f"   {line.strip()}")
                    except Exception as e:
                        logger.error(f"Ошибка чтения лог-файла: {e}")
                    continue
                if user_input.lower() in ['export', 'экспорт']:
                    # Экспортируем метрики в JSON файл
                    try:
                        stats = self.get_performance_stats()
                        export_data = {
                            "export_time": time.strftime("%Y-%m-%d %H:%M:%S"),
                            "performance_stats": stats,
                            "context_info": {
                                "current": self.current_context_length,
                                "safe": self.safe_context_length,
                                "max": self.max_context_length
                            }
                        }
                        filename = f"metrics_export_{int(time.time())}.json"
                        with open(filename, "w", encoding="utf-8") as f:
                            json.dump(export_data, f, ensure_ascii=False, indent=2)
                        logger.info(f"📊 Метрики экспортированы в {filename}")
                    except Exception as e:
                        logger.error(f"Ошибка экспорта метрик: {e}")
                    continue
                if not user_input:
                    continue



            except KeyboardInterrupt:
                logger.info("\n�� Программа прервана пользователем")
                break



    def start_telegram_bot(self) -> bool:
        """Запускает Telegram бота. Возвращает True при успешном старте, иначе False."""
        if not self.telegram_bot_token:
            logger.warning("❌ Telegram Bot токен не указан")
            return False
        
        # Предварительная проверка токена через getMe и редактирование токена в логах
        try:
            redacted = self.telegram_bot_token[:10] + "..." if len(self.telegram_bot_token) > 13 else "***"
            logger.info(f"🔐 Проверяю Telegram токен (redacted: {redacted})")
            resp = requests.get(f"https://api.telegram.org/bot{self.telegram_bot_token}/getMe", timeout=5)
            if resp.status_code != 200:
                logger.error("❌ Ошибка проверки Telegram токена: сервер вернул неуспешный статус")
                return False
            data = resp.json()
            if not data.get("ok"):
                # Не логируем сырой токен
                logger.error(f"❌ Telegram токен отклонен сервером (token: {redacted})")
                return False
        except Exception as e:
            logger.error(f"❌ Ошибка проверки Telegram токена: {e}")
            return False
        
        try:
            # Создаем приложение
            self.telegram_app = Application.builder().token(self.telegram_bot_token).build()
            
            # Добавляем обработчики
            self.telegram_app.add_handler(CommandHandler("start", self._telegram_start))
            self.telegram_app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, self._telegram_text_message))
            self.telegram_app.add_handler(MessageHandler(filters.PHOTO, self._telegram_photo_message))
            self.telegram_app.add_handler(MessageHandler(filters.AUDIO | filters.VOICE, self._telegram_audio_message))
            self.telegram_app.add_handler(MessageHandler(filters.Document.ALL, self._telegram_document_message))
            
            # Запускаем бота в фоне в отдельном потоке
            import threading
            def run_bot():
                loop = None
                try:
                    # Создаем новый event loop для потока
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)
                    from typing import Any, cast
                    coro = self.telegram_app.run_polling(allowed_updates=Update.ALL_TYPES)
                    if coro is not None:
                        loop.run_until_complete(cast(Any, coro))
                except Exception as e:
                    # В веб-режиме логируем тихо
                    if not getattr(self, 'show_images_locally', True):
                        logger.error(f"❌ Ошибка в Telegram боте: {e}")
                    else:
                        logger.debug(f"Telegram bot polling error: {e}")
                finally:
                    if loop is not None:
                        try:
                            loop.close()
                        except Exception:
                            pass
            
            bot_thread = threading.Thread(target=run_bot, daemon=True)
            bot_thread.start()
            
            # Показываем сообщение только в консольном режиме
            if not getattr(self, 'show_images_locally', True):
                logger.info("🤖 Telegram бот запущен в фоновом режиме")
            return True
            
        except Exception as e:
            # В веб-режиме логируем тихо
            if not getattr(self, 'show_images_locally', True):
                logger.error(f"❌ Ошибка запуска Telegram бота: {e}")
            else:
                logger.debug(f"Telegram bot startup error: {e}")
            return False

    async def _safe_reply(self, update: Update, message: str):
        """Безопасная отправка сообщения в Telegram"""
        if update and update.message:
            await update.message.reply_text(message)

    async def _telegram_start(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Обработчик команды /start"""
        if update is None or update.message is None or update.effective_user is None:
            return
        user_id = str(update.effective_user.id)
        # Разрешаем доступ всем пользователям
        
        await self._safe_reply(update,
            "🤖 Привет! Я Нейро - AI оркестратор.\n"
            "Я могу:\n"
            "• Обрабатывать текстовые сообщения\n"
            "• Анализировать изображения\n"
            "• Транскрибировать аудио\n"
            "• Генерировать изображения\n"
            "• Выполнять команды PowerShell\n"
            "• Искать информацию в интернете\n\n"
            "Просто отправьте мне сообщение, изображение или аудио!"
        )

    async def _telegram_text_message(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Обработчик текстовых сообщений"""
        if update is None or update.message is None or update.effective_user is None or update.effective_chat is None:
            return
        user_id = str(update.effective_user.id)
        # Разрешаем доступ всем пользователям
        
        text = update.message.text if update.message and update.message.text else ""
        
        # Проверяем специальные команды OCR
        if any(keyword in text.lower() for keyword in ['ocr', 'распознай текст', 'извлеки текст', 'что написано']):
            # Если есть последнее изображение, применяем к нему OCR
            if hasattr(self, 'last_telegram_image') and self.last_telegram_image:
                await self._safe_reply(update, "🔄 Применяю OCR к последнему изображению...")
                try:
                    if getattr(self, 'use_ocr', False):
                        # Конвертируем base64 в PIL Image
                        from PIL import Image
                        
                        image_data = base64.b64decode(self.last_telegram_image)
                        image = Image.open(io.BytesIO(image_data))
                        
                        # Извлекаем текст принудительно
                        ocr_text, ocr_error = self.extract_text_from_image_object(image)
                        
                        if ocr_text and ocr_text.strip():
                            await update.message.reply_text(f"📖 Извлеченный текст:\n\n{ocr_text.strip()}")
                        elif ocr_error:
                            await update.message.reply_text(f"❌ Ошибка OCR: {ocr_error}")
                        else:
                            await update.message.reply_text("⚠️ Текст в изображении не найден")
                    else:
                        await update.message.reply_text("❌ OCR отключен в системе")
                        
                except Exception as e:
                    await update.message.reply_text(f"❌ Ошибка OCR: {str(e)}")
                return
            else:
                await update.message.reply_text("❌ Нет изображения для OCR. Отправьте сначала изображение.")
                return
        
        await update.message.reply_text("🔄 Обрабатываю ваше сообщение...")

        try:
            # Отправляем в мозг
            ai_response = self.call_brain_model(text or "")

            # Обрабатываем ответ AI
            continue_dialog = self.process_ai_response(ai_response)

            if not continue_dialog:
                # Если диалог завершен, отправляем финальный ответ
                if hasattr(self, 'last_final_response') and self.last_final_response:
                    await update.message.reply_text(self.last_final_response)

                    # Если есть сгенерированное изображение, отправляем его
                    if hasattr(self, 'last_generated_image_b64') and self.last_generated_image_b64:
                        try:
                            # Конвертируем base64 в bytes
                            img_bytes = base64.b64decode(self.last_generated_image_b64)

                            # Отправляем изображение
                            await context.bot.send_photo(
                                chat_id=update.effective_chat.id,
                                photo=img_bytes,
                                caption="🎨 Сгенерированное изображение"
                            )

                            # Очищаем
                            self.last_generated_image_b64 = None

                        except Exception as e:
                            # В веб-режиме логируем тихо
                            if not getattr(self, 'show_images_locally', True):
                                logger.error(f"❌ Ошибка отправки изображения: {e}")
                            else:
                                logger.debug(f"Telegram image send error: {e}")
                            await update.message.reply_text("❌ Ошибка отправки изображения")
                    
                    # Если есть сгенерированный файл, отправляем его
                    if hasattr(self, 'last_generated_file_path') and self.last_generated_file_path:
                        try:
                            # Проверяем что файл существует
                            if os.path.exists(self.last_generated_file_path):
                                # Отправляем файл как документ
                                with open(self.last_generated_file_path, 'rb') as file:
                                    await context.bot.send_document(
                                        chat_id=update.effective_chat.id,
                                        document=file,
                                        filename=self.last_generated_file_name or os.path.basename(self.last_generated_file_path),
                                        caption="📄 Сгенерированный файл"
                                    )
                                
                                # Очищаем
                                self.last_generated_file_path = None
                                self.last_generated_file_name = None
                                
                            else:
                                await update.message.reply_text("❌ Сгенерированный файл не найден")
                                
                        except Exception as e:
                            # В веб-режиме логируем тихо
                            if not getattr(self, 'show_images_locally', True):
                                logger.error(f"❌ Ошибка отправки файла: {e}")
                            else:
                                logger.debug(f"Telegram file send error: {e}")
                            await update.message.reply_text("❌ Ошибка отправки файла")
                else:
                    await update.message.reply_text("✅ Задача выполнена!")
            else:
                # Если диалог продолжается, отправляем промежуточный ответ
                await update.message.reply_text("🔄 Обрабатываю... Пожалуйста, подождите.")

        except Exception as e:
            # Веб-режим: логируем тихо
            if getattr(self, 'show_images_locally', True):
                logger.error(f"❌ Ошибка обработки текстового сообщения: {e}")
            else:
                logger.debug(f"Telegram text message error: {e}")
            await update.message.reply_text(f"❌ Произошла ошибка: {str(e)}")

    async def _telegram_photo_message(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Обработчик фотографий"""
        if update is None or update.message is None or update.effective_user is None or update.effective_chat is None:
            return
        user_id = str(update.effective_user.id)
        # Разрешаем доступ всем пользователям
        
        await update.message.reply_text("🖼️ Обрабатываю изображение...")
        
        try:
            # Получаем фото
            photo = update.message.photo[-1]  # Берем самое большое фото
            file = await context.bot.get_file(photo.file_id)
            
            # Скачиваем фото
            photo_bytes = await file.download_as_bytearray()
            photo_b64 = base64.b64encode(photo_bytes).decode('ascii')
            
            # Анализируем изображение с vision моделью
            vision_desc = self.call_vision_model(photo_b64)
            
            # Применяем умный OCR
            result_message = f"👁️ Описание изображения:\n{vision_desc}"
            
            if getattr(self, 'use_ocr', False):
                try:
                    # Конвертируем байты в PIL Image для OCR
                    from PIL import Image
                    import io
                    
                    image = Image.open(io.BytesIO(photo_bytes))
                    
                    # Проверяем, нужно ли применять OCR
                    should_use_ocr = self.should_use_ocr_on_image(vision_desc)
                    
                    if should_use_ocr:
                        # Извлекаем текст с помощью OCR
                        ocr_text, ocr_error = self.extract_text_from_image_object(image)
                        
                        if ocr_text and ocr_text.strip():
                            result_message += f"\n\n📖 Извлеченный текст:\n{ocr_text.strip()}"
                            await update.message.reply_text("✅ Обнаружен текст на изображении!")
                        elif ocr_error:
                            result_message += f"\n\n⚠️ OCR ошибка: {ocr_error}"
                    else:
                        await update.message.reply_text("ℹ️ Текст на изображении не обнаружен")
                        
                except Exception as ocr_exception:
                    result_message += f"\n\n⚠️ Ошибка OCR: {str(ocr_exception)}"
                    
            # Отправляем полное описание
            await update.message.reply_text(result_message)
            
            # Сохраняем для возможного использования в диалоге
            self.last_telegram_image = photo_b64
            
        except Exception as e:
            # В веб-режиме логируем тихо
            if getattr(self, 'show_images_locally', True):
                logger.error(f"❌ Ошибка обработки фото: {e}")
            else:
                logger.debug(f"Telegram photo processing error: {e}")
            await update.message.reply_text(f"❌ Ошибка обработки изображения: {str(e)}")

    async def _telegram_audio_message(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Обработчик аудио сообщений"""
        if update is None or update.message is None or update.effective_user is None or update.effective_chat is None:
            return
        user_id = str(update.effective_user.id)
        # Разрешаем доступ всем пользователям
        
        await update.message.reply_text("🎵 Обрабатываю аудио...")
        
        try:
            # Получаем аудио
            if update.message.audio:
                audio = update.message.audio
            else:
                audio = update.message.voice

            if audio is None:
                await update.message.reply_text("❌ В сообщении нет аудиофайла")
                return
            
            file = await context.bot.get_file(audio.file_id)
            
            # Скачиваем аудио
            audio_bytes = await file.download_as_bytearray()
            
            # Сохраняем во временный файл
            temp_dir = os.path.join(os.path.dirname(__file__), "temp_audio")
            os.makedirs(temp_dir, exist_ok=True)
            temp_file = os.path.join(temp_dir, f"telegram_audio_{int(time.time())}.ogg")
            
            with open(temp_file, 'wb') as f:
                f.write(audio_bytes)
            
            # Транскрибируем аудио
            transcript = self.transcribe_audio_whisper(temp_file, use_separator=False)
            
            if transcript and not transcript.startswith("[Whisper error]"):
                await update.message.reply_text(f"🎤 Транскрипция аудио:\n{transcript}")
                
                # Сохраняем для возможного использования в диалоге
                self.last_telegram_audio_transcript = transcript
            else:
                await update.message.reply_text("❌ Не удалось распознать аудио")
            
            # Удаляем временный файл
            try:
                os.remove(temp_file)
            except Exception:
                pass
                
        except Exception as e:
            # В веб-режиме логируем тихо
            if getattr(self, 'show_images_locally', True):
                logger.error(f"❌ Ошибка обработки аудио: {e}")
            else:
                logger.debug(f"Telegram audio processing error: {e}")
            await update.message.reply_text(f"❌ Ошибка обработки аудио: {str(e)}")

    async def _telegram_document_message(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Обработчик документов"""
        if update is None or update.message is None or update.effective_user is None or update.effective_chat is None:
            return
        user_id = str(update.effective_user.id)
        # Разрешаем доступ всем пользователям
        
        await update.message.reply_text("📄 Обрабатываю документ...")
        
        try:
            # Получаем документ
            document = update.message.document
            if document is None:
                await update.message.reply_text("❌ В сообщении нет документа")
                return
            
            file_name = document.file_name or "unknown_file"
            file_size = document.file_size
            
            # Проверяем размер файла (максимум 20MB)
            if file_size and file_size > 20 * 1024 * 1024:
                await update.message.reply_text("❌ Файл слишком большой (максимум 20MB)")
                return
            
            file = await context.bot.get_file(document.file_id)
            
            # Определяем тип файла и обработку
            file_lower = file_name.lower()
            
            if file_lower.endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp')):
                # Обрабатываем как изображение с OCR
                await self._process_telegram_image_document(update, file, file_name)
            elif file_lower.endswith(('.docx', '.doc', '.pdf', '.xlsx', '.xls', '.csv', '.txt', '.md', '.rtf', '.json', '.xml', '.html', '.htm')):
                # Обрабатываем как документ
                await self._process_telegram_text_document(update, file, file_name)
            elif file_lower.endswith(('.mp3', '.wav', '.ogg', '.m4a', '.flac', '.aac', '.wma')):
                # Обрабатываем как аудио
                await self._process_telegram_audio_document(update, file, file_name)
            elif file_lower.endswith(('.mp4', '.avi', '.mkv', '.mov', '.wmv', '.flv', '.webm')):
                # Обрабатываем как видео
                await self._process_telegram_video_document(update, file, file_name)
            else:
                await update.message.reply_text(f"❌ Неподдерживаемый тип файла: {file_name}\n\nПоддерживаются:\n• Изображения: JPG, PNG, GIF, BMP, WEBP\n• Документы: DOCX, DOC, PDF, XLSX, XLS, CSV, TXT, MD, RTF, JSON, XML, HTML\n• Аудио: MP3, WAV, OGG, M4A, FLAC, AAC, WMA\n• Видео: MP4, AVI, MKV, MOV, WMV, FLV, WEBM")
                
        except Exception as e:
            # В веб-режиме логируем тихо
            if getattr(self, 'show_images_locally', True):
                logger.error(f"❌ Ошибка обработки документа: {e}")
            else:
                logger.debug(f"Telegram document processing error: {e}")
            await update.message.reply_text(f"❌ Ошибка обработки документа: {str(e)}")

    async def _process_telegram_image_document(self, update: Update, file, file_name: str):
        """Обработка изображений через документы с OCR"""
        try:
            # Скачиваем изображение
            image_bytes = await file.download_as_bytearray()
            image_b64 = base64.b64encode(image_bytes).decode('ascii')
            
            # Анализируем изображение с vision моделью
            vision_desc = self.call_vision_model(image_b64)
            
            result_message = f"🖼️ Анализ изображения '{file_name}':\n\n👁️ Описание:\n{vision_desc}"
            
            # Применяем OCR
            if getattr(self, 'use_ocr', False):
                try:
                    from PIL import Image
                    import io
                    
                    image = Image.open(io.BytesIO(image_bytes))
                    
                    # Всегда применяем OCR для документов (более вероятно содержат текст)
                    ocr_text, ocr_error = self.extract_text_from_image_object(image)
                    
                    if ocr_text and ocr_text.strip():
                        result_message += f"\n\n📖 Извлеченный текст:\n{ocr_text.strip()}"
                        if update.message:
                            await update.message.reply_text("✅ Текст успешно извлечен из документа!")
                    elif ocr_error:
                        result_message += f"\n\n⚠️ OCR ошибка: {ocr_error}"
                    else:
                        result_message += f"\n\n⚠️ Текст в изображении не найден"
                        
                except Exception as ocr_exception:
                    result_message += f"\n\n⚠️ Ошибка OCR: {str(ocr_exception)}"
            else:
                result_message += f"\n\n📖 OCR отключен. Используйте vision описание выше."
            
            if update.message:
                await update.message.reply_text(result_message)
            
        except Exception as e:
            if update.message:
                await update.message.reply_text(f"❌ Ошибка обработки изображения: {str(e)}")

    async def _process_telegram_text_document(self, update: Update, file, file_name: str):
        """Обработка текстовых документов (DOCX, PDF, Excel)"""
        try:
            # Скачиваем документ
            doc_bytes = await file.download_as_bytearray()
            
            # Сохраняем во временный файл
            temp_dir = os.path.join(os.path.dirname(__file__), "temp_docs")
            os.makedirs(temp_dir, exist_ok=True)
            temp_file = os.path.join(temp_dir, f"telegram_doc_{int(time.time())}_{file_name}")
            
            with open(temp_file, 'wb') as f:
                f.write(doc_bytes)
            
            # Обрабатываем документ
            result = self.process_document_request(temp_file)
            
            if update.message:
                await update.message.reply_text(f"📄 Анализ документа '{file_name}':\n\n{result}")
            
            # Удаляем временный файл
            try:
                os.remove(temp_file)
            except Exception:
                pass
                
        except Exception as e:
            if update.message:
                await update.message.reply_text(f"❌ Ошибка обработки документа: {str(e)}")

    async def _process_telegram_audio_document(self, update: Update, file, file_name: str):
        """Обработка аудио файлов через документы"""
        try:
            # Скачиваем аудио
            audio_bytes = await file.download_as_bytearray()
            
            # Сохраняем во временный файл
            temp_dir = os.path.join(os.path.dirname(__file__), "temp_audio")
            os.makedirs(temp_dir, exist_ok=True)
            temp_file = os.path.join(temp_dir, f"telegram_audio_{int(time.time())}_{file_name}")
            
            with open(temp_file, 'wb') as f:
                f.write(audio_bytes)
            
            # Транскрибируем аудио
            transcript = self.transcribe_audio_whisper(temp_file, use_separator=False)
            
            if transcript and not transcript.startswith("[Whisper error]"):
                await self._safe_reply(update, f"🎤 Транскрипция аудио '{file_name}':\n\n{transcript}")
            else:
                await self._safe_reply(update, "❌ Не удалось распознать аудио")
            
            # Удаляем временный файл
            try:
                os.remove(temp_file)
            except Exception:
                pass
                
        except Exception as e:
            await self._safe_reply(update, f"❌ Ошибка обработки аудио: {str(e)}")

    async def _process_telegram_video_document(self, update: Update, file, file_name: str):
        """Обработка видео файлов через документы"""
        try:
            await self._safe_reply(update, "🎬 Загружаю видео и извлекаю кадры для анализа...")
            
            # Скачиваем видео
            video_bytes = await file.download_as_bytearray()
            
            # Сохраняем во временный файл
            temp_dir = os.path.join(os.path.dirname(__file__), "temp_video")
            os.makedirs(temp_dir, exist_ok=True)
            temp_file = os.path.join(temp_dir, f"telegram_video_{int(time.time())}_{file_name}")
            
            with open(temp_file, 'wb') as f:
                f.write(video_bytes)
            
            # Извлекаем кадры для анализа
            frames = self.extract_video_frames(temp_file, fps=1)
            
            if frames:
                await self._safe_reply(update, f"🎬 Анализирую {len(frames)} кадров из видео...")
                
                result_message = f"🎬 Анализ видео '{file_name}':\n\n"
                
                # Анализируем несколько кадров (максимум 3, чтобы не перегружать)
                frames_to_analyze = frames[:3]
                
                for idx, (timecode, b64) in enumerate(frames_to_analyze):
                    if not b64:
                        continue
                    
                    # Vision анализ кадра
                    if getattr(self, 'use_vision', False):
                        vision_desc = self.call_vision_model(b64)
                        result_message += f"🕐 {timecode}: {vision_desc}\n\n"
                        
                        # OCR для кадра
                        if getattr(self, 'use_ocr', False):
                            try:
                                from PIL import Image
                                
                                image_data = base64.b64decode(b64)
                                image = Image.open(io.BytesIO(image_data))
                                
                                # Проверяем, нужно ли применять OCR
                                if self.should_use_ocr_on_image(vision_desc):
                                    ocr_text, ocr_error = self.extract_text_from_image_object(image)
                                    if ocr_text and ocr_text.strip():
                                        result_message += f"📖 Текст в кадре {timecode}: {ocr_text.strip()}\n\n"
                            except Exception:
                                pass  # Игнорируем ошибки OCR для видео
                    
                    # Прогресс для пользователя
                    if idx == 0:
                        await self._safe_reply(update, "🔄 Анализ первого кадра завершен...")
                
                # Проверяем, есть ли аудио дорожка для транскрипции
                try:
                    # Попробуем извлечь аудио и транскрибировать
                    audio_file = temp_file.replace(os.path.splitext(temp_file)[1], '.wav')
                    
                    # Используем ffmpeg для извлечения аудио
                    import subprocess
                    cmd = ['ffmpeg', '-i', temp_file, '-vn', '-acodec', 'pcm_s16le', '-ar', '16000', '-ac', '1', audio_file, '-y']
                    result = subprocess.run(cmd, capture_output=True, text=True)
                    
                    if result.returncode == 0 and os.path.exists(audio_file):
                        await self._safe_reply(update, "🎤 Транскрибирую аудио из видео...")
                        transcript = self.transcribe_audio_whisper(audio_file, use_separator=False)
                        
                        if transcript and not transcript.startswith("[Whisper error]"):
                            result_message += f"🎤 Аудио транскрипция:\n{transcript}\n\n"
                        
                        # Удаляем временный аудио файл
                        try:
                            os.remove(audio_file)
                        except:
                            pass
                            
                except Exception:
                    pass  # Игнорируем ошибки аудио извлечения
                
                # Отправляем полный результат
                if len(result_message) > 4000:  # Telegram лимит ~4096 символов
                    # Разбиваем на части
                    parts = [result_message[i:i+3500] for i in range(0, len(result_message), 3500)]
                    for i, part in enumerate(parts):
                        if i == 0:
                            await self._safe_reply(update, part)
                        else:
                            await self._safe_reply(update, f"(продолжение {i+1}):\n{part}")
                else:
                    await self._safe_reply(update, result_message)
            else:
                await self._safe_reply(update, "❌ Не удалось извлечь кадры из видео")
            
            # Удаляем временный файл
            try:
                os.remove(temp_file)
            except Exception:
                pass
                
        except Exception as e:
            await self._safe_reply(update, f"❌ Ошибка обработки видео: {str(e)}")

    def play_audio_file(self, audio_path: str) -> bool:
        """
        Воспроизводит аудиофайл один раз без зацикливания
        
        Args:
            audio_path: Путь к аудиофайлу
            
        Returns:
            True если воспроизведение запущено успешно, False при ошибке
        """
        try:
            import pygame  # type: ignore
            import time
            
            logger.info(f"🔊 Воспроизводим аудио: {os.path.basename(audio_path)}")
            
            # Инициализируем pygame mixer
            pygame.mixer.init()
            
            # Загружаем и воспроизводим аудио
            try:
                pygame.mixer.music.load(audio_path)
                pygame.mixer.music.play()
                
                # Ждем окончания воспроизведения
                while pygame.mixer.music.get_busy():
                    time.sleep(0.1)
                
                # Останавливаем и закрываем
                pygame.mixer.music.stop()
                pygame.mixer.quit()
                
                logger.info("✅ Аудио воспроизведено через pygame")
                return True
                
            except Exception as e:
                logger.warning(f"⚠️ pygame воспроизведение не удалось: {e}")
                pygame.mixer.quit()
                
                # Fallback: открываем в плеере по умолчанию
                import subprocess
                subprocess.Popen(["start", audio_path], shell=True)
                logger.info("🔄 Открыт в плеере по умолчанию")
                return True
                
        except ImportError:
            logger.warning("⚠️ pygame не установлен, используем fallback")
            # Fallback: открываем в плеере по умолчанию
            import subprocess
            subprocess.Popen(["start", audio_path], shell=True)
            logger.info("🔄 Открыт в плеере по умолчанию")
            return True
            
        except Exception as e:
            logger.error(f"❌ Ошибка воспроизведения аудио: {e}")
            return False

    def text_to_speech(self, text: str, voice: str = "male", language: str = "ru", auto_play: bool = True) -> str:
        """
        Озвучивает текст с помощью gTTS (Google Text-to-Speech)
        
        Args:
            text: Текст для озвучки
            voice: Тип голоса ("male" или "female") - пока не используется в gTTS
            language: Язык текста ("ru", "en", etc.)
            auto_play: Автоматически воспроизвести после создания файла
            
        Returns:
            Путь к созданному аудиофайлу или пустая строка при ошибке
        """
        try:
            from gtts import gTTS
            
            # Создаем папку для сгенерированной речи
            output_dir = os.path.join(os.path.dirname(__file__), "Audio", "generated_speech")
            os.makedirs(output_dir, exist_ok=True)
            
            # Генерируем уникальное имя файла
            timestamp = int(time.time())
            filename = f"tts_{voice}_{language}_{timestamp}.mp3"
            output_path = os.path.join(output_dir, filename)
            
            logger.info(f"🎤 Озвучиваю текст: {text[:100]}...")
            logger.info(f"🔊 Голос: {voice}, Язык: {language}")
            logger.info(f"🌐 Использую Google TTS API")
            
            # Создаем TTS объект
            tts = gTTS(text=text, lang=language, slow=False)
            
            # Сохраняем аудиофайл
            tts.save(output_path)
            
            logger.info(f"✅ Аудиофайл сохранен: {output_path}")
            
            # Автоматически воспроизводим, если включено
            if auto_play:
                self.play_audio_file(output_path)
            
            return output_path
            
        except ImportError:
            logger.error("❌ gTTS не установлен. Установите: pip install gTTS")
            return ""
        except Exception as e:
            logger.error(f"❌ Ошибка озвучки текста: {e}")
            logger.error(f"🔍 Тип ошибки: {type(e).__name__}")
            import traceback
            logger.error(f"📋 Traceback: {traceback.format_exc()}")
            return ""

    def enhance_prompt_with_memory(self, user_message: str, system_prompt: str = "") -> str:
        """
        Улучшает промпт с помощью контекста из памяти
        
        Args:
            user_message: Сообщение пользователя
            system_prompt: Системный промпт
            
        Returns:
            Улучшенный промпт с контекстом
        """
        try:
            # Получаем релевантный контекст
            context = self.get_relevant_context(user_message, max_context_length=1500)
            
            # Получаем предпочтения пользователя
            preferences = self.get_user_preferences(user_message)
            
            # Формируем улучшенный промпт
            enhanced_prompt = system_prompt
            added_content = ""
            
            if context:
                context_section = f"\n\nРЕЛЕВАНТНЫЙ КОНТЕКСТ ИЗ ПРЕДЫДУЩИХ ДИАЛОГОВ:\n{context}\n\nИНСТРУКЦИЯ: Используйте эту информацию только если она релевантна текущему запросу. Не упоминайте источник информации явно."
                enhanced_prompt += context_section
                added_content += context_section
            
            if preferences:
                preferences_section = f"\n\nПРЕДПОЧТЕНИЯ ПОЛЬЗОВАТЕЛЯ:\n{preferences}\n\nИНСТРУКЦИЯ: Учитывайте эти предпочтения при формировании ответа, но не упоминайте их явно."
                enhanced_prompt += preferences_section
                added_content += preferences_section
            
            enhanced_prompt += f"\n\nТЕКУЩИЙ ЗАПРОС ПОЛЬЗОВАТЕЛЯ: {user_message}"
            
            # Подсчитываем только добавленный контент (без системного промпта и запроса пользователя)
            added_length = len(added_content)
            if added_length > 0:
                logger.info(f"📚 Промпт улучшен с помощью памяти (добавлено: {added_length} символов)")
            else:
                logger.info(f"📚 Промпт не улучшен (память пуста)")
            
            # Логируем общую длину промпта для сравнения с токенами
            total_prompt_length = len(enhanced_prompt)
            logger.info(f"📝 Общая длина промпта: {total_prompt_length} символов")
            
            return enhanced_prompt
            
        except Exception as e:
            logger.error(f"❌ Ошибка улучшения промпта: {e}")
            return f"{system_prompt}\n\nТЕКУЩИЙ ЗАПРОС ПОЛЬЗОВАТЕЛЯ: {user_message}"
    
    def auto_save_conversation(self, user_message: str, ai_response: str, 
                              context: str = "", metadata: Optional[Dict[str, Any]] = None):
        """
        Автоматически сохраняет диалог в память
        
        Args:
            user_message: Сообщение пользователя
            ai_response: Ответ ИИ
            context: Дополнительный контекст
            metadata: Дополнительные метаданные
        """
        try:
            # Добавляем базовые метаданные
            if metadata is None:
                metadata = {}
            
            metadata.update({
                "auto_saved": True,
                "response_length": len(ai_response),
                "user_message_length": len(user_message)
            })
            
            # Сохраняем в память
            success = self.add_to_memory(user_message, ai_response, context, metadata)
            
            if success:
                logger.info("💾 Диалог автоматически сохранен в память")
            else:
                logger.warning("⚠️ Не удалось сохранить диалог в память")
                
        except Exception as e:
            logger.error(f"❌ Ошибка автоматического сохранения: {e}")
    
    def extract_preferences_from_response(self, user_message: str, ai_response: str):
        """
        Извлекает предпочтения пользователя из диалога и сохраняет их
        
        Args:
            user_message: Сообщение пользователя
            ai_response: Ответ ИИ
        """
        try:
            # Простые паттерны для извлечения предпочтений
            preference_patterns = [
                r"мне нравится (.+?)(?:\.|$)",
                r"я предпочитаю (.+?)(?:\.|$)",
                r"лучше всего (.+?)(?:\.|$)",
                r"хочу (.+?)(?:\.|$)",
                r"нужно (.+?)(?:\.|$)",
                r"важно (.+?)(?:\.|$)"
            ]
            
            # Ищем предпочтения в сообщении пользователя
            for pattern in preference_patterns:
                matches = re.findall(pattern, user_message.lower())
                for match in matches:
                    if len(match) > 10:  # Минимальная длина предпочтения
                        self.add_user_preference(
                            match.strip(),
                            category="user_preference",
                            metadata={"source": "extracted", "pattern": pattern}
                        )
            
            # Также ищем в ответе ИИ, если там есть подтверждения
            confirmation_patterns = [
                r"понял(?:а)?, что вам нравится (.+?)(?:\.|$)",
                r"запомню, что вы предпочитаете (.+?)(?:\.|$)",
                r"учту ваше предпочтение (.+?)(?:\.|$)"
            ]
            
            for pattern in confirmation_patterns:
                matches = re.findall(pattern, ai_response.lower())
                for match in matches:
                    if len(match) > 10:
                        self.add_user_preference(
                            match.strip(),
                            category="confirmed_preference",
                            metadata={"source": "ai_confirmation", "pattern": pattern}
                        )
                        
        except Exception as e:
            logger.error(f"❌ Ошибка извлечения предпочтений: {e}")

    ### МЕТОДЫ ДЛЯ РАБОТЫ С CHROMADB ###
    
    def add_to_memory(self, user_message: str, ai_response: str, context: str = "", 
                     metadata: Optional[Dict[str, Any]] = None) -> bool:
        self._ensure_chromadb_initialized()
        """
        Добавляет диалог в векторное хранилище памяти
        
        Args:
            user_message: Сообщение пользователя
            ai_response: Ответ ИИ
            context: Дополнительный контекст
            metadata: Дополнительные метаданные
            
        Returns:
            True если успешно добавлено
        """
        self._ensure_chromadb_initialized()
        if self.chromadb_manager is None:
            logger.error("❌ ChromaDB недоступен")
            return False
        try:
            return self.chromadb_manager.add_conversation_memory(
                user_message, ai_response, context, metadata
            )
        except Exception as e:
            logger.error(f"❌ Ошибка добавления в память: {e}")
            return False
    
    def add_user_preference(self, preference_text: str, category: str = "general", 
                           metadata: Optional[Dict[str, Any]] = None) -> bool:
        """
        Добавляет предпочтение пользователя в векторное хранилище
        
        Args:
            preference_text: Текст предпочтения
            category: Категория предпочтения
            metadata: Дополнительные метаданные
            
        Returns:
            True если успешно добавлено
        """
        self._ensure_chromadb_initialized()
        if self.chromadb_manager is None:
            logger.error("❌ ChromaDB недоступен")
            return False
        try:
            return self.chromadb_manager.add_user_preference(
                preference_text, category, metadata
            )
        except Exception as e:
            logger.error(f"❌ Ошибка добавления предпочтения: {e}")
            return False
    
    def get_relevant_context(self, query: str, max_context_length: int = 2000) -> str:
        """
        Получает релевантный контекст из предыдущих диалогов
        
        Args:
            query: Текущий запрос пользователя
            max_context_length: Максимальная длина контекста
            
        Returns:
            Строка с релевантным контекстом
        """
        self._ensure_chromadb_initialized()
        if self.chromadb_manager is None:
            logger.error("❌ ChromaDB недоступен")
            return ""
        try:
            return self.chromadb_manager.get_conversation_context(query, max_context_length)
        except Exception as e:
            logger.error(f"❌ Ошибка получения контекста: {e}")
            return ""
    
    def get_user_preferences(self, query: Optional[str] = None) -> str:
        """
        Получает предпочтения пользователя
        
        Args:
            query: Контекстный запрос (опционально)
            
        Returns:
            Строка с предпочтениями пользователя
        """
        self._ensure_chromadb_initialized()
        if self.chromadb_manager is None:
            logger.error("❌ ChromaDB недоступен")
            return ""
        try:
            return self.chromadb_manager.get_user_preferences_summary(query)
        except Exception as e:
            logger.error(f"❌ Ошибка получения предпочтений: {e}")
            return ""
    
    def search_similar_conversations(self, query: str, n_results: int = 5, 
                                   similarity_threshold: float = 0.7) -> List[Dict[str, Any]]:
        """
        Ищет похожие диалоги в векторном хранилище
        
        Args:
            query: Поисковый запрос
            n_results: Количество результатов
            similarity_threshold: Порог схожести (0-1)
            
        Returns:
            Список найденных диалогов
        """
        self._ensure_chromadb_initialized()
        if self.chromadb_manager is None:
            logger.error("❌ ChromaDB недоступен")
            return []
        try:
            return self.chromadb_manager.search_similar_conversations(
                query, n_results, similarity_threshold
            )
        except Exception as e:
            logger.error(f"❌ Ошибка поиска похожих диалогов: {e}")
            return []
    
    def cleanup_old_memory(self, days_to_keep: int = 30) -> int:
        """
        Удаляет старые записи из памяти
        
        Args:
            days_to_keep: Количество дней для хранения записей
            
        Returns:
            Количество удаленных записей
        """
        self._ensure_chromadb_initialized()
        if self.chromadb_manager is None:
            logger.error("❌ ChromaDB недоступен")
            return 0
        try:
            return self.chromadb_manager.cleanup_old_records(days_to_keep)
        except Exception as e:
            logger.error(f"❌ Ошибка очистки памяти: {e}")
            return 0
    
    def get_memory_stats(self) -> Dict[str, Any]:
        """
        Получает статистику памяти
        
        Returns:
            Словарь со статистикой
        """
        self._ensure_chromadb_initialized()
        if self.chromadb_manager is None:
            logger.error("❌ ChromaDB недоступен")
            return {"error": "ChromaDB недоступен"}
        try:
            return self.chromadb_manager.get_database_stats()
        except Exception as e:
            logger.error(f"❌ Ошибка получения статистики памяти: {e}")
            return {"error": str(e)}
    
    def get_gpu_info(self) -> Dict[str, Any]:
        """
        Получает информацию о GPU для ChromaDB
        
        Returns:
            Словарь с информацией о GPU
        """
        try:
            self._ensure_chromadb_initialized()
            if self.chromadb_manager is None:
                logger.error("❌ ChromaDB недоступен")
                return {"error": "ChromaDB недоступен"}
            return self.chromadb_manager.get_gpu_info()
        except Exception as e:
            logger.error(f"❌ Ошибка получения информации о GPU: {e}")
            return {"error": str(e)}

def ensure_wav(audio_path: str) -> Optional[str]:
    """Конвертирует аудиофайл в WAV формат если он не WAV"""
    try:
        if audio_path.lower().endswith('.wav'):
            return audio_path
        
        # Создаем временный файл
        temp_dir = os.path.join(os.path.dirname(__file__), "Audio", "temp_convert")
        os.makedirs(temp_dir, exist_ok=True)
        wav_path = os.path.join(temp_dir, f"converted_{int(time.time())}.wav")
        
        # Конвертируем через ffmpeg
        cmd = [
            'ffmpeg', '-i', audio_path, '-acodec', 'pcm_s16le', 
            '-ar', '16000', '-ac', '1', wav_path, '-y'
        ]
        subprocess.run(cmd, check=True, capture_output=True)
        
        return wav_path
    except Exception as e:
        logger.error(f"Ошибка конвертации в WAV: {e}")
        return audio_path

def main():
    """Главная функция"""
    parser = argparse.ArgumentParser(description='AI PowerShell Оркестратор')
    parser.add_argument('--web', action='store_true', help='Запустить веб-интерфейс')
    parser.add_argument('--test-startup', action='store_true', help='Тестировать только инициализацию системы')
    args = parser.parse_args()
    
    start_web = args.web
    test_startup = args.test_startup
    
    # Если запущен тест инициализации - выполняем его и выходим
    if test_startup:
        test_startup_initialization()
        return
    
    # Настройка логирования для внешних библиотек при веб-интерфейсе
    if not start_web:
        logging.getLogger('httpx').setLevel(logging.WARNING)
        logging.getLogger('telegram').setLevel(logging.WARNING)
        logging.getLogger('telegram.ext').setLevel(logging.WARNING)
    
    logger.info("Настройка AI PowerShell Оркестратора")
    logger.info("="*50)
    
    # Настройки (можно вынести в конфиг файл)
    LM_STUDIO_URL = "http://localhost:1234"  # URL вашего LM Studio сервера
    GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY", "").strip()  # Ваш Google API ключ
    GOOGLE_CSE_ID = os.getenv("GOOGLE_CSE_ID", "").strip()   # Ваш Google CSE ID
    
    # Telegram Bot настройки
    TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN", "").strip()  # Введите токен вашего бота
    TELEGRAM_ALLOWED_USER_ID = os.getenv("TELEGRAM_ALLOWED_USER_ID", "").strip()  # ID пользователя, которому разрешено использовать бота

    # --- Автоматическое управление инструментами ---
    # Все инструменты по умолчанию выключены для экономии ресурсов
    # Они будут автоматически включаться при необходимости
    use_image_generation = False  # Включается автоматически при генерации изображений
    use_vision = False           # Включается автоматически при анализе изображений
    use_audio = False            # Включается автоматически при обработке аудио
    use_separator = True         # Всегда включен при использовании Whisper (как вы просили)
    use_ocr = False              # Включается автоматически при извлечении текста из изображений

    # Мозг по умолчанию - используем указанную вами модель
    brain_model = "J:/models-LM Studio/mradermacher/Huihui-Qwen3-4B-Thinking-2507-abliterated-GGUF/Huihui-Qwen3-4B-Thinking-2507-abliterated.Q4_K_S.gguf"
    logger.info(f"🧠 Используется модель мозга: {os.path.basename(brain_model)}")
    logger.info("🔧 Инструменты будут автоматически включаться по требованию для экономии ресурсов")

    # Пути к моделям (можно вынести в конфиг)
    vision_model = "moondream2-llamafile"  # Имя vision-модели всегда фиксировано
    whisper_model = "ggerganov/whisper-large-v3-GGUF"

    # Проверяем и запускаем нужные модели
    orchestrator = AIOrchestrator(
        lm_studio_url=LM_STUDIO_URL,
        google_api_key=GOOGLE_API_KEY,
        google_cse_id=GOOGLE_CSE_ID
    )

    # Проверяем, заданы ли Google API настройки
    if not GOOGLE_API_KEY or not GOOGLE_CSE_ID:
        logger.warning("⚠️  ВНИМАНИЕ: Google API ключ или CSE ID не настроены!")
        logger.info("   Поиск в интернете будет недоступен.")
        logger.info("   Для настройки отредактируйте переменные в начале main()")
        logger.info("")

    # Передаем brain_model и настройки в оркестратор
    orchestrator.brain_model = brain_model
    orchestrator.use_separator = use_separator
    orchestrator.use_image_generation = use_image_generation
    orchestrator.use_vision = use_vision
    orchestrator.use_audio = use_audio
    orchestrator.use_ocr = use_ocr
    
    # Передаем настройки Telegram
    orchestrator.telegram_bot_token = TELEGRAM_BOT_TOKEN
    orchestrator.telegram_allowed_user_id = TELEGRAM_ALLOWED_USER_ID
    
    # Запускаем веб-интерфейс если указан флаг --web
    if start_web:
        try:
            # Отключаем локальный показ изображений во всплывающих окнах при веб-режиме
            orchestrator.show_images_locally = False
        except Exception:
            pass
        # Запускаем uvicorn сервер в фоне
        try:
            # subprocess, sys, os уже импортированы в начале файла
            repo_root = os.path.dirname(os.path.abspath(__file__))
            cmd = [
                _sys.executable, "-m", "uvicorn", "webui.server:app",
                "--host", "127.0.0.1", "--port", "8001", "--app-dir", repo_root
            ]
            logger.info(f"🌐 Стартую веб-сервер: {' '.join(cmd)}")
            subprocess.Popen(cmd, cwd=repo_root)
            logger.info("Откройте в браузере: http://127.0.0.1:8001/")
        except Exception as e:
            logger.warning(f"⚠️ Не удалось запустить веб-интерфейс автоматически: {e}")

    # Запускаем Telegram бота если указан токен
    if TELEGRAM_BOT_TOKEN:
        try:
            if start_web:
                logger.info("🤖 Запускаю Telegram бота...")
            tg_started = orchestrator.start_telegram_bot()
            if start_web:
                if tg_started:
                    logger.info("✅ Telegram бот запущен")
                else:
                    logger.info("ℹ️ Telegram бот не запущен (проверьте токен)")
        except Exception as e:
            if start_web:
                logger.error(f"❌ Ошибка запуска Telegram бота: {e}")
            else:
                # В веб-режиме логируем тихо
                logger.debug(f"Telegram bot error: {e}")
    
    def __del__(self):
        """Деструктор для очистки ресурсов плагинов"""
        try:
            if hasattr(self, 'plugin_manager') and self.plugin_manager:
                # Выгружаем все плагины
                for plugin_name in list(self.plugin_manager.loaded_plugins.keys()):
                    self.plugin_manager.unload_plugin(plugin_name)
                logger.info("🔌 Плагины очищены")
        except Exception as e:
            logger.error(f"Ошибка очистки плагинов: {e}")
    
    # Запускаем интерактивный режим
    orchestrator.run_interactive()


def test_startup_initialization():
    """Тестирует инициализацию всех компонентов системы"""
    print("\n" + "="*60)
    print("🧪 ТЕСТ ИНИЦИАЛИЗАЦИИ AI ORCHESTRATOR")
    print("="*60)
    
    total_start_time = time.time()
    
    # Инициализация компонентов
    component_times = {}
    
    # 1. Основной оркестратор
    print("\n📦 Инициализация основного оркестратора...")
    start_time = time.time()
    
    LM_STUDIO_URL = "http://localhost:1234"
    GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY", "").strip()
    GOOGLE_CSE_ID = os.getenv("GOOGLE_CSE_ID", "").strip()
    
    try:
        orchestrator = AIOrchestrator(
            lm_studio_url=LM_STUDIO_URL,
            google_api_key=GOOGLE_API_KEY,
            google_cse_id=GOOGLE_CSE_ID
        )
        component_times["orchestrator"] = time.time() - start_time
        print(f"   ✅ Основной оркестратор: {component_times['orchestrator']:.2f}с")
    except Exception as e:
        component_times["orchestrator"] = time.time() - start_time
        print(f"   ❌ Основной оркестратор: {component_times['orchestrator']:.2f}с - {e}")
        return
    
    # 2. Тестируем ChromaDB
    print("\n🗃️ Тестирование ChromaDB...")
    start_time = time.time()
    
    try:
        # Инициализация ChromaDB через оркестратор
        orchestrator._ensure_chromadb_initialized()
        component_times["chromadb"] = time.time() - start_time
        print(f"   ✅ ChromaDB: {component_times['chromadb']:.2f}с")
        
        # Проверяем работу ChromaDB
        test_memory = orchestrator.add_to_memory(
            "Тестовое сообщение", "Тестовый ответ", "Контекст теста"
        )
        if test_memory:
            print("   ✅ ChromaDB функциональность: OK")
        else:
            print("   ⚠️ ChromaDB функциональность: Ошибка")
            
        component_times["chromadb"] = time.time() - start_time
    except Exception as e:
        component_times["chromadb"] = time.time() - start_time
        print(f"   ❌ ChromaDB: {component_times['chromadb']:.2f}с - {e}")
    
    # 3. Тестируем EasyOCR
    print("\n👁️ Тестирование EasyOCR...")
    start_time = time.time()
    
    try:
        # Сначала проверим, доступен ли EasyOCR как модуль
        try:
            import easyocr  # type: ignore
            easyocr_available = True
        except ImportError:
            easyocr_available = False
        
        if not easyocr_available:
            component_times["easyocr"] = time.time() - start_time
            print(f"   ❌ EasyOCR: {component_times['easyocr']:.2f}с - Модуль не установлен")
            print("   💡 Установите: pip install easyocr")
        elif orchestrator._ensure_ocr_initialized():
            component_times["easyocr"] = time.time() - start_time
            print(f"   ✅ EasyOCR: {component_times['easyocr']:.2f}с")
            
            # Проверим, что OCR reader действительно создан
            if orchestrator.ocr_reader is not None:
                print("   ✅ EasyOCR функциональность: OK")
            else:
                print("   ⚠️ EasyOCR функциональность: Reader не создан")
        else:
            component_times["easyocr"] = time.time() - start_time
            print(f"   ❌ EasyOCR: {component_times['easyocr']:.2f}с - Ошибка инициализации")
    except Exception as e:
        component_times["easyocr"] = time.time() - start_time
        print(f"   ❌ EasyOCR: {component_times['easyocr']:.2f}с - {e}")
    
    # 4. Тестируем модель мозга
    print("\n🧠 Тестирование модели мозга...")
    start_time = time.time()
    
    try:
        brain_model = "J:/models-LM Studio/mradermacher/Huihui-Qwen3-4B-Thinking-2507-abliterated-GGUF/Huihui-Qwen3-4B-Thinking-2507-abliterated.Q4_K_S.gguf"
        orchestrator.brain_model = brain_model
        
        # Проверяем доступность LM Studio
        response = requests.get(f"{LM_STUDIO_URL}/v1/models", timeout=10)
        if response.status_code == 200:
            models = response.json().get("data", [])
            print(f"   📊 Всего моделей в LM Studio: {len(models)}")
            
            # Ищем любые модели, не только загруженные
            loaded_models = [m for m in models if m.get("isLoaded", False)]
            available_models = [m.get("id", "unknown") for m in models]
            
            print(f"   📊 Доступные модели: {available_models}")
            print(f"   📊 Загруженных моделей: {len(loaded_models)}")
            
            if models:  # Если есть любые модели
                component_times["brain_model"] = time.time() - start_time
                print(f"   ✅ Модель мозга: {component_times['brain_model']:.2f}с")
                
                # Тестируем запрос к модели (даже если модель не показывается как загруженная)
                test_response = orchestrator.call_brain_model("Привет! Это тест.")
                if test_response and not test_response.startswith("[Brain error]"):
                    print("   ✅ Тестовый запрос: OK")
                    print(f"   📝 Ответ модели: {test_response[:100]}...")
                else:
                    print(f"   ⚠️ Тестовый запрос: {test_response}")
            else:
                component_times["brain_model"] = time.time() - start_time
                print(f"   ⚠️ Модель мозга: {component_times['brain_model']:.2f}с - Нет моделей в LM Studio")
        else:
            component_times["brain_model"] = time.time() - start_time
            print(f"   ❌ Модель мозга: {component_times['brain_model']:.2f}с - LM Studio недоступен")
    except Exception as e:
        component_times["brain_model"] = time.time() - start_time
        print(f"   ❌ Модель мозга: {component_times['brain_model']:.2f}с - {e}")
    
    # 5. Проверяем фоновый загрузчик
    print("\n🚀 Состояние фонового загрузчика...")
    try:
        loader = get_background_loader()
        loaded = list(loader.loaded_components.keys())
        loading_tasks = list(loader.loading_tasks.keys())
        
        # Показываем только те компоненты, которые еще загружаются
        still_loading = [task for task in loading_tasks if task not in loaded]
        
        print(f"   � Загруженные компоненты: {loaded}")
        if still_loading:
            print(f"   🔄 Еще загружаются: {still_loading}")
        else:
            print(f"   ✅ Все компоненты загружены")
    except Exception as e:
        print(f"   ❌ Фоновый загрузчик: {e}")
    
    # 6. Проверяем плагины
    print("\n🔌 Проверка системы плагинов...")
    try:
        if orchestrator.plugin_manager:
            # Просто проверяем наличие плагинов без обращения к конкретному атрибуту
            print(f"   ✅ Система плагинов: Инициализирована")
        else:
            print("   ⚠️ Система плагинов: Не инициализирована")
    except Exception as e:
        print(f"   ❌ Система плагинов: {e}")
    
    # Итоговая статистика
    total_time = time.time() - total_start_time
    print("\n" + "="*60)
    print("📊 ИТОГОВАЯ СТАТИСТИКА ИНИЦИАЛИЗАЦИИ")
    print("="*60)
    
    # Показываем время каждого компонента с процентом от общего времени
    for component, duration in component_times.items():
        percentage = (duration / total_time * 100) if total_time > 0 else 0
        status = "✅" if duration < 30 else "⚠️" if duration < 60 else "❌"
        print(f"{status} {component:20}: {duration:6.2f}с ({percentage:5.1f}%)")
    
    print(f"\n🕐 Общее время инициализации: {total_time:.2f}с")
    
    if total_time < 10:
        print("🚀 Отлично! Быстрая инициализация")
    elif total_time < 30:
        print("✅ Хорошо! Приемлемое время инициализации")
    elif total_time < 60:
        print("⚠️ Медленно! Требуется оптимизация")
    else:
        print("❌ Очень медленно! Критические проблемы производительности")
    
    print("\n✅ Тест инициализации завершен")


if __name__ == "__main__":
    main()