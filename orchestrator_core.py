#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
AI PowerShell Orchestrator with Google Search Integration
Интегрирует LM Studio/llama.cpp, PowerShell команды и поиск Google

"""

# ============================================================================
# КОНФИГУРАЦИЯ GPU - ВАЖНО! Выполняется ДО импорта библиотек
# ============================================================================
import os
import copy

from config import (
    GPU_CONFIG,
    USE_LLAMA_CPP,
    LLAMA_CPP_MODEL_PATH,
    LLAMA_KV_Q8_DEFAULT,
    LLAMA_CPP_PARAMS as CONFIG_LLAMA_CPP_PARAMS,
    LLAMA_CPP_GENERATION_PARAMS as CONFIG_LLAMA_CPP_GENERATION_PARAMS,
    VISION_MODEL_ID,
    VISION_MODEL_LOAD_ARGS as CONFIG_VISION_MODEL_LOAD_ARGS,
    VISION_GENERATION_PARAMS as CONFIG_VISION_GENERATION_PARAMS,
    VISION_PROMPT_FILENAME,
    VISION_FALLBACK_PROMPT,
    CHROMA_DB_PATH,
    CHROMADB_DEFAULT_COLLECTION_NAME,
    CHROMADB_DEFAULT_COLLECTION_METADATA,
    CHROMADB_BACKGROUND_COLLECTION_NAME,
    CHROMADB_EMBEDDING_MODEL,
    CHROMADB_USE_GPU_BY_DEFAULT,
    CHROMADB_ENABLE_MEMORY,
    DEFAULT_MAX_CONTEXT_LENGTH,
    DEFAULT_SAFE_CONTEXT_LENGTH,
    DEFAULT_MAX_RETRIES,
    AUTO_DISABLE_DELAY_SECONDS,
    DEFAULT_SIMILARITY_THRESHOLD,
    PROMPTS_DIR_NAME,
    PLUGINS_DIR_NAME,
    DEFAULT_LM_STUDIO_URL,
    OUTPUT_DIR_NAME,
    LOG_FILE_NAME,
    FILE_LOG_FORMAT,
    CONSOLE_LOG_LEVEL,
    OCR_AVAILABLE_DEFAULT,
    CHROMADB_AVAILABLE_DEFAULT,
    TORCH_AVAILABLE_DEFAULT,
)

from logging_setup import setup_logging
from prompts import PromptLoader
from llm_services import (
    LlamaCppWrapper,
    LLAMA_CPP_AVAILABLE,
    is_model_running_lm_studio,
    get_model_context_info_lm_studio,
    load_model_lm_studio,
    unload_model_lm_studio,
    ask_qwen_for_prompt,
)
from chromadb_manager import ChromaDBManager, load_chromadb
from image_generation import ModelManager, ImageGenerator
from media_processing import (
    image_to_base64_balanced,
    extract_video_frames,
    get_youtube_cookies_path,
    check_cookies_validity,
    suggest_cookies_update,
    download_youtube_video,
    ensure_wav,
    check_vpn_status,
    get_youtube_info,
    check_youtube_accessibility,
    download_youtube_audio,
    convert_audio_to_wav,
    check_whisper_setup,
    download_whisper_model,
)
from automation import (
    execute_powershell,
    move_mouse,
    left_click,
    right_click,
    scroll,
    mouse_down,
    mouse_up,
    drag_and_drop,
    type_text,
)
from communications import EmailManager, TelegramBotManager
from resource_manager import (
    BackgroundInitializer,
    get_background_loader,
    load_easyocr,
    load_torch,
)
from ui_automation_utils import get_ui_tree_as_text

os.environ['CUDA_VISIBLE_DEVICES'] = GPU_CONFIG.get('cuda_visible_devices', '0')
print(f"🎮 CUDA устройство: {os.environ.get('CUDA_VISIBLE_DEVICES', 'auto')}")
print(f"🚀 Форсируем использование GPU: {GPU_CONFIG.get('force_gpu_message', 'GPU')}")
# Уменьшаем подробный вывод llama.cpp
os.environ.setdefault('LLAMA_LOG_LEVEL', GPU_CONFIG.get('llama_log_level', '40'))
# ============================================================================

LLAMA_CPP_PARAMS = copy.deepcopy(CONFIG_LLAMA_CPP_PARAMS)
LLAMA_CPP_GENERATION_PARAMS = copy.deepcopy(CONFIG_LLAMA_CPP_GENERATION_PARAMS)
VISION_MODEL_LOAD_ARGS = copy.deepcopy(CONFIG_VISION_MODEL_LOAD_ARGS)
VISION_GENERATION_PARAMS = copy.deepcopy(CONFIG_VISION_GENERATION_PARAMS)
LLAMA_KV_Q8 = LLAMA_KV_Q8_DEFAULT

# ============================================================================
# ПЕРЕКЛЮЧАТЕЛЬ БЭКЕНДА и параметры моделей теперь берутся из config.py
# ============================================================================

# Подавляем предупреждения PyTorch для чистого запуска
import warnings
warnings.filterwarnings("ignore", message="expandable_segments not supported on this platform")
warnings.filterwarnings("ignore", message="There is an imbalance between your GPUs")

import requests
import subprocess
import json
import time
import sys as _sys
import os

# Фикс кодировки для Windows терминала (поддержка UTF-8 и эмодзи)
if _sys.platform == 'win32':
    try:
        import codecs
        _sys.stdout = codecs.getwriter('utf-8')(_sys.stdout.buffer, 'strict')
        _sys.stderr = codecs.getwriter('utf-8')(_sys.stderr.buffer, 'strict')
        _sys.stdin = codecs.getreader('utf-8')(_sys.stdin.buffer, 'strict')
    except:
        pass

import base64
import io
import re
import threading
import tempfile
import shutil
import queue
import asyncio
from pathlib import Path
import math
import pyautogui
import mss
import queue
import logging
import argparse
from typing import Dict, Any, List, Union, Optional, TYPE_CHECKING, Tuple, Iterable, cast
import urllib.parse
from PIL import Image
from io import BytesIO
from collections import defaultdict
import asyncio
import telegram
from telegram import Update, Bot
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes
from dotenv import load_dotenv
import concurrent.futures

# На Windows принудительно используем SelectorEventLoopPolicy, чтобы не ломался stdin
if _sys.platform == 'win32':
    try:
        asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())  # type: ignore[attr-defined]
    except Exception as _policy_err:
        print(f"⚠️ Не удалось установить WindowsSelectorEventLoopPolicy: {_policy_err}")

# Импорты для работы с электронной почтой
import imaplib
import smtplib
import email
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from email.mime.base import MIMEBase
from email import encoders
from email.header import decode_header
from email.utils import parseaddr, formataddr
import email.utils
import mimetypes

# Импорт llama-cpp-python (опционально)
try:
    from llama_cpp import Llama
    import llama_cpp as llama_cpp_lib  # type: ignore
    LLAMA_CPP_AVAILABLE = True
except ImportError:
    LLAMA_CPP_AVAILABLE = False
    Llama = None
    llama_cpp_lib = None
    print("⚠️ llama-cpp-python не установлен. Используйте только LM Studio режим или установите: pip install llama-cpp-python")

# Константа для квантизации KV-кэша в Q8 (используем официальное значение, если доступно)
LLAMA_KV_Q8 = getattr(llama_cpp_lib, "GGML_TYPE_Q8_0", LLAMA_KV_Q8) if llama_cpp_lib else LLAMA_KV_Q8
if 'type_k' in LLAMA_CPP_PARAMS:
    LLAMA_CPP_PARAMS['type_k'] = LLAMA_KV_Q8
if 'type_v' in LLAMA_CPP_PARAMS:
    LLAMA_CPP_PARAMS['type_v'] = LLAMA_KV_Q8

# BackgroundInitializer moved to resource_manager.py

# Система плагинов
# Plugin system
try:
    from plugins import PluginManager, PluginError
    PLUGINS_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Система плагинов недоступна: {e}")
    PLUGINS_AVAILABLE = False
    PluginManager = None
    PluginError = Exception

# Помощь статическим анализаторам: явные объявления для опциональных внешних символов
from typing import Any as _Any
chromadb: _Any = None
Settings: _Any = None
SentenceTransformer: _Any = None
torch: _Any = None
_imageio: _Any = None
_pygame: _Any = None

# Импорты для работы с документами
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

try:
    import pandas as pd
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    pd = None
    openpyxl = None

# Импорты для работы с PDF
try:
    import PyPDF2  # type: ignore
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    PyPDF2 = None

# Импорты для генерации файлов
try:
    from reportlab.pdfgen import canvas  # type: ignore
    from reportlab.lib.pagesizes import letter, A4  # type: ignore
    from reportlab.lib.styles import getSampleStyleSheet  # type: ignore
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer  # type: ignore
    from reportlab.lib.units import inch  # type: ignore
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    canvas = None
    _letter = None
    _A4 = None
    getSampleStyleSheet = None
    SimpleDocTemplate = None
    Paragraph = None
    Spacer = None
    _inch = None

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    markdown = None

# Импорты для OCR - теперь ленивые
OCR_AVAILABLE = OCR_AVAILABLE_DEFAULT  # Будем проверять при первом использовании

# Импорты для ChromaDB и векторного поиска - теперь ленивые
CHROMADB_AVAILABLE = CHROMADB_AVAILABLE_DEFAULT  # Будем проверять при первом использовании

# Импорты для Torch - теперь ленивые
TORCH_AVAILABLE = TORCH_AVAILABLE_DEFAULT  # Будем проверять при первом использовании

# Проверки доступности опциональных модулей
try:
    import torch as _torch
    TORCH_AVAILABLE = True
except Exception:
    TORCH_AVAILABLE = False

try:
    import diffusers as _diffusers
    DIFFUSERS_AVAILABLE = True
except Exception:
    DIFFUSERS_AVAILABLE = False

try:
    import imageio as _imageio
    IMAGEIO_AVAILABLE = True
except Exception:
    IMAGEIO_AVAILABLE = False

try:
    import pygame as _pygame
    PYGAME_AVAILABLE = True
except Exception:
    PYGAME_AVAILABLE = False

# Загружаем переменные окружения из .env файла
# override=True - перезаписываем существующие переменные
load_dotenv(override=True)

# Determine if running in web mode (show verbose console logs)
IS_WEB = any(arg == '--web' for arg in _sys.argv)

setup_logging(IS_WEB, LOG_FILE_NAME, FILE_LOG_FORMAT, CONSOLE_LOG_LEVEL)

logger = logging.getLogger(__name__)
if TYPE_CHECKING:
    try:
        import chromadb  # type: ignore
        from chromadb.config import Settings  # type: ignore
    except Exception:
        pass
    # Stubs for optional heavy libraries so Pylance doesn't warn
    try:
        import imageio as _imageio  # type: ignore
    except Exception:
        pass
    try:
        import pygame as _pygame  # type: ignore
    except Exception:
        pass
    try:
        import diffusers as _diffusers  # type: ignore
    except Exception:
        pass
    try:
        import torch as _torch  # type: ignore
    except Exception:
        pass

# image_to_base64_balanced moved to media_processing.py
# ModelManager moved to image_generation.py


class AIOrchestrator:
    # Video/audio methods moved to media_processing.py as standalone functions
    
    def extract_video_frames(self, video_path: str, fps: int = 1) -> list:
        """Wrapper для backwards compatibility."""
        return extract_video_frames(video_path, fps, self.logger)
    def download_youtube_video(self, url: str, out_dir: Optional[str] = None) -> Optional[str]:
        """Wrapper для backwards compatibility."""
        return download_youtube_video(url, out_dir, self.logger)

    def check_vpn_status(self) -> bool:
        """Wrapper для backwards compatibility."""
        return check_vpn_status(self.logger)

    def get_youtube_info(self, url: str) -> dict:
        """Wrapper для backwards compatibility."""
        return get_youtube_info(url, self.logger)

    def check_youtube_accessibility(self, url: str) -> bool:
        """Wrapper для backwards compatibility."""
        return check_youtube_accessibility(url, self.logger)

    def _auto_load_brain_model(self):
        """Автоматически загружает модель мозга при инициализации"""
        try:
            # Режим llama.cpp - прямая загрузка модели
            if self.use_llama_cpp:
                self.logger.info("🔧 Режим: llama-cpp-python (прямое управление)")
                
                if not LLAMA_CPP_AVAILABLE:
                    self.logger.error("❌ llama-cpp-python не установлен!")
                    self.logger.error("💡 Установите: pip install llama-cpp-python")
                    self.logger.error("🔄 Или установите из wheel файла в корне проекта")
                    return
                
                # Создаем обертку llama.cpp
                self.llama_wrapper = LlamaCppWrapper(
                    model_path=self.brain_model,
                    params=LLAMA_CPP_PARAMS,
                    logger_instance=self.logger
                )
                
                # Загружаем модель
                if self.llama_wrapper.load_model():
                    self.brain_model_id = self.llama_wrapper.model_id
                    self.logger.info(f"✅ Модель загружена через llama.cpp: {self.brain_model_id}")
                else:
                    self.logger.error("❌ Не удалось загрузить модель через llama.cpp")
                return
            
            # Режим LM Studio - подключение через HTTP API
            self.logger.info("🔧 Режим: LM Studio (HTTP API)")
            
            # Проверяем, запущена ли модель через прямой запрос к API
            try:
                resp = requests.get(f"{self.lm_studio_url}/v1/models", timeout=10)
                if resp.status_code == 200:
                    data = resp.json()
                    model_loaded = False
                    for m in data.get("data", []):
                        if self.brain_model in m.get("id", "") and m.get("isLoaded", False):
                            model_loaded = True
                            # Сохраняем короткий ID модели для API вызовов
                            self.brain_model_id = m.get("id")
                            self.logger.info(f"✅ Модель мозга уже загружена: {os.path.basename(self.brain_model)} (ID: {self.brain_model_id})")
                            return
                else:
                    self.logger.warning(f"⚠️ Не удалось проверить статус моделей: {resp.status_code}")
            except Exception as e:
                self.logger.warning(f"⚠️ Ошибка проверки статуса моделей: {e}")
            
            # If model not loaded, try to load it
            self.logger.info(f"🧠 Auto-loading brain model: {os.path.basename(self.brain_model)}")
            
            if load_model_lm_studio(self.lm_studio_url, self.brain_model, self.logger):
                self._update_brain_model_id()
            else:
                self.logger.warning("API load failed, trying LM Studio launch...")
                self.launch_model(self.brain_model)
                self._update_brain_model_id()
                
        except Exception as e:
            self.logger.error(f"❌ Ошибка автозагрузки модели мозга: {e}")
    
    def _ensure_vision_model_loaded(self):
        """Гарантирует, что vision-модель в LM Studio загружена перед вызовом."""
        if self.use_llama_cpp:
            # Для режима llama.cpp потребуется отдельный путь загрузки vision моделей
            return

        if self._vision_model_ready or not self.vision_model_id:
            return

        try:
            resp = requests.get(f"{self.lm_studio_url}/v1/models", timeout=10)
            if resp.status_code == 200:
                data = resp.json()
                for model in data.get("data", []):
                    if self.vision_model_id in model.get("id", "") and model.get("isLoaded", False):
                        self._vision_model_ready = True
                        self.logger.info(f"👁️ Vision модель уже загружена: {self.vision_model_id}")
                        return
        except Exception as exc:
            self.logger.debug(f"⚠️ Не удалось проверить статус vision модели: {exc}")

        payload: Dict[str, Any] = {"model": self.vision_model_id, "load": True}
        if self.vision_model_load_args:
            payload["args"] = self.vision_model_load_args

        try:
            self.logger.info(f"👁️ Загружаю vision модель: {self.vision_model_id}")
            resp = requests.post(
                f"{self.lm_studio_url}/v1/models/load",
                json=payload,
                timeout=60
            )
            if resp.status_code == 200:
                self._vision_model_ready = True
                self.logger.info("✅ Vision модель загружена")
            else:
                self.logger.warning(f"⚠️ Не удалось загрузить vision модель ({resp.status_code}): {resp.text}")
        except Exception as exc:
            self.logger.warning(f"⚠️ Ошибка загрузки vision модели: {exc}")

    def _unload_vision_model(self):
        """Unload vision model from LM Studio when idle."""
        if self.use_llama_cpp or not self._vision_model_ready or not self.vision_model_id:
            return

        if unload_model_lm_studio(self.lm_studio_url, self.vision_model_id, self.logger):
            self.logger.info("Vision model unloaded")
        self._vision_model_ready = False

    def _update_brain_model_id(self):
        """Обновляет короткий ID модели мозга из API"""
        try:
            resp = requests.get(f"{self.lm_studio_url}/v1/models", timeout=10)
            if resp.status_code == 200:
                data = resp.json()
                for m in data.get("data", []):
                    if self.brain_model in m.get("id", "") and m.get("isLoaded", False):
                        self.brain_model_id = m.get("id")
                        self.logger.info(f"✅ Обновлен ID модели мозга: {self.brain_model_id}")
                        return
                self.logger.warning("⚠️ Не удалось найти загруженную модель для получения ID")
            else:
                self.logger.warning(f"⚠️ Не удалось получить список моделей для обновления ID: {resp.status_code}")
        except Exception as e:
            self.logger.warning(f"⚠️ Ошибка обновления ID модели мозга: {e}")
    
    def _start_background_loading(self):
        """Запускает фоновую загрузку тяжелых компонентов"""
        loader = get_background_loader()
        
        # Запускаем загрузку EasyOCR в фоне
        loader.start_loading('easyocr', load_easyocr)
        
        # Запускаем загрузку PyTorch в фоне (если нужно)
        loader.start_loading('torch', load_torch)
        
        self.logger.info("🚀 Запущена фоновая загрузка компонентов")
    
    def _ensure_ocr_initialized(self):
        """Обеспечивает инициализацию OCR"""
        if self.ocr_reader is not None:
            return True
            
        loader = get_background_loader()
        ocr_reader = loader.get_component('easyocr', timeout=30)
        
        # Проверяем, что reader не None (успешно загружен)
        if ocr_reader is not None:
            self.ocr_reader = ocr_reader
            self.logger.info("✅ EasyOCR загружен из фонового потока")
            return True
        else:
            # Fallback к синхронной инициализации
            return self._initialize_ocr_sync()
    
    def _initialize_ocr_sync(self):
        """Синхронная инициализация OCR как fallback"""
        try:
            self.logger.info("📖 Синхронная инициализация EasyOCR...")
            import easyocr  # type: ignore
            self.ocr_reader = easyocr.Reader(['ru', 'en'])
            self.logger.info("✅ EasyOCR инициализирован синхронно")
            return True
        except Exception as e:
            self.logger.error(f"❌ Ошибка синхронной инициализации EasyOCR: {e}")
            return False
    
    def _reconnect_brain_model(self):
        """Переподключается к модели мозга, если соединение потеряно"""
        try:
            self.logger.info("🔄 Переподключение к модели мозга...")
            
            # Режим llama.cpp
            if self.use_llama_cpp:
                if self.llama_wrapper:
                    return self.llama_wrapper.reconnect()
                return False
            
            # Режим LM Studio
            # Сначала пытаемся загрузить модель
            self._auto_load_brain_model()
            
            # Ждем немного для загрузки
            time.sleep(3)
            
            # Проверяем, что модель доступна
            response = requests.get(f"{self.lm_studio_url}/v1/models", timeout=10)
            if response.status_code == 200:
                models = response.json().get("data", [])
                if any(self.brain_model in m.get("id", "") for m in models):
                    self.logger.info("✅ Переподключение к модели мозга успешно")
                    return True
            
            self.logger.warning("⚠️ Модель мозга недоступна после переподключения")
            return False
            
        except Exception as e:
            self.logger.error(f"❌ Ошибка переподключения к модели мозга: {e}")
            return False
    
    def _initialize_ocr(self):
        """Инициализирует OCR для распознавания текста на изображениях (теперь ленивая загрузка)"""
        # OCR теперь загружается в фоне, здесь ничего не делаем
        pass

    def _ensure_chromadb_initialized(self):
        """Ленивая инициализация ChromaDB - вызывается только при первом использовании"""
        if not self._chromadb_initialized:
            try:
                self.logger.info("🔄 Инициализирую ChromaDB...")
                self.chromadb_manager = ChromaDBManager(
                    db_path=self._chromadb_config["db_path"],
                    use_gpu=self._chromadb_config["use_gpu"]
                )
                self._chromadb_initialized = True
                self.logger.info("✅ ChromaDB инициализирован")
            except Exception as e:
                self.logger.error(f"❌ Ошибка инициализации ChromaDB: {e}")
                self.chromadb_manager = None
    
    def _check_ffmpeg(self):
        """Проверяет наличие ffmpeg в системе для конвертации аудио"""
        try:
            result = subprocess.run(['ffmpeg', '-version'], capture_output=True, text=True, timeout=10)
            if result.returncode == 0:
                self.logger.info("✅ ffmpeg найден в системе")
            else:
                self.logger.warning("⚠️ ffmpeg найден, но не может быть запущен")
        except (subprocess.CalledProcessError, FileNotFoundError):
            self.logger.warning("⚠️ ffmpeg не найден в системе. Установите ffmpeg для конвертации аудио.")
            self.logger.info("💡 Скачайте с https://ffmpeg.org/download.html")
        except Exception as e:
            self.logger.warning(f"⚠️ Ошибка проверки ffmpeg: {e}")
    
    def is_model_running(self, model_name: str) -> bool:
        """Check if model is running."""
        if self.use_llama_cpp:
            return self.llama_wrapper is not None and self.llama_wrapper.is_loaded()
        return is_model_running_lm_studio(self.lm_studio_url, model_name, self.logger)

    def get_model_context_info(self) -> Dict[str, int]:
        """Get model context information."""
        if self.use_llama_cpp:
            if self.llama_wrapper and self.llama_wrapper.is_loaded():
                return self.llama_wrapper.get_context_info()
            default_ctx = LLAMA_CPP_PARAMS.get("n_ctx", 262144)
            self.logger.warning("Model not loaded, using defaults")
            return {"max_context": default_ctx, "safe_context": int(default_ctx * 0.8)}
        
        # LM Studio mode
        search_terms = ["huihui-qwen3-4b-thinking", "qwen3-4b", "thinking"]
        return get_model_context_info_lm_studio(
            self.lm_studio_url,
            search_terms,
            default_context=262144,
            logger_instance=self.logger,
        )

    def _get_context_info_via_chat(self, model_id: str) -> Optional[Dict[str, int]]:
        """
        Пытается получить информацию о контексте через запрос к модели
        """
        try:
            # Простой запрос для получения информации о модели
            payload = {
                "model": model_id,
                "messages": [{"role": "user", "content": "Hello"}],
                "max_tokens": 1,
                "temperature": 0
            }
            
            resp = requests.post(f"{self.lm_studio_url}/v1/chat/completions", json=payload, timeout=10)
            
            if resp.status_code == 200:
                data = resp.json()
                
                # Проверяем поле stats
                stats = data.get("stats", {})
                if stats:
                    # Ищем информацию о контексте в stats
                    context_length = None
                    if "context_length" in stats:
                        context_length = stats["context_length"]
                    elif "max_context" in stats:
                        context_length = stats["max_context"]
                    elif "max_tokens" in stats:
                        context_length = stats["max_tokens"]
                    
                    if context_length:
                        safe_context = max(context_length // 8, 32768)
                        self.logger.info(f"✅ Найден context_length в stats: {context_length}")
                        return {
                            "max_context": context_length,
                            "safe_context": safe_context
                        }
                
                # Проверяем другие поля на предмет информации о контексте
                for key, value in data.items():
                    if isinstance(value, dict) and ("context" in key.lower() or "token" in key.lower()):
                        self.logger.debug(f"🔍 Проверяем поле {key}: {value}")
                
                self.logger.debug("❌ Информация о контексте не найдена в ответе модели")
                return None
            else:
                self.logger.warning(f"❌ Ошибка запроса к модели: {resp.status_code}")
                return None
                
        except Exception as e:
            self.logger.warning(f"❌ Ошибка получения информации через чат: {e}")
            return None

    def _initialize_dynamic_context(self):
        """
        Инициализирует динамические параметры контекста на основе информации о модели
        """
        try:
            context_info = self.get_model_context_info()
            self.max_context_length = context_info["max_context"]
            self.safe_context_length = context_info["safe_context"]
            self.logger.info(f"📊 Контекст инициализирован: максимум {self.max_context_length:,}, безопасный {self.safe_context_length:,}")
        except Exception as e:
            self.logger.warning(f"⚠️ Ошибка инициализации динамического контекста: {e}")
            # Оставляем значения по умолчанию
            self.max_context_length = 262144
            self.safe_context_length = 32768

    def _initialize_email_config(self):
        """
        Инициализирует конфигурацию почтовой системы
        """
        try:
            # Загружаем настройки почты из переменных окружения
            self.email_config = {
                'gmail': {
                    'smtp_server': os.getenv('GMAIL_SMTP_SERVER', 'smtp.gmail.com'),
                    'smtp_port': int(os.getenv('GMAIL_SMTP_PORT', '587')),
                    'imap_server': os.getenv('GMAIL_IMAP_SERVER', 'imap.gmail.com'),
                    'imap_port': int(os.getenv('GMAIL_IMAP_PORT', '993')),
                    'email': os.getenv('GMAIL_EMAIL'),
                    'app_password': os.getenv('GMAIL_APP_PASSWORD')
                },
                'outlook': {
                    'smtp_server': os.getenv('OUTLOOK_SMTP_SERVER', 'smtp-mail.outlook.com'),
                    'smtp_port': int(os.getenv('OUTLOOK_SMTP_PORT', '587')),
                    'imap_server': os.getenv('OUTLOOK_IMAP_SERVER', 'outlook.office365.com'),
                    'imap_port': int(os.getenv('OUTLOOK_IMAP_PORT', '993')),
                    'email': os.getenv('OUTLOOK_EMAIL'),
                    'app_password': os.getenv('OUTLOOK_APP_PASSWORD')
                },
                'yandex': {
                    'smtp_server': os.getenv('YANDEX_SMTP_SERVER', 'smtp.yandex.ru'),
                    'smtp_port': int(os.getenv('YANDEX_SMTP_PORT', '587')),
                    'imap_server': os.getenv('YANDEX_IMAP_SERVER', 'imap.yandex.ru'),
                    'imap_port': int(os.getenv('YANDEX_IMAP_PORT', '993')),
                    'email': os.getenv('YANDEX_EMAIL'),
                    'app_password': os.getenv('YANDEX_APP_PASSWORD')
                },
                'mail_ru': {
                    'smtp_server': os.getenv('MAIL_RU_SMTP_SERVER', 'smtp.mail.ru'),
                    'smtp_port': int(os.getenv('MAIL_RU_SMTP_PORT', '587')),
                    'imap_server': os.getenv('MAIL_RU_IMAP_SERVER', 'imap.mail.ru'),
                    'imap_port': int(os.getenv('MAIL_RU_IMAP_PORT', '993')),
                    'email': os.getenv('MAIL_RU_EMAIL'),
                    'app_password': os.getenv('MAIL_RU_APP_PASSWORD')
                }
            }

            # Проверяем какие провайдеры настроены
            self.available_email_providers = []
            for provider, config in self.email_config.items():
                if config['email'] and config['app_password']:
                    self.available_email_providers.append(provider)
            
            if self.available_email_providers:
                self.logger.info(f"📧 Почтовая система инициализирована. Доступные провайдеры: {', '.join(self.available_email_providers)}")
            else:
                self.logger.warning("⚠️ Почтовые провайдеры не настроены. Проверьте переменные окружения.")
            
            # Инициализируем EmailManager
            self.email_manager = EmailManager(self.email_config)
                
        except Exception as e:
            self.logger.error(f"❌ Ошибка инициализации почтовой системы: {e}")
            self.email_config = {}
            self.available_email_providers = []
            self.email_manager = None

    def send_email(self, provider: str, to_email: str, subject: str, body: str, attachments: Optional[List[str]] = None, reply_to: Optional[str] = None):
        """Wrapper для EmailManager.send_email"""
        if not self.email_manager:
            return "❌ Почтовая система не инициализирована"
        return self.email_manager.send_email(provider, to_email, subject, body, attachments, reply_to)

    def get_emails(self, provider: str, folder: str = 'INBOX', limit: int = 10, search_criteria: str = 'ALL'):
        """Wrapper для EmailManager.get_emails"""
        if not self.email_manager:
            return "❌ Почтовая система не инициализирована"
        return self.email_manager.get_emails(provider, folder, limit, search_criteria)

    def reply_to_email(self, provider: str, original_email_id: str, reply_text: str, attachments: Optional[List[str]] = None):
        """Wrapper для EmailManager.reply_to_email"""
        if not self.email_manager:
            return "❌ Почтовая система не инициализирована"
        return self.email_manager.reply_to_email(provider, original_email_id, reply_text, attachments)

    def search_emails(self, provider: str, query: str, folder: str = 'INBOX', limit: int = 20):
        """Wrapper для EmailManager.search_emails"""
        if not self.email_manager:
            return "❌ Почтовая система не инициализирована"
        return self.email_manager.search_emails(provider, query, folder, limit)

    def _trim_context_if_needed(self):
        """
        Обрезает контекст если он превышает безопасные лимиты
        Использует self.current_context_length (total_tokens) для проверки
        """
        if self.current_context_length > self.max_context_length:
            # Критический лимит - агрессивная обрезка
            self.conversation_history = self.conversation_history[-2:]  # Оставляем только 2 последних сообщения
            self.logger.warning(f"Критическое превышение контекста ({self.current_context_length:,} > {self.max_context_length:,}) - агрессивная обрезка истории")
        elif self.current_context_length > self.safe_context_length:
            # Превышение безопасного лимита - аккуратная обрезка
            self.conversation_history = self.conversation_history[-5:]  # Оставляем только 5 последних сообщений
            self.logger.warning(f"Превышение безопасного контекста ({self.current_context_length:,} > {self.safe_context_length:,}) - аккуратная обрезка истории")
        elif self.current_context_length > self.safe_context_length * 0.8:
            # Приближение к безопасному лимиту - профилактическая обрезка
            self.conversation_history = self.conversation_history[-10:]  # Оставляем только 10 последних сообщений
            self.logger.info(f"Приближение к безопасному лимиту ({self.current_context_length:,} > {self.safe_context_length * 0.8:,}) - профилактическая обрезка истории")

    def launch_model(self, model_path: str):
        """
        Запускает модель через LM Studio (локально, subprocess)
        """
        try:
            # threading уже импортирован в начале файла
            lmstudio_exe = os.getenv("LMSTUDIO_EXE", r"C:\Program Files\LM Studio\LM Studio.exe")
            self.logger.info(f"Запускаю модель: {model_path}")
            threading.Thread(target=lambda: os.system(f'"{lmstudio_exe}" --model "{model_path}"'), daemon=True).start()
        except Exception as e:
            self.logger.error(f"Ошибка запуска модели: {e}")

    def ask_qwen(self, question: str) -> Optional[str]:
        """Wrapper для ask_qwen_for_prompt"""
        model_id = self.brain_model_id if hasattr(self, 'brain_model_id') and self.brain_model_id else self.brain_model
        return ask_qwen_for_prompt(
            self.lm_studio_url,
            model_id,
            question,
            self.logger
        )

    def get_youtube_cookies_path(self) -> Optional[str]:
        """
        Получает путь к файлу cookies для YouTube
        Возвращает путь к файлу или None если файл не найден
        """
        cookies_file = "youtube_cookies.txt"
        
        # Сначала ищем в текущей рабочей директории
        cookies_path = os.path.join(os.getcwd(), cookies_file)
        if os.path.exists(cookies_path) and os.path.getsize(cookies_path) > 0:
            self.logger.info(f"🍪 Найден файл cookies в рабочей директории: {cookies_file}")
            return cookies_path
        
        # Затем ищем в директории скрипта
        cookies_path = os.path.join(os.path.dirname(__file__), cookies_file)
        if os.path.exists(cookies_path) and os.path.getsize(cookies_path) > 0:
            self.logger.info(f"🍪 Найден файл cookies в директории скрипта: {cookies_file}")
            return cookies_path
        
        # Если файл не найден нигде
        self.logger.info(f"ℹ️ Файл cookies не найден: {cookies_file}")
        return None

    def check_cookies_validity(self, cookies_path: str) -> bool:
        """
        Проверяет валидность файла cookies
        """
        try:
            with open(cookies_path, 'r', encoding='utf-8') as f:
                content = f.read()
                
            # Проверяем базовую структуру
            if not content.strip():
                return False
                
            # Проверяем наличие YouTube доменов
            youtube_domains = ['youtube.com', '.youtube.com', 'google.com', '.google.com']
            has_youtube = any(domain in content for domain in youtube_domains)
            
            if not has_youtube:
                self.logger.warning("⚠️ В файле cookies не найдены домены YouTube")
                return False
                
            # Проверяем формат (должен содержать табуляции)
            if '\t' not in content:
                self.logger.warning("⚠️ Неверный формат файла cookies (отсутствуют табуляции)")
                return False
                
            self.logger.info("✅ Файл cookies валиден")
            return True
            
        except Exception as e:
            self.logger.error(f"❌ Ошибка проверки cookies: {e}")
            return False

    def suggest_cookies_update(self):
        """
        Предлагает пользователю обновить cookies
        """
        self.logger.info("💡 Для улучшения работы с YouTube рекомендуется:")
        self.logger.info("   1. Запустить: python extract_chrome_cookies.py")
        self.logger.info("   2. Закрыть Chrome перед извлечением")
        self.logger.info("   3. Войти в YouTube через VPN")
        self.logger.info("   4. Cookies обновляются каждые 2-3 месяца")

    def generate_image_stable_diffusion(self, prompt: str, negative_prompt: str, params: dict) -> Optional[str]:
        """Wrapper для ImageGenerator.generate_image"""
        if not self.image_generator:
            self.logger.error("❌ Генератор изображений не инициализирован")
            return None
            
        # Генерация
        output_path = self.image_generator.generate_image(prompt, negative_prompt, params)
        if not output_path:
            return None
            
        # Апскейл
        if self.image_generator._is_realesrgan_available():
             self.logger.info("🔍 Модель RealESRGAN найдена, применяю апскейл...")
             upscaled_path = self.image_generator.upscale_image_realesrgan(output_path)
             if upscaled_path:
                 output_path = upscaled_path
                 self.logger.info("✨ Использую увеличенное изображение")
        
        # Выгружаем pipeline
        self.image_generator._unload_current_pipeline()
        
        # Конвертация в base64
        try:
            with open(output_path, "rb") as f:
                return base64.b64encode(f.read()).decode("ascii")
        except Exception as e:
            self.logger.error(f"❌ Ошибка чтения файла изображения: {e}")
            return None

    # _unload_current_pipeline moved to ImageGenerator
    # _is_realesrgan_available moved to ImageGenerator

    # upscale_image_realesrgan moved to ImageGenerator
    
    # _upscale_image_alternative moved to ImageGenerator
    # _install_realesrgan_dependencies moved to ImageGenerator

    def generate_video_stable_diffusion(self, prompt: str, negative_prompt: str, params: dict) -> Optional[str]:
        """Wrapper для ImageGenerator.generate_video"""
        if not self.image_generator:
            self.logger.error("❌ Генератор изображений не инициализирован")
            return None
        return self.image_generator.generate_video(prompt, negative_prompt, params)

    # _add_dynamic_elements moved to ImageGenerator
    
    # _install_diffusers_dependencies moved to ImageGenerator

    def show_image_base64_temp(self, b64img: str):
        """Показать изображение из base64 через универсальный метод Windows"""
        try:
            # В веб-режиме отключаем всплывающее окно показа
            if not getattr(self, 'show_images_locally', True):
                return
            
            # Создаем временный файл для показа
            import tempfile
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                tmp_path = tmp_file.name
                img_data = base64.b64decode(b64img)
                tmp_file.write(img_data)
            
            # Открываем через универсальный метод Windows
            try:
                subprocess.run(["start", tmp_path], shell=True, check=True)
                self.logger.info("🖼️ Изображение автоматически открыто")
            except Exception as e:
                self.logger.warning(f"⚠️ Не удалось открыть изображение: {e}")
                
        except Exception as e:
            self.logger.error(f"Ошибка показа изображения: {e}")

    def find_new_audio(self) -> Optional[str]:
        """Находит новый аудиофайл для обработки"""
        audio_extensions = ['.mp3', '.wav', '.m4a', '.flac', '.ogg', '.aac']
        
        # Ищем в папке Audio
        audio_dir = os.path.join(self.base_dir, 'Audio')
        if os.path.exists(audio_dir):
            for file in os.listdir(audio_dir):
                if any(file.lower().endswith(ext) for ext in audio_extensions):
                    # Проверяем, что файл не помечен как обработанный
                    if '.used' not in file and not file.endswith('.used'):
                        file_path = os.path.join(audio_dir, file)
                        # Проверяем, что файл действительно существует и не пустой
                        if os.path.exists(file_path) and os.path.getsize(file_path) > 0:
                            return file_path
        
        return ""

    def mark_audio_used(self, audio_path: str):
        """Удаляет аудиофайл после обработки"""
        try:
            if os.path.exists(audio_path):
                # Удаляем файл полностью
                os.remove(audio_path)
                self.logger.info(f"✅ Аудиофайл удален после обработки: {os.path.basename(audio_path)}")
        except Exception as e:
            self.logger.error(f"❌ Ошибка при удалении аудиофайла: {e}")

    def transcribe_audio_whisper(self, audio_path: str, lang: str = "ru", use_separator: bool = True) -> Optional[str]:
        """
        Распознаёт аудио через whisper-cli. Если use_separator=True, предварительно выделяет вокал через audio-separator.
        Возвращает текст транскрипта (выводит только один раз при получении).
        """
        start_time = time.time()
        
        # Автоматически включаем audio модель при необходимости
        if not getattr(self, 'use_audio', False):
            self.logger.info("🔧 Автоматически включаю audio модель")
            self.use_audio = True
            # Запускаем таймер автоматического выключения
            self.auto_disable_tools("audio")
        
        # Проверяем и загружаем whisper модель если нужно
        if not self.check_whisper_setup():
            return "[Whisper error] Проблемы с настройкой Whisper. Проверьте наличие whisper-cli.exe и модели."
        
        try:
            base_dir = os.path.dirname(os.path.abspath(__file__))
            exe_path = os.path.join(base_dir, "Release", "whisper-cli.exe")
            model_path = os.path.join(base_dir, "models", "whisper-large-v3-q8_0.gguf")
            
            # Проверяем существование whisper-cli.exe
            if not os.path.exists(exe_path):
                return "[Whisper error] Не найден whisper-cli.exe в папке Release"
            
            # Проверяем существование модели
            if not os.path.exists(model_path):
                return "[Whisper error] Не найдена модель whisper в папке models"
            
            audio_for_whisper = audio_path
            
            # Используем audio separator если включен
            if use_separator:
                try:
                    from audio_separator.separator import Separator
                    self.logger.info("🎵 Использую audio-separator для выделения вокала...")
                    out_dir = os.path.join(base_dir, "separated")
                    os.makedirs(out_dir, exist_ok=True)
                    separator = Separator(output_dir=out_dir)
                    separator.load_model(model_name='UVR-MDX-NET-Inst_HQ_3')
                    output_files = separator.separate(audio_path)
                    vocals_path = None
                    for file_path in output_files:
                        if '(Vocals)' in os.path.basename(file_path):
                            vocals_path = file_path  # audio-separator возвращает полный путь
                            self.logger.info(f"[SUCCESS] Вокал найден: {vocals_path}")
                            break
                    if not vocals_path:
                        self.logger.warning("⚠️ Не удалось найти файл с голосом после разделения дорожек, использую оригинал")
                    else:
                        audio_for_whisper = vocals_path
                except ImportError:
                    self.logger.warning("⚠️ Не установлена библиотека audio-separator. Пытаюсь установить автоматически...")
                    try:
                        import subprocess
                        subprocess.run([_sys.executable, "-m", "pip", "install", "audio-separator"], 
                                     capture_output=True, check=True)
                        self.logger.info("✅ audio-separator успешно установлен")
                        # Повторно пытаемся импортировать
                        from audio_separator.separator import Separator
                        self.logger.info("🎵 Использую audio-separator для выделения вокала...")
                        out_dir = os.path.join(base_dir, "separated")
                        os.makedirs(out_dir, exist_ok=True)
                        separator = Separator(output_dir=out_dir)
                        separator.load_model(model_name='UVR-MDX-NET-Inst_HQ_3')
                        output_files = separator.separate(audio_path)
                        vocals_path = None
                        for file_path in output_files:
                            if '(Vocals)' in os.path.basename(file_path):
                                vocals_path = file_path  # audio-separator возвращает полный путь
                                self.logger.info(f"[SUCCESS] Вокал найден: {vocals_path}")
                                break
                        if not vocals_path:
                            self.logger.warning("⚠️ Не удалось найти файл с голосом после разделения дорожек, использую оригинал")
                        else:
                            audio_for_whisper = vocals_path
                    except Exception as install_error:
                        self.logger.warning(f"⚠️ Не удалось установить audio-separator: {install_error}")
                        self.logger.info("ℹ️ Продолжаю без разделения дорожек")
                except Exception as e:
                    self.logger.warning(f"⚠️ Ошибка audio-separator: {e}, использую оригинал")
            
            # Конвертируем аудио в WAV формат для Whisper (если это не уже WAV)
            if not audio_for_whisper.lower().endswith('.wav'):
                wav_path = self.convert_audio_to_wav(audio_for_whisper)
                if wav_path:
                    audio_for_whisper = wav_path
                    self.logger.info(f"✅ Аудио конвертировано в WAV: {os.path.basename(wav_path)}")
                else:
                    self.logger.warning("⚠️ Не удалось конвертировать в WAV, использую оригинал")
            else:
                self.logger.info("✅ Аудио уже в WAV формате")
            
            # Переименовать используемый файл в .used.расширение
            base_used, ext_used = os.path.splitext(audio_for_whisper)
            used_path = base_used + ".used" + ext_used
            try:
                if os.path.exists(audio_for_whisper):
                    os.rename(audio_for_whisper, used_path)
                    self.logger.info(f"✅ Аудиофайл переименован в: {os.path.basename(used_path)}")
                else:
                    self.logger.warning(f"⚠️ Аудиофайл не найден для переименования: {audio_for_whisper}")
                    used_path = audio_for_whisper
            except Exception as e:
                self.logger.error(f"Ошибка переименования аудио после whisper: {e}")
                # Если не удалось переименовать, используем оригинальный файл
                used_path = audio_for_whisper
            
            # Теперь используем used_path для whisper
            cmd = [exe_path, "--model", model_path]
            if lang:
                cmd += ["--language", lang]
            cmd.append(used_path)
            self.logger.info(f"[INFO] Запуск Whisper: {' '.join(cmd)}")
            import subprocess
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=300, encoding="utf-8", errors="replace")
            transcript = result.stdout.strip() if result.stdout else ""
            if transcript:
                self.logger.info("\n=== ТРАНСКРИПТ АУДИО ===\n" + transcript)
                return self._wrap_info_if_needed(transcript, source="audio")
            
            # Очистка временных файлов если был separator
            if use_separator and 'separated' in audio_for_whisper:
                try:
                    separated_dir = os.path.dirname(audio_for_whisper)
                    if os.path.exists(separated_dir):
                        shutil.rmtree(separated_dir)
                        self.logger.info("🧹 Временные файлы audio-separator очищены")
                except Exception as e:
                    self.logger.warning(f"⚠️ Не удалось очистить временные файлы: {e}")
            
            err = result.stderr.strip() if result.stderr else ""
            return f"[Whisper error] Не удалось получить транскрипт. STDERR: {err}"
        except Exception as e:
            error_msg = f"Исключение whisper-cli: {str(e)}"
            self.logger.error(error_msg)
            return f"[Whisper error] {error_msg}"
        finally:
            # Записываем метрику производительности
            response_time = time.time() - start_time
            self.add_performance_metric("whisper_transcription", response_time)
            self.logger.info(f"🎤 Whisper обработал за {response_time:.2f} сек")

    def convert_audio_to_wav(self, audio_path: str) -> Optional[str]:
        """Wrapper для backwards compatibility."""
        return convert_audio_to_wav(audio_path)

    def check_whisper_setup(self) -> bool:
        """Wrapper для backwards compatibility."""
        return check_whisper_setup()

    def download_whisper_model(self) -> bool:
        """Wrapper для backwards compatibility."""
        return download_whisper_model()

    def download_youtube_audio(self, url: str, out_dir: Optional[str] = None) -> str:
        """Wrapper для backwards compatibility."""
        return download_youtube_audio(url, out_dir)
    def find_new_image(self) -> str:
        """
        Находит первое новое изображение (png или jpg) в папке Photos, игнорируя файлы с .used перед расширением
        """
        photos_dir = os.path.join(os.path.dirname(__file__), "Photos")
        if not os.path.exists(photos_dir):
            return ""
        for fname in os.listdir(photos_dir):
            lower = fname.lower()
            if lower.endswith(('.png', '.jpg', '.jpeg')) and '.used' not in os.path.splitext(lower)[0]:
                return os.path.join(photos_dir, fname)
        return ""

    def mark_image_used(self, image_path: str):
        """
        Переименовывает изображение, чтобы нейросеть его больше не использовала
        """
        if not image_path:
            return
        base, ext = os.path.splitext(image_path)
        new_path = base + ".used" + ext
        try:
            os.rename(image_path, new_path)
        except Exception as e:
            self.logger.error(f"Ошибка переименования изображения: {e}")
    def extract_think_content(self, text: str) -> Optional[str]:
        """
        Извлекает содержимое из блока <think> или альтернативных маркеров размышлений.
        Поддерживает форматы:
        - <think>...</think> 
        - <|begin_of_thought|>...<|end_of_thought|>
        - BEGIN_OF_THOUGHT...END_OF_THOUGHT

        Args:
            text: Входной текст для поиска блока размышлений

        Returns:
            Optional[str]: Содержимое блока размышлений или None если блок не найден
        """
        # Паттерны для поиска (case-insensitive)
        patterns = [
            r'<think>(.*?)</think>',
            r'<\|begin_of_thought\|>(.*?)<\|end_of_thought\|>',
            r'BEGIN_OF_THOUGHT(.*?)END_OF_THOUGHT'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
            if match:
                content = match.group(1).strip()
                if content:  # Проверяем что контент не пустой
                    return content
        
        return None

    def extract_first_json(self, text: str, allow_json_in_think: bool = False) -> str:
        """
        Извлекает первый корректный JSON-блок из текста.
        
        Args:
            text: Входной текст для поиска JSON
            allow_json_in_think: Искать ли JSON внутри блока think
            
        Returns:
            str: Найденный JSON или исходный текст если JSON не найден
        """
        # Сначала ищем think-блок и удаляем его из текста если он есть
        think_content = None
        clean_text = text
        
        patterns = [
            r'<think>.*?</think>',
            r'<\|begin_of_thought\|>.*?<\|end_of_thought\|>',
            r'BEGIN_OF_THOUGHT.*?END_OF_THOUGHT'
        ]
        
        for pattern in patterns:
            think_match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
            if think_match:
                think_content = self.extract_think_content(think_match.group(0))
                clean_text = re.sub(pattern, '', text, flags=re.DOTALL | re.IGNORECASE).strip()
                break
        
        # Сначала пробуем найти чистый JSON (без обрамлений)
        json_in_text = self._extract_json_from_text(clean_text)
        if json_in_text:
            return json_in_text
        
        # Если не нашли, пробуем искать JSON с обрамлением и удалять его
        json_with_wrapper = re.search(r'```(?:json)?\s*(.*?)\s*```', clean_text, re.DOTALL)
        if json_with_wrapper:
            potential_json = json_with_wrapper.group(1).strip()
            json_in_wrapper = self._extract_json_from_text(potential_json)
            if json_in_wrapper:
                return json_in_wrapper
            
        # Если разрешено и есть think-контент - ищем JSON там
        if allow_json_in_think and think_content:
            json_in_think = self._extract_json_from_text(think_content)
            if json_in_think:
                return json_in_think
        
        return text  # если не найдено, вернуть исходное
    
    def _extract_json_from_text(self, text: str) -> str:
        """
        Вспомогательная функция для извлечения JSON из текста
        """
        # re уже импортирован в начале файла
        stack = []
        start = None
        
        for i, c in enumerate(text):
            if c == '{':
                if not stack:
                    start = i
                stack.append('{')
            elif c == '}':
                if stack:
                    stack.pop()
                    if not stack and start is not None:
                        return text[start:i+1]
        
        return ""

    def _smart_json_parse(self, s: str):
        """Умный парсер JSON с несколькими попытками автокоррекции.

        Возвращает tuple (data_or_none, fixes_list).
        """
        logger.info(f"🔍 Парсинг JSON: {s[:200]}...")
        try:
            return json.loads(s), []
        except Exception as e:
            fixes = [f"Первый парсинг не удался: {e}"]

        # Попытка закрыть скобки
        open_braces = s.count('{')
        close_braces = s.count('}')
        if open_braces > close_braces:
            s2 = s + '}' * (open_braces - close_braces)
            fixes.append(f"Добавлено {open_braces - close_braces} }} для баланса скобок")
            try:
                return json.loads(s2), fixes
            except Exception:
                pass
        elif close_braces > open_braces:
            s2 = re.sub(r'}+$', '', s)
            fixes.append(f"Удалены лишние закрывающие скобки")
            try:
                return json.loads(s2), fixes
            except Exception:
                pass

        # Заменяем одинарные кавычки на двойные если это безопасно
        if "'" in s and '"' not in s:
            s2 = s.replace("'", '"')
            try:
                return json.loads(s2), fixes+['Заменены одинарные кавычки на двойные']
            except Exception:
                pass

        # Удаляем лишние запятые перед закрывающей скобкой
        s3 = re.sub(r',\s*([}\]])', r'\1', s)
        try:
            return json.loads(s3), fixes+['Удалены лишние запятые']
        except Exception:
            pass

        # Оборачиваем ключи в кавычки (грубая попытка)
        s4 = re.sub(r'([,{]\s*)([a-zA-Z0-9_]+)\s*:', r'\1"\2":', s3)
        try:
            return json.loads(s4), fixes+['Добавлены кавычки к ключам']
        except Exception:
            pass

        # Исправляем незакрытые строки
        s5 = re.sub(r'([^\"])\s*$', r'\1"', s4)
        try:
            return json.loads(s5), fixes+['Исправлены незакрытые строки']
        except Exception:
            pass

        # Финальная попытка с очисткой от непечатаемых символов
        s6 = re.sub(r'[^\x20-\x7E]', '', s5)
        try:
            return json.loads(s6), fixes+['Очищены непечатаемые символы']
        except Exception as e2:
            fixes.append(f"Не удалось распарсить даже после исправлений: {e2}")

        return None, fixes
    def __init__(self, lm_studio_url: str = DEFAULT_LM_STUDIO_URL, 
                 google_api_key: str = "", google_cse_id: str = ""):
        """
        Инициализация оркестратора
        
        Args:
            lm_studio_url: URL сервера LM Studio (используется только если USE_LLAMA_CPP = False)
            google_api_key: API ключ Google Custom Search
            google_cse_id: ID поисковой системы Google CSE
        """
        # Определяем режим работы: llama.cpp или LM Studio
        self.use_llama_cpp = USE_LLAMA_CPP
        
        # Инициализация для режима LM Studio
        self.lm_studio_url = lm_studio_url.rstrip("/")
        
        # Инициализация для режима llama.cpp
        self.llama_wrapper: Optional[LlamaCppWrapper] = None
        
        self.google_api_key = google_api_key
        self.google_cse_id = google_cse_id
        # unify logger usage for instance methods
        self.logger = logger
        self.conversation_history: List[Dict[str, Any]] = []
        self.brain_model = LLAMA_CPP_MODEL_PATH
        self.brain_model_id = None  # Короткий ID модели для API вызовов
        self.use_separator = True  # По умолчанию True, чтобы убрать предупреждение Pylance
        self.use_image_generation = False  # По умолчанию отключена генерация изображений
        # Тумблеры функционала (визуал и аудио)
        self.use_vision = False
        self.use_audio = False
        self.use_ocr = False  # По умолчанию отключен OCR
        # Управление локальным показом изображений (для веб-режима можно отключить)
        self.show_images_locally = True
    # Конфигурация vision-модели (Moondream2)
        self.vision_model_id = VISION_MODEL_ID
        self.vision_model_load_args = dict(VISION_MODEL_LOAD_ARGS)
        self.vision_generation_params = dict(VISION_GENERATION_PARAMS)
        self._vision_model_ready = False
        # Хранилище последнего сгенерированного изображения (base64) и ответа
        self.last_generated_image_b64 = None
        self.last_final_response = ""
        
        # Хранилище последнего сгенерированного файла для Telegram
        self.last_generated_file_path = None
        self.last_generated_file_name = None
        
        # Динамическое управление контекстом
        self.max_context_length = DEFAULT_MAX_CONTEXT_LENGTH
        self.safe_context_length = DEFAULT_SAFE_CONTEXT_LENGTH
        self.current_context_length = 0    # Текущий размер контекста
        
        # Метрики производительности
        self.performance_metrics = []  # Список метрик производительности
        
        # Счетчик попыток для предотвращения зацикливания
        self.retry_count = 0
        self.max_retries = DEFAULT_MAX_RETRIES
        
        # Постоянная голосовая запись
        self.continuous_recording = False
        self.audio_queue = queue.Queue()
        self.recording_thread = None
        
        # Таймеры для автоматического выключения инструментов
        self.tool_timers = {}
        self.auto_disable_delay = AUTO_DISABLE_DELAY_SECONDS
        
        # Автоматически запускаем модель мозга при инициализации
        self._auto_load_brain_model()
        
        # Инициализируем динамические параметры контекста после загрузки модели
        self._initialize_dynamic_context()
        
        # Инициализируем базовую директорию
        self.base_dir = os.path.dirname(os.path.abspath(__file__))
        
        # Запускаем фоновую загрузку тяжелых компонентов
        self._start_background_loading()
        
        # Инициализируем ChromaDB для векторного хранилища (ленивая инициализация)
        self.chromadb_manager = None
        self._chromadb_initialized = False
        chroma_path = CHROMA_DB_PATH
        if not os.path.isabs(chroma_path):
            chroma_path = os.path.join(self.base_dir, os.path.normpath(chroma_path))
        self._chromadb_config = {
            "db_path": chroma_path,
            "use_gpu": CHROMADB_USE_GPU_BY_DEFAULT
        }
        
        # OCR будет инициализирован в фоне
        self.ocr_reader = None
        
        # Проверяем наличие ffmpeg для конвертации аудио
        self._check_ffmpeg()
        
        # Telegram Bot настройки
        self.telegram_bot_token = ""
        self.telegram_allowed_user_id = ""
        
        # Инициализируем систему плагинов
        # Initialize plugin system
        self.plugin_manager = None
        if PLUGINS_AVAILABLE and PluginManager is not None:
            try:
                self.plugin_manager = PluginManager(plugins_dir=PLUGINS_DIR_NAME)
                self.plugin_manager.load_all_plugins(orchestrator=self)
                logger.info("✅ Система плагинов инициализирована")
            except Exception as e:
                logger.error(f"❌ Ошибка инициализации плагинов: {e}")
                self.plugin_manager = None
        else:
            logger.warning("⚠️ Система плагинов недоступна")
        
        # Инициализируем почтовые настройки
        self._initialize_email_config()

        # Инициализируем загрузчик промптов
        self.prompt_loader = PromptLoader(self.base_dir)
        
        # Инициализируем менеджер моделей и LoRA
        self.model_manager = ModelManager(self.base_dir)
        
        # Инициализируем генератор изображений
        self.image_generator = ImageGenerator(self.model_manager)
        
        # Загружаем базовый системный промпт из файла
        self.system_prompt = self.prompt_loader.load_base_prompt()

    # ------------------------------------------------------------------
    # Вспомогательные утилиты для безопасной передачи данных в модель
    # ------------------------------------------------------------------
    def _wrap_info_block(self, content: str, source: str = "") -> str:
        """Оборачивает произвольный текст в тег <INFO> для защиты от jailbreak.

        Args:
            content: Текст, который нужно передать в модель.
            source:  Необязательный маркер источника (file/audio/video).
        """
        if content is None:
            content = ""
        source_clean = re.sub(r"[^a-zA-Z0-9_-]+", "", source or "")
        source_attr = f" source=\"{source_clean}\"" if source_clean else ""
        return f"<INFO{source_attr}>\n{content}\n</INFO>"

    def _wrap_info_if_needed(self, content: str, source: str = "") -> str:
        """Оборачивает текст в <INFO>, если он ещё не обёрнут."""
        if not content:
            return content
        if "<INFO" in content:
            return content
        return self._wrap_info_block(content, source)

    def _strip_info_tags(self, content: str, highlight: bool = True) -> str:
        """Удаляет <INFO> теги из финального ответа и подсвечивает содержимое.

        По умолчанию подсвечивает ANSI-жёлтым. Если терминал не поддерживает ANSI,
        текст останется без цвета, но теги всё равно будут убраны.
        """
        if not content:
            return content

        def _repl(match: re.Match) -> str:
            inner = match.group(1).strip()
            if not inner:
                return ""
            if not highlight:
                return inner
            return f"\033[43m{inner}\033[0m"

        return re.sub(r"<INFO[^>]*>(.*?)</INFO>", _repl, content, flags=re.IGNORECASE | re.DOTALL)

    def list_folder_contents(self, folder_name: str) -> str:
        """
        Получение списка файлов в указанной папке
        
        Args:
            folder_name: Имя папки (Audio, Photos, Video, Excel, Docx, PDF)
        
        Returns:
            Строка со списком файлов или сообщение об ошибке
        """
        try:
            folder_path = os.path.join(self.base_dir, folder_name)
            
            if not os.path.exists(folder_path):
                return f"Папка {folder_name} не существует"
            
            files = os.listdir(folder_path)
            if not files:
                return f"Папка {folder_name} пуста"
            
            # Группируем файлы по типу
            file_types = {
                'Документы DOCX': [],
                'Таблицы Excel': [],
                'Документы PDF': [],
                'Изображения': [],
                'Аудио': [],
                'Видео': [],
                'Другие': []
            }
            
            for file in files:
                file_lower = file.lower()
                if file_lower.endswith(('.docx', '.doc')):
                    file_types['Документы DOCX'].append(file)
                elif file_lower.endswith(('.xlsx', '.xls', '.csv')):
                    file_types['Таблицы Excel'].append(file)
                elif file_lower.endswith('.pdf'):
                    file_types['Документы PDF'].append(file)
                elif file_lower.endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp')):
                    file_types['Изображения'].append(file)
                elif file_lower.endswith(('.mp3', '.wav', '.ogg', '.flac', '.m4a')):
                    file_types['Аудио'].append(file)
                elif file_lower.endswith(('.mp4', '.avi', '.mkv', '.mov', '.wmv')):
                    file_types['Видео'].append(file)
                else:
                    file_types['Другие'].append(file)
            
            result = f"Содержимое папки {folder_name}:\n"
            for file_type, file_list in file_types.items():
                if file_list:
                    result += f"\n{file_type}:\n"
                    for file in sorted(file_list):
                        result += f"  - {file}\n"
            
            return result
            
        except Exception as e:
            return f"Ошибка при чтении папки {folder_name}: {str(e)}"

    def extract_docx_content(self, file_path: str) -> tuple[str, str]:
        """
        Извлечение текста из DOCX файла
        
        Args:
            file_path: Путь к DOCX файлу
        
        Returns:
            Кортеж (текст, сообщение_об_ошибке)
        """
        try:
            from docx import Document
            
            if not os.path.exists(file_path):
                return "", f"Файл {file_path} не найден"
            
            doc = Document(file_path)
            text_content = []
            
            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            full_text = '\n'.join(text_content)
            
            if not full_text.strip():
                return "", "Документ не содержит текста"
            
            return full_text, ""
            
        except ImportError:
            return "", "Библиотека python-docx не установлена. Установите: pip install python-docx"
        except Exception as e:
            return "", f"Ошибка при чтении DOCX файла: {str(e)}"

    def extract_excel_content(self, file_path: str) -> tuple[str, str]:
        """
        Извлечение данных из Excel файла
        
        Args:
            file_path: Путь к Excel файлу
        
        Returns:
            Кортеж (данные_в_текстовом_формате, сообщение_об_ошибке)
        """
        try:
            import pandas as pd
            
            if not os.path.exists(file_path):
                return "", f"Файл {file_path} не найден"
            
            # Читаем все листы Excel файла
            excel_data = pd.read_excel(file_path, sheet_name=None)
            
            content_parts = []
            
            for sheet_name, df in excel_data.items():
                content_parts.append(f"=== Лист: {sheet_name} ===\n")
                
                # Проверяем, есть ли данные
                if df.empty:
                    content_parts.append("Лист пуст\n")
                    continue
                
                # Конвертируем DataFrame в текстовое представление
                content_parts.append(df.to_string(index=False))
                content_parts.append("\n")
            
            full_content = '\n'.join(content_parts)
            
            if not full_content.strip():
                return "", "Excel файл не содержит данных"
            
            return full_content, ""
            
        except ImportError:
            return "", "Библиотеки pandas/openpyxl не установлены. Установите: pip install pandas openpyxl"
        except Exception as e:
            return "", f"Ошибка при чтении Excel файла: {str(e)}"

    def extract_pdf_content(self, file_path: str) -> tuple[str, str]:
        """
        Извлечение текста из PDF файла
        
        Args:
            file_path: Путь к PDF файлу
        
        Returns:
            Кортеж (текст, сообщение_об_ошибке)
        """
        try:
            if not PDF_AVAILABLE or PyPDF2 is None:
                return "", "Библиотека PyPDF2 не установлена. Установите: pip install PyPDF2"
            
            if not os.path.exists(file_path):
                return "", f"Файл {file_path} не найден"
            
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Проверяем, есть ли страницы
                if len(pdf_reader.pages) == 0:
                    return "", "PDF файл не содержит страниц"
                
                # Извлекаем текст со всех страниц
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"=== Страница {page_num} ===\n{page_text.strip()}")
                    except Exception as e:
                        text_content.append(f"=== Страница {page_num} ===\n[Ошибка извлечения текста: {str(e)}]")
            
            full_text = '\n\n'.join(text_content)
            
            if not full_text.strip():
                return "", "PDF файл не содержит извлекаемого текста"
            
            return full_text, ""
            
        except Exception as e:
            return "", f"Ошибка при чтении PDF файла: {str(e)}"

    def rag_process_large_content(self, content: str, max_tokens: int = 4000) -> str:
        """
        RAG-обработка больших документов с разделением на части
        
        Args:
            content: Содержимое документа
            max_tokens: Максимальное количество токенов на часть
        
        Returns:
            Обработанное содержимое (сжатое или разделенное)
        """
        try:
            # Приблизительная оценка токенов (1 токен ≈ 4 символа для русского текста)
            estimated_tokens = len(content) // 4
            
            if estimated_tokens <= max_tokens:
                return content
            
            # Если документ слишком большой, разделяем на части
            chunk_size = max_tokens * 4  # Размер части в символах
            chunks = []
            
            # Разделяем по предложениям, чтобы сохранить смысл
            sentences = content.split('.')
            current_chunk = ""
            
            for sentence in sentences:
                if len(current_chunk + sentence) < chunk_size:
                    current_chunk += sentence + "."
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                        current_chunk = sentence + "."
                    else:
                        # Если одно предложение больше chunk_size, берем как есть
                        chunks.append(sentence.strip())
            
            # Добавляем последнюю часть
            if current_chunk:
                chunks.append(current_chunk.strip())
            
            # Возвращаем первые 3 части с указанием общего количества
            if len(chunks) <= 3:
                if len(chunks) > 1:
                    return f"=== ДОКУМЕНТ РАЗДЕЛЕН НА {len(chunks)} ЧАСТЕЙ ===\n\n" + '\n\n=== ЧАСТЬ ДОКУМЕНТА ===\n\n'.join(chunks)
                else:
                    return chunks[0] if chunks else content[:max_tokens * 4]
            else:
                result = f"=== ДОКУМЕНТ РАЗДЕЛЕН НА {len(chunks)} ЧАСТЕЙ ===\n\n"
                result += '\n\n=== ЧАСТЬ ДОКУМЕНТА ===\n\n'.join(chunks[:3])
                result += f"\n\n[ПОКАЗАНЫ ПЕРВЫЕ 3 ЧАСТИ ИЗ {len(chunks)}. СПРОСИТЕ, ЕСЛИ НУЖНО БОЛЬШЕ ИНФОРМАЦИИ]"
                return result
                
        except Exception as e:
            logger.error(f"Ошибка RAG обработки: {e}")
            return content[:max_tokens * 4] + "\n\n[СОДЕРЖИМОЕ ОБРЕЗАНО ДУЕ К РАЗМЕРУ]"

    def process_document_request(self, file_path: str) -> str:
        """
        Обработка запроса на работу с документом
        
        Args:
            file_path: Путь к файлу
        
        Returns:
            Обработанное содержимое документа
        """
        try:
            file_lower = file_path.lower()
            
            if file_lower.endswith(('.docx', '.doc')):
                content, error = self.extract_docx_content(file_path)
                if error:
                    return f"Ошибка при обработке DOCX: {error}"
                
                # RAG обработка для больших документов
                processed_content = self.rag_process_large_content(content)
                wrapped = self._wrap_info_if_needed(processed_content, source="docx")
                return f"Содержимое DOCX документа:\n\n{wrapped}"
                
            elif file_lower.endswith(('.xlsx', '.xls')):
                content, error = self.extract_excel_content(file_path)
                if error:
                    return f"Ошибка при обработке Excel: {error}"
                
                # RAG обработка для больших таблиц
                processed_content = self.rag_process_large_content(content)
                wrapped = self._wrap_info_if_needed(processed_content, source="excel")
                return f"Содержимое Excel файла:\n\n{wrapped}"
                
            elif file_lower.endswith('.pdf'):
                content, error = self.extract_pdf_content(file_path)
                if error:
                    return f"Ошибка при обработке PDF: {error}"
                
                # RAG обработка для больших PDF документов
                processed_content = self.rag_process_large_content(content)
                wrapped = self._wrap_info_if_needed(processed_content, source="pdf")
                return f"Содержимое PDF документа:\n\n{wrapped}"
                
            elif file_lower.endswith('.csv'):
                # Для CSV используем pandas
                try:
                    import pandas as pd
                    df = pd.read_csv(file_path)
                    content = df.to_string(index=False)
                    processed_content = self.rag_process_large_content(content)
                    wrapped = self._wrap_info_if_needed(processed_content, source="csv")
                    return f"Содержимое CSV файла:\n\n{wrapped}"
                except Exception as e:
                    return f"Ошибка при чтении CSV: {str(e)}"
                    
            elif file_lower.endswith(('.txt', '.md')):
                # Текстовые файлы
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    processed_content = self.rag_process_large_content(content)
                    wrapped = self._wrap_info_if_needed(processed_content, source="text")
                    file_type = "Markdown" if file_lower.endswith('.md') else "текстового"
                    return f"Содержимое {file_type} файла:\n\n{wrapped}"
                except UnicodeDecodeError:
                    # Попробуем другие кодировки
                    for encoding in ['cp1251', 'latin1']:
                        try:
                            with open(file_path, 'r', encoding=encoding) as f:
                                content = f.read()
                            processed_content = self.rag_process_large_content(content)
                            wrapped = self._wrap_info_if_needed(processed_content, source="text")
                            return f"Содержимое текстового файла (кодировка {encoding}):\n\n{wrapped}"
                        except:
                            continue
                    return f"Ошибка: не удалось определить кодировку файла {file_path}"
                except Exception as e:
                    return f"Ошибка при чтении текстового файла: {str(e)}"
                    
            elif file_lower.endswith(('.rtf')):
                # RTF файлы
                try:
                    # Простое извлечение текста из RTF (базовое)
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    # Удаляем RTF команды (очень упрощенно)
                    import re
                    content = re.sub(r'\\[a-z]+\d*', '', content)  # Убираем команды типа \par, \b1 и т.д.
                    content = re.sub(r'[{}]', '', content)  # Убираем фигурные скобки
                    content = content.strip()
                    processed_content = self.rag_process_large_content(content)
                    wrapped = self._wrap_info_if_needed(processed_content, source="rtf")
                    return f"Содержимое RTF файла:\n\n{wrapped}"
                except Exception as e:
                    return f"Ошибка при обработке RTF: {str(e)}"
                    
            elif file_lower.endswith('.json'):
                # JSON файлы
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    # Преобразуем JSON в читаемый формат
                    content = json.dumps(data, indent=2, ensure_ascii=False)
                    processed_content = self.rag_process_large_content(content)
                    wrapped = self._wrap_info_if_needed(processed_content, source="json")
                    return f"Содержимое JSON файла:\n\n{wrapped}"
                except json.JSONDecodeError as e:
                    return f"Ошибка в формате JSON: {str(e)}"
                except Exception as e:
                    return f"Ошибка при обработке JSON: {str(e)}"
                    
            elif file_lower.endswith(('.xml', '.html', '.htm')):
                # XML/HTML файлы
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    
                    # Для HTML попробуем извлечь только текст
                    if file_lower.endswith(('.html', '.htm')):
                        try:
                            from bs4 import BeautifulSoup
                            soup = BeautifulSoup(content, 'html.parser')
                            # Удаляем скрипты и стили
                            for script in soup(["script", "style"]):
                                script.decompose()
                            content = soup.get_text()
                            # Очищаем лишние пробелы
                            content = '\n'.join(line.strip() for line in content.splitlines() if line.strip())
                        except ImportError:
                            # Если BeautifulSoup не установлен, используем простое удаление тегов
                            import re
                            content = re.sub('<[^<]+?>', '', content)
                    
                    processed_content = self.rag_process_large_content(content)
                    wrapped = self._wrap_info_if_needed(processed_content, source="html" if file_lower.endswith(('.html', '.htm')) else "xml")
                    file_type = "HTML" if file_lower.endswith(('.html', '.htm')) else "XML"
                    return f"Содержимое {file_type} файла:\n\n{wrapped}"
                except Exception as e:
                    return f"Ошибка при обработке XML/HTML: {str(e)}"
                    
            else:
                return f"Неподдерживаемый формат файла: {file_path}"
                
        except Exception as e:
            return f"Ошибка при обработке документа: {str(e)}"

    def generate_docx_file(self, content: str, filename: str) -> str:
        """
        Генерация DOCX файла
        
        Args:
            content: Содержимое документа
            filename: Полное имя файла (с расширением или без)
        
        Returns:
            Сообщение о результате создания файла
        """
        try:
            if not DOCX_AVAILABLE or Document is None:
                return "Ошибка: Библиотека python-docx не установлена"
            
            # Создаем документ
            doc = Document()
            
            # Разделяем контент на параграфы
            paragraphs = content.split('\n')
            
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    # Проверяем, является ли строка заголовком (начинается с #)
                    if paragraph_text.strip().startswith('#'):
                        # Убираем # и создаем заголовок
                        title_text = paragraph_text.strip().lstrip('#').strip()
                        heading = doc.add_heading(title_text, level=1)
                    else:
                        doc.add_paragraph(paragraph_text.strip())
            
            # Создаем папку если не существует
            os.makedirs(os.path.join(self.base_dir, OUTPUT_DIR_NAME), exist_ok=True)
            
            # Сохраняем файл - убираем расширение если есть и добавляем .docx
            base_name = filename.replace('.docx', '').replace('.doc', '')
            output_path = os.path.join(self.base_dir, OUTPUT_DIR_NAME, f"{base_name}.docx")
            doc.save(output_path)
            
            return f"Документ успешно создан: {output_path}"
            
        except Exception as e:
            return f"Ошибка при создании DOCX файла: {str(e)}"

    def generate_excel_file(self, content: str, filename: str) -> str:
        """
        Генерация Excel файла
        
        Args:
            content: Содержимое в формате таблицы (разделители - табуляция или запятые)
            filename: Полное имя файла (с расширением или без)
        
        Returns:
            Сообщение о результате создания файла
        """
        try:
            if not EXCEL_AVAILABLE or pd is None:
                return "Ошибка: Библиотеки pandas/openpyxl не установлены"
            
            # Пытаемся разобрать контент как табличные данные
            lines = content.strip().split('\n')
            if not lines:
                return "Ошибка: Пустое содержимое для Excel файла"
            
            # Определяем разделитель (табуляция или запятая)
            delimiter = '\t' if '\t' in lines[0] else ','
            
            # Создаем DataFrame
            import io
            data_string = '\n'.join(lines)
            df = pd.read_csv(io.StringIO(data_string), delimiter=delimiter)
            
            # Сохраняем файл - убираем расширение если есть и добавляем .xlsx
            base_name = filename.replace('.xlsx', '').replace('.xls', '')
            output_path = os.path.join(self.base_dir, OUTPUT_DIR_NAME, f"{base_name}.xlsx")
            df.to_excel(output_path, index=False)
            
            return f"Excel файл успешно создан: {output_path}"
            
        except Exception as e:
            return f"Ошибка при создании Excel файла: {str(e)}"

    def generate_markdown_file(self, content: str, filename: str) -> str:
        """
        Генерация Markdown файла
        
        Args:
            content: Содержимое в формате Markdown
            filename: Полное имя файла (с расширением или без)
        
        Returns:
            Сообщение о результате создания файла
        """
        try:
            # Сохраняем файл - убираем расширение если есть и добавляем .md
            base_name = filename.replace('.md', '').replace('.markdown', '')
            output_path = os.path.join(self.base_dir, OUTPUT_DIR_NAME, f"{base_name}.md")
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return f"Markdown файл успешно создан: {output_path}"
            
        except Exception as e:
            return f"Ошибка при создании Markdown файла: {str(e)}"

    def generate_pdf_file(self, content: str, filename: str) -> str:
        """
        Генерация PDF файла
        
        Args:
            content: Содержимое документа
            filename: Полное имя файла (с расширением или без)
        
        Returns:
            Сообщение о результате создания файла
        """
        try:
            if not REPORTLAB_AVAILABLE:
                return "Ошибка: Библиотека reportlab не установлена"
            
            # Убираем расширение если есть и добавляем .pdf
            base_name = filename.replace('.pdf', '')
            os.makedirs(os.path.join(self.base_dir, OUTPUT_DIR_NAME), exist_ok=True)
            output_path = os.path.join(self.base_dir, OUTPUT_DIR_NAME, f"{base_name}.pdf")
            
            # Создаем PDF документ - импорт во время выполнения
            from reportlab.lib.pagesizes import A4  # type: ignore
            from reportlab.lib.styles import getSampleStyleSheet  # type: ignore
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer  # type: ignore
            
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Разделяем контент на параграфы
            paragraphs = content.split('\n')
            
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    # Проверяем, является ли строка заголовком
                    if paragraph_text.strip().startswith('#'):
                        title_text = paragraph_text.strip().lstrip('#').strip()
                        p = Paragraph(title_text, styles['Heading1'])
                    else:
                        p = Paragraph(paragraph_text.strip(), styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 12))
            
            # Строим документ
            doc.build(story)
            
            return f"PDF файл успешно создан: {output_path}"
            
        except Exception as e:
            return f"Ошибка при создании PDF файла: {str(e)}"

    def generate_txt_file(self, content: str, filename: str) -> str:
        """
        Генерация простого текстового файла
        
        Args:
            content: Содержимое файла
            filename: Полное имя файла (с расширением или без)
        
        Returns:
            Сообщение о результате создания файла
        """
        try:
            # Убираем расширение если есть и добавляем .txt
            base_name = filename.replace('.txt', '')
            output_path = os.path.join(self.base_dir, OUTPUT_DIR_NAME, f"{base_name}.txt")
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return f"Текстовый файл успешно создан: {output_path}"
            
        except Exception as e:
            return f"Ошибка при создании текстового файла: {str(e)}"

    def generate_json_file(self, content: str, filename: str) -> str:
        """
        Генерация JSON файла
        
        Args:
            content: Содержимое файла в формате JSON (строка)
            filename: Полное имя файла (с расширением или без)
        
        Returns:
            Сообщение о результате создания файла
        """
        try:
            # Убираем расширение если есть и добавляем .json
            base_name = filename.replace('.json', '')
            output_path = os.path.join(self.base_dir, OUTPUT_DIR_NAME, f"{base_name}.json")
            
            # Проверяем валидность JSON
            import json
            try:
                json.loads(content)  # Проверяем что content - валидный JSON
            except json.JSONDecodeError:
                # Если не JSON, оборачиваем в кавычки как строку
                content = json.dumps(content, ensure_ascii=False, indent=2)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return f"JSON файл успешно создан: {output_path}"
            
        except Exception as e:
            return f"Ошибка при создании JSON файла: {str(e)}"

    def generate_csv_file(self, content: str, filename: str) -> str:
        """
        Генерация CSV файла
        
        Args:
            content: Содержимое файла в формате CSV (строка с разделителями)
            filename: Полное имя файла (с расширением или без)
        
        Returns:
            Сообщение о результате создания файла
        """
        try:
            # Убираем расширение если есть и добавляем .csv
            base_name = filename.replace('.csv', '')
            output_path = os.path.join(self.base_dir, OUTPUT_DIR_NAME, f"{base_name}.csv")
            
            with open(output_path, 'w', encoding='utf-8', newline='') as f:
                f.write(content)
            
            return f"CSV файл успешно создан: {output_path}"
            
        except Exception as e:
            return f"Ошибка при создании CSV файла: {str(e)}"

    def generate_html_file(self, content: str, filename: str) -> str:
        """
        Генерация HTML файла
        
        Args:
            content: Содержимое файла в формате HTML
            filename: Полное имя файла (с расширением или без)
        
        Returns:
            Сообщение о результате создания файла
        """
        try:
            # Убираем расширение если есть и добавляем .html
            base_name = filename.replace('.html', '').replace('.htm', '')
            output_path = os.path.join(self.base_dir, OUTPUT_DIR_NAME, f"{base_name}.html")
            
            # Если контент не содержит HTML структуру, добавляем базовую
            if not content.strip().lower().startswith('<!doctype') and not content.strip().lower().startswith('<html'):
                content = f"""<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{base_name}</title>
</head>
<body>
{content}
</body>
</html>"""
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return f"HTML файл успешно создан: {output_path}"
            
        except Exception as e:
            return f"Ошибка при создании HTML файла: {str(e)}"

    def generate_xml_file(self, content: str, filename: str) -> str:
        """
        Генерация XML файла
        
        Args:
            content: Содержимое файла в формате XML
            filename: Полное имя файла (с расширением или без)
        
        Returns:
            Сообщение о результате создания файла
        """
        try:
            # Убираем расширение если есть и добавляем .xml
            base_name = filename.replace('.xml', '')
            output_path = os.path.join(self.base_dir, OUTPUT_DIR_NAME, f"{base_name}.xml")
            
            # Если контент не содержит XML декларацию, добавляем
            if not content.strip().startswith('<?xml'):
                content = f'<?xml version="1.0" encoding="UTF-8"?>\n{content}'
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return f"XML файл успешно создан: {output_path}"
            
        except Exception as e:
            return f"Ошибка при создании XML файла: {str(e)}"

    def generate_bat_file(self, content: str, filename: str) -> str:
        """
        Создает .bat файл с командами для Windows
        
        Args:
            content: Содержимое bat файла (команды)
            filename: Имя файла (с расширением .bat или без)
            
        Returns:
            Сообщение о результате создания файла
        """
        try:
            # Обеспечиваем правильное расширение
            if not filename.lower().endswith('.bat'):
                filename += '.bat'
            
            # Путь для сохранения в папку output
            output_path = os.path.join(os.getcwd(), OUTPUT_DIR_NAME, filename)
            
            # Убеждаемся что папка output существует
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Добавляем @echo off в начало, если его нет
            if not content.strip().startswith('@echo off'):
                content = '@echo off\n' + content.strip()
            
            # Добавляем pause в конец, если его нет
            if not content.strip().endswith('pause'):
                content = content.strip() + '\npause'
            
            with open(output_path, 'w', encoding='cp1251') as f:  # cp1251 для Windows bat файлов
                f.write(content)
            
            self.logger.info(f"📄 BAT файл создан: {filename}")
            return f"BAT файл успешно создан: {output_path}"
            
        except Exception as e:
            return f"Ошибка при создании BAT файла: {str(e)}"

    def run_bat_file(self, file_path: str, working_dir: Optional[str] = None) -> Dict[str, Any]:
        """
        Запускает .bat файл
        
        Args:
            file_path: Путь к .bat файлу
            working_dir: Рабочая директория для выполнения (опционально)
            
        Returns:
            Результат выполнения с выводом и кодом возврата
        """
        try:
            # Разрешаем путь к файлу
            resolved_path = self.resolve_path(file_path)
            
            if not os.path.exists(resolved_path):
                return {
                    "success": False,
                    "error": f"BAT файл не найден: {resolved_path}",
                    "output": "",
                    "return_code": -1
                }
            
            # Определяем рабочую директорию
            if working_dir is None:
                working_dir = os.path.dirname(resolved_path)
            else:
                working_dir = self.resolve_path(working_dir)
            
            self.logger.info(f"🚀 Запускаю BAT файл: {os.path.basename(resolved_path)}")
            
            # Запускаем bat файл
            result = subprocess.run(
                [resolved_path],
                cwd=working_dir,
                capture_output=True,
                text=True,
                encoding='cp1251',  # Кодировка для Windows
                timeout=300  # Таймаут 5 минут
            )
            
            success = result.returncode == 0
            
            if success:
                self.logger.info(f"✅ BAT файл выполнен успешно")
            else:
                self.logger.warning(f"⚠️ BAT файл завершился с кодом: {result.returncode}")
            
            return {
                "success": success,
                "output": result.stdout,
                "error": result.stderr,
                "return_code": result.returncode,
                "working_dir": working_dir
            }
            
        except subprocess.TimeoutExpired:
            return {
                "success": False,
                "error": "Превышено время ожидания выполнения BAT файла (5 минут)",
                "output": "",
                "return_code": -1
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"Ошибка при выполнении BAT файла: {str(e)}",
                "output": "",
                "return_code": -1
            }

    def generate_file(self, content: str, filename: str, file_format: str) -> bool:
        """
        Универсальный метод генерации файлов
        
        Args:
            content: Содержимое файла
            filename: Полное имя файла с расширением
            file_format: Формат файла (docx, excel, md, pdf, txt, json, csv, html, xml, bat)
        
        Returns:
            True если файл создан успешно, False иначе
        """
        try:
            # Создаем папку output если её нет
            output_dir = os.path.join(self.base_dir, OUTPUT_DIR_NAME)
            os.makedirs(output_dir, exist_ok=True)
            
            format_lower = file_format.lower()
            
            if format_lower in ['docx', 'doc', 'word']:
                result = self.generate_docx_file(content, filename)
                return "успешно создан" in result.lower()
            elif format_lower in ['excel', 'xlsx', 'xls']:
                result = self.generate_excel_file(content, filename)
                return "успешно создан" in result.lower()
            elif format_lower in ['md', 'markdown']:
                result = self.generate_markdown_file(content, filename)
                return "успешно создан" in result.lower()
            elif format_lower in ['pdf']:
                result = self.generate_pdf_file(content, filename)
                return "успешно создан" in result.lower()
            elif format_lower in ['txt', 'text']:
                result = self.generate_txt_file(content, filename)
                return "успешно создан" in result.lower()
            elif format_lower in ['json']:
                result = self.generate_json_file(content, filename)
                return "успешно создан" in result.lower()
            elif format_lower in ['csv']:
                result = self.generate_csv_file(content, filename)
                return "успешно создан" in result.lower()
            elif format_lower in ['html', 'htm']:
                result = self.generate_html_file(content, filename)
                return "успешно создан" in result.lower()
            elif format_lower in ['xml']:
                result = self.generate_xml_file(content, filename)
                return "успешно создан" in result.lower()
            elif format_lower in ['bat', 'batch']:
                result = self.generate_bat_file(content, filename)
                return "успешно создан" in result.lower()
            else:
                logger.error(f"Неподдерживаемый формат файла: {file_format}")
                return False
                
        except Exception as e:
            logger.error(f"Ошибка при генерации файла: {str(e)}")
            return False

    def extract_text_from_image(self, image_path: str) -> Tuple[str, str]:
        """
        Извлекает текст из изображения с помощью OCR
        
        Args:
            image_path: Путь к изображению
        
        Returns:
            Tuple[str, str]: (extracted_text, error_message)
        """
        try:
            if not self._ensure_ocr_initialized():
                return "", "OCR не может быть инициализирован"
            
            if not os.path.exists(image_path):
                return "", f"Файл {image_path} не найден"
            
            logger.info(f"📖 Извлекаю текст из изображения: {image_path}")
            
            # Выполняем OCR (дополнительная проверка для типизации)
            if self.ocr_reader is None:
                return "", "OCR reader не инициализирован после попытки загрузки"
            results = self.ocr_reader.readtext(image_path)
            
            if not results:
                return "", "Текст на изображении не обнаружен"
            
            # Объединяем весь найденный текст
            extracted_text = []
            for (bbox, text, confidence) in results:
                # Фильтруем результаты с низкой уверенностью
                try:
                    confidence_value = float(confidence)
                except (TypeError, ValueError):
                    confidence_value = 0.0
                if confidence_value > 0.3:  # Порог уверенности 30%
                    extracted_text.append(text.strip())
            
            if not extracted_text:
                return "", "Текст не обнаружен с достаточной уверенностью"
            
            final_text = '\n'.join(extracted_text)
            logger.info(f"✅ Извлечено {len(extracted_text)} блоков текста")
            
            return final_text, ""
            
        except Exception as e:
            error_msg = f"Ошибка OCR: {str(e)}"
            logger.error(error_msg)
            return "", error_msg

    def extract_text_from_image_object(self, image_obj) -> Tuple[str, str]:
        """
        Извлекает текст из PIL.Image объекта с помощью OCR
        
        Args:
            image_obj: PIL.Image объект
        
        Returns:
            Tuple[str, str]: (extracted_text, error_message)
        """
        try:
            if not self._ensure_ocr_initialized():
                return "", "OCR не может быть инициализирован"
            
            import numpy as np
            
            # Конвертируем PIL Image в numpy array для EasyOCR
            image_array = np.array(image_obj)
            
            logger.info(f"📖 Извлекаю текст из изображения (объект)")
            
            # Выполняем OCR (дополнительная проверка для типизации)
            if self.ocr_reader is None:
                return "", "OCR reader не инициализирован после попытки загрузки"
            results = self.ocr_reader.readtext(image_array)
            
            if not results:
                return "", "Текст на изображении не обнаружен"
            
            # Объединяем весь найденный текст
            extracted_text = []
            for (bbox, text, confidence) in results:
                # Фильтруем результаты с низкой уверенностью
                try:
                    confidence_value = float(confidence)
                except (TypeError, ValueError):
                    confidence_value = 0.0
                if confidence_value > 0.3:  # Порог уверенности 30%
                    extracted_text.append(text.strip())
            
            if not extracted_text:
                return "", "Текст не обнаружен с достаточной уверенностью"
            
            final_text = '\n'.join(extracted_text)
            logger.info(f"✅ Извлечено {len(extracted_text)} блоков текста")
            
            return final_text, ""
            
        except Exception as e:
            error_msg = f"Ошибка OCR объекта: {str(e)}"
            logger.error(error_msg)
            return "", error_msg

    def should_use_ocr_on_image(self, vision_description: str) -> bool:
        """
        Определяет, нужно ли применять OCR на основе описания изображения.
        Теперь использует интеллектуальную логику вместо простого поиска ключевых слов.
        
        Args:
            vision_description: Описание изображения от vision модели
            
        Returns:
            bool: True если нужно применить OCR, False если нет
        """
        return self.should_use_ocr_intelligently(vision_description)

    def process_image_with_smart_ocr(self, image_path: str, vision_description: str = "", force_ocr: bool = False) -> Tuple[str, str, str]:
        """
        Обрабатывает изображение с умным применением OCR
        
        Args:
            image_path: Путь к изображению
            vision_description: Описание от vision модели (опционально)
            force_ocr: Принудительно применить OCR
        
        Returns:
            Tuple[str, str, str]: (extracted_text, description, error_message)
        """
        extracted_text = ""
        ocr_error = ""
        
        # Решаем, применять ли OCR
        should_ocr = force_ocr or self.should_use_ocr_on_image(vision_description)
        
        if should_ocr:
            logger.info("🔍 Применяю OCR к изображению")
            text, error = self.extract_text_from_image(image_path)
            if text:
                extracted_text = f"\n\n[Текст с изображения]:\n{text}"
                logger.info(f"✅ Извлечен текст: {len(text)} символов")
            elif error:
                ocr_error = error
                logger.warning(f"⚠️ OCR не удался: {error}")
        else:
            logger.info("⏭️ OCR не применяется - текст не обнаружен в описании")
        
        # Объединяем описание и извлеченный текст
        full_description = vision_description
        if extracted_text:
            full_description += extracted_text
        
        return extracted_text, full_description, ocr_error

    def auto_disable_tools(self, tool_name: Optional[str] = None):
        """Автоматически выключает инструмент через заданное время после использования"""
        import threading
        import time
        import gc
        
        def disable_tool(tool_name):
            time.sleep(self.auto_disable_delay)
            
            if tool_name == 'image_generation':
                if hasattr(self, 'use_image_generation') and self.use_image_generation:
                    self.use_image_generation = False
                    
                    # Выгружаем pipeline из памяти
                    if hasattr(self, 'current_pipeline') and self.current_pipeline is not None:
                        try:
                            # Освобождаем GPU память
                            if hasattr(self.current_pipeline, 'to'):
                                self.current_pipeline.to('cpu')
                            del self.current_pipeline
                            self.current_pipeline = None
                        except Exception as e:
                            self.logger.warning(f"⚠️ Ошибка при выгрузке pipeline: {e}")
                    
                    # Принудительная очистка GPU памяти
                    try:
                        import torch
                        if torch.cuda.is_available():
                            torch.cuda.empty_cache()
                            torch.cuda.synchronize()
                    except Exception as e:
                        self.logger.warning(f"⚠️ Ошибка при очистке GPU памяти: {e}")
                    
                    # Сборка мусора
                    gc.collect()
                    self.logger.info(f"🔧 Автоматически выключил {tool_name} и освободил память")
                    
            elif tool_name == 'vision':
                if hasattr(self, 'use_vision') and self.use_vision:
                    self.use_vision = False
                    
                    # Выгружаем vision модели если они есть
                    vision_attrs = ['vision_model', 'vision_processor', 'vision_pipeline']
                    for attr in vision_attrs:
                        if hasattr(self, attr):
                            try:
                                model = getattr(self, attr)
                                if model is not None and hasattr(model, 'to'):
                                    model.to('cpu')
                                delattr(self, attr)
                            except Exception as e:
                                self.logger.warning(f"⚠️ Ошибка при выгрузке {attr}: {e}")
                    
                    # Очистка GPU памяти
                    try:
                        import torch
                        if torch.cuda.is_available():
                            torch.cuda.empty_cache()
                    except Exception:
                        pass
                    
                    gc.collect()
                    self.logger.info(f"🔧 Автоматически выключил {tool_name} и освободил память")
                    self._unload_vision_model()
                    
            elif tool_name == 'audio':
                if hasattr(self, 'use_audio') and self.use_audio:
                    self.use_audio = False
                    
                    # Выгружаем audio модели если они есть
                    audio_attrs = ['whisper_model', 'audio_model', 'tts_model']
                    for attr in audio_attrs:
                        if hasattr(self, attr):
                            try:
                                model = getattr(self, attr)
                                if model is not None and hasattr(model, 'to'):
                                    model.to('cpu')
                                delattr(self, attr)
                            except Exception as e:
                                self.logger.warning(f"⚠️ Ошибка при выгрузке {attr}: {e}")
                    
                    # Очистка GPU памяти
                    try:
                        import torch
                        if torch.cuda.is_available():
                            torch.cuda.empty_cache()
                    except Exception:
                        pass
                    
                    gc.collect()
                    self.logger.info(f"🔧 Автоматически выключил {tool_name} и освободил память")
        
        # Если указан конкретный инструмент, запускаем таймер только для него
        if tool_name:
            # Отменяем предыдущий таймер если он есть
            if tool_name in self.tool_timers and self.tool_timers[tool_name].is_alive():
                self.tool_timers[tool_name].cancel() if hasattr(self.tool_timers[tool_name], 'cancel') else None
                
            timer = threading.Thread(target=disable_tool, args=(tool_name,), daemon=True)
            self.tool_timers[tool_name] = timer
            timer.start()
            self.logger.info(f"⏰ Запустил таймер автоматического выключения для {tool_name}")
        else:
            # Запускаем таймеры для всех активных инструментов
            for tool in ['image_generation', 'vision', 'audio']:
                if tool not in self.tool_timers or not self.tool_timers[tool].is_alive():
                    timer = threading.Thread(target=disable_tool, args=(tool,), daemon=True)
                    self.tool_timers[tool] = timer
                    timer.start()
                    self.logger.info(f"⏰ Запустил таймер автоматического выключения для {tool}")
                
    def _log(self, message: str, level: str = "INFO"):
        """Логирование с временной меткой в файл и консоль"""
        timestamp = time.strftime("%H:%M:%S")
        formatted_message = f"[{timestamp}] {level}: {message}"
        
        # Логируем в файл
        if level == "ERROR":
            logger.error(message)
        elif level == "WARNING":
            logger.warning(message)
        else:
            logger.info(message)
        
        # Выводим в консоль
        print(formatted_message)
    
    def get_context_info(self) -> str:
        """Возвращает информацию о текущем использовании контекста"""
        return f"Контекст: {self.current_context_length:,} токенов / {self.safe_context_length:,} (безопасный) / {self.max_context_length:,} (максимум)"

    def add_performance_metric(self, action: str, response_time: float, context_length: int = 0):
        """Добавляет метрику производительности"""
        metric = {
            "timestamp": time.time(),
            "action": action,
            "response_time": response_time,
            "context_length": context_length
        }
        self.performance_metrics.append(metric)
        
        # Ограничиваем количество метрик
        if len(self.performance_metrics) > 100:
            self.performance_metrics.pop(0)
    
    def get_performance_stats(self) -> Dict[str, Any]:
        """Возвращает статистику производительности"""
        if not self.performance_metrics:
            return {"total_actions": 0, "avg_response_time": 0, "recent_metrics": []}
        
        total_actions = len(self.performance_metrics)
        avg_response_time = sum(m["response_time"] for m in self.performance_metrics) / total_actions
        recent_metrics = self.performance_metrics[-10:]  # Последние 10 метрик
        
        return {
            "total_actions": total_actions,
            "avg_response_time": round(avg_response_time, 3),
            "recent_metrics": recent_metrics
        }

    def take_screenshot(self) -> str:
        """
        Делает скриншот основного монитора и возвращает base64
        """
        try:
            # mss уже импортирован в начале файла
            with mss.mss() as sct:
                # Скриншот основного монитора
                monitor = sct.monitors[1]  # 0 - все мониторы, 1 - основной
                screenshot = sct.grab(monitor)
                # Конвертируем в PIL Image
                img = Image.frombytes("RGB", screenshot.size, screenshot.rgb)
                # Сжимаем для экономии места
                img.thumbnail((1280, 720), Image.Resampling.LANCZOS)
                # Конвертируем в base64
                buf = BytesIO()
                img.save(buf, format="PNG", optimize=True)
                return base64.b64encode(buf.getvalue()).decode("ascii")
        except ImportError:
            logger.warning("mss не установлен, используем pyautogui")
            try:
                # pyautogui уже импортирован в начале файла
                screenshot = pyautogui.screenshot()
                screenshot.thumbnail((1280, 720), Image.Resampling.LANCZOS)
                buf = BytesIO()
                screenshot.save(buf, format="PNG", optimize=True)
                return base64.b64encode(buf.getvalue()).decode("ascii")
            except ImportError:
                logger.error("pyautogui не установлен")
                return ""
        except Exception as e:
            logger.error(f"Ошибка создания скриншота: {e}")
            return ""

    def resolve_path(self, path: str) -> str:
        """
        Разрешает относительный путь к изображению
        """
        if os.path.isabs(path):
            return path
        
        # Проверяем в базовой директории
        full_path = os.path.join(self.base_dir, path)
        if os.path.exists(full_path):
            return full_path
            
        # Проверяем в папке Photos
        photos_path = os.path.join(self.base_dir, "Photos", path)
        if os.path.exists(photos_path):
            return photos_path
            
        # Проверяем в папке Images
        images_path = os.path.join(self.base_dir, "Images", path)
        if os.path.exists(images_path):
            return images_path
            
        return path  # Возвращаем исходный путь если не найден

    def analyze_image_with_vision(self, image_path: str) -> str:
        """
        Анализирует изображение с помощью vision модели
        """
        try:
            # Конвертируем изображение в base64
            image_b64 = image_to_base64_balanced(image_path)
            if not image_b64:
                return "Ошибка: не удалось загрузить изображение"
            
            # Автоматически включаем vision если не включен
            if not getattr(self, 'use_vision', False):
                self.logger.info("🔍 Автоматически включаю анализ изображений...")
                self.use_vision = True
                self.auto_disable_tools("vision")
            
            # Вызываем vision модель
            return self.call_vision_model(image_b64)
            
        except Exception as e:
            self.logger.error(f"Ошибка анализа изображения: {e}")
            return f"Ошибка анализа изображения: {str(e)}"

    # Mouse/Keyboard methods moved to automation.py
    # move_mouse, left_click, right_click, scroll, mouse_down, mouse_up, drag_and_drop, type_text

    def start_continuous_recording(self) -> bool:
        """Запуск постоянной голосовой записи."""
        if self.continuous_recording:
            logger.info("Постоянная голосовая запись уже активна")
            return True

        try:
            self.continuous_recording = True
            self.recording_thread = threading.Thread(target=self._continuous_recording_worker, daemon=True)
            self.recording_thread.start()
            logger.info("Постоянная голосовая запись запущена")
            return True
        except Exception as e:
            self.continuous_recording = False
            logger.error(f"Ошибка запуска постоянной голосовой записи: {e}")
            return False

    def stop_continuous_recording(self) -> bool:
        """Остановка постоянной голосовой записи."""
        try:
            if not self.continuous_recording:
                logger.info("Постоянная голосовая запись уже остановлена")
                return True

            self.continuous_recording = False
            if self.recording_thread:
                self.recording_thread.join(timeout=2)
            logger.info("Постоянная голосовая запись остановлена")
            return True
        except Exception as e:
            logger.error(f"Ошибка остановки постоянной голосовой записи: {e}")
            return False

    def _continuous_recording_worker(self):
        """Воркер для постоянной голосовой записи (заглушка - нужна реализация через веб-интерфейс)"""
        # Эта функция будет вызываться из веб-интерфейса через API
        while self.continuous_recording:
            try:
                # Проверяем очередь на наличие аудиочанков
                if not self.audio_queue.empty():
                    audio_data = self.audio_queue.get_nowait()
                    # Обрабатываем аудио
                    self._process_audio_chunk(audio_data)
                time.sleep(0.1)
            except Exception as e:
                logger.error(f"Ошибка в continuous recording worker: {e}")

    def _process_audio_chunk(self, audio_source: Union[bytes, str], lang: str = "ru", cleanup_source: bool = False) -> Optional[Dict[str, Any]]:
        """Обработка чанка аудио из постоянной записи или веб-API.

        Поддерживает как байты аудио, так и путь к уже сохраненному временному файлу.
        """
        temp_file: Optional[str] = None
        source_path: Optional[str] = None
        try:
            if isinstance(audio_source, (bytes, bytearray)):
                temp_dir = os.path.join(os.path.dirname(__file__), "temp_audio")
                os.makedirs(temp_dir, exist_ok=True)
                temp_file = os.path.join(temp_dir, f"chunk_{int(time.time())}.wav")
                with open(temp_file, 'wb') as f:
                    f.write(audio_source)
                source_path = temp_file
            elif isinstance(audio_source, str):
                source_path = audio_source
            else:
                logger.error("Неподдерживаемый тип аудио для _process_audio_chunk")
                return None

            if not source_path or not os.path.exists(source_path):
                logger.error(f"Аудиоисточник не найден: {source_path}")
                return None
            
            # Распознаём аудио
            transcript = self.transcribe_audio_whisper(source_path, lang=lang, use_separator=False)
            
            if transcript and not transcript.startswith("[Whisper error]"):
                result_payload: Dict[str, Any] = {
                    "text": transcript,
                    "continue": False,
                    "response": ""
                }

                # Проверяем, содержит ли текст команду или имя "Алиса"
                if self._is_valid_command(transcript):
                    logger.info(f"Получена команда из голоса: {transcript}")
                    # Делаем скриншот для контекста
                    screenshot_b64 = self.take_screenshot()
                    vision_desc = ""
                    if screenshot_b64:
                        vision_desc = self.call_vision_model(screenshot_b64)
                    
                    # Формируем запрос для мозга
                    safe_transcript = self._wrap_info_if_needed(transcript, source="audio")
                    brain_input = f"[Скриншот экрана]: {vision_desc}\n\nГолосовая команда: {safe_transcript}"
                    
                    # Отправляем в мозг
                    ai_response = self.call_brain_model(brain_input)
                    continue_dialog = self.process_ai_response(ai_response)
                    result_payload["continue"] = continue_dialog
                    result_payload["response"] = getattr(self, 'last_final_response', '')
                    return result_payload
                else:
                    # Игнорируем бессмысленные фразы
                    result_payload["response"] = transcript
                    return result_payload

            logger.warning("Не удалось получить валидную транскрипцию аудиочанка")
            return None
                
        except Exception as e:
            logger.error(f"Ошибка обработки аудиочанка: {e}")
            return None
        finally:
            cleanup_targets = []
            if temp_file:
                cleanup_targets.append(temp_file)
            if cleanup_source and isinstance(audio_source, str):
                cleanup_targets.append(audio_source)

            for cleanup_target in cleanup_targets:
                try:
                    if cleanup_target and os.path.exists(cleanup_target):
                        os.remove(cleanup_target)
                except Exception as cleanup_error:
                    logger.debug(f"Не удалось удалить временный аудиофайл {cleanup_target}: {cleanup_error}")

    def _is_valid_command(self, text: str) -> bool:
        """Всегда возвращает True — фильтрация отключена, все команды проходят к нейросети"""
        return True

    def call_vision_model(self, image_base64: str) -> str:
        """
        Отправка изображения только в vision-модель ("глаза")
        Возвращает описание изображения (текст).
        """
        start_time = time.time()
        
        # Автоматически включаем vision модель при необходимости
        if not getattr(self, 'use_vision', False):
            logger.info("🔧 Автоматически включаю vision модель")
            self.use_vision = True
            # Запускаем таймер автоматического выключения
            self.auto_disable_tools("vision")
        
        try:
            self._ensure_vision_model_loaded()

            vision_prompt = ""
            if getattr(self, "prompt_loader", None) is not None:
                try:
                    vision_prompt = self.prompt_loader.load_vision_prompt()
                except Exception as exc:
                    logger.debug(f"Vision prompt load error: {exc}")

            messages: List[Dict[str, Any]] = []
            if vision_prompt:
                messages.append({
                    "role": "system",
                    "content": [{"type": "text", "text": vision_prompt}]
                })

            user_content: List[Dict[str, Any]] = []
            if vision_prompt:
                user_content.append({
                    "type": "text",
                    "text": "Проанализируй изображение согласно инструкциям."
                })
            user_content.append({
                "type": "image_url",
                "image_url": {"url": f"data:image/png;base64,{image_base64}"}
            })

            messages.append({"role": "user", "content": user_content})

            payload: Dict[str, Any] = {
                "model": self.vision_model_id or "moondream2-llamafile",
                "messages": messages,
            }
            for key, value in self.vision_generation_params.items():
                if value is not None:
                    payload[key] = value
            logger.info("Отправляю изображение в vision-модель (глаза)")
            response = requests.post(
                f"{self.lm_studio_url}/v1/chat/completions",
                json=payload,
                headers={"Content-Type": "application/json"}
            )
            if response.status_code == 200:
                result = response.json()
                return result["choices"][0]["message"]["content"].strip()
            else:
                error_msg = f"Ошибка vision-модели: {response.status_code} - {response.text}"
                logger.error(error_msg)
                return f"[Vision error] {error_msg}"
        except Exception as e:
            error_msg = f"Исключение vision: {str(e)}"
            logger.error(error_msg)
            return f"[Vision error] {error_msg}"
        finally:
            # Записываем метрику производительности
            response_time = time.time() - start_time
            self.add_performance_metric("vision_processing", response_time)
            logger.info(f"👁️ Vision обработал за {response_time:.2f} сек")

    def call_brain_model(self, user_message: str, vision_desc: str = "") -> str:
        """
        Отправка текстового запроса (и опционально описания изображения) в "мозг" (текстовая модель)
        Поддерживает как LM Studio, так и llama.cpp режимы
        """
        start_time = time.time()
        # Инициализируем переменные для использования в except блоке
        processed_message = user_message
        messages = []

        logger.info(f"🧠 call_brain_model старт, длина сообщения {len(user_message)} символов")
        
        try:
            # Проверяем, является ли сообщение командой для загрузки модуля
            if self.prompt_loader.is_module_command(user_message.strip()):
                module_content = self.prompt_loader.load_module(user_message.strip())
                logger.info(f"📚 Загружен модуль для команды: {user_message.strip()}")
                return module_content
            
            # Обрабатываем сообщение плагинами (hook on_message_received)
            processed_message = user_message
            if self.plugin_manager:
                processed_message = self.plugin_manager.call_hook_message_received(user_message, self)
            
            # Улучшаем промпт с помощью памяти ChromaDB
            enhanced_system_prompt = self.enhance_prompt_with_memory(processed_message, self.system_prompt)
            
            # Добавляем информацию о доступных плагинах в системный промпт
            if self.plugin_manager:
                plugin_info = self._get_plugin_info_for_prompt()
                if plugin_info:
                    enhanced_system_prompt += f"\n\n{plugin_info}"
            
            messages: List[Dict[str, Any]] = [
                {"role": "system", "content": enhanced_system_prompt}
            ]
            
            # Добавляем историю разговора с управлением контекстом
            messages.extend(self.conversation_history)
            
            # Добавляем описание изображения, если есть
            if vision_desc:
                messages.append({"role": "user", "content": vision_desc})
            # Добавляем основной запрос пользователя
            messages.append({"role": "user", "content": user_message})
            
            # Динамическое управление контекстом - подсчитываем длину
            total_length = sum(len(str(msg.get("content", ""))) for msg in messages)
            self.current_context_length = total_length
            
            # Обрезаем контекст если необходимо (используем предварительную оценку)
            if total_length > self.safe_context_length:
                self.conversation_history = self.conversation_history[-5:]  # Оставляем только 5 последних сообщений
                logger.warning(f"Превышение безопасного контекста ({total_length:,} символов > {self.safe_context_length:,}) - аккуратная обрезка истории")
            
            # Пересобираем сообщения после обрезки
            messages = [{"role": "system", "content": enhanced_system_prompt}]
            messages.extend(self.conversation_history)
            if vision_desc:
                messages.append({"role": "user", "content": vision_desc})
            messages.append({"role": "user", "content": user_message})
            
            logger.info(f"Отправляю запрос в мозг: {user_message[:100]}...")
            
            # ===================================================================
            # РЕЖИМ LLAMA.CPP - прямая работа с моделью через Python API
            # ===================================================================
            if self.use_llama_cpp:
                if not self.llama_wrapper or not self.llama_wrapper.is_loaded():
                    error_msg = "Модель llama.cpp не загружена"
                    logger.error(f"❌ {error_msg}")
                    return f"[Brain error] {error_msg}"
                
                try:
                    # Добавляем детальное логирование
                    start_total = time.time()
                    logger.info(f"🚀 Вызываю llama.cpp для генерации...")
                    logger.info(f"📝 Параметры: messages={len(messages)}, temp={LLAMA_CPP_GENERATION_PARAMS['temperature']}, max_tokens={LLAMA_CPP_GENERATION_PARAMS['max_tokens']}")
                    
                    # Показываем размер каждого сообщения для диагностики
                    total_chars = sum(len(str(msg.get('content', ''))) for msg in messages)
                    logger.info(f"📊 Размер всех сообщений: {total_chars} символов, {len(messages)} сообщений")
                    for i, msg in enumerate(messages):
                        content_len = len(str(msg.get('content', '')))
                        logger.info(f"  Сообщение {i+1}: роль={msg.get('role', '?')}, длина={content_len} символов")
                    
                    # Засекаем время вызова
                    start_call = time.time()
                    logger.info("⏱️ НАЧАЛО ВЫЗОВА create_chat_completion...")
                    
                    # Вызываем llama.cpp для генерации с параметрами из начала файла
                    result = self.llama_wrapper.create_chat_completion(  # type: ignore
                        messages=messages,  # type: ignore
                        temperature=LLAMA_CPP_GENERATION_PARAMS['temperature'],
                        max_tokens=LLAMA_CPP_GENERATION_PARAMS['max_tokens'],
                        top_p=LLAMA_CPP_GENERATION_PARAMS['top_p'],
                        top_k=LLAMA_CPP_GENERATION_PARAMS['top_k'],
                        repeat_penalty=LLAMA_CPP_GENERATION_PARAMS['repeat_penalty'],
                        stream=LLAMA_CPP_GENERATION_PARAMS['stream']
                    )
                    
                    end_call = time.time()
                    call_duration = end_call - start_call
                    logger.info(f"⏱️ КОНЕЦ ВЫЗОВА create_chat_completion: {call_duration:.2f} секунд")
                    logger.info("✅ Генерация завершена успешно")
                    
                    ai_response = result["choices"][0]["message"]["content"].strip()
                    
                    # Извлекаем информацию о токенах
                    usage_info = result.get("usage", {})
                    prompt_tokens = usage_info.get("prompt_tokens", 0)
                    completion_tokens = usage_info.get("completion_tokens", 0)
                    total_tokens = usage_info.get("total_tokens", 0)
                    
                    if total_tokens > 0:
                        self.current_context_length = total_tokens
                        logger.info(f"📊 Токены (llama.cpp): prompt={prompt_tokens}, completion={completion_tokens}, total={total_tokens}")
                        self._trim_context_if_needed()
                    
                    # Добавляем в историю разговора
                    if ai_response and ai_response != "{}":
                        self.conversation_history.append({"role": "user", "content": processed_message})
                        self.conversation_history.append({"role": "assistant", "content": ai_response})
                        self.auto_save_conversation(processed_message, ai_response, vision_desc)
                        self.extract_preferences_from_response(processed_message, ai_response)
                    
                    # Обрабатываем ответ плагинами
                    final_response = ai_response
                    if self.plugin_manager:
                        final_response = self.plugin_manager.call_hook_response_generated(ai_response, self)
                    
                    return final_response
                    
                except Exception as llama_error:
                    error_msg = f"Ошибка llama.cpp: {str(llama_error)}"
                    logger.error(f"❌ {error_msg}")
                    
                    # Попытка переподключения
                    logger.info("🔄 Попытка переподключения к модели llama.cpp...")
                    if self._reconnect_brain_model():
                        try:
                            logger.info("🔄 Повторный запрос после переподключения...")
                            result = self.llama_wrapper.create_chat_completion(
                                messages=messages,
                                temperature=LLAMA_CPP_GENERATION_PARAMS['temperature'],
                                max_tokens=LLAMA_CPP_GENERATION_PARAMS['max_tokens'],
                                top_p=LLAMA_CPP_GENERATION_PARAMS['top_p'],
                                top_k=LLAMA_CPP_GENERATION_PARAMS['top_k'],
                                repeat_penalty=LLAMA_CPP_GENERATION_PARAMS['repeat_penalty'],
                                stream=LLAMA_CPP_GENERATION_PARAMS['stream']
                            )
                            ai_response = result["choices"][0]["message"]["content"].strip()
                            logger.info("✅ Повторный запрос после переподключения успешен")
                            
                            if ai_response and ai_response != "{}":
                                self.conversation_history.append({"role": "user", "content": processed_message})
                                self.conversation_history.append({"role": "assistant", "content": ai_response})
                                self.auto_save_conversation(processed_message, ai_response, vision_desc)
                                self.extract_preferences_from_response(processed_message, ai_response)
                            
                            final_response = ai_response
                            if self.plugin_manager:
                                final_response = self.plugin_manager.call_hook_response_generated(ai_response, self)
                            
                            return final_response
                        except Exception as retry_e:
                            logger.error(f"❌ Ошибка повторного запроса: {retry_e}")
                    
                    return f"[Brain error] {error_msg}"
            
            # ===================================================================
            # РЕЖИМ LM STUDIO - работа через HTTP API
            # ===================================================================
            payload = {
                "model": self.brain_model_id if hasattr(self, 'brain_model_id') and self.brain_model_id else self.brain_model,
                "messages": messages,
                "temperature": LLAMA_CPP_GENERATION_PARAMS['temperature'],
                "max_tokens": LLAMA_CPP_GENERATION_PARAMS['max_tokens'],
                "stream": LLAMA_CPP_GENERATION_PARAMS['stream']
            }
            
            response = requests.post(
                f"{self.lm_studio_url}/v1/chat/completions",
                json=payload,
                headers={"Content-Type": "application/json"}
            )
            if response.status_code == 200:
                result = response.json()
                ai_response = result["choices"][0]["message"]["content"].strip()
                
                # Извлекаем информацию о токенах из ответа модели
                usage_info = result.get("usage", {})
                prompt_tokens = usage_info.get("prompt_tokens", 0)
                completion_tokens = usage_info.get("completion_tokens", 0)
                total_tokens = usage_info.get("total_tokens", 0)
                
                # Обновляем текущий размер контекста на основе total_tokens
                if total_tokens > 0:
                    self.current_context_length = total_tokens
                    logger.info(f"📊 Токены (LM Studio): prompt={prompt_tokens}, completion={completion_tokens}, total={total_tokens}")
                    
                    # Обрезаем контекст на основе реальных токенов
                    self._trim_context_if_needed()
                
                # Добавляем в историю разговора (если ответ не пустой)
                if ai_response and ai_response != "{}":
                    self.conversation_history.append({"role": "user", "content": processed_message})
                    self.conversation_history.append({"role": "assistant", "content": ai_response})
                    
                    # Автоматически сохраняем диалог в ChromaDB
                    self.auto_save_conversation(processed_message, ai_response, vision_desc)
                    
                    # Извлекаем предпочтения пользователя из диалога
                    self.extract_preferences_from_response(processed_message, ai_response)
                
                # Обрабатываем ответ плагинами (hook on_response_generated)
                final_response = ai_response
                if self.plugin_manager:
                    final_response = self.plugin_manager.call_hook_response_generated(ai_response, self)
                
                return final_response
            else:
                error_msg = f"Ошибка brain-модели: {response.status_code} - {response.text}"
                logger.error(error_msg)
                
                # Попытка переподключения при ошибке
                if response.status_code in [404, 500, 503]:
                    logger.info("🔄 Попытка переподключения к модели мозга...")
                    if self._reconnect_brain_model():
                        # Повторяем запрос после переподключения
                        try:
                            retry_response = requests.post(
                                f"{self.lm_studio_url}/v1/chat/completions",
                                json=payload,
                                headers={"Content-Type": "application/json"}
                            )
                            if retry_response.status_code == 200:
                                result = retry_response.json()
                                ai_response = result["choices"][0]["message"]["content"].strip()
                                logger.info("✅ Повторный запрос после переподключения успешен")
                                
                                # Добавляем в историю разговора
                                if ai_response and ai_response != "{}":
                                    self.conversation_history.append({"role": "user", "content": processed_message})
                                    self.conversation_history.append({"role": "assistant", "content": ai_response})
                                    self.auto_save_conversation(processed_message, ai_response, vision_desc)
                                    self.extract_preferences_from_response(processed_message, ai_response)
                                
                                # Обрабатываем ответ плагинами
                                final_response = ai_response
                                if self.plugin_manager:
                                    final_response = self.plugin_manager.call_hook_response_generated(ai_response, self)
                                
                                return final_response
                        except Exception as retry_e:
                            logger.error(f"❌ Ошибка повторного запроса: {retry_e}")
                
                return f"[Brain error] {error_msg}"
        except Exception as e:
            error_msg = f"Исключение brain: {str(e)}"
            logger.error(error_msg)
            
            # Попытка переподключения при исключении (может быть связано с обрывом соединения)
            if "Connection" in str(e) or "timeout" in str(e).lower() or "refused" in str(e).lower():
                logger.info("🔄 Попытка переподключения к модели мозга из-за проблем соединения...")
                if self._reconnect_brain_model():
                    # Повторяем запрос после переподключения
                    try:
                        if self.use_llama_cpp:
                            result = self.llama_wrapper.create_chat_completion(  # type: ignore
                                messages=messages,  # type: ignore
                                temperature=LLAMA_CPP_GENERATION_PARAMS['temperature'],
                                max_tokens=LLAMA_CPP_GENERATION_PARAMS['max_tokens'],
                                top_p=LLAMA_CPP_GENERATION_PARAMS['top_p'],
                                top_k=LLAMA_CPP_GENERATION_PARAMS['top_k'],
                                repeat_penalty=LLAMA_CPP_GENERATION_PARAMS['repeat_penalty'],
                                stream=LLAMA_CPP_GENERATION_PARAMS['stream']
                            )
                        else:
                            payload = {
                                "model": self.brain_model_id if hasattr(self, 'brain_model_id') and self.brain_model_id else self.brain_model,
                                "messages": messages,
                                "temperature": LLAMA_CPP_GENERATION_PARAMS['temperature'],
                                "max_tokens": LLAMA_CPP_GENERATION_PARAMS['max_tokens'],
                                "stream": LLAMA_CPP_GENERATION_PARAMS['stream']
                            }
                            retry_response = requests.post(
                                f"{self.lm_studio_url}/v1/chat/completions",
                                json=payload,
                                headers={"Content-Type": "application/json"}
                            )
                            result = retry_response.json() if retry_response.status_code == 200 else None
                        
                        if result:
                            ai_response = result["choices"][0]["message"]["content"].strip()
                            logger.info("✅ Повторный запрос после переподключения успешен")
                            
                            # Добавляем в историю разговора
                            if ai_response and ai_response != "{}":
                                self.conversation_history.append({"role": "user", "content": processed_message})
                                self.conversation_history.append({"role": "assistant", "content": ai_response})
                                self.auto_save_conversation(processed_message, ai_response, vision_desc)
                                self.extract_preferences_from_response(processed_message, ai_response)
                            
                            # Обрабатываем ответ плагинами
                            final_response = ai_response
                            if self.plugin_manager:
                                final_response = self.plugin_manager.call_hook_response_generated(ai_response, self)
                            
                            return final_response
                    except Exception as retry_e:
                        logger.error(f"❌ Ошибка повторного запроса после исключения: {retry_e}")
            
            return f"[Brain error] {error_msg}"
        finally:
            # Записываем метрику производительности
            response_time = time.time() - start_time
            self.add_performance_metric("brain_response", response_time, self.current_context_length)
            logger.info(f"🧠 Мозг ответил за {response_time:.2f} сек")

    # execute_powershell moved to automation.py

    def google_search(self, query: str, num_results: int = 10) -> List[Dict[str, str]]:
        """
        Поиск в Google Custom Search API
        
        Args:
            query: Поисковый запрос
            num_results: Количество результатов для парсинга (по умолчанию 10)
            
        Returns:
            Список результатов поиска
        """
        try:
            if not self.google_api_key or not self.google_cse_id:
                return [{"error": "Google API ключ или CSE ID не настроены"}]
            
            logger.info(f"Выполняю поиск Google: {query}")
            
            # Кодируем запрос для URL
            encoded_query = urllib.parse.quote(query)
            
            # Формируем URL для Google Custom Search API (максимум 10 результатов)
            url = f"https://www.googleapis.com/customsearch/v1?key={self.google_api_key}&cx={self.google_cse_id}&q={encoded_query}&num=10"
            
            response = requests.get(url, timeout=10)
            
            if response.status_code == 200:
                data = response.json()
                
                if "items" not in data:
                    return [{"error": "Результаты не найдены"}]
                
                # Берем первые num_results результатов для парсинга (максимум 10)
                actual_results = min(num_results, 10)
                search_results = []
                for i, item in enumerate(data["items"][:actual_results]):
                    result = {
                        "title": item.get("title", ""),
                        "url": item.get("link", ""),
                        "snippet": item.get("snippet", "")
                    }
                    
                    # Пытаемся получить содержимое страницы
                    try:
                        page_response = requests.get(result["url"], timeout=5, headers={
                            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
                        })
                        if page_response.status_code == 200:
                            # Берем первые 2000 символов текста для полного анализа
                            content = page_response.text[:2000]
                            result["content"] = content
                        else:
                            result["content"] = "Не удалось получить содержимое страницы"
                    except:
                        result["content"] = "Ошибка при получении содержимого страницы"
                    
                    search_results.append(result)
                    logger.info(f"Получен результат {i+1}: {result['title']}")
                
                logger.info(f"Поиск завершен: найдено {len(search_results)} результатов")
                return search_results
            else:
                error_msg = f"Ошибка Google Search API: {response.status_code}"
                logger.error(error_msg)
                return [{"error": error_msg}]
                
        except Exception as e:
            error_msg = f"Ошибка поиска Google: {str(e)}"
            logger.error(error_msg)
            return [{"error": error_msg}]

    def process_ai_response(self, ai_response: str) -> bool:
        """Light wrapper to keep `process_ai_response` simple for static analysers.

        Реализует прокси к подробной реализации `_process_ai_response_impl`.
        Это уменьшает сложность видимой функции для Pylance.
        """
        return self._process_ai_response_impl(ai_response)

    # --- Разделённые обработчики действий (уменьшают сложность основного метода) ---
    def _is_english_simple(self, s: str) -> bool:
        if not s:
            return False
        allowed_chars = 0
        total_chars = len(s)
        for c in s:
            if (c.isascii() and (c.isalpha() or c.isdigit()) or
                c in '.,;:-_=+!@#$%^&*()[]{}<>?/\\|`~\'\" ' or
                c.isspace()):
                allowed_chars += 1
        return allowed_chars / total_chars > 0.85

    def _handle_powershell(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        command = action_data.get("command", "")
        description = action_data.get("description", "")
        logger.info(f"\n🔧 ВЫПОЛНЕНИЕ КОМАНДЫ: {description}")
        logger.info(f"📝 Команда: {command}")
        result = execute_powershell(command, self.logger)
        if result["success"]:
            feedback = f"Команда '{command}' выполнена успешно. Результат: {result['output']}"
        else:
            feedback = f"Команда '{command}' завершилась с ошибкой: {result.get('error','') }"
        follow_up = self.call_brain_model(feedback)
        return follow_up

    def _handle_search(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        query = action_data.get("query", "")
        description = action_data.get("description", "")
        logger.info(f"\n🔍 ПОИСК В ИНТЕРНЕТЕ: {description}")
        logger.info(f"🔎 Запрос: {query}")
        search_results = self.google_search(query)
        results_text = "Результаты поиска:\n"
        for i, result in enumerate(search_results, 1):
            if "error" in result:
                results_text += f"{i}. Ошибка: {result['error']}\n"
            else:
                results_text += f"{i}. {result['title']}\n"
                results_text += f"   URL: {result['url']}\n"
                results_text += f"   Описание: {result['snippet']}\n"
                results_text += f"   Содержимое: {result.get('content', '')}\n\n"
        logger.info("✅ Поиск завершен")
        follow_up = self.call_brain_model(f"Результаты поиска по запросу '{query}': {results_text}")
        return follow_up

    def _handle_take_screenshot(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        logger.info(f"\n📸 СОЗДАНИЕ СКРИНШОТА")
        if not getattr(self, 'use_vision', False):
            logger.info("🔧 Автоматически включаю vision модель")
            self.use_vision = True
            self.auto_disable_tools("vision")
        screenshot_b64 = self.take_screenshot()
        if screenshot_b64:
            vision_desc = self.call_vision_model(screenshot_b64)
            feedback = f"Скриншот экрана получен. Описание от vision-модели: {vision_desc}"
        else:
            feedback = "Не удалось создать скриншот"
        follow_up = self.call_brain_model(feedback)
        return follow_up

    def _handle_move_mouse(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        x = action_data.get("x", 0)
        y = action_data.get("y", 0)
        description = action_data.get("description", "")
        logger.info(f"\n🖱️ ПЕРЕМЕЩЕНИЕ МЫШИ: {description}")
        result = move_mouse(x, y)
        feedback = f"Мышь перемещена в координаты ({x}, {y})" if result.get("success") else f"Ошибка перемещения мыши: {result.get('error','') }"
        follow_up = self.call_brain_model(feedback)
        return follow_up

    def _handle_left_click(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        x = action_data.get("x", 0)
        y = action_data.get("y", 0)
        result = left_click(x, y)
        feedback = f"Клик ЛКМ выполнен в координатах ({x}, {y})" if result.get("success") else f"Ошибка клика: {result.get('error','') }"
        follow_up = self.call_brain_model(feedback)
        return follow_up

    def _handle_right_click(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        x = action_data.get("x", 0)
        y = action_data.get("y", 0)
        result = right_click(x, y)
        feedback = f"Клик ПКМ выполнен в координатах ({x}, {y})" if result.get("success") else f"Ошибка клика ПКМ: {result.get('error','') }"
        follow_up = self.call_brain_model(feedback)
        return follow_up

    def _handle_scroll(self, action: str, action_data: Dict[str, Any]) -> Union[bool, str]:
        pixels = action_data.get("pixels", 100)
        if action == "scroll_down":
            pixels = -pixels
        result = scroll(pixels)
        feedback = f"Прокрутка выполнена: {result.get('message','') }" if result.get("success") else f"Ошибка прокрутки: {result.get('error','') }"
        follow_up = self.call_brain_model(feedback)
        return follow_up

    def _handle_mouse_down(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        x = action_data.get("x", 0)
        y = action_data.get("y", 0)
        result = mouse_down(x, y)
        feedback = f"ЛКМ зажата в координатах ({x}, {y})" if result.get("success") else f"Ошибка зажатия ЛКМ: {result.get('error','') }"
        follow_up = self.call_brain_model(feedback)
        return follow_up

    def _handle_mouse_up(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        x = action_data.get("x", 0)
        y = action_data.get("y", 0)
        result = mouse_up(x, y)
        feedback = f"ЛКМ отпущена в координатах ({x}, {y})" if result.get("success") else f"Ошибка отпускания ЛКМ: {result.get('error','') }"
        follow_up = self.call_brain_model(feedback)
        return follow_up

    def _handle_drag_and_drop(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        x1 = action_data.get("x1", 0)
        y1 = action_data.get("y1", 0)
        x2 = action_data.get("x2", 0)
        y2 = action_data.get("y2", 0)
        result = drag_and_drop(x1, y1, x2, y2)
        feedback = f"Перетаскивание выполнено от ({x1}, {y1}) к ({x2}, {y2})" if result.get("success") else f"Ошибка перетаскивания: {result.get('error','') }"
        follow_up = self.call_brain_model(feedback)
        return follow_up

    def _handle_type_text(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        text = action_data.get("text", "")
        result = type_text(text)
        feedback = f"Текст введён: {text}" if result.get("success") else f"Ошибка ввода текста: {result.get('error','') }"
        follow_up = self.call_brain_model(feedback)
        return follow_up

    def _handle_generate_image(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        if not getattr(self, 'use_image_generation', False):
            logger.info("🔧 Автоматически включаю генерацию изображений")
            self.use_image_generation = True
            self.auto_disable_tools("image_generation")
        if not getattr(self, 'use_image_generation', False):
            logger.error("❌ Генерация изображений отключена")
            follow_up = self.call_brain_model("Генерация изображений отключена. Предложи альтернативный способ помочь пользователю.")
            return follow_up

        description = action_data.get("description", "")
        text = action_data.get("text", "")
        style = action_data.get("style", "")
        negative_prompt = action_data.get("negative_prompt", "")
        prompt = text.strip() if text else (description or "")
        if style and prompt:
            prompt += f", {style}"

        params = {}
        if not isinstance(params, dict):
            params = {}

        if not prompt:
            follow_up = self.call_brain_model("Нейросеть вернула пустой промт для генерации изображения. Попроси её создать описание.")
            return follow_up

        if not self._is_english_simple(prompt):
            follow_up = self.call_brain_model(f"Нейросеть вернула промт не на английском языке: {prompt}. Попроси её создать промт на английском.")
            return follow_up

        neg = negative_prompt.strip() if negative_prompt else ""
        if not neg or not self._is_english_simple(neg):
            neg = "(worst quality, low quality, normal quality:1.4)"
            self.logger.info(f"⚠️ Используется fallback negative_prompt: {neg}")
        else:
            self.logger.info(f"✅ Используется negative_prompt из JSON: {neg}")
        
        self.logger.info(f"🎨 Передаем в генерацию - prompt: {prompt[:50]}..., negative_prompt: {neg}")

        # default params and validation (kept simple here)
        default_params = {"seed": -1, "steps": 30, "width": 1024, "height": 1024, "cfg": 4.0}
        gen_params = default_params.copy()

        img_b64 = self.generate_image_stable_diffusion(prompt, neg, gen_params)
        if img_b64:
            self.last_generated_image_b64 = img_b64
            self.show_image_base64_temp(img_b64)
            final_msg = f"✅ Изображение успешно сгенерировано по вашему описанию: {description}\n🎨 Использованный промт: {prompt}"
            self.last_final_response = final_msg
            logger.info(final_msg)
            return False
        else:
            follow_up = self.call_brain_model(f"Не удалось сгенерировать изображение по описанию: {description}.")
            return follow_up

    def _handle_generate_video(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        if not getattr(self, 'use_image_generation', False):
            logger.info("🔧 Автоматически включаю генерацию изображений")
            self.use_image_generation = True
            self.auto_disable_tools("image_generation")
        if not getattr(self, 'use_image_generation', False):
            follow_up = self.call_brain_model("Генерация изображений отключена. Предложи альтернативный способ помочь пользователю.")
            return follow_up

        description = action_data.get("description", "")
        text = action_data.get("text", "")
        style = action_data.get("style", "")
        negative_prompt = action_data.get("negative_prompt", "")
        prompt = text.strip() if text else (description or "")
        if style and prompt:
            prompt += f", {style}"

        if not prompt:
            follow_up = self.call_brain_model("Нейросеть вернула пустой промт для генерации видео. Попроси её создать описание.")
            return follow_up

        # negative prompt handling
        neg = negative_prompt.strip() if negative_prompt else ""
        fallback_negative = "(worst quality, low quality, normal quality:1.4), (deformed, distorted, disfigured:1.3), poorly drawn, bad anatomy, text, watermark"
        if not neg or not self._is_english_simple(neg):
            neg = fallback_negative
            self.logger.info(f"⚠️ Используется fallback negative_prompt для видео: {neg}")
        else:
            self.logger.info(f"✅ Используется negative_prompt из JSON для видео: {neg}")
        
        self.logger.info(f"🎬 Передаем в генерацию видео - prompt: {prompt[:50]}..., negative_prompt: {neg[:50]}...")

        # default video params and validation
        default_params = {"seed": -1, "steps": 20, "width": 512, "height": 512, "cfg": 7.0, "num_frames": 24, "fps": 8}
        params = action_data.get("params", {}) or {}
        gen_params = default_params.copy()
        if isinstance(params, dict):
            for key, value in params.items():
                if key in default_params:
                    try:
                        if key in ["seed", "steps", "width", "height", "num_frames", "fps"]:
                            gen_params[key] = int(value)
                        elif key == "cfg":
                            gen_params[key] = float(value)
                        else:
                            gen_params[key] = value
                    except (ValueError, TypeError):
                        logger.warning(f"⚠️ Неверный параметр {key}={value}, используется значение по умолчанию")

        # basic bounds
        if gen_params["steps"] < 1 or gen_params["steps"] > 100:
            gen_params["steps"] = 20
        if gen_params["width"] < 64 or gen_params["width"] > 2048:
            gen_params["width"] = 512
        if gen_params["height"] < 64 or gen_params["height"] > 2048:
            gen_params["height"] = 512

        video_path = self.generate_video_stable_diffusion(prompt, neg, gen_params)
        if video_path:
            final_msg = f"✅ Видео успешно сгенерировано по вашему описанию: {description}\n📁 Путь к видео: {video_path}"
            self.last_final_response = final_msg
            logger.info(final_msg)
            return False
        else:
            follow_up = self.call_brain_model(f"Не удалось сгенерировать видео по описанию: {description}.")
            return follow_up

    def _handle_speak(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        text_to_speak = action_data.get("text", "")
        voice = action_data.get("voice", "male")
        language = action_data.get("language", "ru")
        if not text_to_speak:
            follow_up = self.call_brain_model("Текст для озвучки пустой. Укажите текст в поле 'text'.")
            return follow_up
        audio_path = self.text_to_speech(text_to_speak, voice, language)
        if audio_path:
            follow_up = self.call_brain_model(f"Текст успешно озвучен: {text_to_speak}. Аудиофайл сохранен: {os.path.basename(audio_path)}")
        else:
            follow_up = self.call_brain_model(f"Не удалось озвучить текст: {text_to_speak}.")
        return follow_up

    def _handle_response(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        # Поддерживаем и "text" и "content" для совместимости
        raw_content = action_data.get("text", action_data.get("content", ""))
        # Удаляем служебные <INFO> теги перед выводом пользователю и подсвечиваем содержимое
        display_content = self._strip_info_tags(raw_content, highlight=True)
        self.last_final_response = display_content
        logger.info(f"\n🤖 ФИНАЛЬНЫЙ ОТВЕТ:")
        logger.info(display_content)
        
        # Если есть сгенерированный файл, уведомляем об этом в интерактивном режиме
        if (hasattr(self, 'last_generated_file_path') and self.last_generated_file_path and 
            getattr(self, 'show_images_locally', True)):
            logger.info(f"\n📄 Создан файл: {self.last_generated_file_name}")
            logger.info(f"📂 Расположение: {self.last_generated_file_path}")
            # Очищаем после показа в интерактивном режиме
            self.last_generated_file_path = None
            self.last_generated_file_name = None
        
        return False

    def _handle_list_files(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для просмотра содержимого папок
        """
        folder = action_data.get("folder", "")
        description = action_data.get("description", f"Просмотр содержимого папки {folder}")
        
        logger.info(f"\n📁 ПРОСМОТР ПАПКИ: {description}")
        logger.info(f"📂 Папка: {folder}")
        
        result = self.list_folder_contents(folder)
        
        logger.info(f"📋 Результат:\n{result}")
        
        follow_up = self.call_brain_model(f"Содержимое папки '{folder}': {result}")
        return follow_up

    def _handle_process_document(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для обработки документов (DOCX, Excel, PDF)
        """
        file_path = action_data.get("file_path", "")
        description = action_data.get("description", f"Обработка документа {file_path}")
        
        logger.info(f"\n📄 ОБРАБОТКА ДОКУМЕНТА: {description}")
        logger.info(f"📁 Файл: {file_path}")
        
        # Проверяем, является ли путь относительным и добавляем базовую директорию
        if not os.path.isabs(file_path):
            # Проверяем, не начинается ли путь уже с папки
            if (file_path.startswith("Docx/") or file_path.startswith("Docx\\") or
                file_path.startswith("Excel/") or file_path.startswith("Excel\\") or
                file_path.startswith("PDF/") or file_path.startswith("PDF\\")):
                full_path = os.path.join(self.base_dir, file_path)
            else:
                # Определяем папку по расширению файла
                file_lower = file_path.lower()
                if file_lower.endswith(('.docx', '.doc')):
                    full_path = os.path.join(self.base_dir, "Docx", file_path)
                elif file_lower.endswith(('.xlsx', '.xls', '.csv')):
                    full_path = os.path.join(self.base_dir, "Excel", file_path)
                elif file_lower.endswith('.pdf'):
                    full_path = os.path.join(self.base_dir, "PDF", file_path)
                else:
                    full_path = os.path.join(self.base_dir, file_path)
        else:
            full_path = file_path
        
        result = self.process_document_request(full_path)
        
        logger.info(f"📋 Результат обработки:\n{result[:500]}...")
        
        follow_up = self.call_brain_model(f"Результат обработки документа '{file_path}': {result}")
        return follow_up

    def _handle_generate_file(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для генерации файлов (DOCX, Excel, PDF, Markdown)
        """
        content = action_data.get("content", "")
        filename = action_data.get("filename", "")
        file_path_param = action_data.get("file_path") or action_data.get("path") or ""
        file_type = action_data.get("file_type", "").lower()
        description = action_data.get("description", f"Генерация файла {filename or file_path_param}")
        
        logger.info(f"\n📝 ГЕНЕРАЦИЯ ФАЙЛА: {description}")
        logger.info(f"📁 Имя файла: {filename}")
        if file_path_param:
            logger.info(f"📂 Путь из JSON: {file_path_param}")
        logger.info(f"📄 Тип файла: {file_type}")
        
        if not content:
            follow_up = self.call_brain_model("Ошибка: не указано содержимое для генерации файла")
            return follow_up
        
        output_path = None
        if not filename:
            if file_path_param:
                normalized_path = os.path.expanduser(file_path_param)
                filename = os.path.basename(normalized_path) or ""
                logger.info(f"🔄 Использую имя из пути: {filename}")
                if os.path.isabs(normalized_path):
                    output_path = normalized_path
                else:
                    output_path = os.path.join(self.base_dir, OUTPUT_DIR_NAME, file_path_param)
            if not filename:
                follow_up = self.call_brain_model("Ошибка: не указано имя файла")
                return follow_up
        if output_path is None:
            output_path = os.path.join(self.base_dir, OUTPUT_DIR_NAME, filename)
        output_path = os.path.expanduser(output_path)
        logger.info(f"🛣️ Путь сохранения: {output_path}")
        
        try:
            success = self.generate_file(content, output_path, file_type)
            if success:
                logger.info(f"✅ Файл успешно создан: {output_path}")
                
                # Сохраняем информацию о последнем сгенерированном файле для Telegram
                self.last_generated_file_path = output_path
                self.last_generated_file_name = filename
                
                follow_up = self.call_brain_model(f"Файл '{filename}' успешно создан в папке output")
            else:
                logger.error(f"❌ Ошибка при создании файла: {output_path}")
                follow_up = self.call_brain_model(f"Ошибка при создании файла '{filename}'")
        except Exception as e:
            logger.error(f"❌ Исключение при создании файла: {e}")
            follow_up = self.call_brain_model(f"Ошибка при создании файла '{filename}': {str(e)}")
        
        return follow_up

    def _handle_extract_text(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для извлечения текста из изображений с помощью OCR
        """
        image_path = action_data.get("image_path", "")
        description = action_data.get("description", f"Извлечение текста из {image_path}")
        force_ocr = action_data.get("force_ocr", False)
        
        logger.info(f"\n📖 ИЗВЛЕЧЕНИЕ ТЕКСТА OCR: {description}")
        logger.info(f"📁 Изображение (исходный путь): {image_path}")
        
        # Проверяем, является ли путь относительным и добавляем базовую директорию
        if not os.path.isabs(image_path):
            logger.info(f"🔍 Путь относительный, проверяю логику...")
            # Проверяем, не начинается ли путь уже с папки Photos
            if image_path.startswith("Photos/") or image_path.startswith("Photos\\"):
                logger.info(f"🔍 Путь начинается с Photos/, используем как есть")
                full_path = os.path.join(self.base_dir, image_path)
            else:
                logger.info(f"🔍 Путь не начинается с Photos/, определяю по расширению")
                # Определяем папку по расширению файла
                file_lower = image_path.lower()
                if file_lower.endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp')):
                    logger.info(f"🔍 Это изображение, добавляю папку Photos")
                    full_path = os.path.join(self.base_dir, "Photos", image_path)
                else:
                    logger.info(f"🔍 Не изображение, добавляю к base_dir")
                    full_path = os.path.join(self.base_dir, image_path)
        else:
            logger.info(f"🔍 Путь абсолютный, используем как есть")
            full_path = image_path
        
        logger.info(f"🔍 Полный путь к файлу: {full_path}")
        
        # Сначала получим описание от vision модели если доступна
        vision_description = ""
        if getattr(self, 'use_vision', False):
            try:
                # Конвертируем изображение в base64 для vision модели
                with open(full_path, 'rb') as img_file:
                    img_b64 = base64.b64encode(img_file.read()).decode('ascii')
                    vision_description = self.call_vision_model(img_b64)
                    logger.info(f"🔍 Vision описание: {vision_description[:100]}...")
            except Exception as e:
                logger.warning(f"⚠️ Не удалось получить vision описание: {e}")
        
        # Применяем умный OCR
        extracted_text, full_description, ocr_error = self.process_image_with_smart_ocr(
            full_path, vision_description, force_ocr
        )
        
        if extracted_text:
            wrapped_text = self._wrap_info_if_needed(extracted_text, source="image_ocr")
            result_text = f"Извлеченный текст из изображения '{image_path}':\n{wrapped_text}"
        elif ocr_error:
            result_text = f"Ошибка при обработке изображения '{image_path}':\n\n{ocr_error}\n\nОписание изображения (если доступно):\n{vision_description}"
        else:
            result_text = f"Описание изображения '{image_path}':\n\n{vision_description}"
        
        logger.info(f"📋 Результат OCR:\n{result_text}")
        
        follow_up = self.call_brain_model(result_text)
        return follow_up

    def _handle_analyze_image(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для анализа изображения с возможностью интеллектуального OCR.
        
        Args:
            action_data: Данные действия с полями:
                - image_path: Путь к изображению  
                - check_for_text: Нужно ли проверять наличие текста (bool)
                - description: Описание задачи (optional)
        
        Returns:
            str: Follow-up для модели с результатами анализа
        """
        try:
            image_path = action_data.get("image_path", "").strip()
            check_for_text = action_data.get("check_for_text", False)
            description = action_data.get("description", "")
            
            if not image_path:
                logger.error("❌ Не указан путь к изображению")
                return self.call_brain_model("❌ Ошибка: не указан путь к изображению")
                
            # Разрешаем путь
            full_path = self.resolve_path(image_path)
            if not os.path.exists(full_path):
                logger.error(f"❌ Изображение не найдено: {full_path}")
                return self.call_brain_model(f"❌ Изображение не найдено: {image_path}")
            
            logger.info(f"🔍 Анализирую изображение: {image_path}")
            if description:
                logger.info(f"📝 Задача: {description}")
            
            # Получаем vision описание
            vision_description = ""
            try:
                vision_description = self.analyze_image_with_vision(full_path)
                logger.info(f"👁️ Vision описание: {vision_description}")
            except Exception as e:
                logger.warning(f"⚠️ Не удалось получить vision описание: {e}")
                vision_description = f"Ошибка анализа изображения: {e}"
            
            # Если запрошена проверка текста, принимаем решение об OCR
            extracted_text = ""
            ocr_info = ""
            
            if check_for_text:
                logger.info("🔍 Проверяю наличие текста на изображении...")
                
                # Интеллектуальное решение об OCR на основе vision описания
                should_use_ocr = self.should_use_ocr_intelligently(vision_description, description)
                
                if should_use_ocr:
                    logger.info("✅ Применяю OCR для извлечения текста")
                    try:
                        extracted_text = self.extract_text_from_image(full_path)
                        if extracted_text:
                            ocr_info = f"\n\n📋 Извлеченный текст:\n{extracted_text}"
                        else:
                            ocr_info = "\n\n❌ OCR не обнаружил текст на изображении"
                    except Exception as e:
                        logger.warning(f"⚠️ Ошибка OCR: {e}")
                        ocr_info = f"\n\n❌ Ошибка при извлечении текста: {e}"
                else:
                    logger.info("❌ OCR не требуется для данного изображения")
                    ocr_info = "\n\n📝 OCR не применялся - на изображении не обнаружено значимого текста"
            
            # Формируем результат анализа
            result_text = f"Анализ изображения '{image_path}':\n\n📸 Описание:\n{vision_description}{ocr_info}"
            
            logger.info(f"✅ Анализ завершен")
            
            follow_up = self.call_brain_model(result_text)
            return follow_up
            
        except Exception as e:
            logger.error(f"❌ Ошибка при анализе изображения: {e}")
            return self.call_brain_model(f"❌ Ошибка при анализе изображения: {e}")

    def _handle_inspect_ui(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для инспекции UI приложений
        """
        window_name = action_data.get("window_name", "")
        max_depth = action_data.get("max_depth", 5)
        description = action_data.get("description", f"Инспекция UI окна '{window_name}'")
        
        logger.info(f"\n🔍 ИНСПЕКЦИЯ UI: {description}")
        logger.info(f"🪟 Окно: {window_name}")
        
        if not window_name:
            # Если имя окна не указано, пробуем активное окно, но лучше предупредить
            logger.warning("⚠️ Имя окна не указано, будет использовано активное окно")
        
        ui_tree = get_ui_tree_as_text(window_name, max_depth)
        
        # Ограничиваем размер вывода, чтобы не перегрузить контекст
        if len(ui_tree) > 10000:
            ui_tree = ui_tree[:10000] + "\n... (truncated)"
            
        logger.info(f"📋 Результат UI Tree (первые 500 символов):\n{ui_tree[:500]}...")
        
        follow_up = self.call_brain_model(f"Результат инспекции UI окна '{window_name}':\n\n{ui_tree}")
        return follow_up

    def _handle_plugin_action(self, action: str, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для действий плагинов.
        
        Args:
            action: Строка действия в формате "plugin:plugin_name:action_name"
            action_data: Данные действия
        
        Returns:
            str: Follow-up для модели с результатами выполнения
        """
        try:
            # Проверяем доступность системы плагинов
            if not self.plugin_manager:
                logger.error("❌ Система плагинов недоступна")
                return self.call_brain_model("❌ Система плагинов недоступна")
            
            # Парсим action в формате "plugin:plugin_name:action_name"
            parts = action.split(":", 2)
            if len(parts) != 3 or parts[0] != "plugin":
                logger.error(f"❌ Неверный формат действия плагина: {action}")
                return self.call_brain_model(f"❌ Неверный формат действия плагина. Ожидается 'plugin:plugin_name:action_name'")
            
            plugin_name = parts[1]
            plugin_action = parts[2]
            plugin_data = action_data.get("data", {})
            
            logger.info(f"🔌 Выполняется действие плагина: {plugin_name}.{plugin_action}")
            
            # Выполняем действие плагина
            result = self.plugin_manager.execute_plugin_action(
                plugin_name=plugin_name,
                action=plugin_action,
                data=plugin_data,
                orchestrator=self
            )
            
            # Формируем ответ
            result_text = f"Результат выполнения плагина '{plugin_name}', действие '{plugin_action}':\n\n{result}"
            
            logger.info(f"✅ Действие плагина выполнено успешно")
            
            follow_up = self.call_brain_model(result_text)
            return follow_up
            
        except PluginError as e:
            logger.error(f"❌ Ошибка плагина: {e}")
            return self.call_brain_model(f"❌ Ошибка плагина: {e}")
        except Exception as e:
            logger.error(f"❌ Неожиданная ошибка при выполнении действия плагина: {e}")
            return self.call_brain_model(f"❌ Ошибка при выполнении действия плагина: {e}")

    def _handle_get_help(self, action: str, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для команд get_*_help - загружает соответствующий модуль.
        
        Args:
            action: Команда типа "get_image_generation_help"
            action_data: Данные действия
        
        Returns:
            str: Follow-up для модели с загруженным модулем
        """
        try:
            logger.info(f"📚 Загружаю модуль для команды: {action}")
            
            # Загружаем модуль через PromptLoader
            module_content = self.prompt_loader.load_module(action)
            
            if module_content is None:
                logger.warning(f"❌ Модуль для команды {action} не найден")
                return self.call_brain_model(f"❌ Модуль для команды {action} не найден. Доступные команды: {', '.join(self.prompt_loader.module_commands.keys())}")
            
            logger.info(f"✅ Модуль загружен, размер: {len(module_content)} символов")
            
            # Отправляем загруженный модуль как контекст для модели
            follow_up_prompt = f"""
Загружен модуль по запросу {action}:

{module_content}

Теперь ты можешь использовать эту информацию для ответа пользователю или выполнения соответствующих действий.
"""
            
            follow_up = self.call_brain_model(follow_up_prompt)
            return follow_up
            
        except Exception as e:
            logger.error(f"❌ Ошибка при загрузке модуля {action}: {e}")
            return self.call_brain_model(f"❌ Ошибка при загрузке модуля {action}: {e}")

    def _handle_send_email(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для отправки email
        """
        provider = action_data.get("provider", "")
        to_email = action_data.get("to_email", "")
        subject = action_data.get("subject", "")
        body = action_data.get("body", "")
        attachments = action_data.get("attachments", [])
        description = action_data.get("description", f"Отправка письма на {to_email}")
        
        logger.info(f"\n📧 ОТПРАВКА EMAIL: {description}")
        logger.info(f"📨 Провайдер: {provider}")
        logger.info(f"📧 Получатель: {to_email}")
        logger.info(f"📝 Тема: {subject}")
        
        if not provider:
            # Используем первый доступный провайдер
            if self.available_email_providers:
                provider = self.available_email_providers[0]
                logger.info(f"🔧 Автоматически выбран провайдер: {provider}")
            else:
                result = "❌ Почтовые провайдеры не настроены"
                logger.error(result)
                follow_up = self.call_brain_model(result)
                return follow_up
        
        result = self.send_email(provider, to_email, subject, body, attachments)
        logger.info(f"📧 Результат отправки: {result}")
        
        follow_up = self.call_brain_model(f"Результат отправки письма: {result}")
        return follow_up

    def _handle_get_emails(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для получения списка писем
        """
        provider = action_data.get("provider", "")
        folder = action_data.get("folder", "INBOX")
        limit = action_data.get("limit", 10)
        search_criteria = action_data.get("search_criteria", "ALL")
        description = action_data.get("description", f"Получение писем из {folder}")
        
        logger.info(f"\n📧 ПОЛУЧЕНИЕ ПИСЕМ: {description}")
        logger.info(f"📨 Провайдер: {provider}")
        logger.info(f"📁 Папка: {folder}")
        logger.info(f"🔢 Лимит: {limit}")
        
        if not provider:
            # Используем первый доступный провайдер
            if self.available_email_providers:
                provider = self.available_email_providers[0]
                logger.info(f"🔧 Автоматически выбран провайдер: {provider}")
            else:
                result = "❌ Почтовые провайдеры не настроены"
                logger.error(result)
                follow_up = self.call_brain_model(result)
                return follow_up
        
        emails = self.get_emails(provider, folder, limit, search_criteria)
        
        if isinstance(emails, list):
            # Форматируем список писем для отправки в модель
            emails_summary = f"Получено {len(emails)} писем из папки '{folder}':\n\n"
            for i, email_info in enumerate(emails, 1):
                emails_summary += f"{i}. От: {email_info.get('from', 'Неизвестно')}\n"
                emails_summary += f"   Тема: {email_info.get('subject', 'Без темы')}\n"
                emails_summary += f"   Дата: {email_info.get('date', 'Неизвестно')}\n"
                emails_summary += f"   ID: {email_info.get('id', 'Неизвестно')}\n"
                
                # Обрезаем текст письма до 200 символов для краткости
                body = email_info.get('body', '')
                if len(body) > 200:
                    body = body[:200] + "..."
                emails_summary += f"   Текст: {body}\n\n"
            
            logger.info(f"📧 Получено {len(emails)} писем")
            follow_up = self.call_brain_model(emails_summary)
        else:
            # Ошибка получения писем
            logger.error(f"❌ Ошибка: {emails}")
            follow_up = self.call_brain_model(f"Ошибка получения писем: {emails}")
        
        return follow_up

    def _handle_reply_email(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для ответа на письмо
        """
        provider = action_data.get("provider", "")
        original_email_id = action_data.get("original_email_id", "")
        reply_text = action_data.get("reply_text", "")
        attachments = action_data.get("attachments", [])
        description = action_data.get("description", f"Ответ на письмо {original_email_id}")
        
        logger.info(f"\n📧 ОТВЕТ НА ПИСЬМО: {description}")
        logger.info(f"📨 Провайдер: {provider}")
        logger.info(f"🆔 ID оригинального письма: {original_email_id}")
        
        if not provider:
            # Используем первый доступный провайдер
            if self.available_email_providers:
                provider = self.available_email_providers[0]
                logger.info(f"🔧 Автоматически выбран провайдер: {provider}")
            else:
                result = "❌ Почтовые провайдеры не настроены"
                logger.error(result)
                follow_up = self.call_brain_model(result)
                return follow_up
        
        result = self.reply_to_email(provider, original_email_id, reply_text, attachments)
        logger.info(f"📧 Результат ответа: {result}")
        
        follow_up = self.call_brain_model(f"Результат ответа на письмо: {result}")
        return follow_up

    def _handle_search_emails(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для поиска писем
        """
        provider = action_data.get("provider", "")
        query = action_data.get("query", "")
        folder = action_data.get("folder", "INBOX")
        limit = action_data.get("limit", 20)
        description = action_data.get("description", f"Поиск писем: {query}")
        
        logger.info(f"\n🔍 ПОИСК ПИСЕМ: {description}")
        logger.info(f"📨 Провайдер: {provider}")
        logger.info(f"🔎 Запрос: {query}")
        logger.info(f"📁 Папка: {folder}")
        
        if not provider:
            # Используем первый доступный провайдер
            if self.available_email_providers:
                provider = self.available_email_providers[0]
                logger.info(f"🔧 Автоматически выбран провайдер: {provider}")
            else:
                result = "❌ Почтовые провайдеры не настроены"
                logger.error(result)
                follow_up = self.call_brain_model(result)
                return follow_up
        
        emails = self.search_emails(provider, query, folder, limit)
        
        if isinstance(emails, list):
            # Форматируем результаты поиска
            search_summary = f"Найдено {len(emails)} писем по запросу '{query}':\n\n"
            for i, email_info in enumerate(emails, 1):
                search_summary += f"{i}. От: {email_info.get('from', 'Неизвестно')}\n"
                search_summary += f"   Тема: {email_info.get('subject', 'Без темы')}\n"
                search_summary += f"   Дата: {email_info.get('date', 'Неизвестно')}\n"
                search_summary += f"   ID: {email_info.get('id', 'Неизвестно')}\n"
                
                # Обрезаем текст письма до 200 символов
                body = email_info.get('body', '')
                if len(body) > 200:
                    body = body[:200] + "..."
                search_summary += f"   Текст: {body}\n\n"
            
            logger.info(f"🔍 Найдено {len(emails)} писем")
            follow_up = self.call_brain_model(search_summary)
        else:
            # Ошибка поиска
            logger.error(f"❌ Ошибка поиска: {emails}")
            follow_up = self.call_brain_model(f"Ошибка поиска писем: {emails}")
        
        return follow_up

    def _handle_run_bat_file(self, action_data: Dict[str, Any]) -> Union[bool, str]:
        """
        Обработчик для запуска BAT файлов
        """
        file_path = action_data.get("file_path", "")
        working_dir = action_data.get("working_dir", None)
        description = action_data.get("description", f"Запуск BAT файла: {file_path}")
        
        logger.info(f"\n🚀 ЗАПУСК BAT ФАЙЛА: {description}")
        logger.info(f"📄 Файл: {file_path}")
        if working_dir:
            logger.info(f"📁 Рабочая директория: {working_dir}")
        
        if not file_path:
            error_msg = "❌ Не указан путь к BAT файлу"
            logger.error(error_msg)
            follow_up = self.call_brain_model(f"Ошибка: {error_msg}")
            return follow_up
        
        try:
            result = self.run_bat_file(file_path, working_dir)
            
            if result["success"]:
                # Успешное выполнение
                output_info = f"✅ BAT файл выполнен успешно!\n"
                output_info += f"📄 Файл: {os.path.basename(file_path)}\n"
                output_info += f"📁 Рабочая директория: {result.get('working_dir', 'не указана')}\n"
                output_info += f"🔢 Код возврата: {result.get('return_code', 0)}\n"
                
                if result.get("output"):
                    output_info += f"\n📝 Вывод:\n{result['output']}"
                
                logger.info("✅ BAT файл выполнен успешно")
                follow_up = self.call_brain_model(output_info)
            else:
                # Ошибка выполнения
                error_info = f"❌ Ошибка выполнения BAT файла!\n"
                error_info += f"📄 Файл: {os.path.basename(file_path)}\n"
                error_info += f"🔢 Код возврата: {result.get('return_code', -1)}\n"
                error_info += f"❌ Ошибка: {result.get('error', 'Неизвестная ошибка')}\n"
                
                if result.get("output"):
                    error_info += f"\n📝 Вывод:\n{result['output']}"
                
                logger.error(f"❌ Ошибка выполнения BAT файла: {result.get('error')}")
                follow_up = self.call_brain_model(error_info)
                
        except Exception as e:
            error_msg = f"❌ Исключение при запуске BAT файла: {str(e)}"
            logger.error(error_msg)
            follow_up = self.call_brain_model(f"Ошибка: {error_msg}")
        
        return follow_up

    def _get_plugin_info_for_prompt(self) -> str:
        """
        Формирует информацию о доступных плагинах для добавления в системный промпт.
        
        Returns:
            str: Текст с информацией о плагинах
        """
        if not self.plugin_manager:
            return ""
        
        try:
            loaded_plugins = self.plugin_manager.get_loaded_plugins()
            if not loaded_plugins:
                return ""
            
            plugin_info_parts = ["ДОСТУПНЫЕ ПЛАГИНЫ:"]
            
            for plugin_name, plugin in loaded_plugins.items():
                try:
                    info = plugin.get_plugin_info()
                    actions = plugin.get_available_actions()
                    
                    plugin_desc = f"\n🔌 {info.get('name', plugin_name)} v{info.get('version', '1.0')}"
                    plugin_desc += f"\n   Описание: {info.get('description', 'Нет описания')}"
                    plugin_desc += f"\n   Автор: {info.get('author', 'Неизвестен')}"
                    plugin_desc += f"\n   Действия: {', '.join(actions)}"
                    plugin_desc += f"\n   Формат вызова: plugin:{plugin_name}:action_name"
                    
                    plugin_info_parts.append(plugin_desc)
                except Exception as e:
                    logger.warning(f"Ошибка получения информации о плагине {plugin_name}: {e}")
                    plugin_info_parts.append(f"\n🔌 {plugin_name} (ошибка получения информации)")
            
            return "\n".join(plugin_info_parts)
            
        except Exception as e:
            logger.error(f"Ошибка формирования информации о плагинах: {e}")
            return ""

    def should_use_ocr_intelligently(self, vision_description: str, task_description: str = "") -> bool:
        """
        Интеллектуальное решение о необходимости OCR на основе контекста.
        
        Args:
            vision_description: Описание изображения от vision модели
            task_description: Описание задачи пользователя
            
        Returns:
            bool: True если OCR нужен, False если нет
        """
        vision_lower = vision_description.lower()
        task_lower = task_description.lower()
        
        # Явные индикаторы необходимости OCR
        text_indicators = [
            "text", "writing", "words", "letters", "document", "page", "book", "sign", 
            "label", "caption", "title", "heading", "paragraph", "sentence", "line",
            "screen", "display", "interface", "menu", "button", "dialog", "window",
            "newspaper", "article", "magazine", "poster", "banner", "billboard",
            "form", "table", "chart", "graph", "spreadsheet", "invoice", "receipt",
            "card", "certificate", "license", "passport", "id", "ticket",
            "текст", "надпись", "слова", "буквы", "документ", "страница", "книга",
            "вывеска", "подпись", "заголовок", "строка", "экран", "интерфейс",
            "меню", "кнопка", "окно", "газета", "статья", "плакат", "форма",
            "таблица", "график", "чек", "карта", "сертификат", "билет"
        ]
        
        # Индикаторы отсутствия текста
        no_text_indicators = [
            "landscape", "nature", "animal", "person", "face", "building", "car",
            "food", "flower", "tree", "sky", "mountain", "water", "art", "painting",
            "photo", "picture", "image", "scenery", "portrait", "selfie",
            "пейзаж", "природа", "животное", "человек", "лицо", "здание", "машина",
            "еда", "цветок", "дерево", "небо", "гора", "вода", "искусство", "картина",
            "фото", "изображение", "портрет"
        ]
        
        # Проверяем явные запросы OCR в задаче
        if any(keyword in task_lower for keyword in ["текст", "text", "прочита", "read", "извлеч", "extract"]):
            return True
        
        # Считаем индикаторы в vision описании
        text_score = sum(1 for indicator in text_indicators if indicator in vision_lower)
        no_text_score = sum(1 for indicator in no_text_indicators if indicator in vision_lower)
        
        # Решение на основе баланса индикаторов
        if text_score >= 2:  # Несколько индикаторов текста
            return True
        elif text_score >= 1 and no_text_score == 0:  # Есть индикатор текста, нет противопоказаний
            return True
        elif no_text_score >= 2:  # Явно не текстовое изображение
            return False
        else:
            # Граничный случай - OCR может быть полезен
            return text_score > no_text_score

    def _process_ai_response_impl(self, ai_response: str) -> bool:
        """
        Подробная реализация обработки ответа AI (перенесена из original `process_ai_response`).

        Args:
            ai_response: JSON ответ от AI

        Returns:
            True если нужно продолжать диалог, False если завершить
        """
        # Итеративный обработчик цепочек follow_up: избегаем рекурсии.
        next_input: Optional[str] = ai_response
        attempts = 0
        while next_input is not None and attempts <= self.max_retries:
            attempts += 1
            try:
                json_str = self.extract_first_json(next_input, allow_json_in_think=True)
                if not json_str or json_str == next_input:
                    logger.info("💬 Модель вернула текстовый ответ без JSON")
                    if len(next_input.strip()) > 5 and not next_input.strip().startswith('{'):
                        logger.info("💬 Использую текстовый ответ как финальный")
                        clean_answer = re.sub(r'<think>.*?</think>', '', next_input, flags=re.DOTALL | re.IGNORECASE).strip()
                        if not clean_answer:
                            clean_answer = next_input.strip()
                        self.last_final_response = clean_answer
                        logger.info("\n🤖 ФИНАЛЬНЫЙ ОТВЕТ (текст):")
                        logger.info(clean_answer)
                        return False
                    think_content = self.extract_think_content(next_input)
                    if think_content:
                        json_in_think = self._extract_json_from_text(think_content)
                        if json_in_think:
                            logger.info("🔍 Найден JSON внутри <think> блока, использую его")
                            json_str = json_in_think
                        else:
                            feedback = "Модель вернула неполный ответ. Пожалуйста, сформулируй полный ответ или действие."
                            next_input = self.call_brain_model(feedback)
                            continue
                    else:
                        feedback = "Модель вернула неполный ответ. Пожалуйста, сформулируй полный ответ или действие."
                        next_input = self.call_brain_model(feedback)
                        continue

                action_data, fixes = self._smart_json_parse(json_str)
                if fixes:
                    logger.warning(f"⚠️ Исправления JSON: {'; '.join(fixes)}")
                if not action_data:
                    logger.error(f"❌ Не удалось распарсить JSON даже после исправлений:\n{json_str}")
                    self.retry_count += 1
                    if self.retry_count > self.max_retries:
                        logger.warning(f"🔄 Достигнут лимит попыток ({self.max_retries}), завершаю диалог")
                        self.retry_count = 0
                        self.last_final_response = "Извините, возникла проблема с обработкой запроса. Попробуйте переформулировать вопрос."
                        return False
                    think_content = self.extract_think_content(next_input)
                    if think_content:
                        logger.info("💭 Обнаружен блок размышлений - модель размышляет, но не генерирует действие")
                        logger.info(f"💭 Содержимое: {think_content[:200]}...")
                        if len(think_content) > 20 and any(word in think_content.lower() for word in ['привет', 'hello', 'здравствуй']):
                            logger.info("💭 Обнаружено приветствие в размышлениях, завершаю диалог")
                            self.retry_count = 0
                            self.last_final_response = "Привет! Я Нейро, ваш AI-помощник. Чем могу помочь?"
                            return False
                        if self.retry_count >= 2 and 'правил' in think_content.lower():
                            logger.info("💭 Модель зацикливается на правилах, принудительно завершаю")
                            self.retry_count = 0
                            self.last_final_response = "Привет! Как дела? Чем могу помочь?"
                            return False
                        next_input = self.call_brain_model("Пожалуйста, дай простой дружелюбный ответ или конкретное действие. Не анализируй правила.")
                        continue
                    else:
                        next_input = self.call_brain_model("Пожалуйста, дай простой ответ или конкретное действие.")
                        continue

                if 'action' not in action_data:
                    logger.warning("⚠️ В JSON отсутствует ключ 'action', добавляю action: 'unknown'")
                    action_data['action'] = 'unknown'

                if action_data == {} or (len(action_data) == 1 and 'action' in action_data and action_data['action'] == 'unknown'):
                    logger.info("💭 Модель вернула пустой JSON - возможно, она размышляет")
                    think_content = self.extract_think_content(next_input)
                    if think_content and len(think_content) > 20:
                        logger.info("💭 Использую содержимое размышлений как ответ")
                        self.last_final_response = think_content
                        logger.info("\n🤖 ФИНАЛЬНЫЙ ОТВЕТ (из размышлений):")
                        logger.info(think_content)
                        return False
                    next_input = self.call_brain_model("Модель вернула пустой JSON. Пожалуйста, сформулируй конкретный ответ или действие.")
                    continue

                action = action_data.get("action")
                
                # Проверка на пустой ответ "..."
                if action == "response":
                    text_content = action_data.get("text", action_data.get("content", "")).strip()
                    if text_content == "..." or text_content == ".." or not text_content:
                        logger.warning("⚠️ Модель вернула пустой ответ '...', прошу перегенерировать")
                        next_input = self.call_brain_model("Ты вернул пустой ответ '...'. Пожалуйста, дай полноценный текстовый ответ.")
                        continue

                self.retry_count = 0

                # Вызов соответствующего хендлера и обработка результата (следующий ввод или завершение)
                handler_result = None
                if action == "powershell":
                    handler_result = self._handle_powershell(action_data)
                elif action == "search":
                    handler_result = self._handle_search(action_data)
                elif action == "take_screenshot":
                    handler_result = self._handle_take_screenshot(action_data)
                elif action == "move_mouse":
                    handler_result = self._handle_move_mouse(action_data)
                elif action == "left_click":
                    handler_result = self._handle_left_click(action_data)
                elif action == "right_click":
                    handler_result = self._handle_right_click(action_data)
                elif action in ["scroll_up", "scroll_down"]:
                    handler_result = self._handle_scroll(action, action_data)
                elif action == "mouse_down":
                    handler_result = self._handle_mouse_down(action_data)
                elif action == "mouse_up":
                    handler_result = self._handle_mouse_up(action_data)
                elif action == "drag_and_drop":
                    handler_result = self._handle_drag_and_drop(action_data)
                elif action == "type_text":
                    handler_result = self._handle_type_text(action_data)
                elif action == "generate_image":
                    handler_result = self._handle_generate_image(action_data)
                elif action == "generate_video":
                    handler_result = self._handle_generate_video(action_data)
                elif action == "speak":
                    handler_result = self._handle_speak(action_data)
                elif action == "list_files":
                    handler_result = self._handle_list_files(action_data)
                elif action == "process_document":
                    handler_result = self._handle_process_document(action_data)
                elif action == "generate_file":
                    handler_result = self._handle_generate_file(action_data)
                elif action == "extract_text":
                    handler_result = self._handle_extract_text(action_data)
                elif action == "analyze_image":
                    handler_result = self._handle_analyze_image(action_data)
                elif action == "inspect_ui":
                    handler_result = self._handle_inspect_ui(action_data)
                elif action == "response":
                    handler_result = self._handle_response(action_data)
                elif action == "send_email":
                    handler_result = self._handle_send_email(action_data)
                elif action == "get_emails":
                    handler_result = self._handle_get_emails(action_data)
                elif action == "reply_email":
                    handler_result = self._handle_reply_email(action_data)
                elif action == "search_emails":
                    handler_result = self._handle_search_emails(action_data)
                elif action == "run_bat_file":
                    handler_result = self._handle_run_bat_file(action_data)
                elif action.startswith("plugin:"):
                    handler_result = self._handle_plugin_action(action, action_data)
                elif action.startswith("get_") and action.endswith("_help"):
                    handler_result = self._handle_get_help(action, action_data)
                else:
                    logger.warning(f"❓ Неизвестное действие: {action}")
                    return False

                # Обработка результата хендлера: False => завершить, str => новый ввод для итерации
                if handler_result is False:
                    return False
                elif isinstance(handler_result, str):
                    next_input = handler_result
                    continue
                else:
                    # Если хендлер вернул None или неожиданный тип — завершаем
                    logger.error("❌ Хендлер вернул некорректный результат, завершаю")
                    return False

            except json.JSONDecodeError as e:
                logger.error(f"❌ Ошибка парсинга JSON ответа AI: {e}")
                logger.info(f"📝 Ответ AI: {next_input}")
                return False
            except Exception as e:
                logger.error(f"❌ Ошибка обработки ответа AI: {str(e)}")
                return False

        logger.warning("🔄 Превышен цикл попыток обработки follow_up или достигнут лимит. Завершаю.")
        return False

    def _read_user_input(self, prompt_text: str) -> str:
        """Безопасное чтение ввода пользователя с поддержкой Windows-консоли."""
        if os.name != 'nt':
            # На Unix-подобных системах стандартный input работает корректно
            return input(prompt_text).strip()

        import msvcrt  # type: ignore

        print(prompt_text, end='', flush=True)
        buffer: List[str] = []

        while True:
            ch = msvcrt.getwch()

            # Enter / возврат
            if ch in ('\r', '\n'):
                print()
                return ''.join(buffer).strip()

            # Ctrl+C
            if ch == '\x03':
                print()
                raise KeyboardInterrupt()

            # Ctrl+Z (EOF)
            if ch == '\x1a':
                print()
                return ''.join(buffer).strip()

            # Backspace
            if ch == '\b':
                if buffer:
                    buffer.pop()
                    # Стираем символ в консоли
                    print('\b \b', end='', flush=True)
                continue

            # Игнорируем служебные клавиши (стрелки и т.п.)
            if ch in ('\x00', '\xe0'):
                # пропускаем следующий код клавиши
                msvcrt.getwch()
                continue

            buffer.append(ch)
            print(ch, end='', flush=True)

    def run_interactive(self):
        """Запуск интерактивного режима (глаза, аудио, мозг)"""
        # Показываем сообщения только в консольном режиме
        if getattr(self, 'show_images_locally', True):
            logger.info("🚀 AI PowerShell Оркестратор запущен!")
            logger.info("💡 Используйте новые команды для работы с файлами:")
            logger.info("   - list_files: просмотр содержимого папок (Audio, Photos, Video, Excel, Docx, PDF)")
            logger.info("   - process_document: обработка всех типов документов (DOCX, DOC, PDF, XLSX, XLS, CSV, TXT, MD, RTF)")
            logger.info("   - extract_text: OCR распознавание текста с изображений (с поддержкой русского и английского)")
            logger.info("   - generate_file: создание файлов (DOCX, Excel, PDF, Markdown) в папке output")
            logger.info("📱 Telegram бот поддерживает:")
            logger.info("   - Изображения: автоматический OCR + vision анализ")
            logger.info("   - Документы: полная обработка всех поддерживаемых форматов")
            logger.info("   - Аудио: транскрипция через Whisper")
            logger.info("   - Видео: анализ кадров + OCR + извлечение аудио")
            logger.info(f"🧠 Модель: {os.path.basename(self.brain_model)}")
            logger.info(f"📊 {self.get_context_info()}")
            logger.info("💻 Доступные команды: 'stats' (метрики), 'reset' (сброс), 'logs' (логи), 'export' (экспорт), 'memory' (память), 'gpu' (видеокарта), 'search' (поиск), 'preferences' (предпочтения), 'cleanup' (очистка), 'unload' (выгрузка моделей), 'exit' (выход)")
            logger.info("="*60)

        vision_desc = ""
        audio_text = ""
        while True:
            try:
                # АВТОМАТИЧЕСКАЯ ОБРАБОТКА ФАЙЛОВ ОТКЛЮЧЕНА - ТЕПЕРЬ ТОЛЬКО ПО ЗАПРОСУ
                # Автоматический поиск и обработка изображений/аудио удалена
                # Используйте новые действия: list_files и process_document
                
                vision_desc = ""
                audio_text = ""

                # 3. Запрашиваем у пользователя текстовый вопрос
                prompt_text = "\n👤 Ваш вопрос (или Enter для пропуска, либо вставьте ссылку на YouTube): "
                auto_question = os.getenv("AUTO_QUESTION")
                try:
                    if auto_question:
                        user_input = auto_question.strip()
                        logger.info(f"🤖 AUTO_QUESTION активен, используем запрос: {user_input}")
                        # Отключаем AUTO_QUESTION после первого использования, чтобы избежать зацикливания
                        os.environ.pop("AUTO_QUESTION", None)
                        auto_question = None
                    elif getattr(self, 'show_images_locally', True) or not IS_WEB:
                        logger.info("⌛ Жду ввод пользователя...")
                        logger.debug("➡️ Ожидание input() начато")
                        user_input = self._read_user_input(prompt_text)
                        logger.debug("✅ input() завершён")
                        logger.debug(f"[DEBUG] Получен ввод: {repr(user_input)}")
                        logger.debug(f"📝 Сырой ввод (repr): {repr(user_input)}")
                        logger.info(f"📥 Ввод пользователя: {user_input}")
                    else:
                        # В веб-режиме локальный ввод отключен
                        logger.debug("🌐 Веб-режим: локальный ввод отключен, ожидаю события из интерфейса")
                        user_input = ""
                except EOFError:
                    # Если ввод из файла/pipe, используем пустую строку
                    user_input = ""
                    if getattr(self, 'show_images_locally', True) or not IS_WEB:
                        logger.info("📝 Ввод из файла/pipe, продолжаю...")
                if user_input.lower() in ['exit', 'quit', 'выход']:
                    if getattr(self, 'show_images_locally', True):
                        logger.info("👋 До свидания!")
                    break
                if user_input.lower() in ['stats', 'метрики', 'статистика']:
                    # Показываем метрики производительности только в консольном режиме
                    if getattr(self, 'show_images_locally', True):
                        stats = self.get_performance_stats()
                        logger.info("\n📊 МЕТРИКИ ПРОИЗВОДИТЕЛЬНОСТИ:")
                        logger.info(f"   Всего действий: {stats['total_actions']}")
                        logger.info(f"   Среднее время ответа: {stats['avg_response_time']} сек")
                        if stats['recent_metrics']:
                            logger.info("   Последние действия:")
                            for metric in stats['recent_metrics'][-5:]:  # Показываем последние 5
                                timestamp = time.strftime("%H:%M:%S", time.localtime(metric['timestamp']))
                                logger.info(f"     [{timestamp}] {metric['action']}: {metric['response_time']:.2f} сек")
                        logger.info(f"   {self.get_context_info()}")
                    continue
                if user_input.lower() in ['reset', 'сброс', 'очистить']:
                    # Сбрасываем метрики и историю
                    self.performance_metrics.clear()
                    self.conversation_history.clear()
                    self.current_context_length = 0
                    if getattr(self, 'show_images_locally', True):
                        logger.info("🔄 Метрики и история сброшены")
                    continue
                if user_input.lower() in ['logs', 'логи']:
                    # Показываем последние записи из лог-файла только в консольном режиме
                    if getattr(self, 'show_images_locally', True):
                        try:
                            with open("ai_orchestrator.log", "r", encoding="utf-8") as f:
                                lines = f.readlines()
                                logger.info("\n📝 ПОСЛЕДНИЕ ЗАПИСИ В ЛОГЕ:")
                                for line in lines[-10:]:  # Показываем последние 10 строк
                                    logger.info(f"   {line.strip()}")
                        except Exception as e:
                            logger.error(f"Ошибка чтения лог-файла: {e}")
                    continue
                if user_input.lower() in ['export', 'экспорт']:
                    # Экспортируем метрики в JSON файл только в консольном режиме
                    if getattr(self, 'show_images_locally', True):
                        try:
                            stats = self.get_performance_stats()
                            export_data = {
                                "export_time": time.strftime("%Y-%m-%d %H:%M:%S"),
                                "performance_stats": stats,
                                "context_info": {
                                    "current": self.current_context_length,
                                    "safe": self.safe_context_length,
                                    "max": self.max_context_length
                                }
                            }
                            filename = f"metrics_export_{int(time.time())}.json"
                            with open(filename, "w", encoding="utf-8") as f:
                                json.dump(export_data, f, ensure_ascii=False, indent=2)
                            logger.info(f"📊 Метрики экспортированы в {filename}")
                        except Exception as e:
                            logger.error(f"Ошибка экспорта метрик: {e}")
                    continue
                if user_input.lower() in ['memory', 'память', 'mem']:
                    # Показываем статистику памяти ChromaDB
                    if getattr(self, 'show_images_locally', True):
                        try:
                            stats = self.get_memory_stats()
                            logger.info("\n🧠 СТАТИСТИКА ПАМЯТИ CHROMADB:")
                            if "error" not in stats:
                                logger.info(f"   Всего записей: {stats['total_records']}")
                                logger.info(f"   Диалогов: {stats['conversations']}")
                                logger.info(f"   Предпочтений: {stats['preferences']}")
                                logger.info(f"   Путь к БД: {stats['database_path']}")
                                logger.info(f"   Модель эмбеддингов: {stats['embedding_model']}")
                            else:
                                logger.error(f"   Ошибка получения статистики: {stats['error']}")
                        except Exception as e:
                            logger.error(f"Ошибка получения статистики памяти: {e}")
                    continue
                if user_input.lower() in ['gpu', 'видеокарта', 'gpuinfo']:
                    # Показываем информацию о GPU для ChromaDB
                    if getattr(self, 'show_images_locally', True):
                        try:
                            gpu_info = self.get_gpu_info()
                            logger.info("\n🎮 ИНФОРМАЦИЯ О GPU ДЛЯ CHROMADB:")
                            if "error" not in gpu_info:
                                logger.info(f"   GPU доступен: {'Да' if gpu_info['gpu_available'] else 'Нет'}")
                                if gpu_info['gpu_available']:
                                    logger.info(f"   Название GPU: {gpu_info['gpu_name']}")
                                    logger.info(f"   Память GPU: {gpu_info['gpu_memory']:.1f} GB")
                                logger.info(f"   Используемое устройство: {gpu_info['device_used']}")
                            else:
                                logger.error(f"   Ошибка получения информации о GPU: {gpu_info['error']}")
                        except Exception as e:
                            logger.error(f"Ошибка получения информации о GPU: {e}")
                    continue
                if user_input.lower() in ['cleanup', 'очистка', 'clean']:
                    # Очищаем старые записи из памяти
                    if getattr(self, 'show_images_locally', True):
                        try:
                            days = input("🗑️ Введите количество дней для хранения записей (по умолчанию 30): ").strip()
                            days_to_keep = int(days) if days.isdigit() else 30
                            deleted_count = self.cleanup_old_memory(days_to_keep)
                            logger.info(f"🧹 Удалено {deleted_count} старых записей из памяти")
                        except Exception as e:
                            logger.error(f"Ошибка очистки памяти: {e}")
                    continue
                if user_input.lower() in ['unload', 'выгрузка', 'unload_models']:
                    # Принудительная выгрузка всех моделей кроме мозга
                    logger.info("🔧 Принудительная выгрузка всех моделей...")
                    try:
                        import gc
                        import torch
                        
                        # Выключаем все системы
                        self.use_image_generation = False
                        self.use_vision = False
                        self.use_audio = False
                        
                        # Выгружаем pipeline для генерации изображений
                        if hasattr(self, 'current_pipeline') and self.current_pipeline is not None:
                            try:
                                if hasattr(self.current_pipeline, 'to'):
                                    self.current_pipeline.to('cpu')
                                del self.current_pipeline
                                self.current_pipeline = None
                                logger.info("✅ Pipeline для генерации изображений выгружен")
                            except Exception as e:
                                logger.warning(f"⚠️ Ошибка при выгрузке pipeline: {e}")
                        
                        # Выгружаем другие модели
                        model_attrs = ['vision_model', 'vision_processor', 'vision_pipeline', 
                                     'whisper_model', 'audio_model', 'tts_model']
                        for attr in model_attrs:
                            if hasattr(self, attr):
                                try:
                                    model = getattr(self, attr)
                                    if model is not None and hasattr(model, 'to'):
                                        model.to('cpu')
                                    delattr(self, attr)
                                    logger.info(f"✅ {attr} выгружен")
                                except Exception as e:
                                    logger.warning(f"⚠️ Ошибка при выгрузке {attr}: {e}")
                        
                        # Очистка GPU памяти
                        if torch.cuda.is_available():
                            torch.cuda.empty_cache()
                            torch.cuda.synchronize()
                            logger.info("✅ GPU память очищена")
                        
                        # Сборка мусора
                        gc.collect()
                        logger.info("✅ Все модели выгружены, память освобождена")
                        
                    except Exception as e:
                        logger.error(f"Ошибка при выгрузке моделей: {e}")
                    continue
                if user_input.lower() in ['search', 'поиск', 'find']:
                    # Поиск похожих диалогов в памяти
                    if getattr(self, 'show_images_locally', True):
                        try:
                            query = input("🔍 Введите поисковый запрос: ").strip()
                            if query:
                                results = self.search_similar_conversations(query, n_results=3)
                                if results:
                                    logger.info(f"\n🔍 НАЙДЕНО {len(results)} ПОХОЖИХ ДИАЛОГОВ:")
                                    for i, result in enumerate(results, 1):
                                        logger.info(f"   {i}. Схожесть: {result['similarity']:.2f}")
                                        logger.info(f"      ID: {result['id']}")
                                        logger.info(f"      Текст: {result['document'][:100]}...")
                                        logger.info(f"      Время: {time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(result['metadata']['timestamp']))}")
                                        logger.info("")
                                else:
                                    logger.info("🔍 Похожих диалогов не найдено")
                        except Exception as e:
                            logger.error(f"Ошибка поиска: {e}")
                    continue
                if user_input.lower() in ['preferences', 'предпочтения', 'prefs']:
                    # Показываем предпочтения пользователя
                    if getattr(self, 'show_images_locally', True):
                        try:
                            query = input("👤 Введите контекст для поиска предпочтений (или Enter для всех): ").strip()
                            preferences = self.get_user_preferences(query if query else None)
                            if preferences:
                                logger.info(f"\n👤 ПРЕДПОЧТЕНИЯ ПОЛЬЗОВАТЕЛЯ:")
                                logger.info(preferences)
                            else:
                                logger.info("👤 Предпочтения не найдены")
                        except Exception as e:
                            logger.error(f"Ошибка получения предпочтений: {e}")
                    continue
                if not user_input:
                    continue

                # 4. Перехват: если есть YouTube-ссылка, скачиваем видео и аудио, обрабатываем аудио, затем формируем brain_input: вопрос пользователя (без ссылки, вместо неё название ролика), затем текст из аудио, и только после этого отправляем brain_input в мозг
                # re уже импортирован в начале файла
                yt_url_match = re.search(r'https?://(?:www\.)?(?:youtube\.com|youtu\.be)/\S+', user_input)
                yt_processed = False  # Флаг для отслеживания обработки YouTube
                if yt_url_match:
                    yt_url = yt_url_match.group(0)
                    logger.info("🔗 Обнаружена ссылка на YouTube, скачиваю видео...")
                    
                    # Проверяем VPN статус перед скачиванием
                    if not self.check_vpn_status():
                        logger.warning("⚠️ VPN может не работать корректно. YouTube может быть недоступен.")
                        user_input_no_url = re.sub(r'https?://\S+', '[YouTube недоступен - проверьте VPN]', user_input).strip()
                        brain_input = f"{user_input_no_url}\n\n[ОШИБКА]: Не удалось скачать YouTube видео. Проверьте VPN соединение."
                        ai_response = self.call_brain_model(brain_input)
                        continue
                    
                    # Проверяем доступность YouTube ссылки
                    if not self.check_youtube_accessibility(yt_url):
                        logger.error("❌ YouTube ссылка недоступна. Проверьте VPN или ссылку.")
                        user_input_no_url = re.sub(r'https?://\S+', '[YouTube недоступен]', user_input).strip()
                        brain_input = f"{user_input_no_url}\n\n[ОШИБКА]: YouTube ссылка недоступна. Проверьте VPN или ссылку."
                        ai_response = self.call_brain_model(brain_input)
                        continue
                    
                    # Получаем информацию о видео
                    video_info = self.get_youtube_info(yt_url)
                    if video_info.get('success'):
                        video_title = video_info['title']
                        logger.info(f"📹 Название видео: {video_title}")
                    else:
                        video_title = "Неизвестное видео"
                        logger.warning(f"⚠️ Не удалось получить название: {video_info.get('error', 'Неизвестная ошибка')}")
                    yt_video = self.download_youtube_video(yt_url)
                    if yt_video:
                        logger.info(f"✅ Видео скачано: {yt_video}")
                        # video_title уже получен выше из get_youtube_info
                    else:
                        logger.error("❌ Не удалось скачать видео с YouTube")
                        video_title = "Видео не скачано"
                    logger.info("🔗 Скачиваю аудиодорожку для транскрипции...")
                    yt_audio = self.download_youtube_audio(yt_url)
                    if yt_audio:
                        logger.info(f"✅ Аудио скачано: {yt_audio}")
                        # Определяем язык по названию видео, а не по имени файла
                        if "english" in video_title.lower() or "eng" in video_title.lower():
                            lang = "en"
                        elif "рус" in video_title.lower() or "russian" in video_title.lower():
                            lang = "ru"
                        else:
                            # Автоматически определяем язык по названию видео
                            # Для Rick Astley и подобных - английский
                            if any(word in video_title.lower() for word in ["rick", "astley", "never", "gonna", "give", "you", "up"]):
                                lang = "en"
                                logger.info("🌐 Автоматически определен язык: английский (по названию видео)")
                            else:
                                # По умолчанию используем русский
                                lang = "ru"
                                logger.info("🌐 Используется язык по умолчанию: русский")
                        audio_text = self.transcribe_audio_whisper(yt_audio, lang=lang, use_separator=getattr(self, 'use_separator', True))
                        # --- VISION ПОКАДРОВО ---
                        vision_frames_desc = ""
                        if getattr(self, 'use_vision', False) and yt_video:
                            logger.info("🖼️ Извлекаю кадры из видео для vision...")
                            frames = self.extract_video_frames(yt_video, fps=1)
                            frame_results = []  # [(timecode, desc)]
                            for idx, (timecode, b64) in enumerate(frames):
                                if not b64:
                                    continue
                                vision_prompt = "Describe absolutely everything you see in the image, including all small details, their positions, and any visible text. Be as detailed as possible."
                                desc = self.call_vision_model(b64 + "\n" + vision_prompt)
                                
                                # Добавляем OCR к кадру видео
                                if getattr(self, 'use_ocr', False):
                                    try:
                                        # Декодируем base64 для OCR
                                        from PIL import Image
                                        
                                        image_data = base64.b64decode(b64)
                                        image = Image.open(io.BytesIO(image_data))
                                        
                                        # Извлекаем текст с помощью OCR
                                        ocr_text, ocr_error = self.extract_text_from_image_object(image)
                                        if ocr_text and ocr_text.strip():
                                            desc += f"\n[OCR TEXT]: {ocr_text.strip()}"
                                            logger.info(f"[OCR][{timecode}] Извлечен текст: {ocr_text.strip()}")
                                        elif ocr_error:
                                            logger.warning(f"[OCR][{timecode}] Ошибка OCR: {ocr_error}")
                                    except Exception as e:
                                        logger.warning(f"[OCR][{timecode}] Ошибка OCR: {e}")
                                
                                frame_results.append((timecode, desc))
                                logger.info(f"[VISION][{timecode}] {desc}")
                            # Группировка одинаковых описаний
                            # collections.defaultdict уже импортирован в начале файла
                            grouped = []  # [(list_of_timecodes, desc)]
                            prev_desc = None
                            prev_times = []
                            for i, (tc, desc) in enumerate(frame_results):
                                if prev_desc is None:
                                    prev_desc = desc
                                    prev_times = [tc]
                                elif desc == prev_desc:
                                    prev_times.append(tc)
                                else:
                                    grouped.append((prev_times, prev_desc))
                                    prev_desc = desc
                                    prev_times = [tc]
                            if prev_times:
                                grouped.append((prev_times, prev_desc))
                            # Формируем текст с диапазонами и списками
                            def format_timecodes(times):
                                def tc_to_sec(tc):
                                    h, m, s = map(int, tc.strip('[]').split(':'))
                                    return h*3600 + m*60 + s
                                if len(times) == 1:
                                    return times[0]
                                secs = [tc_to_sec(t) for t in times]
                                sorted_pairs = sorted(zip(secs, times))
                                secs_sorted, times_sorted = zip(*sorted_pairs)
                                # Проверяем, есть ли диапазон подряд
                                ranges = []
                                start = end = secs_sorted[0]
                                start_tc = times_sorted[0]
                                for i in range(1, len(secs_sorted)):
                                    if secs_sorted[i] == end + 1:
                                        end = secs_sorted[i]
                                    else:
                                        if start != end:
                                            ranges.append(f"[{start_tc}-{times_sorted[i-1]}]")
                                        else:
                                            ranges.append(f"{start_tc}")
                                        start = end = secs_sorted[i]
                                        start_tc = times_sorted[i]
                                if start != end:
                                    ranges.append(f"[{start_tc}-{times_sorted[-1]}]")
                                else:
                                    ranges.append(f"{start_tc}")
                                return ', '.join(ranges)
                            for times, desc in grouped:
                                vision_frames_desc += f"{format_timecodes(times)}: {desc}\n"
                        # Формируем brain_input: вопрос пользователя (без ссылки, вместо неё название ролика), ответы vision по кадрам, затем текст из аудио (с таймкодами)
                        user_input_no_url = re.sub(r'https?://\S+', f'[Видео]: {video_title}' if video_title else '', user_input).strip()
                        brain_input = ""
                        brain_input += user_input_no_url
                        if vision_frames_desc:
                            wrapped_frames = self._wrap_info_if_needed(vision_frames_desc, source="video_frames")
                            brain_input += f"\n[Покадровое описание видео]:\n{wrapped_frames}"
                        if audio_text:
                            wrapped_audio = self._wrap_info_if_needed(audio_text, source="audio")
                            brain_input += f"\n[Текст из аудио]:\n{wrapped_audio}"
                        # Отправляем в мозг и обрабатываем ответ
                        ai_response = self.call_brain_model(brain_input)
                        # Показываем информацию о контексте
                        logger.info(f"📊 {self.get_context_info()}")
                        continue_dialog = self.process_ai_response(ai_response)
                        if not continue_dialog:
                            logger.info("\n" + "="*60)
                        yt_processed = True  # Отмечаем, что YouTube обработан
                        continue  # пропускаем стандартный brain_input ниже
                    else:
                        logger.error("❌ Не удалось скачать аудио с YouTube")

                # 5. Собираем итоговый запрос для мозга (только если не было YouTube обработки)
                if not yt_processed:  # Используем флаг вместо yt_url_match
                    brain_input = ""
                    if vision_desc:
                        brain_input += f"[Описание изображения]:\n{vision_desc}\n"
                    if audio_text:
                        wrapped_audio = self._wrap_info_if_needed(audio_text, source="audio")
                        brain_input += f"[Текст из аудио]:\n{wrapped_audio}\n"
                    brain_input += user_input

                    logger.info(f"🧠 Формирую запрос к модели (длина {len(brain_input)} символов)")
                    logger.debug(f"brain_input (truncated): {brain_input[:200]}")

                    # 6. Отправляем запрос в мозг
                    ai_response = self.call_brain_model(brain_input)
                    # Показываем информацию о контексте
                    logger.info(f"📊 {self.get_context_info()}")
                    # Обрабатываем ответ AI
                    continue_dialog = self.process_ai_response(ai_response)
                    if not continue_dialog:
                        logger.info("\n" + "="*60)

            except Exception as e:
                if isinstance(e, KeyboardInterrupt):
                    logger.info("\n👋 Программа прервана пользователем")
                    break
                logger.error(f"❌ Неожиданная ошибка: {str(e)}")
                # Убираем дублирование обработки аудио - это уже делается выше
                audio_text = ""

                # 3. Запрашиваем у пользователя текстовый вопрос
                try:
                    user_input = input("\n👤 Ваш вопрос (или Enter для пропуска, либо вставьте ссылку на YouTube): ").strip()
                    logger.info(f"📥 Ввод пользователя (после ошибки): {user_input}")
                except EOFError:
                    # Если ввод из файла/pipe, используем пустую строку
                    user_input = ""
                    logger.info("📝 Ввод из файла/pipe, продолжаю...")
                if user_input.lower() in ['exit', 'quit', 'выход']:
                    logger.info("👋 До свидания!")
                    break
                if user_input.lower() in ['stats', 'метрики', 'статистика']:
                    # Показываем метрики производительности
                    stats = self.get_performance_stats()
                    logger.info("\n📊 МЕТРИКИ ПРОИЗВОДИТЕЛЬНОСТИ:")
                    logger.info(f"   Всего действий: {stats['total_actions']}")
                    logger.info(f"   Среднее время ответа: {stats['avg_response_time']} сек")
                    if stats['recent_metrics']:
                        logger.info("   Последние действия:")
                        for metric in stats['recent_metrics'][-5:]:  # Показываем последние 5
                            timestamp = time.strftime("%H:%M:%S", time.localtime(metric['timestamp']))
                            logger.info(f"     [{timestamp}] {metric['action']}: {metric['response_time']:.2f} сек")
                    logger.info(f"   {self.get_context_info()}")
                    continue
                if user_input.lower() in ['reset', 'сброс', 'очистить']:
                    # Сбрасываем метрики и историю
                    self.performance_metrics.clear()
                    self.conversation_history.clear()
                    self.current_context_length = 0
                    logger.info("🔄 Метрики и история сброшены")
                    continue
                if user_input.lower() in ['logs', 'логи']:
                    # Показываем последние записи из лог-файла
                    try:
                        with open("ai_orchestrator.log", "r", encoding="utf-8") as f:
                            lines = f.readlines()
                            logger.info("\n📝 ПОСЛЕДНИЕ ЗАПИСИ В ЛОГЕ:")
                            for line in lines[-10:]:  # Показываем последние 10 строк
                                logger.info(f"   {line.strip()}")
                    except Exception as e:
                        logger.error(f"Ошибка чтения лог-файла: {e}")
                    continue
                if user_input.lower() in ['export', 'экспорт']:
                    # Экспортируем метрики в JSON файл
                    try:
                        stats = self.get_performance_stats()
                        export_data = {
                            "export_time": time.strftime("%Y-%m-%d %H:%M:%S"),
                            "performance_stats": stats,
                            "context_info": {
                                "current": self.current_context_length,
                                "safe": self.safe_context_length,
                                "max": self.max_context_length
                            }
                        }
                        filename = f"metrics_export_{int(time.time())}.json"
                        with open(filename, "w", encoding="utf-8") as f:
                            json.dump(export_data, f, ensure_ascii=False, indent=2)
                        logger.info(f"📊 Метрики экспортированы в {filename}")
                    except Exception as e:
                        logger.error(f"Ошибка экспорта метрик: {e}")
                    continue
                if not user_input:
                    continue



            except KeyboardInterrupt:
                logger.info("\n�� Программа прервана пользователем")
                break



    def start_telegram_bot(self) -> bool:
        """Запускает Telegram бота. Возвращает True при успешном старте, иначе False."""
        if not self.telegram_bot_token:
            logger.warning("❌ Telegram Bot токен не указан")
            return False
        
        # Предварительная проверка токена через getMe и редактирование токена в логах
        try:
            redacted = self.telegram_bot_token[:10] + "..." if len(self.telegram_bot_token) > 13 else "***"
            logger.info(f"🔐 Проверяю Telegram токен (redacted: {redacted})")
            resp = requests.get(f"https://api.telegram.org/bot{self.telegram_bot_token}/getMe", timeout=5)
            if resp.status_code != 200:
                logger.error("❌ Ошибка проверки Telegram токена: сервер вернул неуспешный статус")
                return False
            data = resp.json()
            if not data.get("ok"):
                # Не логируем сырой токен
                logger.error(f"❌ Telegram токен отклонен сервером (token: {redacted})")
                return False
        except Exception as e:
            logger.error(f"❌ Ошибка проверки Telegram токена: {e}")
            return False
        
        try:
            # Создаем приложение
            self.telegram_app = Application.builder().token(self.telegram_bot_token).build()
            
            # Добавляем обработчики
            self.telegram_app.add_handler(CommandHandler("start", self._telegram_start))
            self.telegram_app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, self._telegram_text_message))
            self.telegram_app.add_handler(MessageHandler(filters.PHOTO, self._telegram_photo_message))
            self.telegram_app.add_handler(MessageHandler(filters.AUDIO | filters.VOICE, self._telegram_audio_message))
            self.telegram_app.add_handler(MessageHandler(filters.Document.ALL, self._telegram_document_message))
            
            # Запускаем бота в фоне в отдельном потоке
            import threading
            def run_bot():
                loop = None
                try:
                    # Создаем новый event loop для потока
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)
                    from typing import Any, cast
                    coro = self.telegram_app.run_polling(allowed_updates=Update.ALL_TYPES)
                    if coro is not None:
                        loop.run_until_complete(cast(Any, coro))
                except Exception as e:
                    # В веб-режиме логируем тихо
                    if not getattr(self, 'show_images_locally', True):
                        logger.error(f"❌ Ошибка в Telegram боте: {e}")
                    else:
                        logger.debug(f"Telegram bot polling error: {e}")
                finally:
                    if loop is not None:
                        try:
                            loop.close()
                        except Exception:
                            pass
            
            bot_thread = threading.Thread(target=run_bot, daemon=True)
            bot_thread.start()
            
            # Показываем сообщение только в консольном режиме
            if not getattr(self, 'show_images_locally', True):
                logger.info("🤖 Telegram бот запущен в фоновом режиме")
            return True
            
        except Exception as e:
            # В веб-режиме логируем тихо
            if not getattr(self, 'show_images_locally', True):
                logger.error(f"❌ Ошибка запуска Telegram бота: {e}")
            else:
                logger.debug(f"Telegram bot startup error: {e}")
            return False

    async def _safe_reply(self, update: Update, message: str):
        """Безопасная отправка сообщения в Telegram"""
        if update and update.message:
            await update.message.reply_text(message)

    async def _telegram_start(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Обработчик команды /start"""
        if update is None or update.message is None or update.effective_user is None:
            return
        user_id = str(update.effective_user.id)
        # Разрешаем доступ всем пользователям
        
        await self._safe_reply(update,
            "🤖 Привет! Я Нейро - AI оркестратор.\n"
            "Я могу:\n"
            "• Обрабатывать текстовые сообщения\n"
            "• Анализировать изображения\n"
            "• Транскрибировать аудио\n"
            "• Генерировать изображения\n"
            "• Выполнять команды PowerShell\n"
            "• Искать информацию в интернете\n\n"
            "Просто отправьте мне сообщение, изображение или аудио!"
        )

    async def _telegram_text_message(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Обработчик текстовых сообщений"""
        if update is None or update.message is None or update.effective_user is None or update.effective_chat is None:
            return
        user_id = str(update.effective_user.id)
        # Разрешаем доступ всем пользователям
        
        text = update.message.text if update.message and update.message.text else ""
        
        # Проверяем специальные команды OCR
        if any(keyword in text.lower() for keyword in ['ocr', 'распознай текст', 'извлеки текст', 'что написано']):
            # Если есть последнее изображение, применяем к нему OCR
            if hasattr(self, 'last_telegram_image') and self.last_telegram_image:
                await self._safe_reply(update, "🔄 Применяю OCR к последнему изображению...")
                try:
                    if getattr(self, 'use_ocr', False):
                        # Конвертируем base64 в PIL Image
                        from PIL import Image
                        
                        image_data = base64.b64decode(self.last_telegram_image)
                        image = Image.open(io.BytesIO(image_data))
                        
                        # Извлекаем текст принудительно
                        ocr_text, ocr_error = self.extract_text_from_image_object(image)
                        
                        if ocr_text and ocr_text.strip():
                            await update.message.reply_text(f"📖 Извлеченный текст:\n\n{ocr_text.strip()}")
                        elif ocr_error:
                            await update.message.reply_text(f"❌ Ошибка OCR: {ocr_error}")
                        else:
                            await update.message.reply_text("⚠️ Текст в изображении не найден")
                    else:
                        await update.message.reply_text("❌ OCR отключен в системе")
                        
                except Exception as e:
                    await update.message.reply_text(f"❌ Ошибка OCR: {str(e)}")
                return
            else:
                await update.message.reply_text("❌ Нет изображения для OCR. Отправьте сначала изображение.")
                return
        
        await update.message.reply_text("🔄 Обрабатываю ваше сообщение...")

        try:
            # Отправляем в мозг
            ai_response = self.call_brain_model(text or "")

            # Обрабатываем ответ AI
            continue_dialog = self.process_ai_response(ai_response)

            if not continue_dialog:
                # Если диалог завершен, отправляем финальный ответ
                if hasattr(self, 'last_final_response') and self.last_final_response:
                    await update.message.reply_text(self.last_final_response)

                    # Если есть сгенерированное изображение, отправляем его
                    if hasattr(self, 'last_generated_image_b64') and self.last_generated_image_b64:
                        try:
                            # Конвертируем base64 в bytes
                            img_bytes = base64.b64decode(self.last_generated_image_b64)

                            # Отправляем изображение
                            await context.bot.send_photo(
                                chat_id=update.effective_chat.id,
                                photo=img_bytes,
                                caption="🎨 Сгенерированное изображение"
                            )

                            # Очищаем
                            self.last_generated_image_b64 = None

                        except Exception as e:
                            # В веб-режиме логируем тихо
                            if not getattr(self, 'show_images_locally', True):
                                logger.error(f"❌ Ошибка отправки изображения: {e}")
                            else:
                                logger.debug(f"Telegram image send error: {e}")
                            await update.message.reply_text("❌ Ошибка отправки изображения")
                    
                    # Если есть сгенерированный файл, отправляем его
                    if hasattr(self, 'last_generated_file_path') and self.last_generated_file_path:
                        try:
                            # Проверяем что файл существует
                            if os.path.exists(self.last_generated_file_path):
                                # Отправляем файл как документ
                                with open(self.last_generated_file_path, 'rb') as file:
                                    await context.bot.send_document(
                                        chat_id=update.effective_chat.id,
                                        document=file,
                                        filename=self.last_generated_file_name or os.path.basename(self.last_generated_file_path),
                                        caption="📄 Сгенерированный файл"
                                    )
                                
                                # Очищаем
                                self.last_generated_file_path = None
                                self.last_generated_file_name = None
                                
                            else:
                                await update.message.reply_text("❌ Сгенерированный файл не найден")
                                
                        except Exception as e:
                            # В веб-режиме логируем тихо
                            if not getattr(self, 'show_images_locally', True):
                                logger.error(f"❌ Ошибка отправки файла: {e}")
                            else:
                                logger.debug(f"Telegram file send error: {e}")
                            await update.message.reply_text("❌ Ошибка отправки файла")
                else:
                    await update.message.reply_text("✅ Задача выполнена!")
            else:
                # Если диалог продолжается, отправляем промежуточный ответ
                await update.message.reply_text("🔄 Обрабатываю... Пожалуйста, подождите.")

        except Exception as e:
            # Веб-режим: логируем тихо
            if getattr(self, 'show_images_locally', True):
                logger.error(f"❌ Ошибка обработки текстового сообщения: {e}")
            else:
                logger.debug(f"Telegram text message error: {e}")
            await update.message.reply_text(f"❌ Произошла ошибка: {str(e)}")

    async def _telegram_photo_message(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Обработчик фотографий"""
        if update is None or update.message is None or update.effective_user is None or update.effective_chat is None:
            return
        user_id = str(update.effective_user.id)
        # Разрешаем доступ всем пользователям
        
        await update.message.reply_text("🖼️ Обрабатываю изображение...")
        
        try:
            # Получаем фото
            photo = update.message.photo[-1]  # Берем самое большое фото
            file = await context.bot.get_file(photo.file_id)
            
            # Скачиваем фото
            photo_bytes = await file.download_as_bytearray()
            photo_b64 = base64.b64encode(photo_bytes).decode('ascii')
            
            # Анализируем изображение с vision моделью
            vision_desc = self.call_vision_model(photo_b64)
            
            # Применяем умный OCR
            result_message = f"👁️ Описание изображения:\n{vision_desc}"
            
            if getattr(self, 'use_ocr', False):
                try:
                    # Конвертируем байты в PIL Image для OCR
                    from PIL import Image
                    import io
                    
                    image = Image.open(io.BytesIO(photo_bytes))
                    
                    # Проверяем, нужно ли применять OCR
                    should_use_ocr = self.should_use_ocr_on_image(vision_desc)
                    
                    if should_use_ocr:
                        # Извлекаем текст с помощью OCR
                        ocr_text, ocr_error = self.extract_text_from_image_object(image)
                        
                        if ocr_text and ocr_text.strip():
                            result_message += f"\n\n📖 Извлеченный текст:\n{ocr_text.strip()}"
                            await update.message.reply_text("✅ Обнаружен текст на изображении!")
                        elif ocr_error:
                            result_message += f"\n\n⚠️ OCR ошибка: {ocr_error}"
                    else:
                        await update.message.reply_text("ℹ️ Текст на изображении не обнаружен")
                        
                except Exception as ocr_exception:
                    result_message += f"\n\n⚠️ Ошибка OCR: {str(ocr_exception)}"
                    
            # Отправляем полное описание
            await update.message.reply_text(result_message)
            
            # Сохраняем для возможного использования в диалоге
            self.last_telegram_image = photo_b64
            
        except Exception as e:
            # В веб-режиме логируем тихо
            if getattr(self, 'show_images_locally', True):
                logger.error(f"❌ Ошибка обработки фото: {e}")
            else:
                logger.debug(f"Telegram photo processing error: {e}")
            await update.message.reply_text(f"❌ Ошибка обработки изображения: {str(e)}")

    async def _telegram_audio_message(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Обработчик аудио сообщений"""
        if update is None or update.message is None or update.effective_user is None or update.effective_chat is None:
            return
        user_id = str(update.effective_user.id)
        # Разрешаем доступ всем пользователям
        
        await update.message.reply_text("🎵 Обрабатываю аудио...")
        
        try:
            # Получаем аудио
            if update.message.audio:
                audio = update.message.audio
            else:
                audio = update.message.voice

            if audio is None:
                await update.message.reply_text("❌ В сообщении нет аудиофайла")
                return
            
            file = await context.bot.get_file(audio.file_id)
            
            # Скачиваем аудио
            audio_bytes = await file.download_as_bytearray()
            
            # Сохраняем во временный файл
            temp_dir = os.path.join(os.path.dirname(__file__), "temp_audio")
            os.makedirs(temp_dir, exist_ok=True)
            temp_file = os.path.join(temp_dir, f"telegram_audio_{int(time.time())}.ogg")
            
            with open(temp_file, 'wb') as f:
                f.write(audio_bytes)
            
            # Транскрибируем аудио
            transcript = self.transcribe_audio_whisper(temp_file, use_separator=False)
            
            if transcript and not transcript.startswith("[Whisper error]"):
                await update.message.reply_text(f"🎤 Транскрипция аудио:\n{transcript}")
                
                # Сохраняем для возможного использования в диалоге
                self.last_telegram_audio_transcript = transcript
            else:
                await update.message.reply_text("❌ Не удалось распознать аудио")
            
            # Удаляем временный файл
            try:
                os.remove(temp_file)
            except Exception:
                pass
                
        except Exception as e:
            # В веб-режиме логируем тихо
            if getattr(self, 'show_images_locally', True):
                logger.error(f"❌ Ошибка обработки аудио: {e}")
            else:
                logger.debug(f"Telegram audio processing error: {e}")
            await update.message.reply_text(f"❌ Ошибка обработки аудио: {str(e)}")

    async def _telegram_document_message(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Обработчик документов"""
        if update is None or update.message is None or update.effective_user is None or update.effective_chat is None:
            return
        user_id = str(update.effective_user.id)
        # Разрешаем доступ всем пользователям
        
        await update.message.reply_text("📄 Обрабатываю документ...")
        
        try:
            # Получаем документ
            document = update.message.document
            if document is None:
                await update.message.reply_text("❌ В сообщении нет документа")
                return
            
            file_name = document.file_name or "unknown_file"
            file_size = document.file_size
            
            # Проверяем размер файла (максимум 20MB)
            if file_size and file_size > 20 * 1024 * 1024:
                await update.message.reply_text("❌ Файл слишком большой (максимум 20MB)")
                return
            
            file = await context.bot.get_file(document.file_id)
            
            # Определяем тип файла и обработку
            file_lower = file_name.lower()
            
            if file_lower.endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp')):
                # Обрабатываем как изображение с OCR
                await self._process_telegram_image_document(update, file, file_name)
            elif file_lower.endswith(('.docx', '.doc', '.pdf', '.xlsx', '.xls', '.csv', '.txt', '.md', '.rtf', '.json', '.xml', '.html', '.htm')):
                # Обрабатываем как документ
                await self._process_telegram_text_document(update, file, file_name)
            elif file_lower.endswith(('.mp3', '.wav', '.ogg', '.m4a', '.flac', '.aac', '.wma')):
                # Обрабатываем как аудио
                await self._process_telegram_audio_document(update, file, file_name)
            elif file_lower.endswith(('.mp4', '.avi', '.mkv', '.mov', '.wmv', '.flv', '.webm')):
                # Обрабатываем как видео
                await self._process_telegram_video_document(update, file, file_name)
            else:
                await update.message.reply_text(f"❌ Неподдерживаемый тип файла: {file_name}\n\nПоддерживаются:\n• Изображения: JPG, PNG, GIF, BMP, WEBP\n• Документы: DOCX, DOC, PDF, XLSX, XLS, CSV, TXT, MD, RTF, JSON, XML, HTML\n• Аудио: MP3, WAV, OGG, M4A, FLAC, AAC, WMA\n• Видео: MP4, AVI, MKV, MOV, WMV, FLV, WEBM")
                
        except Exception as e:
            # В веб-режиме логируем тихо
            if getattr(self, 'show_images_locally', True):
                logger.error(f"❌ Ошибка обработки документа: {e}")
            else:
                logger.debug(f"Telegram document processing error: {e}")
            await update.message.reply_text(f"❌ Ошибка обработки документа: {str(e)}")

    async def _process_telegram_image_document(self, update: Update, file, file_name: str):
        """Обработка изображений через документы с OCR"""
        try:
            # Скачиваем изображение
            image_bytes = await file.download_as_bytearray()
            image_b64 = base64.b64encode(image_bytes).decode('ascii')
            
            # Анализируем изображение с vision моделью
            vision_desc = self.call_vision_model(image_b64)
            
            result_message = f"🖼️ Анализ изображения '{file_name}':\n\n👁️ Описание:\n{vision_desc}"
            
            # Применяем OCR
            if getattr(self, 'use_ocr', False):
                try:
                    from PIL import Image
                    import io
                    
                    image = Image.open(io.BytesIO(image_bytes))
                    
                    # Всегда применяем OCR для документов (более вероятно содержат текст)
                    ocr_text, ocr_error = self.extract_text_from_image_object(image)
                    
                    if ocr_text and ocr_text.strip():
                        result_message += f"\n\n📖 Извлеченный текст:\n{ocr_text.strip()}"
                        if update.message:
                            await update.message.reply_text("✅ Текст успешно извлечен из документа!")
                    elif ocr_error:
                        result_message += f"\n\n⚠️ OCR ошибка: {ocr_error}"
                    else:
                        result_message += f"\n\n⚠️ Текст в изображении не найден"
                        
                except Exception as ocr_exception:
                    result_message += f"\n\n⚠️ Ошибка OCR: {str(ocr_exception)}"
            else:
                result_message += f"\n\n📖 OCR отключен. Используйте vision описание выше."
            
            if update.message:
                await update.message.reply_text(result_message)
            
        except Exception as e:
            if update.message:
                await update.message.reply_text(f"❌ Ошибка обработки изображения: {str(e)}")

    async def _process_telegram_text_document(self, update: Update, file, file_name: str):
        """Обработка текстовых документов (DOCX, PDF, Excel)"""
        try:
            # Скачиваем документ
            doc_bytes = await file.download_as_bytearray()
            
            # Сохраняем во временный файл
            temp_dir = os.path.join(os.path.dirname(__file__), "temp_docs")
            os.makedirs(temp_dir, exist_ok=True)
            temp_file = os.path.join(temp_dir, f"telegram_doc_{int(time.time())}_{file_name}")
            
            with open(temp_file, 'wb') as f:
                f.write(doc_bytes)
            
            # Обрабатываем документ
            result = self.process_document_request(temp_file)
            
            if update.message:
                await update.message.reply_text(f"📄 Анализ документа '{file_name}':\n\n{result}")
            
            # Удаляем временный файл
            try:
                os.remove(temp_file)
            except Exception:
                pass
                
        except Exception as e:
            if update.message:
                await update.message.reply_text(f"❌ Ошибка обработки документа: {str(e)}")

    async def _process_telegram_audio_document(self, update: Update, file, file_name: str):
        """Обработка аудио файлов через документы"""
        try:
            # Скачиваем аудио
            audio_bytes = await file.download_as_bytearray()
            
            # Сохраняем во временный файл
            temp_dir = os.path.join(os.path.dirname(__file__), "temp_audio")
            os.makedirs(temp_dir, exist_ok=True)
            temp_file = os.path.join(temp_dir, f"telegram_audio_{int(time.time())}_{file_name}")
            
            with open(temp_file, 'wb') as f:
                f.write(audio_bytes)
            
            # Транскрибируем аудио
            transcript = self.transcribe_audio_whisper(temp_file, use_separator=False)
            
            if transcript and not transcript.startswith("[Whisper error]"):
                await self._safe_reply(update, f"🎤 Транскрипция аудио '{file_name}':\n\n{transcript}")
            else:
                await self._safe_reply(update, "❌ Не удалось распознать аудио")
            
            # Удаляем временный файл
            try:
                os.remove(temp_file)
            except Exception:
                pass
                
        except Exception as e:
            await self._safe_reply(update, f"❌ Ошибка обработки аудио: {str(e)}")

    async def _process_telegram_video_document(self, update: Update, file, file_name: str):
        """Обработка видео файлов через документы"""
        # Заглушка для видео
        await self._safe_reply(update, f"🎥 Видео '{file_name}' получено, но обработка видео пока не реализована.")

    def play_audio_file(self, audio_path: str) -> bool:
        """
        Воспроизводит аудиофайл один раз без зацикливания
        
        Args:
            audio_path: Путь к аудиофайлу
            
        Returns:
            True если воспроизведение запущено успешно, False при ошибке
        """
        try:
            import pygame  # type: ignore
            import time
            
            logger.info(f"🔊 Воспроизводим аудио: {os.path.basename(audio_path)}")
            
            # Инициализируем pygame mixer
            pygame.mixer.init()
            
            # Загружаем и воспроизводим аудио
            try:
                pygame.mixer.music.load(audio_path)
                pygame.mixer.music.play()
                
                # Ждем окончания воспроизведения
                while pygame.mixer.music.get_busy():
                    time.sleep(0.1)
                
                # Останавливаем и закрываем
                pygame.mixer.music.stop()
                pygame.mixer.quit()
                
                logger.info("✅ Аудио воспроизведено через pygame")
                return True
                
            except Exception as e:
                logger.warning(f"⚠️ pygame воспроизведение не удалось: {e}")
                pygame.mixer.quit()
                
                # Fallback: открываем в плеере по умолчанию
                import subprocess
                subprocess.Popen(["start", audio_path], shell=True)
                logger.info("🔄 Открыт в плеере по умолчанию")
                return True
                
        except ImportError:
            logger.warning("⚠️ pygame не установлен, используем fallback")
            # Fallback: открываем в плеере по умолчанию
            import subprocess
            subprocess.Popen(["start", audio_path], shell=True)
            logger.info("🔄 Открыт в плеере по умолчанию")
            return True
            
        except Exception as e:
            logger.error(f"❌ Ошибка воспроизведения аудио: {e}")
            return False

    def text_to_speech(self, text: str, voice: str = "male", language: str = "ru", auto_play: bool = True) -> str:
        """
        Озвучивает текст с помощью gTTS (Google Text-to-Speech)
        
        Args:
            text: Текст для озвучки
            voice: Тип голоса ("male" или "female") - пока не используется в gTTS
            language: Язык текста ("ru", "en", etc.)
            auto_play: Автоматически воспроизвести после создания файла
            
        Returns:
            Путь к созданному аудиофайлу или пустая строка при ошибке
        """
        try:
            from gtts import gTTS
            
            # Создаем папку для сгенерированной речи
            output_dir = os.path.join(os.path.dirname(__file__), "Audio", "generated_speech")
            os.makedirs(output_dir, exist_ok=True)
            
            # Генерируем уникальное имя файла
            timestamp = int(time.time())
            filename = f"tts_{voice}_{language}_{timestamp}.mp3"
            output_path = os.path.join(output_dir, filename)
            
            logger.info(f"🎤 Озвучиваю текст: {text[:100]}...")
            logger.info(f"🔊 Голос: {voice}, Язык: {language}")
            logger.info(f"🌐 Использую Google TTS API")
            
            # Создаем TTS объект
            tts = gTTS(text=text, lang=language, slow=False)
            
            # Сохраняем аудиофайл
            tts.save(output_path)
            
            logger.info(f"✅ Аудиофайл сохранен: {output_path}")
            
            # Автоматически воспроизводим, если включено
            if auto_play:
                self.play_audio_file(output_path)
            
            return output_path
            
        except ImportError:
            logger.error("❌ gTTS не установлен. Установите: pip install gTTS")
            return ""
        except Exception as e:
            logger.error(f"❌ Ошибка озвучки текста: {e}")
            logger.error(f"🔍 Тип ошибки: {type(e).__name__}")
            import traceback
            logger.error(f"📋 Traceback: {traceback.format_exc()}")
            return ""

    def enhance_prompt_with_memory(self, user_message: str, system_prompt: str = "") -> str:
        """
        Улучшает промпт с помощью контекста из памяти
        
        Args:
            user_message: Сообщение пользователя
            system_prompt: Системный промпт
            
        Returns:
            Улучшенный промпт с контекстом
        """
        try:
            # Получаем релевантный контекст
            context = self.get_relevant_context(user_message, max_context_length=1500)
            
            # Получаем предпочтения пользователя
            preferences = self.get_user_preferences(user_message)
            
            # Формируем улучшенный промпт
            enhanced_prompt = system_prompt
            added_content = ""
            
            if context:
                context_section = f"\n\nРЕЛЕВАНТНЫЙ КОНТЕКСТ ИЗ ПРЕДЫДУЩИХ ДИАЛОГОВ:\n{context}\n\nИНСТРУКЦИЯ: Используйте эту информацию только если она релевантна текущему запросу. Не упоминайте источник информации явно."
                enhanced_prompt += context_section
                added_content += context_section
            
            if preferences:
                preferences_section = f"\n\nПРЕДПОЧТЕНИЯ ПОЛЬЗОВАТЕЛЯ:\n{preferences}\n\nИНСТРУКЦИЯ: Учитывайте эти предпочтения при формировании ответа, но не упоминайте их явно."
                enhanced_prompt += preferences_section
                added_content += preferences_section
            
            # Подсчитываем только добавленный контент (без системного промпта и запроса пользователя)
            added_length = len(added_content)
            if added_length > 0:
                logger.info(f"📚 Промпт улучшен с помощью памяти (добавлено: {added_length} символов)")
            else:
                logger.info(f"📚 Промпт не улучшен (память пуста)")
            
            # Логируем общую длину промпта для сравнения с токенами
            total_prompt_length = len(enhanced_prompt)
            logger.info(f"📝 Общая длина промпта: {total_prompt_length} символов")
            
            return enhanced_prompt
            
        except Exception as e:
            logger.error(f"❌ Ошибка улучшения промпта: {e}")
            return f"{system_prompt}\n\nТЕКУЩИЙ ЗАПРОС ПОЛЬЗОВАТЕЛЯ: {user_message}"
    
    def auto_save_conversation(self, user_message: str, ai_response: str, 
                              context: str = "", metadata: Optional[Dict[str, Any]] = None):
        """
        Автоматически сохраняет диалог в память
        
        Args:
            user_message: Сообщение пользователя
            ai_response: Ответ ИИ
            context: Дополнительный контекст
            metadata: Дополнительные метаданные
        """
        try:
            # Добавляем базовые метаданные
            if metadata is None:
                metadata = {}
            
            metadata.update({
                "auto_saved": True,
                "response_length": len(ai_response),
                "user_message_length": len(user_message)
            })
            
            # Сохраняем в память
            success = self.add_to_memory(user_message, ai_response, context, metadata)
            
            if success:
                logger.info("💾 Диалог автоматически сохранен в память")
            else:
                logger.warning("⚠️ Не удалось сохранить диалог в память")
                
        except Exception as e:
            logger.error(f"❌ Ошибка автоматического сохранения: {e}")
    
    def extract_preferences_from_response(self, user_message: str, ai_response: str):
        """
        Извлекает предпочтения пользователя из диалога и сохраняет их
        
        Args:
            user_message: Сообщение пользователя
            ai_response: Ответ ИИ
        """
        try:
            # Простые паттерны для извлечения предпочтений
            preference_patterns = [
                r"мне нравится (.+?)(?:\.|$)",
                r"я предпочитаю (.+?)(?:\.|$)",
                r"лучше всего (.+?)(?:\.|$)",
                r"хочу (.+?)(?:\.|$)",
                r"нужно (.+?)(?:\.|$)",
                r"важно (.+?)(?:\.|$)"
            ]
            
            # Ищем предпочтения в сообщении пользователя
            for pattern in preference_patterns:
                matches = re.findall(pattern, user_message.lower())
                for match in matches:
                    if len(match) > 10:  # Минимальная длина предпочтения
                        self.add_user_preference(
                            match.strip(),
                            category="user_preference",
                            metadata={"source": "extracted", "pattern": pattern}
                        )
            
            # Также ищем в ответе ИИ, если там есть подтверждения
            confirmation_patterns = [
                r"понял(?:а)?, что вам нравится (.+?)(?:\.|$)",
                r"запомню, что вы предпочитаете (.+?)(?:\.|$)",
                r"учту ваше предпочтение (.+?)(?:\.|$)"
            ]
            
            for pattern in confirmation_patterns:
                matches = re.findall(pattern, ai_response.lower())
                for match in matches:
                    if len(match) > 10:
                        self.add_user_preference(
                            match.strip(),
                            category="confirmed_preference",
                            metadata={"source": "ai_confirmation", "pattern": pattern}
                        )
                        
        except Exception as e:
            logger.error(f"❌ Ошибка извлечения предпочтений: {e}")

    ### МЕТОДЫ ДЛЯ РАБОТЫ С CHROMADB ###
    
    def add_to_memory(self, user_message: str, ai_response: str, context: str = "", 
                     metadata: Optional[Dict[str, Any]] = None) -> bool:
        self._ensure_chromadb_initialized()
        """
        Добавляет диалог в векторное хранилище памяти
        
        Args:
            user_message: Сообщение пользователя
            ai_response: Ответ ИИ
            context: Дополнительный контекст
            metadata: Дополнительные метаданные
            
        Returns:
            True если успешно добавлено
        """
        self._ensure_chromadb_initialized()
        if not CHROMADB_ENABLE_MEMORY:
            logger.debug("ChromaDB memory storage is disabled by configuration")
            return False
        if self.chromadb_manager is None:
            logger.error("❌ ChromaDB недоступен")
            return False
        try:
            return self.chromadb_manager.add_conversation_memory(
                user_message, ai_response, context, metadata
            )
        except Exception as e:
            logger.error(f"❌ Ошибка добавления в память: {e}")
            return False
    
    def add_user_preference(self, preference_text: str, category: str = "general", 
                           metadata: Optional[Dict[str, Any]] = None) -> bool:
        """
        Добавляет предпочтение пользователя в векторное хранилище
        
        Args:
            preference_text: Текст предпочтения
            category: Категория предпочтения
            metadata: Дополнительные метаданные
        
        Returns:
            True если успешно добавлено
        """
        if not CHROMADB_ENABLE_MEMORY:
            logger.debug("ChromaDB memory storage is disabled, skipping preference save")
            return False
        self._ensure_chromadb_initialized()
        if self.chromadb_manager is None:
            logger.error("❌ ChromaDB недоступен")
            return False
        try:
            return self.chromadb_manager.add_user_preference(
                preference_text, category, metadata
            )
        except Exception as e:
            logger.error(f"❌ Ошибка добавления предпочтения: {e}")
            return False
    
    def get_relevant_context(self, query: str, max_context_length: int = 2000) -> str:
        """
        Получает релевантный контекст из предыдущих диалогов
        
        Args:
            query: Текущий запрос пользователя
            max_context_length: Максимальная длина контекста
            
        Returns:
            Строка с релевантным контекстом
        """
        self._ensure_chromadb_initialized()
        if self.chromadb_manager is None:
            logger.error("❌ ChromaDB недоступен")
            return ""
        try:
            return self.chromadb_manager.get_conversation_context(query, max_context_length)
        except Exception as e:
            logger.error(f"❌ Ошибка получения контекста: {e}")
            return ""
    
    def get_user_preferences(self, query: Optional[str] = None) -> str:
        """
        Получает предпочтения пользователя
        
        Args:
            query: Контекстный запрос (опционально)
            
        Returns:
            Строка с предпочтениями пользователя
        """
        self._ensure_chromadb_initialized()
        if self.chromadb_manager is None:
            logger.error("❌ ChromaDB недоступен")
            return ""
        try:
            return self.chromadb_manager.get_user_preferences_summary(query)
        except Exception as e:
            logger.error(f"❌ Ошибка получения предпочтений: {e}")
            return ""
    
    def search_similar_conversations(self, query: str, n_results: int = 5, 
                                   similarity_threshold: float = 0.7) -> List[Dict[str, Any]]:
        """
        Ищет похожие диалоги в векторном хранилище
        
        Args:
            query: Поисковый запрос
            n_results: Количество результатов
            similarity_threshold: Порог схожести (0-1)
            
        Returns:
            Список найденных диалогов
        """
        self._ensure_chromadb_initialized()
        if self.chromadb_manager is None:
            logger.error("❌ ChromaDB недоступен")
            return []
        try:
            return self.chromadb_manager.search_similar_conversations(
                query, n_results, similarity_threshold
            )
        except Exception as e:
            logger.error(f"❌ Ошибка поиска похожих диалогов: {e}")
            return []
    
    def cleanup_old_memory(self, days_to_keep: int = 30) -> int:
        """
        Удаляет старые записи из памяти
        
        Args:
            days_to_keep: Количество дней для хранения записей
            
        Returns:
            Количество удаленных записей
        """
        self._ensure_chromadb_initialized()
        if self.chromadb_manager is None:
            logger.error("❌ ChromaDB недоступен")
            return 0
        try:
            return self.chromadb_manager.cleanup_old_records(days_to_keep)
        except Exception as e:
            logger.error(f"❌ Ошибка очистки памяти: {e}")
            return 0
    
    def get_memory_stats(self) -> Dict[str, Any]:
        """
        Получает статистику памяти
        
        Returns:
            Словарь со статистикой
        """
        self._ensure_chromadb_initialized()
        if self.chromadb_manager is None:
            logger.error("❌ ChromaDB недоступен")
            return {"error": "ChromaDB недоступен"}
        try:
            return self.chromadb_manager.get_database_stats()
        except Exception as e:
            logger.error(f"❌ Ошибка получения статистики памяти: {e}")
            return {"error": str(e)}
    
    def get_gpu_info(self) -> Dict[str, Any]:
        """
        Получает информацию о GPU для ChromaDB
        
        Returns:
            Словарь с информацией о GPU
        """
        try:
            self._ensure_chromadb_initialized()
            if self.chromadb_manager is None:
                logger.error("❌ ChromaDB недоступен")
                return {"error": "ChromaDB недоступен"}
            return self.chromadb_manager.get_gpu_info()
        except Exception as e:
            logger.error(f"❌ Ошибка получения информации о GPU: {e}")
            return {"error": str(e)}

# ensure_wav moved to media_processing.py

